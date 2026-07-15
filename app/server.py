from __future__ import annotations

import datetime as dt
import gzip
import hashlib
import importlib.util
import io
import json
import os
import re
import sqlite3
import subprocess
import sys
import threading
import time
import urllib.error
import urllib.request
import uuid
import webbrowser
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from functools import lru_cache
from html import unescape
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path, PurePosixPath
from urllib.parse import parse_qs, quote, quote_plus, urlencode, urljoin, urlparse

try:
    import jwt
except ImportError:  # pragma: no cover - only required when Supabase auth is enabled.
    jwt = None


APP_DIR = Path(__file__).parent.absolute()


def load_env_files() -> None:
    for env_path in [APP_DIR / ".env.local", APP_DIR / ".env"]:
        if not env_path.exists():
            continue
        for line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


load_env_files()


def env_path(name: str, default: Path) -> Path:
    value = os.environ.get(name)
    return Path(value).expanduser() if value else default


PUBLIC_DIR = APP_DIR / "public"
DATA_DIR = env_path("JOB_ASSISTANT_DATA_DIR", APP_DIR / "data")
WORKSPACE_DIR = env_path("JOB_ASSISTANT_WORKSPACE_DIR", APP_DIR / "workspace")
DB_PATH = DATA_DIR / "career_copilot.sqlite"
PROFILE_PATH = DATA_DIR / "profile.json"
USER_CONTEXT_PATH = DATA_DIR / "user_context.json"
APPLY_ASSIST_DIR = DATA_DIR / "apply-assist"
BROWSER_PROFILE_DIR = DATA_DIR / "browser-profile"
RESUME_UPLOAD_DIR = DATA_DIR / "resumes"
DEFAULT_RESUME_PDF = RESUME_UPLOAD_DIR / "active-resume.pdf"
RESUME_PATH = env_path("JOB_ASSISTANT_RESUME", DEFAULT_RESUME_PDF)
REFERENCE_RESUME_DIR = DEFAULT_RESUME_PDF.parent
PROFILE_PHOTO_PATH = REFERENCE_RESUME_DIR / "profile-photo.jpg"

APP_HOST = os.environ.get("JOB_ASSISTANT_HOST", "127.0.0.1")
APP_PORT = int(os.environ.get("PORT") or os.environ.get("JOB_ASSISTANT_PORT", "8787"))
NOTION_VERSION = os.environ.get("NOTION_VERSION", "2022-06-28")
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1-mini")
LOCAL_USER_ID = os.environ.get("JOB_ASSISTANT_LOCAL_USER_ID", "local-owner")
REQUEST_CONTEXT = threading.local()
INITIALIZED_DB_PATHS: set[str] = set()
DB_INIT_LOCK = threading.RLock()
CLOUD_STATE_LOADED: set[str] = set()
CLOUD_STATE_BUCKET_READY: set[str] = set()
CLOUD_STATE_LOCK = threading.RLock()
CLOUD_STATE_USER_LOCKS: dict[str, threading.RLock] = {}

DATE_FMT = "%Y-%m-%d"

STATUS_VALUES = {
    "New",
    "Recommended",
    "Apply Queue",
    "Drafted",
    "Applied",
    "Watch",
    "Dropped",
    "Follow Up",
    "Interview",
    "Rejected",
    "Offer",
    "Closed",
}

AI_RESUME_HEADLINE = "AI-Enabled Product, UX, and Service Design Candidate"
SCAN_THREADS: dict[str, threading.Thread] = {}
SCAN_THREADS_LOCK = threading.Lock()

AI_PROFILE_SIGNAL = (
    "I consistently use AI-assisted research synthesis, prompt-based ideation, scenario exploration, "
    "UX writing iteration, JD/capability matching, and workflow automation to accelerate evidence-led design decisions."
)

NOTION_APPLICATION_STATUSES = {
    "Apply Queue",
    "Drafted",
    "Applied",
    "Follow Up",
    "Interview",
    "Rejected",
    "Offer",
}

WATCH_COMPANIES = [
    ("ByteDance", "Company Site", "https://jobs.bytedance.com/en/position", "Product, design, operations, data, AI"),
    ("TikTok", "Company Site", "https://careers.tiktok.com/", "Product, design, trust and safety, operations"),
    ("Shopee", "Company Site", "https://careers.shopee.sg/", "Product, UX, business, operations"),
    ("Lazada", "Company Site", "https://www.lazada.com/en/careers/", "Product, UX, commercial, operations"),
    ("Sea", "Company Site", "https://www.sea.com/careers", "Graduate, product, corporate, design"),
    ("Grab", "Company Site", "https://www.grab.careers/", "Product, design, analytics, operations"),
    ("GovTech", "Company Site", "https://www.tech.gov.sg/careers/", "Design, product, digital services"),
    ("DBS", "Company Site", "https://www.dbs.com/careers/default.page", "Graduate, innovation, product, UX"),
    ("PDD", "Company Site", "https://careers.pddglobalhr.com/campus/intern", "Internship, product, platform and AI roles; require explicit Singapore location"),
    ("Tencent", "Company Site", "https://careers.tencent.com/en-us/home.html", "Product, design, technology"),
]

REGION_CONFIGS = {
    "SG": {
        "code": "SG",
        "label": "Singapore",
        "default_city": "Singapore",
        "cities": ["Singapore"],
        "default_sources": ["LinkedIn", "InternSG", "Indeed", "JobStreet", "Company Site"],
        "search_location": "Singapore",
        "indeed_host": "sg.indeed.com",
        "daily_copy": "Singapore roles with local eligibility checks.",
    },
    "CN": {
        "code": "CN",
        "label": "China Mainland",
        "default_city": "Shanghai",
        "cities": ["Shanghai", "Beijing", "Shenzhen", "Hangzhou", "Guangzhou"],
        "default_sources": ["LinkedIn", "Mainland Public Search", "Company Site"],
        "search_location": "China",
        "indeed_host": "",
        "daily_copy": "Mainland China roles, city fit first.",
    },
    "HK": {
        "code": "HK",
        "label": "Hong Kong",
        "default_city": "Hong Kong",
        "cities": ["Hong Kong", "Kowloon", "New Territories"],
        "default_sources": ["LinkedIn", "JobsDB", "Company Site"],
        "search_location": "Hong Kong",
        "indeed_host": "hk.indeed.com",
        "daily_copy": "Hong Kong roles, Greater China fit first.",
    },
}

COMPANY_CATALOG = [
    {
        "region": "SG",
        "company": "TikTok",
        "source": "Company Site",
        "url": "https://careers.tiktok.com/",
        "focus": "Product, design, trust and safety, operations",
        "company_type": "Internet",
        "city_tags": ["Singapore"],
        "priority": 95,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "Shopee",
        "source": "Company Site",
        "url": "https://careers.shopee.sg/",
        "focus": "Product, UX, business, operations",
        "company_type": "E-commerce",
        "city_tags": ["Singapore"],
        "priority": 92,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "Grab",
        "source": "Company Site",
        "url": "https://www.grab.careers/",
        "focus": "Product, design, analytics, operations",
        "company_type": "Mobility / Fintech",
        "city_tags": ["Singapore"],
        "priority": 90,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "GovTech",
        "source": "Company Site",
        "url": "https://www.tech.gov.sg/careers/",
        "focus": "Design, product, digital services",
        "company_type": "Public Digital",
        "city_tags": ["Singapore"],
        "priority": 88,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "DBS",
        "source": "Company Site",
        "url": "https://www.dbs.com/careers/default.page",
        "focus": "Graduate, innovation, product, UX",
        "company_type": "Financial Services",
        "city_tags": ["Singapore"],
        "priority": 86,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "Sea",
        "source": "Company Site",
        "url": "https://www.sea.com/careers",
        "focus": "Graduate, product, corporate, design",
        "company_type": "Internet",
        "city_tags": ["Singapore"],
        "priority": 84,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "Lazada",
        "source": "Company Site",
        "url": "https://www.lazada.com/en/careers/",
        "focus": "Product, UX, commercial, operations",
        "company_type": "E-commerce",
        "city_tags": ["Singapore"],
        "priority": 82,
        "default_watch": True,
    },
    {
        "region": "SG",
        "company": "IKEA Singapore",
        "source": "Company Site",
        "url": "https://jobs.ikea.com/en/location/singapore-jobs/22908/1880251/2",
        "focus": "Retail experience, service design, customer operations, visual merchandising",
        "company_type": "Retail / Service Design",
        "city_tags": ["Singapore"],
        "tags": ["本地服务品牌", "体验设计"],
        "language_signal": "English first",
        "recommend_reason": "强服务体验场景，适合把 service design / customer journey 讲成业务价值。",
        "priority": 80,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "MUJI Singapore",
        "source": "Company Site",
        "url": "https://www.muji.com/sg/",
        "focus": "Retail operations, customer experience, visual merchandising, brand service",
        "company_type": "Retail / Lifestyle",
        "city_tags": ["Singapore"],
        "tags": ["本地服务品牌", "日系品牌"],
        "language_signal": "English first",
        "recommend_reason": "适合关注品牌体验、门店服务和生活方式零售相关机会。",
        "priority": 76,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "foodpanda Singapore",
        "source": "Company Site",
        "url": "https://careers.foodpanda.com/singapore",
        "focus": "Product operations, commercial, logistics, marketing, analytics",
        "company_type": "Food Delivery / Local Services",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "运营"],
        "language_signal": "English first",
        "recommend_reason": "本地生活服务场景强，适合 product ops、growth、service experience。",
        "priority": 85,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "POP MART Singapore",
        "source": "Company Site",
        "url": "https://www.popmart.com/sg",
        "focus": "Retail operations, brand, community, IP, visual merchandising",
        "company_type": "Consumer / IP Retail",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "消费品牌"],
        "language_signal": "Chinese-friendly likely",
        "recommend_reason": "中国品牌出海场景，适合讲跨文化用户、IP 社群和零售体验。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Changi Airport Group",
        "source": "Company Site",
        "url": "https://www.changiairport.com/en/careers/job-opportunities.html",
        "focus": "Internships, customer experience, digital product, service operations",
        "company_type": "Travel / Service Experience",
        "city_tags": ["Singapore"],
        "tags": ["本地标杆", "实习"],
        "language_signal": "English first",
        "recommend_reason": "服务设计和体验运营场景非常强，适合用作品集讲 journey / touchpoint。",
        "priority": 87,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "ByteDance",
        "source": "Company Site",
        "url": "https://jobs.bytedance.com/en/position",
        "focus": "AI product, design, content platform, operations, trust and safety",
        "company_type": "Internet / AI",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "AI"],
        "language_signal": "Chinese-friendly likely",
        "recommend_reason": "新加坡岗位多，中文和中国互联网背景可能是加分项。",
        "priority": 94,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Tencent Singapore",
        "source": "Company Site",
        "url": "https://careers.tencent.com/en-us/home.html",
        "focus": "Product, design, gaming, cloud, international operations",
        "company_type": "Internet / Gaming",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "产品"],
        "language_signal": "Chinese-friendly likely",
        "recommend_reason": "适合关注游戏、云服务和国际产品运营方向。",
        "priority": 81,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Trip.com Group",
        "source": "Company Site",
        "url": "https://careers.trip.com/",
        "focus": "Travel product, customer experience, operations, content, growth",
        "company_type": "Travel Tech",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "出海"],
        "language_signal": "Chinese-friendly likely",
        "recommend_reason": "旅游平台和跨境用户场景适合讲多语言体验、服务流程和增长。",
        "priority": 79,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Ant International",
        "source": "Company Site",
        "url": "https://www.ant-intl.com/en/job-search",
        "focus": "Payments, product operations, growth, risk, customer experience",
        "company_type": "Fintech / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "Fintech", "出海"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "支付和跨境商业场景强，适合产品运营、增长和风控相关实习；中文背景可能是加分项。",
        "priority": 86,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Alibaba Cloud Singapore",
        "source": "Company Site",
        "url": "https://www.alibabagroup.com/careers",
        "focus": "Cloud, AI, product operations, ecosystem marketing, customer success",
        "company_type": "Cloud / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "AI", "出海"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "云与企业服务出海场景，适合讲 AI 产品、B2B 用户和区域市场理解。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Lark",
        "source": "Company Site",
        "url": "https://www.larksuite.com/site/job",
        "focus": "B2B SaaS, product, customer success, UX research, APAC growth",
        "company_type": "B2B SaaS / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "B2B SaaS", "产品"],
        "language_signal": "Chinese-friendly likely",
        "recommend_reason": "协作工具和 APAC 客户场景清晰，适合 AI product、UX research 和 customer success。",
        "priority": 85,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Huawei Singapore",
        "source": "Company Site",
        "url": "https://career.huawei.com/reccampportal/portal5/index.html",
        "focus": "Cloud, AI, telecom, product marketing, solution operations",
        "company_type": "Telecom / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "AI", "本地企业客户"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "企业科技和区域业务场景多，适合关注云、AI、解决方案和产品运营岗位。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Xiaomi Singapore",
        "source": "Company Site",
        "url": "https://career.mi.com/",
        "focus": "AIoT, e-commerce, marketing, product operations, retail experience",
        "company_type": "Consumer Tech / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "消费科技", "出海"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "消费科技和 AIoT 出海场景，适合产品运营、市场和用户体验相关机会。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "OPPO Singapore",
        "source": "Company Site",
        "url": "https://careers.oppo.com/",
        "focus": "Consumer technology, retail operations, marketing, product experience",
        "company_type": "Consumer Tech / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "消费科技"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "适合关注手机生态、门店体验、市场和区域运营相关机会。",
        "priority": 75,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "vivo Singapore",
        "source": "Company Site",
        "url": "https://www.vivo.com/sg/about-vivo/career/jobs",
        "focus": "Consumer technology, marketing, retail operations, product experience",
        "company_type": "Consumer Tech / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "消费科技"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "手机和消费电子场景，适合市场、运营、用户体验和渠道相关机会。",
        "priority": 74,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "SHEIN Singapore",
        "source": "Company Site",
        "url": "https://careers.shein.com/",
        "focus": "E-commerce, merchandising, logistics, growth, content, operations",
        "company_type": "E-commerce / Greater China",
        "city_tags": ["Singapore"],
        "tags": ["中文友好概率较高", "电商", "出海"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "跨境电商和供应链场景强，适合增长、内容、运营和体验相关岗位。",
        "priority": 80,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Singtel",
        "source": "Company Site",
        "url": "https://www.singtel.com/about-us/careers",
        "focus": "Internships, digital product, data, AI, customer experience, NCS/Nxera ecosystem",
        "company_type": "Telco / Local Anchor",
        "city_tags": ["Singapore"],
        "tags": ["本地标杆", "实习", "AI"],
        "language_signal": "English first",
        "recommend_reason": "新加坡本地大型雇主，实习、数据、数字化和 NCS 生态机会较多。",
        "priority": 86,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "PatSnap",
        "source": "Company Site",
        "url": "https://www.patsnap.com/careers/",
        "focus": "AI product, UX, research, data, product engineering",
        "company_type": "AI / IP Intelligence",
        "city_tags": ["Singapore"],
        "tags": ["AI", "设计与产品"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "AI + 知识工作场景强，和 AI product / UX research 匹配度高。",
        "priority": 86,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Hypotenuse AI",
        "source": "Company Site",
        "url": "https://careers.hypotenuse.ai/",
        "focus": "Generative AI product, content workflow, growth, design",
        "company_type": "AI Startup",
        "city_tags": ["Singapore"],
        "tags": ["AI", "高潜力初创"],
        "language_signal": "English first",
        "recommend_reason": "生成式 AI 产品场景直接，适合 AI product 和内容工作流经验。",
        "priority": 83,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "WIZ.AI",
        "source": "Company Site",
        "url": "https://www.wiz.ai/pages/join-us.html",
        "focus": "Conversational AI, product, customer success, operations",
        "company_type": "AI Startup",
        "city_tags": ["Singapore"],
        "tags": ["AI", "中文友好概率较高"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "对话式 AI 和企业服务场景，适合讲 AI workflow 与用户研究。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "ADVANCE.AI",
        "source": "Company Site",
        "url": "https://advanceai.tech/about-us/careers/",
        "focus": "AI, fintech, risk products, data operations, product management",
        "company_type": "AI / Fintech",
        "city_tags": ["Singapore"],
        "tags": ["AI", "Fintech"],
        "language_signal": "English first",
        "recommend_reason": "AI 风控和金融科技场景，适合产品、运营和数据相关实习。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "AI Singapore",
        "source": "Company Site",
        "url": "https://aisingapore.org/home/careers/",
        "focus": "AI programmes, research translation, product, education, community",
        "company_type": "AI / Public Programme",
        "city_tags": ["Singapore"],
        "tags": ["AI", "本地标杆"],
        "language_signal": "English first",
        "recommend_reason": "本地 AI 生态核心组织，适合关注 AI 产品、项目和研究转化机会。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "AiChat",
        "source": "Company Site",
        "url": "https://www.aichat.com/careers/",
        "focus": "Conversational commerce, chatbot product, customer success, marketing",
        "company_type": "AI Startup",
        "city_tags": ["Singapore"],
        "tags": ["AI", "产品运营"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "AI chatbot 与商业转化场景，适合 AI operations / product ops。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Horizon Labs",
        "source": "Company Site",
        "url": "https://www.horizon-labs.co/",
        "focus": "AI product, software, automation, startup generalist roles",
        "company_type": "AI Startup",
        "city_tags": ["Singapore"],
        "tags": ["AI", "高潜力初创"],
        "language_signal": "English first",
        "recommend_reason": "适合关注小团队 AI 产品和泛产品/运营机会。",
        "priority": 77,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Fabrica Robotics",
        "source": "Company Site",
        "url": "https://www.fabricarobotics.com/careers",
        "focus": "Robotics, product, operations, hardware-software experience",
        "company_type": "Robotics Startup",
        "city_tags": ["Singapore"],
        "tags": ["AI", "硬件体验"],
        "language_signal": "English first",
        "recommend_reason": "机器人和实体体验场景，适合把交互、原型和服务流程结合起来讲。",
        "priority": 76,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "k-ID",
        "source": "Company Site",
        "url": "https://www.k-id.com/careers",
        "focus": "Trust and safety, product, youth digital experience, policy operations",
        "company_type": "Safety Tech Startup",
        "city_tags": ["Singapore"],
        "tags": ["高潜力初创", "产品"],
        "language_signal": "English first",
        "recommend_reason": "信任安全和年轻用户体验场景，适合 UX research / product ops。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "BeeX",
        "source": "Company Site",
        "url": "https://www.beex.sg/careers",
        "focus": "Robotics, maritime, product operations, design-adjacent systems",
        "company_type": "Robotics Startup",
        "city_tags": ["Singapore"],
        "tags": ["高潜力初创", "硬件体验"],
        "language_signal": "English first",
        "recommend_reason": "适合关注机器人、海事科技和复杂系统体验相关机会。",
        "priority": 75,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Flowmingo AI",
        "source": "Company Site",
        "url": "https://flowmingo.ai/careers",
        "focus": "AI hiring platform, product UX, growth, marketing, operations",
        "company_type": "AI Startup",
        "city_tags": ["Singapore", "Remote"],
        "tags": ["AI", "高潜力初创", "实习"],
        "language_signal": "English first",
        "recommend_reason": "AI 招聘产品和候选人体验场景直接，适合讲 AI-native UX、增长和招聘流程自动化。",
        "priority": 89,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "X0PA AI",
        "source": "Company Site",
        "url": "https://x0pa.com/career/",
        "focus": "Responsible AI hiring, product, customer success, data, HR tech",
        "company_type": "AI / HR Tech",
        "city_tags": ["Singapore"],
        "tags": ["AI", "B2B SaaS"],
        "language_signal": "English first",
        "recommend_reason": "AI 招聘与评估场景清晰，适合产品运营、UX research 和 B2B AI 方向。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "ViSenze",
        "source": "Company Site",
        "url": "https://apply.workable.com/visenze/?lng=en",
        "focus": "Visual AI, commerce search, recommendation, product, data",
        "company_type": "AI / Commerce Search",
        "city_tags": ["Singapore"],
        "tags": ["AI", "电商"],
        "language_signal": "English first",
        "recommend_reason": "视觉搜索和推荐系统场景强，适合 AI product、UX 和数据相关岗位。",
        "priority": 83,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "PixCap",
        "source": "Company Site",
        "url": "https://jobs.surgeahead.com/jobs/pixcap",
        "focus": "3D design tools, AI design workflow, product, frontend, creative tooling",
        "company_type": "Design / AI Tooling",
        "city_tags": ["Singapore", "Remote"],
        "tags": ["设计与产品", "AI", "高潜力初创"],
        "language_signal": "English first",
        "recommend_reason": "3D 与 AI 创作工具，和产品设计、前端体验、创意工具方向贴合。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "ShopBack",
        "source": "Company Site",
        "url": "https://jobs.lever.co/shopback-2",
        "focus": "Consumer growth, rewards, payments, data, marketing, product",
        "company_type": "Consumer / Fintech",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "实习", "增长"],
        "language_signal": "English first",
        "recommend_reason": "新加坡成长起来的消费平台，数据、增长、产品和实习岗位较多。",
        "priority": 88,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Syfe",
        "source": "Company Site",
        "url": "https://syfe.careers-page.com/",
        "focus": "WealthTech, operations, product, investment, analytics, internships",
        "company_type": "WealthTech",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "实习"],
        "language_signal": "English first",
        "recommend_reason": "数字财富管理产品，适合运营、分析、产品和金融科技实习方向。",
        "priority": 83,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Endowus",
        "source": "Company Site",
        "url": "https://endowus.com/careers",
        "focus": "Wealth management, product, operations, customer experience, content",
        "company_type": "WealthTech",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "本地标杆"],
        "language_signal": "English first",
        "recommend_reason": "财富科技和复杂服务体验强，适合产品、内容、运营和客户旅程方向。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "StashAway",
        "source": "Company Site",
        "url": "https://www.stashaway.sg/careers",
        "focus": "Digital wealth, product, growth, client experience, analytics",
        "company_type": "WealthTech",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "增长"],
        "language_signal": "English first",
        "recommend_reason": "数字投顾和用户教育场景清晰，适合 UX、内容、增长和产品运营。",
        "priority": 81,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Carousell Group",
        "source": "Company Site",
        "url": "https://careers.smartrecruiters.com/carousellgroup",
        "focus": "Marketplace, trust and safety, consumer product, data, operations",
        "company_type": "Marketplace",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "产品"],
        "language_signal": "English first",
        "recommend_reason": "本地 marketplace 场景丰富，适合讲交易体验、信任安全和用户增长。",
        "priority": 85,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Airwallex",
        "source": "Company Site",
        "url": "https://jobs.ashbyhq.com/airwallex",
        "focus": "Fintech infrastructure, payments, risk, product, operations, AI enablement",
        "company_type": "Fintech / Payments",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "产品", "AI"],
        "language_signal": "English first",
        "recommend_reason": "跨境支付和金融基础设施岗位多，产品、风控、运营和数据方向都可看。",
        "priority": 86,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Trust Bank",
        "source": "Company Site",
        "url": "https://trustbank.sg/careers/",
        "focus": "Digital banking, product, customer experience, operations, marketing",
        "company_type": "Digital Bank",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "本地标杆"],
        "language_signal": "English first",
        "recommend_reason": "数字银行服务体验完整，适合产品、运营、客户旅程和品牌增长方向。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Nansen",
        "source": "Company Site",
        "url": "https://job-boards.greenhouse.io/nansen",
        "focus": "Onchain analytics, AI/data product, research, product design, growth",
        "company_type": "Crypto / Data Intelligence",
        "city_tags": ["Singapore", "Remote"],
        "tags": ["AI", "数据", "高潜力初创"],
        "language_signal": "English first",
        "recommend_reason": "链上数据和 AI 情报产品场景强，适合数据产品、研究和增长方向。",
        "priority": 81,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Transcelestial",
        "source": "Company Site",
        "url": "https://transcelestial.com/careers/",
        "focus": "Space laser communications, product operations, systems, hardware-software experience",
        "company_type": "DeepTech",
        "city_tags": ["Singapore"],
        "tags": ["高潜力初创", "硬件体验"],
        "language_signal": "English first",
        "recommend_reason": "深科技和复杂系统场景，适合讲系统体验、运营流程和跨团队协作。",
        "priority": 77,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Una Brands",
        "source": "Company Site",
        "url": "https://www.una-brands.com/careers",
        "focus": "E-commerce brands, growth, operations, corporate development, marketplace",
        "company_type": "E-commerce Aggregator",
        "city_tags": ["Singapore"],
        "tags": ["出海", "电商", "中文友好概率较高"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "电商品牌运营和 APAC 场景，适合增长、运营、内容和跨境业务。",
        "priority": 76,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Sleek",
        "source": "Company Site",
        "url": "https://apply.workable.com/careers-at-sleek/?lng=en",
        "focus": "Business automation, fintech, operations, growth, AI-native content",
        "company_type": "B2B SaaS / Fintech",
        "city_tags": ["Singapore", "Remote"],
        "tags": ["Fintech", "B2B SaaS", "实习"],
        "language_signal": "English first",
        "recommend_reason": "中小企业自动化和 fintech 场景，适合产品运营、内容、增长和实习方向。",
        "priority": 80,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Aspire",
        "source": "Company Site",
        "url": "https://aspireapp.com/careers",
        "focus": "B2B fintech, spend management, operations, product, partnerships",
        "company_type": "Fintech / B2B",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "高潜力初创"],
        "language_signal": "English first",
        "recommend_reason": "B2B 金融操作系统，适合产品、运营、策略和客户旅程方向。",
        "priority": 81,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Circles",
        "source": "Company Site",
        "url": "https://circles.co/careers",
        "focus": "Digital telco, customer experience, product, growth, operations",
        "company_type": "Digital Telco",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "产品"],
        "language_signal": "English first",
        "recommend_reason": "数字电信和订阅体验场景，适合产品运营、服务体验和增长方向。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Doctor Anywhere",
        "source": "Company Site",
        "url": "https://careers.smartrecruiters.com/DoctorAnywhere",
        "focus": "HealthTech, data, product, operations, patient experience",
        "company_type": "HealthTech",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "服务体验"],
        "language_signal": "English first",
        "recommend_reason": "医疗服务和数字健康场景，适合服务设计、运营和数据体验方向。",
        "priority": 79,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Razer",
        "source": "Company Site",
        "url": "https://razer.wd3.myworkdayjobs.com/Careers",
        "focus": "Gaming hardware, software, product, data, marketing, retail experience",
        "company_type": "Gaming / Consumer Tech",
        "city_tags": ["Singapore"],
        "tags": ["产品", "消费科技"],
        "language_signal": "English first",
        "recommend_reason": "游戏硬件和生态产品丰富，适合产品、数据、市场和体验相关岗位。",
        "priority": 80,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Canva",
        "source": "Company Site",
        "url": "https://www.lifeatcanva.com/en/",
        "focus": "Design tools, content, product design, research, marketing, AI creative workflow",
        "company_type": "Design Platform",
        "city_tags": ["Singapore", "Remote"],
        "tags": ["设计与产品", "AI"],
        "language_signal": "English first",
        "recommend_reason": "设计工具和 AI 创作工作流高度贴合 UX/product design 和内容方向。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Wise",
        "source": "Company Site",
        "url": "https://wise.jobs/",
        "focus": "Payments, content design, product operations, analytics, compliance experience",
        "company_type": "Fintech / Payments",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "内容设计"],
        "language_signal": "English first",
        "recommend_reason": "跨境金融产品复杂度高，适合 content design、运营、合规体验和产品方向。",
        "priority": 83,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Crypto.com",
        "source": "Company Site",
        "url": "https://jobs.lever.co/crypto",
        "focus": "Web3, security, product, operations, growth, compliance",
        "company_type": "Crypto / Fintech",
        "city_tags": ["Singapore", "Hong Kong"],
        "tags": ["Fintech", "Web3"],
        "language_signal": "English first",
        "recommend_reason": "Web3 金融和安全运营岗位多，适合产品、运营、合规和增长方向。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "NodeFlair",
        "source": "Company Site",
        "url": "https://nodeflair.com/careers",
        "focus": "Tech jobs platform, product, growth, content, community, data",
        "company_type": "Career Tech",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "产品", "增长"],
        "language_signal": "English first",
        "recommend_reason": "技术招聘和薪资数据产品，适合产品运营、内容增长和求职生态相关经验。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Funding Societies",
        "source": "Company Site",
        "url": "https://apply.workable.com/fundingsocieties/",
        "focus": "SME fintech, lending, operations, risk, customer experience, product",
        "company_type": "Fintech / SME Lending",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "本地平台"],
        "language_signal": "English first",
        "recommend_reason": "东南亚 SME 金融场景清晰，适合产品、运营、风控和客户旅程方向。",
        "priority": 83,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "YouTrip",
        "source": "Company Site",
        "url": "https://apply.workable.com/youtrip/?lng=en",
        "focus": "Travel fintech, multi-currency wallet, product, growth, operations, marketing",
        "company_type": "Fintech / Travel",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "实习", "增长"],
        "language_signal": "English first",
        "recommend_reason": "消费金融和旅行场景强，适合产品运营、增长、数据和市场实习方向。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "PropertyGuru Group",
        "source": "Company Site",
        "url": "https://propertyguru.wd105.myworkdayjobs.com/PropertyGuru",
        "focus": "PropTech, marketplace, AI transformation, product, UX, fulfilment strategy",
        "company_type": "PropTech / Marketplace",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "产品", "AI"],
        "language_signal": "English first",
        "recommend_reason": "本地房产平台和 marketplace 场景扎实，适合讲搜索、推荐、交易和 AI 转型。",
        "priority": 85,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Ninja Van",
        "source": "Company Site",
        "url": "https://www.ninjavan.co/en-sg/company/careers",
        "focus": "Logistics, business operations, product ops, data, service recovery",
        "company_type": "Logistics / E-commerce Infrastructure",
        "city_tags": ["Singapore"],
        "tags": ["本地平台", "运营", "服务体验"],
        "language_signal": "English first",
        "recommend_reason": "物流和履约体验复杂，适合 product ops、service design、数据运营方向。",
        "priority": 81,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "bolttech",
        "source": "Company Site",
        "url": "https://bolttech.io/careers/",
        "focus": "Insurtech, platform partnerships, product, operations, customer experience",
        "company_type": "Insurtech",
        "city_tags": ["Singapore", "Hong Kong"],
        "tags": ["Fintech", "平台"],
        "language_signal": "English first",
        "recommend_reason": "保险科技和平台合作场景，适合产品、商业运营和客户体验相关岗位。",
        "priority": 79,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Accredify",
        "source": "Company Site",
        "url": "https://www.accredify.io/careers",
        "focus": "Verifiable credentials, trust infrastructure, product, customer success, operations",
        "company_type": "TrustTech / SaaS",
        "city_tags": ["Singapore"],
        "tags": ["B2B SaaS", "本地初创"],
        "language_signal": "English first",
        "recommend_reason": "数字凭证和信任基础设施场景清楚，适合 B2B 产品、运营和客户成功。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "EPOS",
        "source": "Company Site",
        "url": "https://www.epos.com.sg/careers-at-epos/",
        "focus": "POS SaaS, AI product, product design, merchant operations, growth",
        "company_type": "B2B SaaS / Local Commerce",
        "city_tags": ["Singapore"],
        "tags": ["AI", "产品", "实习"],
        "language_signal": "English first",
        "recommend_reason": "本地商户 SaaS 和 AI 产品岗位明确，适合 AI product、产品设计和运营实习。",
        "priority": 83,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Ebb & Flow Group",
        "source": "Company Site",
        "url": "https://www.ebbflowgroup.com/careers",
        "focus": "F&B brands, guest experience, brand marketing, operations, service design",
        "company_type": "F&B / Experience Brand",
        "city_tags": ["Singapore"],
        "tags": ["服务体验", "品牌增长"],
        "language_signal": "English first",
        "recommend_reason": "如果想做服务体验、品牌和内容增长，这类线下体验品牌很适合练案例。",
        "priority": 74,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "NCS",
        "source": "Company Site",
        "url": "https://www.ncs.co/careers/",
        "focus": "AI services, public digital, consulting, data, product delivery, UX",
        "company_type": "Tech Services / AI",
        "city_tags": ["Singapore"],
        "tags": ["AI", "本地标杆", "公共数字"],
        "language_signal": "English first",
        "recommend_reason": "AI 与公共数字服务岗位多，适合关注咨询、产品交付、数据和体验设计。",
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Moomoo Singapore",
        "source": "Company Site",
        "url": "https://apply.workable.com/moomoo/",
        "focus": "Investment app, fintech, product, marketing, customer operations, data",
        "company_type": "Fintech / Brokerage",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "中文友好概率较高"],
        "language_signal": "Chinese-friendly possible",
        "recommend_reason": "中资背景金融科技产品，适合产品运营、内容、市场和用户增长方向。",
        "priority": 82,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "Holmusk",
        "source": "Company Site",
        "url": "https://www.holmusk.com/careers",
        "focus": "Health data, AI analytics, research, product, clinical operations",
        "company_type": "HealthTech / AI Data",
        "city_tags": ["Singapore", "Remote"],
        "tags": ["AI", "HealthTech", "数据"],
        "language_signal": "English first",
        "recommend_reason": "医疗数据和 AI 分析场景，适合研究、数据产品和复杂行业 UX。",
        "priority": 79,
        "default_watch": False,
    },
    {
        "region": "SG",
        "company": "M-DAQ",
        "source": "Company Site",
        "url": "https://www.m-daq.com/careers",
        "focus": "Cross-border payments, FX, product, compliance, client experience",
        "company_type": "Fintech / Payments",
        "city_tags": ["Singapore"],
        "tags": ["Fintech", "B2B"],
        "language_signal": "English first",
        "recommend_reason": "跨境支付和外汇产品复杂度高，适合产品、客户体验、合规和运营方向。",
        "priority": 78,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "ByteDance",
        "source": "Company Site",
        "url": "https://jobs.bytedance.com/en/position",
        "focus": "AI product, design, operations, content platform",
        "company_type": "Internet / AI",
        "city_tags": ["Shanghai", "Beijing", "Shenzhen", "Hangzhou"],
        "priority": 96,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "Alibaba",
        "source": "Company Site",
        "url": "https://talent.alibaba.com/",
        "focus": "Product, design, commerce, cloud, AI",
        "company_type": "Internet / Cloud",
        "city_tags": ["Hangzhou", "Shanghai", "Beijing", "Shenzhen"],
        "priority": 94,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "Tencent",
        "source": "Company Site",
        "url": "https://careers.tencent.com/en-us/home.html",
        "focus": "Product, design, gaming, social, cloud",
        "company_type": "Internet",
        "city_tags": ["Shenzhen", "Shanghai", "Beijing", "Guangzhou"],
        "priority": 93,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "Xiaohongshu",
        "source": "Company Site",
        "url": "https://job.xiaohongshu.com/",
        "focus": "Product, UX, content, growth, community",
        "company_type": "Consumer Tech",
        "city_tags": ["Shanghai", "Beijing"],
        "priority": 91,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "Meituan",
        "source": "Company Site",
        "url": "https://zhaopin.meituan.com/",
        "focus": "Product operations, service design, growth, local services",
        "company_type": "Local Services",
        "city_tags": ["Beijing", "Shanghai", "Shenzhen"],
        "priority": 89,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "Baidu",
        "source": "Company Site",
        "url": "https://talent.baidu.com/",
        "focus": "AI product, search, cloud, UX research",
        "company_type": "AI / Search",
        "city_tags": ["Beijing", "Shanghai", "Shenzhen"],
        "priority": 86,
        "default_watch": False,
    },
    {
        "region": "CN",
        "company": "Huawei",
        "source": "Company Site",
        "url": "https://career.huawei.com/",
        "focus": "Product, UX, service, cloud, device ecosystem",
        "company_type": "Technology",
        "city_tags": ["Shenzhen", "Shanghai", "Beijing", "Hangzhou"],
        "priority": 84,
        "default_watch": False,
    },
    {
        "region": "HK",
        "company": "HSBC",
        "source": "Company Site",
        "url": "https://www.hsbc.com/careers",
        "focus": "Graduate, digital product, service design, fintech",
        "company_type": "Financial Services",
        "city_tags": ["Hong Kong"],
        "priority": 95,
        "default_watch": False,
    },
    {
        "region": "HK",
        "company": "HKEX",
        "source": "Company Site",
        "url": "https://www.hkexgroup.com/Careers",
        "focus": "Digital product, data, market operations",
        "company_type": "Exchange / Finance",
        "city_tags": ["Hong Kong"],
        "priority": 92,
        "default_watch": False,
    },
    {
        "region": "HK",
        "company": "Cathay",
        "source": "Company Site",
        "url": "https://careers.cathaypacific.com/",
        "focus": "Customer experience, service design, digital product",
        "company_type": "Travel / Service",
        "city_tags": ["Hong Kong"],
        "priority": 90,
        "default_watch": False,
    },
    {
        "region": "HK",
        "company": "WeLab",
        "source": "Company Site",
        "url": "https://www.welab.co/en/careers/",
        "focus": "Fintech product, UX, growth, operations",
        "company_type": "Fintech",
        "city_tags": ["Hong Kong"],
        "priority": 88,
        "default_watch": False,
    },
    {
        "region": "HK",
        "company": "Mox",
        "source": "Company Site",
        "url": "https://mox.com/careers/",
        "focus": "Digital banking, product, design, operations",
        "company_type": "Fintech",
        "city_tags": ["Hong Kong"],
        "priority": 86,
        "default_watch": False,
    },
    {
        "region": "HK",
        "company": "Animoca Brands",
        "source": "Company Site",
        "url": "https://www.animocabrands.com/careers",
        "focus": "Product, growth, digital experiences, Web3",
        "company_type": "Digital Entertainment",
        "city_tags": ["Hong Kong"],
        "priority": 82,
        "default_watch": False,
    },
]

TARGET_QUERIES = [
    "ai product intern",
    "ai ux intern",
    "generative ai intern",
    "ux design intern",
    "internship with conversion Singapore",
    "graduate internship Singapore sponsorship",
    "Mandarin Chinese product intern Singapore",
    "product design intern",
    "service design intern",
    "user research intern",
    "product management intern",
    "product operations intern",
    "marketing intern",
    "content intern",
    "design graduate",
    "product analyst intern",
]

AI_TARGET_QUERIES = [
    "ai product intern",
    "ai product manager intern",
    "ai ux intern",
    "ai user research intern",
    "generative ai intern",
    "ai operations intern",
    "ai project intern",
    "ai data operations intern",
    "machine learning product intern",
    "automation intern",
]

STARTUP_OPPORTUNITY_QUERIES = [
    "startup intern",
    "ai startup intern",
    "venture builder intern",
    "founders office intern",
    "startup innovation intern",
]

STARTUP_OPPORTUNITY_SIGNALS = [
    "startup",
    "start-up",
    "venture",
    "founder",
    "innovation",
    "sginnovate",
    "ai",
    "artificial intelligence",
    "automation",
    "agentic",
    "fintech",
    "incubator",
    "accelerator",
]

SOURCE_LIMITS = {
    "LinkedIn": 24,
    "InternSG": 24,
    "Indeed": 12,
    "JobStreet": 24,
    "Watched Companies": 24,
    "Company Site": 80,
    "Google Jobs": 24,
    "MyCareersFuture": 18,
    "Careers@Gov": 30,
    "Internship.sg": 24,
    "Cultjobs": 18,
    "Startup Channels": 18,
    "AI Startup ATS": 60,
    "LinkedIn AI": 18,
    "InternSG AI": 12,
    "Indeed AI": 8,
    "JobStreet AI": 12,
}

SG_AI_STARTUP_ATS_BOARDS = [
    ("Temus", "https://job-boards.greenhouse.io/temus", "AI, product, UX and digital transformation internships"),
    ("Workato", "https://job-boards.greenhouse.io/workato", "AI integration, product analytics and marketing internships"),
    ("k-ID", "https://jobs.ashbyhq.com/k-ID", "Product, HCI, UX, engineering and operations internships"),
    ("Motion Ventures", "https://job-boards.greenhouse.io/motionventures", "Venture, startup ecosystem and strategy internships"),
    ("Menlo", "https://jobs.ashbyhq.com/menlo", "Applied AI, product engineering and research fellowships"),
    ("Manus AI", "https://jobs.ashbyhq.com/manusai", "AI product, growth and go-to-market internships"),
    ("Simular", "https://jobs.ashbyhq.com/Simular", "AI startup operations and product internships"),
    ("Plaud", "https://jobs.ashbyhq.com/Plaud", "AI hardware, product and global growth internships"),
    ("Bifrost AI", "https://jobs.ashbyhq.com/Bifrost", "Physical AI, product and strategy internships"),
    ("Dexmate", "https://jobs.ashbyhq.com/dexmate", "Physical AI and robotics internships"),
    ("Flagright", "https://jobs.ashbyhq.com/flagright.com", "Fintech AI, product and marketing internships"),
    ("Venti Technologies", "https://jobs.ashbyhq.com/GoVenti", "Autonomous systems, AI and operations internships"),
    ("Coinhako", "https://jobs.ashbyhq.com/Coinhako", "Singapore fintech and operations internships"),
    ("ShopBack", "https://jobs.lever.co/shopback-2", "Data, product, growth and commerce internships"),
    ("Carousell Group", "https://careers.smartrecruiters.com/carousellgroup", "Marketplace, operations, product and graduate opportunities"),
    ("YouTrip", "https://apply.workable.com/youtrip", "Fintech product, community, growth and business internships"),
    ("StraitsX", "https://job-boards.greenhouse.io/straitsx", "Data and AI, product and fintech engineering internships"),
    ("MoneySmart", "https://job-boards.greenhouse.io/moneysmart", "Product design, UX and fintech internships"),
    ("Razer", "https://razer.wd3.myworkdayjobs.com/Careers", "AI, gaming, community, product and operations internships"),
    ("Circles", "https://circles.wd103.myworkdayjobs.com/en-US/Circles", "AI product, digital telco, growth and strategy early-career roles"),
    ("Wise", "https://careers.smartrecruiters.com/Wise", "Fintech analytics, product, operations and marketing internships"),
    ("Bosch Singapore", "https://careers.smartrecruiters.com/BoschGroup", "AI, connected services, product and digital transformation internships"),
    ("We. Singapore", "https://job-boards.greenhouse.io/wesingapore", "Innovation, AI-enabled content, design and communications internships"),
    ("Carta", "https://job-boards.greenhouse.io/carta", "Fintech, startup strategy, product and operations internships"),
    ("Marshall Wace", "https://job-boards.greenhouse.io/mwinternshipprogram", "Technology, AI and product internships with graduate conversion paths"),
]

EMPLOYMENT_PRIORITY_VALUES = {"internship", "full_time", "both", "unspecified"}
EMPLOYMENT_TYPE_VALUES = {"Internship", "Full-time", "Graduate", "Contract", "Unknown"}
SALARY_PERIODS = {"monthly", "yearly", "daily", "hourly", "unknown"}
LIMITED_SCAN_SOURCES = {"Indeed", "JobStreet"}
REGION_CURRENCIES = {"SG": "SGD", "CN": "CNY", "HK": "HKD"}
MIN_RESUME_DIRECTION_SCORE = 0.24
WORK_AUTH_OPTIONS = {
    "SG": [
        {"value": "Student Pass", "label": "Student Pass"},
        {"value": "Singapore work eligibility to be confirmed", "label": "待确认"},
        {"value": "EP/S Pass sponsorship needed", "label": "需要工签支持"},
        {"value": "Citizen / PR", "label": "Citizen / PR"},
    ],
    "CN": [
        {"value": "China mainland work eligibility to be confirmed", "label": "待确认"},
        {"value": "Mainland China citizen", "label": "中国大陆身份"},
        {"value": "Work permit / visa support needed", "label": "需要工作许可支持"},
    ],
    "HK": [
        {"value": "Hong Kong work eligibility to be confirmed", "label": "待确认"},
        {"value": "IANG / student visa", "label": "IANG / 学生签"},
        {"value": "Work visa sponsorship needed", "label": "需要工签支持"},
        {"value": "HK permanent resident", "label": "HK PR"},
    ],
}
JOB_TYPE_OPTIONS = [
    {"value": "Internship", "label": "Internship"},
    {"value": "Graduate", "label": "Graduate"},
    {"value": "Full-time", "label": "Full-time"},
    {"value": "Contract", "label": "Contract"},
]
EMPLOYMENT_PRIORITY_OPTIONS = [
    {"value": "internship", "label": "实习为主"},
    {"value": "full_time", "label": "正式工为主"},
    {"value": "both", "label": "都考虑"},
    {"value": "unspecified", "label": "暂不确定"},
]
CAREER_GOAL_VALUES = {"sg_internship_to_fulltime", "experience_first", "full_time_sg", "explore"}
PRIORITY_LEVEL_VALUES = {"high", "medium", "low", "unspecified"}
LANGUAGE_PREFERENCE_VALUES = {"chinese_friendly", "bilingual", "english_ok", "unspecified"}
COMPANY_GROUP_VALUES = {
    "greater_china",
    "sg_anchor",
    "ai_startup",
    "product_design",
    "fintech",
    "service_brand",
}
CAREER_GOAL_OPTIONS = [
    {"value": "sg_internship_to_fulltime", "label": "实习到留新加坡"},
    {"value": "experience_first", "label": "先累积经验"},
    {"value": "full_time_sg", "label": "直接找正式工"},
    {"value": "explore", "label": "先探索"},
]
PRIORITY_LEVEL_OPTIONS = [
    {"value": "high", "label": "优先"},
    {"value": "medium", "label": "加权"},
    {"value": "low", "label": "不强求"},
    {"value": "unspecified", "label": "暂不确定"},
]
LANGUAGE_PREFERENCE_OPTIONS = [
    {"value": "chinese_friendly", "label": "中文友好优先"},
    {"value": "bilingual", "label": "中英双语都可"},
    {"value": "english_ok", "label": "英文为主也可以"},
    {"value": "unspecified", "label": "暂不确定"},
]
USER_JOB_TAG_OPTIONS = [
    {"value": "internship", "label": "实习", "category": "岗位类型"},
    {"value": "graduate", "label": "Graduate", "category": "岗位类型"},
    {"value": "full_time", "label": "正式工", "category": "岗位类型"},
    {"value": "contract", "label": "合同/兼职", "category": "岗位类型"},
    {"value": "conversion_strong", "label": "明确可转正", "category": "留新路径"},
    {"value": "conversion_possible", "label": "可能可转正", "category": "留新路径"},
    {"value": "conversion_none", "label": "明确无转正", "category": "留新路径"},
    {"value": "visa_possible", "label": "工签可能", "category": "留新路径"},
    {"value": "visa_unclear", "label": "工签待确认", "category": "留新路径"},
    {"value": "visa_unlikely", "label": "工签风险", "category": "留新路径"},
    {"value": "chinese_friendly", "label": "中文友好可能", "category": "语言环境"},
    {"value": "english_first", "label": "英文为主", "category": "语言环境"},
    {"value": "company_greater_china", "label": "大中华背景", "category": "公司类型"},
    {"value": "company_sg_anchor", "label": "新加坡本地大厂", "category": "公司类型"},
    {"value": "company_ai_startup", "label": "AI/高潜力初创", "category": "公司类型"},
    {"value": "company_product_design", "label": "产品/设计型公司", "category": "公司类型"},
    {"value": "company_fintech", "label": "Fintech", "category": "公司类型"},
    {"value": "company_service_brand", "label": "服务体验品牌", "category": "公司类型"},
    {"value": "ai_related", "label": "AI 相关", "category": "方向标签"},
    {"value": "product_related", "label": "Product", "category": "方向标签"},
    {"value": "ux_related", "label": "UX / Research", "category": "方向标签"},
    {"value": "operations_related", "label": "Ops / Project", "category": "方向标签"},
    {"value": "marketing_related", "label": "Marketing / Content", "category": "方向标签"},
    {"value": "salary_match", "label": "薪资匹配", "category": "薪资"},
    {"value": "salary_unknown", "label": "薪资未知", "category": "薪资"},
    {"value": "salary_low", "label": "薪资偏低", "category": "薪资"},
    {"value": "source_official", "label": "官网 / ATS", "category": "来源"},
    {"value": "source_linkedin", "label": "LinkedIn", "category": "来源"},
    {"value": "source_internsg", "label": "InternSG", "category": "来源"},
    {"value": "source_internship_sg", "label": "Internship.sg", "category": "来源"},
    {"value": "source_jobstreet", "label": "JobStreet", "category": "来源"},
    {"value": "source_indeed", "label": "Indeed", "category": "来源"},
    {"value": "source_mycareersfuture", "label": "MyCareersFuture", "category": "来源"},
    {"value": "source_cultjobs", "label": "Cultjobs 创意岗位", "category": "来源"},
    {"value": "source_startup", "label": "创业与 AI 机会", "category": "来源"},
    {"value": "source_google_jobs", "label": "Google Jobs", "category": "来源"},
    {"value": "fresh_today", "label": "今天新发现", "category": "新鲜度"},
    {"value": "fresh_recent", "label": "近期岗位", "category": "新鲜度"},
    {"value": "fresh_stale", "label": "较早岗位", "category": "新鲜度"},
    {"value": "high_experience", "label": "年限偏高", "category": "风险"},
]
USER_JOB_TAG_VALUES = {item["value"] for item in USER_JOB_TAG_OPTIONS}
USER_JOB_TAG_LABELS = {item["value"]: item["label"] for item in USER_JOB_TAG_OPTIONS}
SALARY_PERIOD_OPTIONS = [
    {"value": "monthly", "label": "月薪"},
    {"value": "yearly", "label": "年薪"},
    {"value": "daily", "label": "日薪"},
    {"value": "hourly", "label": "时薪"},
]
SALARY_BAND_OPTIONS = {
    "SG": {
        "monthly": [
            {"value": "", "label": "先不填"},
            {"value": "800", "label": "SGD 800+"},
            {"value": "1200", "label": "SGD 1,200+"},
            {"value": "1800", "label": "SGD 1,800+"},
            {"value": "2500", "label": "SGD 2,500+"},
            {"value": "3500", "label": "SGD 3,500+"},
        ],
        "yearly": [
            {"value": "", "label": "先不填"},
            {"value": "30000", "label": "SGD 30k+"},
            {"value": "45000", "label": "SGD 45k+"},
            {"value": "60000", "label": "SGD 60k+"},
        ],
        "daily": [
            {"value": "", "label": "先不填"},
            {"value": "80", "label": "SGD 80+/天"},
            {"value": "120", "label": "SGD 120+/天"},
            {"value": "180", "label": "SGD 180+/天"},
        ],
        "hourly": [
            {"value": "", "label": "先不填"},
            {"value": "10", "label": "SGD 10+/时"},
            {"value": "15", "label": "SGD 15+/时"},
            {"value": "25", "label": "SGD 25+/时"},
        ],
    },
    "CN": {
        "monthly": [
            {"value": "", "label": "先不填"},
            {"value": "3000", "label": "CNY 3,000+"},
            {"value": "6000", "label": "CNY 6,000+"},
            {"value": "10000", "label": "CNY 10,000+"},
            {"value": "15000", "label": "CNY 15,000+"},
        ],
        "yearly": [
            {"value": "", "label": "先不填"},
            {"value": "80000", "label": "CNY 80k+"},
            {"value": "150000", "label": "CNY 150k+"},
            {"value": "250000", "label": "CNY 250k+"},
        ],
    },
    "HK": {
        "monthly": [
            {"value": "", "label": "先不填"},
            {"value": "8000", "label": "HKD 8,000+"},
            {"value": "12000", "label": "HKD 12,000+"},
            {"value": "18000", "label": "HKD 18,000+"},
            {"value": "25000", "label": "HKD 25,000+"},
        ],
        "yearly": [
            {"value": "", "label": "先不填"},
            {"value": "180000", "label": "HKD 180k+"},
            {"value": "300000", "label": "HKD 300k+"},
            {"value": "450000", "label": "HKD 450k+"},
        ],
    },
}
SCAN_SOURCE_MODES = {
    "LinkedIn（含 AI 关键词）": "primary",
    "InternSG（含 AI 关键词）": "primary",
    "MyCareersFuture": "supplemental",
    "Careers@Gov": "primary",
    "Internship.sg": "supplemental",
    "Cultjobs": "primary",
    "新加坡科技与 AI ATS": "primary",
    "Google Jobs": "primary",
    "Indeed": "supplemental",
    "JobStreet": "supplemental",
    "关注公司公开来源": "supplemental",
    "创业与 AI 机会": "supplemental",
    "公司官网": "company",
    "JobsDB": "primary",
    "LinkedIn": "primary",
    "Mainland Public Search": "primary",
}

SCAN_SOURCE_NAME_ALIASES = {
    "Glints / NodeFlair / Startups": "创业与 AI 机会",
    "新加坡 AI 初创 ATS": "新加坡科技与 AI ATS",
}
RETIRED_AUTO_SCAN_SOURCES: set[str] = set()

COMPANY_ALIAS_OVERRIDES = {
    "advance.ai": ["ADVANCE.AI", "Advance AI", "Advance Intelligence Group", "advanceai"],
    "ai singapore": ["AI Singapore", "AISG"],
    "aichat": ["AiChat", "AI Chat"],
    "alibaba cloud singapore": ["Alibaba Cloud Singapore", "Alibaba", "Alibaba Cloud", "AliCloud"],
    "ant international": ["Ant International", "Ant Group", "Alipay+", "Alipay Labs", "Ant"],
    "bytedance": ["ByteDance", "TikTok"],
    "changi airport group": ["Changi Airport Group", "Changi Airport", "CAG"],
    "dbs": ["DBS", "DBS Bank"],
    "foodpanda singapore": ["foodpanda Singapore", "foodpanda"],
    "grab": ["Grab", "Grab Singapore"],
    "horizon labs": ["Horizon Labs", "Horizon Labs SG"],
    "hypotenuse ai": ["Hypotenuse AI", "Hypotenuse"],
    "huawei singapore": ["Huawei Singapore", "Huawei", "Huawei Technologies"],
    "ikea singapore": ["IKEA Singapore", "IKEA"],
    "k-id": ["k-ID", "k ID", "kID"],
    "lazada": ["Lazada", "Lazada Singapore"],
    "muji singapore": ["MUJI Singapore", "MUJI"],
    "oppo singapore": ["OPPO Singapore", "OPPO", "Sinoppel"],
    "patsnap": ["PatSnap", "Patsnap"],
    "pdd": ["PDD", "Pinduoduo", "Temu"],
    "pop mart singapore": ["POP MART Singapore", "POP MART", "Popmart"],
    "sea": ["Sea", "Sea Group", "Shopee", "Garena"],
    "shopee": ["Shopee", "Shopee Singapore", "SeaMoney"],
    "shein singapore": ["SHEIN Singapore", "SHEIN", "SHEIN Group"],
    "singtel": ["Singtel", "Singtel Group", "Singapore Telecommunications"],
    "tencent singapore": ["Tencent Singapore", "Tencent"],
    "trip.com group": ["Trip.com Group", "Trip.com", "Ctrip"],
    "vivo singapore": ["vivo Singapore", "vivo"],
    "wiz.ai": ["WIZ.AI", "WIZ AI", "WIZ HOLDINGS", "WIZ HOLDINGS PTE LTD", "Wiz Holdings"],
    "xiaomi singapore": ["Xiaomi Singapore", "Xiaomi", "Xiaomi Technology"],
    "lark": ["Lark", "Lark Suite", "Lark APAC", "ByteDance Lark"],
    "flowmingo ai": ["Flowmingo AI", "Flowmingo", "Featurii"],
    "x0pa ai": ["X0PA AI", "X0PA"],
    "visenze": ["ViSenze", "ViSenze AI", "Rezolve Ai"],
    "pixcap": ["PixCap", "Pixcap"],
    "shopback": ["ShopBack", "ShopBack Pay"],
    "syfe": ["Syfe"],
    "endowus": ["Endowus"],
    "stashaway": ["StashAway"],
    "carousell group": ["Carousell Group", "Carousell"],
    "airwallex": ["Airwallex"],
    "trust bank": ["Trust Bank", "Trust Singapore", "Trust"],
    "nansen": ["Nansen", "Nansen.ai"],
    "transcelestial": ["Transcelestial", "Transcelestial Technologies"],
    "una brands": ["Una Brands"],
    "sleek": ["Sleek"],
    "aspire": ["Aspire", "Aspire App"],
    "circles": ["Circles", "Circles.Life", "Circles Life"],
    "doctor anywhere": ["Doctor Anywhere"],
    "razer": ["Razer"],
    "canva": ["Canva"],
    "wise": ["Wise", "Wise Payments"],
    "crypto.com": ["Crypto.com", "Crypto Com"],
    "nodeflair": ["NodeFlair", "Nodeflair"],
    "funding societies": ["Funding Societies", "Funding Societies | Modalku", "Modalku"],
    "youtrip": ["YouTrip", "YouTrip Singapore"],
    "propertyguru group": ["PropertyGuru Group", "PropertyGuru"],
    "ninja van": ["Ninja Van", "NinjaVan"],
    "bolttech": ["bolttech", "Bolttech"],
    "accredify": ["Accredify"],
    "epos": ["EPOS", "EPOS Pte Ltd", "Epos"],
    "ebb & flow group": ["Ebb & Flow Group", "Ebb Flow", "EBB & FLOW PTE. LTD."],
    "ncs": ["NCS", "NCS Group", "NCS PTE. LTD."],
    "moomoo singapore": ["Moomoo Singapore", "Moomoo", "Futu", "Moomoo Financial Singapore"],
    "holmusk": ["Holmusk"],
    "m-daq": ["M-DAQ", "M-DAQ Global", "M DAQ"],
}

COMPANY_SCAN_ROLE_PATTERN = re.compile(
    r"\b(intern|internship|graduate|associate|junior|entry level|ux|ui|user research|design|designer|"
    r"product|operations|analyst|content|marketing|growth|customer success|project|programme|program|ai|"
    r"machine learning|data|research|service|software|engineer|engineering|developer|development)\b",
    flags=re.I,
)
COMPANY_SCAN_TITLE_ROLE_PATTERN = re.compile(
    r"\b(intern|internship|graduate|trainee|apprentice|fellow|manager|designer|engineer|developer|analyst|researcher|"
    r"scientist|specialist|associate|executive|coordinator|assistant|consultant|architect|officer|director|lead|"
    r"head|owner|sales|operations|recruiter|technician|support|planner|controller|producer|editor|writer|strategist)\b",
    flags=re.I,
)
COMPANIES_REQUIRE_EXPLICIT_LOCATION = {"pdd"}

COMPANY_CAREER_LINK_PATTERN = re.compile(
    r"(career|job|jobs|opening|openings|join-us|join us|position|positions|vacanc|work-with-us|greenhouse|lever|ashby|workday|smartrecruiters|workable|bamboohr|careers-page|surgeahead)",
    flags=re.I,
)

COMPANY_SCAN_PAGE_CAP = 5
COMPANY_SCAN_PER_COMPANY_CAP = 6
COMPANY_SCAN_TIMEOUT = 10
COMPANY_SCAN_SOURCE_LABELS = {
    "Company Site": "官网 / ATS",
    "Company Site / ATS": "官网 / ATS",
    "LinkedIn": "LinkedIn 匹配",
    "JobStreet": "JobStreet 匹配",
    "Google Jobs": "Google Jobs 匹配",
}

AI_EXPLICIT_KEYWORDS = [
    "ai",
    "artificial intelligence",
    "generative ai",
    "genai",
    "gen ai",
    "machine learning",
    "large language model",
    "llm",
    "ai transformation",
    "ai data",
    "ai product",
    "ai operations",
    "ai-enabled",
    "ai powered",
    "ai-powered",
    "agentic",
    "chatbot",
    "prompt",
    "nlp",
    "computer vision",
]

AI_DOMAIN_KEYWORDS = [
    "generative ai",
    "genai",
    "gen ai",
    "machine learning",
    "large language model",
    "llm",
    "ai transformation",
    "ai data",
    "ai product",
    "ai operations",
    "ai agent",
    "agentic",
    "chatbot",
    "prompt",
    "nlp",
    "computer vision",
    "focused on artificial intelligence",
    "ai-powered platform",
]

AI_ADJACENT_KEYWORDS = [
    "automation",
    "workflow automation",
    "data orchestration",
    "data annotation",
    "model evaluation",
    "human moderation",
    "trust and safety",
]

AI_ROLE_ANCHORS = [
    "product",
    "design",
    "designer",
    "ux",
    "user research",
    "research",
    "service",
    "operations",
    "project",
    "marketing",
    "content",
    "analyst",
    "data",
    "strategy",
    "innovation",
    "intern",
    "graduate",
    "associate",
]

AI_NOISE_PATTERNS = [
    r"use ai to assess how you fit",
    r"get ai-powered advice",
    r"ai-powered advice",
    r"see how you compare to .*? applicants",
    r"tailor my resume.*?sign in",
    r"to support an efficient and fair hiring process.*?applicant privacy notice",
    r"we may use artificial intelligence \(?ai\)? tools to help .*?recruiters",
    r"we may use .*?artificial intelligence .*?(hiring|recruitment|application|applicant).*?(decision|process|review)",
]

CAREER_DIRECTIONS = [
    {
        "id": "ai-product",
        "label": "AI Product",
        "keywords": ["ai", "genai", "ai product", "llm", "generative ai", "ai agent", "chatbot", "prompt", "automation", "workflow"],
        "evidence": ["ai-assisted", "prompt", "workflow automation", "jd/capability matching", "scenario exploration"],
        "gaps": ["LLM evaluation metrics", "AI product launch evidence", "technical product specs"],
    },
    {
        "id": "ux-product-design",
        "label": "UX/Product Design",
        "keywords": ["ux", "product design", "figma", "prototype", "interaction", "experience design", "usability"],
        "evidence": ["figma", "prototype", "user journey", "visual design", "human-centred"],
        "gaps": ["portfolio case evidence for shipped UI", "interaction metrics", "design system examples"],
    },
    {
        "id": "user-research",
        "label": "User Research",
        "keywords": ["user research", "ux research", "interview", "survey", "usability", "insight", "qualitative"],
        "evidence": ["user research", "interview", "journey", "research synthesis", "service blueprint"],
        "gaps": ["quantitative research methods", "research repository examples", "sample-size evidence"],
    },
    {
        "id": "service-design",
        "label": "Service Design",
        "keywords": ["service design", "service blueprint", "customer journey", "touchpoint", "healthcare", "public service"],
        "evidence": ["service design", "service blueprint", "healthcare", "spatial flow", "stakeholder"],
        "gaps": ["measured service outcomes", "stakeholder implementation evidence", "business impact"],
    },
    {
        "id": "product-ops",
        "label": "Product Ops",
        "keywords": ["product operations", "operations", "process", "workflow", "data operations", "project management"],
        "evidence": ["workflow automation", "on-site operations", "process", "coordination", "documentation"],
        "gaps": ["dashboard metrics", "SQL/data workflow evidence", "cross-functional operating cadence"],
    },
    {
        "id": "growth-content",
        "label": "Growth/Content",
        "keywords": ["growth", "marketing", "content", "campaign", "community", "copywriting", "conversion"],
        "evidence": ["content creation", "visual design", "ux writing", "scenario exploration", "portfolio"],
        "gaps": ["conversion metrics", "campaign results", "audience segmentation evidence"],
    },
]

DIRECTION_GENERIC_KEYWORDS = {
    "ai-product": {"ai", "genai", "generative ai", "llm", "automation", "workflow"},
    "ux-product-design": {"prototype"},
    "user-research": {"interview", "survey", "insight"},
    "service-design": {"healthcare", "public service"},
    "product-ops": {"operations", "process", "workflow"},
    "growth-content": {"growth", "content", "community", "conversion"},
}

NON_TARGET_FUNCTION_TITLE_PATTERN = re.compile(
    r"\b(finance|financial|accounting|accountant|audit|tax|legal|counsel|compliance|risk|"
    r"human resources|hr|people|talent|recruit(?:er|ing)?|procurement|supply chain|"
    r"investment|portfolio|pre[- ]?sales|sales|business development|account management|"
    r"partner success|customer success|data science|data scientist|data engineer(?:ing)?|"
    r"machine learning engineer|ml engineer|ai engineer|artificial intelligence engineer|"
    r"software engineer(?:ing)?|software developer|back[- ]?end|front[- ]?end|full[- ]?stack|"
    r"devops|site reliability|sre|cloud engineer|platform engineer|security engineer|"
    r"cybersecurity|quantitative|robotics engineer|robot learning|ai\s*/\s*data|data intern|"
    r"autonomous vehicle|integration\s*(?:&|and)\s*validation|software tools|"
    r"engineering intern(?:ship)?|repair|maintenance)\b",
    flags=re.I,
)

PURE_TECHNICAL_TITLE_PATTERN = re.compile(
    r"\b(data science|data scientist|data engineer(?:ing)?|machine learning engineer|ml engineer|"
    r"ai engineer|artificial intelligence engineer|software engineer(?:ing)?|software developer|"
    r"back[- ]?end|front[- ]?end|full[- ]?stack|devops|site reliability|sre|cloud engineer|"
    r"platform engineer|security engineer|cybersecurity|quantitative|robotics engineer|"
    r"robot learning|ai\s*/\s*data|data intern|autonomous vehicle|"
    r"integration\s*(?:&|and)\s*validation|software tools|engineering intern(?:ship)?)\b",
    flags=re.I,
)

PRODUCT_FACING_TITLE_PATTERN = re.compile(
    r"\b(product|ux|ui|user research|design|research|growth|content|marketing|operations|"
    r"project|program(?:me)?|strategy|innovation)\b",
    flags=re.I,
)

DIRECTION_CATEGORIES = {
    "ai-product": "AI 与产品",
    "ux-product-design": "设计与体验",
    "user-research": "研究与洞察",
    "service-design": "研究与洞察",
    "product-ops": "运营与商业",
    "growth-content": "运营与商业",
}

FALLBACK_PROFILE_TEXT = """
Early-career product, UX, and service design candidate.
Service design, UX research, user journeys, service blueprints, design innovation,
product design, experience design, Figma, visual design, prototyping,
content creation, operations, workflow documentation, and AI-assisted research synthesis.
"""


def today() -> str:
    return dt.date.today().strftime(DATE_FMT)


def now_iso() -> str:
    return dt.datetime.now().replace(microsecond=0).isoformat()


def truthy_env(name: str) -> bool:
    return (os.environ.get(name) or "").strip().lower() in {"1", "true", "yes", "on"}


def supabase_auth_configured() -> bool:
    return bool(os.environ.get("SUPABASE_URL") and supabase_anon_key())


def auth_required() -> bool:
    return truthy_env("JOB_ASSISTANT_REQUIRE_AUTH") or supabase_auth_configured()


def auth_config_payload() -> dict:
    return {
        "auth_required": auth_required(),
        "supabase_url": os.environ.get("SUPABASE_URL", ""),
        "supabase_anon_key": supabase_anon_key(),
    }


class AuthError(ValueError):
    pass


def safe_user_id(user_id: str | None) -> str:
    value = (user_id or LOCAL_USER_ID).strip() or LOCAL_USER_ID
    cleaned = re.sub(r"[^A-Za-z0-9_.@-]+", "-", value).strip(".-")
    if cleaned:
        return cleaned[:120]
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:32]


def request_user_id() -> str:
    return getattr(REQUEST_CONTEXT, "user_id", LOCAL_USER_ID) or LOCAL_USER_ID


def scoped_storage_enabled() -> bool:
    return auth_required() or request_user_id() != LOCAL_USER_ID


def current_data_dir() -> Path:
    if scoped_storage_enabled():
        return DATA_DIR / "users" / safe_user_id(request_user_id())
    return DATA_DIR


def current_workspace_dir() -> Path:
    if scoped_storage_enabled():
        return WORKSPACE_DIR / "users" / safe_user_id(request_user_id())
    return WORKSPACE_DIR


def current_db_path() -> Path:
    return current_data_dir() / "career_copilot.sqlite" if scoped_storage_enabled() else DB_PATH


def current_profile_path() -> Path:
    return current_data_dir() / "profile.json" if scoped_storage_enabled() else PROFILE_PATH


def current_user_context_path() -> Path:
    return current_data_dir() / "user_context.json" if scoped_storage_enabled() else USER_CONTEXT_PATH


def current_apply_assist_dir() -> Path:
    return current_data_dir() / "apply-assist" if scoped_storage_enabled() else APPLY_ASSIST_DIR


def current_browser_profile_dir() -> Path:
    return current_data_dir() / "browser-profile" if scoped_storage_enabled() else BROWSER_PROFILE_DIR


def current_resume_upload_dir() -> Path:
    return current_data_dir() / "resumes" if scoped_storage_enabled() else RESUME_UPLOAD_DIR


def current_default_resume_pdf() -> Path:
    return current_resume_upload_dir() / "active-resume.pdf"


def current_resume_path() -> Path:
    configured = os.environ.get("JOB_ASSISTANT_RESUME")
    return Path(configured).expanduser() if configured else current_default_resume_pdf()


def current_profile_photo_path() -> Path:
    return current_resume_upload_dir() / "profile-photo.jpg"


def current_notion_config_path() -> Path:
    return current_data_dir() / "notion_config.json"


def supabase_base_url() -> str:
    return (os.environ.get("SUPABASE_URL") or "").rstrip("/")


def supabase_anon_key() -> str:
    return os.environ.get("SUPABASE_ANON_KEY") or os.environ.get("SUPABASE_PUBLISHABLE_KEY", "")


def supabase_service_role_key() -> str:
    return os.environ.get("SUPABASE_SERVICE_ROLE_KEY") or os.environ.get("SUPABASE_SECRET_KEY", "")


def cloud_state_bucket() -> str:
    return os.environ.get("SUPABASE_STORAGE_BUCKET", "job-assistant-users")


def cloud_state_enabled() -> bool:
    if not (supabase_base_url() and supabase_service_role_key()):
        return False
    return auth_required() or truthy_env("JOB_ASSISTANT_CLOUD_STATE")


def cloud_state_user_key(user_id: str | None = None) -> str:
    return safe_user_id(user_id or request_user_id())


def cloud_state_object_path(user_id: str | None = None) -> str:
    return f"{cloud_state_user_key(user_id)}/state.zip"


def cloud_state_lock_for(user_id: str | None = None) -> threading.RLock:
    key = cloud_state_user_key(user_id)
    with CLOUD_STATE_LOCK:
        lock = CLOUD_STATE_USER_LOCKS.get(key)
        if lock is None:
            lock = threading.RLock()
            CLOUD_STATE_USER_LOCKS[key] = lock
        return lock


def supabase_storage_request(
    method: str,
    path: str,
    data: bytes | None = None,
    headers: dict | None = None,
    tolerate_404: bool = False,
) -> tuple[int, bytes]:
    base = supabase_base_url()
    service_key = supabase_service_role_key()
    if not base or not service_key:
        raise ValueError("Supabase Storage is not configured.")
    request = urllib.request.Request(
        f"{base}/storage/v1{path}",
        data=data,
        method=method,
        headers={
            "apikey": service_key,
            "Authorization": f"Bearer {service_key}",
            **(headers or {}),
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            return response.status, response.read()
    except urllib.error.HTTPError as exc:
        body = exc.read()
        detail = body.decode("utf-8", errors="ignore")
        if tolerate_404 and (exc.code == 404 or '"statusCode":"404"' in detail or '"statusCode":404' in detail):
            return 404, body
        raise ValueError(f"Supabase Storage error {exc.code}: {detail}") from exc


def ensure_cloud_state_bucket() -> None:
    if not cloud_state_enabled():
        return
    bucket = cloud_state_bucket()
    if bucket in CLOUD_STATE_BUCKET_READY:
        return
    with CLOUD_STATE_LOCK:
        if bucket in CLOUD_STATE_BUCKET_READY:
            return
        status, _ = supabase_storage_request("GET", f"/bucket/{quote_plus(bucket)}", tolerate_404=True)
        if status == 404:
            payload = json.dumps({"id": bucket, "name": bucket, "public": False}).encode("utf-8")
            supabase_storage_request(
                "POST",
                "/bucket",
                data=payload,
                headers={"Content-Type": "application/json"},
            )
        CLOUD_STATE_BUCKET_READY.add(bucket)


def user_state_paths() -> list[tuple[str, Path]]:
    return [
        ("data", current_data_dir()),
        ("workspace", current_workspace_dir()),
    ]


def should_skip_cloud_path(path: Path) -> bool:
    ignored_parts = {"browser-profile", "__pycache__"}
    return any(part in ignored_parts for part in path.parts) or path.name in {".DS_Store"}


def build_user_state_archive() -> bytes:
    ensure_dirs()
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for prefix, root in user_state_paths():
            if not root.exists():
                continue
            for path in root.rglob("*"):
                if should_skip_cloud_path(path):
                    continue
                relative = path.relative_to(root)
                archive_name = PurePosixPath(prefix, *relative.parts).as_posix()
                if path.is_dir():
                    archive.writestr(f"{archive_name}/", b"")
                    continue
                try:
                    archive.write(path, archive_name)
                except FileNotFoundError:
                    continue
    return buffer.getvalue()


def restore_user_state_archive(payload: bytes) -> bool:
    if not payload:
        return False
    ensure_dirs()
    roots = {prefix: root.resolve() for prefix, root in user_state_paths()}
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        for info in archive.infolist():
            posix_path = PurePosixPath(info.filename)
            parts = posix_path.parts
            if len(parts) < 2 or parts[0] not in roots:
                continue
            target_root = roots[parts[0]]
            target = (target_root / Path(*parts[1:])).resolve()
            if target != target_root and target_root not in target.parents:
                continue
            if should_skip_cloud_path(target):
                continue
            if info.is_dir():
                target.mkdir(parents=True, exist_ok=True)
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(info) as source, target.open("wb") as output:
                output.write(source.read())
    return True


def ensure_cloud_state_loaded() -> None:
    if not cloud_state_enabled() or not scoped_storage_enabled():
        return
    user_key = cloud_state_user_key()
    if user_key in CLOUD_STATE_LOADED:
        return
    with cloud_state_lock_for(user_key):
        if user_key in CLOUD_STATE_LOADED:
            return
        ensure_cloud_state_bucket()
        object_path = quote_plus(cloud_state_object_path(user_key)).replace("%2F", "/")
        status, payload = supabase_storage_request(
            "GET",
            f"/object/{quote_plus(cloud_state_bucket())}/{object_path}",
            tolerate_404=True,
        )
        if status != 404:
            restore_user_state_archive(payload)
        CLOUD_STATE_LOADED.add(user_key)


def sync_cloud_state(reason: str = "") -> bool:
    if not cloud_state_enabled() or not scoped_storage_enabled():
        return False
    user_key = cloud_state_user_key()
    with cloud_state_lock_for(user_key):
        ensure_cloud_state_loaded()
        ensure_cloud_state_bucket()
        payload = build_user_state_archive()
        object_path = quote_plus(cloud_state_object_path(user_key)).replace("%2F", "/")
        supabase_storage_request(
            "POST",
            f"/object/{quote_plus(cloud_state_bucket())}/{object_path}",
            data=payload,
            headers={
                "Content-Type": "application/zip",
                "x-upsert": "true",
                "Cache-Control": "no-store",
            },
        )
        CLOUD_STATE_LOADED.add(user_key)
        return True


def safe_sync_cloud_state(reason: str = "") -> None:
    try:
        sync_cloud_state(reason)
    except Exception as exc:
        print(f"Cloud state sync failed ({reason or 'request'}): {exc}", file=sys.stderr)


def should_sync_after_request(method: str, path: str) -> bool:
    if not path.startswith("/api/") or is_public_api_path(path):
        return False
    if method in {"POST", "PUT", "DELETE"} and path != "/api/open-path":
        return True
    return method == "GET" and path == "/api/report/today"


@contextmanager
def request_user_context(user_id: str):
    previous = getattr(REQUEST_CONTEXT, "user_id", None)
    REQUEST_CONTEXT.user_id = user_id or LOCAL_USER_ID
    try:
        yield
    finally:
        if previous is None:
            try:
                delattr(REQUEST_CONTEXT, "user_id")
            except AttributeError:
                pass
        else:
            REQUEST_CONTEXT.user_id = previous


@contextmanager
def db_initialization_context():
    previous = getattr(REQUEST_CONTEXT, "initializing_db", False)
    REQUEST_CONTEXT.initializing_db = True
    try:
        yield
    finally:
        REQUEST_CONTEXT.initializing_db = previous


def is_public_api_path(path: str) -> bool:
    return path in {"/api/health", "/api/auth/config", "/api/profile-options"}


def user_id_from_bearer_token(handler: SimpleHTTPRequestHandler) -> str:
    header = handler.headers.get("Authorization", "")
    if not header.startswith("Bearer "):
        raise AuthError("请先登录。")
    token = header.split(" ", 1)[1].strip()
    if not token:
        raise AuthError("请先登录。")
    secret = os.environ.get("SUPABASE_JWT_SECRET")
    if secret:
        if jwt is None:
            raise AuthError("服务端缺少 PyJWT，请重新安装依赖。")
        try:
            payload = jwt.decode(token, secret, algorithms=["HS256"], options={"verify_aud": False})
        except Exception as exc:
            raise AuthError("登录已过期，请重新登录。") from exc
        user_id = payload.get("sub")
        if not user_id:
            raise AuthError("登录信息无效，请重新登录。")
        return str(user_id)
    return user_id_from_supabase_auth(token)


def user_id_from_supabase_auth(token: str) -> str:
    base = supabase_base_url()
    anon_key = supabase_anon_key()
    if not base or not anon_key:
        raise AuthError("服务端还没有配置 Supabase 登录服务。")
    request = urllib.request.Request(
        f"{base}/auth/v1/user",
        headers={
            "apikey": anon_key,
            "Authorization": f"Bearer {token}",
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=20) as response:
            payload = json.loads(response.read().decode("utf-8") or "{}")
    except Exception as exc:
        raise AuthError("登录已过期，请重新登录。") from exc
    user_id = payload.get("id") or payload.get("sub")
    if not user_id:
        raise AuthError("登录信息无效，请重新登录。")
    return str(user_id)


def user_id_for_request(handler: SimpleHTTPRequestHandler, path: str) -> str:
    if not path.startswith("/api/") or is_public_api_path(path):
        return LOCAL_USER_ID
    if not auth_required():
        return LOCAL_USER_ID
    return user_id_from_bearer_token(handler)


def scan_thread_key(scan_run_id: int, user_id: str | None = None) -> str:
    return f"{safe_user_id(user_id or request_user_id())}:{scan_run_id}"


def ensure_dirs() -> None:
    for path in [
        current_data_dir(),
        current_apply_assist_dir(),
        current_browser_profile_dir(),
        current_resume_upload_dir(),
        current_workspace_dir(),
        current_workspace_dir() / "drafts",
        current_workspace_dir() / "applications",
        current_workspace_dir() / "reports",
    ]:
        path.mkdir(parents=True, exist_ok=True)


def health_payload() -> dict:
    ensure_dirs()
    db_path = current_db_path()
    with sqlite3.connect(db_path) as conn:
        conn.execute("select 1").fetchone()
    return {
        "ok": True,
        "app": "job-assistant",
        "time": now_iso(),
        "database": str(db_path),
        "storage": "scoped" if scoped_storage_enabled() else "local",
        "cloud_state": "enabled" if cloud_state_enabled() else "disabled",
        "cloud_bucket": cloud_state_bucket() if cloud_state_enabled() else "",
    }


@contextmanager
def get_db():
    ensure_dirs()
    ensure_cloud_state_loaded()
    db_path = current_db_path()
    if not getattr(REQUEST_CONTEXT, "initializing_db", False) and str(db_path) not in INITIALIZED_DB_PATHS:
        with DB_INIT_LOCK:
            if str(db_path) not in INITIALIZED_DB_PATHS:
                setup_db()
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def ensure_column(conn: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    existing = {row["name"] for row in conn.execute(f"pragma table_info({table})").fetchall()}
    if column not in existing:
        conn.execute(f"alter table {table} add column {column} {definition}")


def table_columns(conn: sqlite3.Connection, table: str) -> set[str]:
    return {row["name"] for row in conn.execute(f"pragma table_info({table})").fetchall()}


def migrate_watch_companies_table(conn: sqlite3.Connection) -> None:
    columns = table_columns(conn, "watch_companies")
    schema_row = conn.execute(
        "select sql from sqlite_master where type='table' and name='watch_companies'"
    ).fetchone()
    schema_sql = (schema_row["sql"] or "").lower() if schema_row else ""
    needs_rebuild = "region" not in columns or "company text not null unique" in schema_sql
    if not needs_rebuild:
        ensure_column(conn, "watch_companies", "region", "text not null default 'SG'")
        ensure_column(conn, "watch_companies", "city_tags_json", "text not null default '[]'")
        ensure_column(conn, "watch_companies", "aliases_json", "text not null default '[]'")
        ensure_column(conn, "watch_companies", "company_type", "text not null default 'Company'")
        ensure_column(conn, "watch_companies", "user_added", "integer not null default 0")
        ensure_column(conn, "watch_companies", "priority", "integer not null default 50")
        ensure_column(conn, "watch_companies", "notes", "text not null default ''")
        ensure_column(conn, "watch_companies", "last_scan_status", "text not null default 'not_scanned'")
        ensure_column(conn, "watch_companies", "last_scan_note", "text not null default ''")
        ensure_column(conn, "watch_companies", "last_jobs_found", "integer not null default 0")
        conn.execute(
            "create unique index if not exists idx_watch_companies_region_company on watch_companies(region, company)"
        )
        return

    conn.execute("alter table watch_companies rename to watch_companies_old")
    conn.execute(
        """
        create table watch_companies (
            id integer primary key autoincrement,
            company text not null,
            source text not null,
            url text not null,
            focus text not null,
            region text not null default 'SG',
            city_tags_json text not null default '["Singapore"]',
            aliases_json text not null default '[]',
            company_type text not null default 'Company',
            user_added integer not null default 0,
            priority integer not null default 50,
            notes text not null default '',
            last_checked_at text,
            last_scan_status text not null default 'not_scanned',
            last_scan_note text not null default '',
            last_jobs_found integer not null default 0,
            status text not null default 'Watch',
            unique(region, company)
        )
        """
    )
    old_columns = table_columns(conn, "watch_companies_old")
    select_columns = [
        "id",
        "company",
        "source",
        "url",
        "focus",
        "'SG' as region",
        "'[\"Singapore\"]' as city_tags_json",
        "aliases_json" if "aliases_json" in old_columns else "'[]' as aliases_json",
        "'Company' as company_type",
        "0 as user_added",
        "50 as priority",
        "'' as notes",
        "last_checked_at" if "last_checked_at" in old_columns else "null as last_checked_at",
        "last_scan_status" if "last_scan_status" in old_columns else "'not_scanned' as last_scan_status",
        "last_scan_note" if "last_scan_note" in old_columns else "'' as last_scan_note",
        "last_jobs_found" if "last_jobs_found" in old_columns else "0 as last_jobs_found",
        "status" if "status" in old_columns else "'Watch' as status",
    ]
    conn.execute(
        f"""
        insert or ignore into watch_companies(
            id, company, source, url, focus, region, city_tags_json,
            aliases_json, company_type, user_added, priority, notes, last_checked_at,
            last_scan_status, last_scan_note, last_jobs_found, status
        )
        select {", ".join(select_columns)}
        from watch_companies_old
        """
    )
    conn.execute("drop table watch_companies_old")
    conn.execute(
        "create unique index if not exists idx_watch_companies_region_company on watch_companies(region, company)"
    )


def seed_default_watch_companies(conn: sqlite3.Connection) -> None:
    for company, source, url, focus in WATCH_COMPANIES:
        conn.execute(
            """
            insert into watch_companies(
                company, source, url, focus, region, city_tags_json,
                aliases_json, company_type, user_added, priority, notes, status
            )
            values(?, ?, ?, ?, 'SG', '["Singapore"]', ?, 'Company', 0, 50, '', 'Watch')
            on conflict(region, company) do update set
                source=excluded.source,
                url=excluded.url,
                focus=excluded.focus,
                aliases_json=excluded.aliases_json
            """,
            (company, source, url, focus, json.dumps(COMPANY_ALIAS_OVERRIDES.get(company.lower(), []), ensure_ascii=False)),
        )
    for item in COMPANY_CATALOG:
        if not item.get("default_watch"):
            continue
        conn.execute(
            """
            insert into watch_companies(
                company, source, url, focus, region, city_tags_json,
                aliases_json, company_type, user_added, priority, notes, status
            )
            values(?, ?, ?, ?, ?, ?, ?, ?, 0, ?, '', 'Watch')
            on conflict(region, company) do update set
                source=excluded.source,
                url=excluded.url,
                focus=excluded.focus,
                city_tags_json=excluded.city_tags_json,
                aliases_json=excluded.aliases_json,
                company_type=excluded.company_type,
                priority=max(watch_companies.priority, excluded.priority)
            """,
            (
                item["company"],
                item["source"],
                item["url"],
                item["focus"],
                item["region"],
                json.dumps(item.get("city_tags") or [], ensure_ascii=False),
                json.dumps(item.get("aliases") or COMPANY_ALIAS_OVERRIDES.get(item["company"].lower(), []), ensure_ascii=False),
                item.get("company_type") or "Company",
                int(item.get("priority") or 50),
            ),
        )


def setup_db() -> None:
    ensure_dirs()
    db_path = current_db_path()
    with DB_INIT_LOCK:
        with db_initialization_context():
            with get_db() as conn:
                conn.executescript(
                    """
            create table if not exists jobs (
                id integer primary key autoincrement,
                company text not null,
                position text not null,
                name text not null,
                source text not null,
                url text not null unique,
                external_job_id text,
                location text not null default 'Singapore',
                job_type text,
                employment_type text not null default 'Unknown',
                conversion_opportunity integer not null default 0,
                salary_min real,
                salary_max real,
                salary_currency text,
                salary_period text,
                salary_text text,
                salary_fit text not null default 'unknown',
                conversion_signal text not null default 'unknown',
                visa_sponsorship_signal text not null default 'unknown',
                language_signal text not null default 'unknown',
                pathway_score real not null default 0,
                pathway_evidence_json text not null default '[]',
                application_deadline text,
                jd_text text not null,
                jd_hash text not null,
                score real not null default 0,
                status text not null default 'New',
                decision text,
                eligibility_flags text not null default '[]',
                match_notes text not null default '',
                found_date text not null,
                batch_date text,
                recommended_date text,
                applied_date text,
                last_followup_at text,
                followup_count integer not null default 0,
                last_checked_at text not null,
                notion_page_id text,
                resume_path text,
                cover_letter_path text,
                jd_cn_text text,
                created_at text not null,
                updated_at text not null
            );

            create table if not exists applications (
                id integer primary key autoincrement,
                job_id integer not null,
                status text not null default 'Drafted',
                resume_path text,
                cover_letter_path text,
                submitted_at text,
                submission_mode text not null default 'human_confirmed',
                custom_questions_json text not null default '[]',
                notes text not null default '',
                created_at text not null,
                updated_at text not null,
                foreign key(job_id) references jobs(id)
            );

            create table if not exists daily_reports (
                date text primary key,
                searched_count integer not null default 0,
                recommended_count integer not null default 0,
                drafted_count integer not null default 0,
                apply_queue_count integer not null default 0,
                applied_count integer not null default 0,
                watch_count integer not null default 0,
                drop_count integer not null default 0,
                failures_json text not null default '[]',
                report_markdown_path text,
                updated_at text not null
            );

            create table if not exists watch_companies (
                id integer primary key autoincrement,
                company text not null unique,
                source text not null,
                url text not null,
                focus text not null,
                last_checked_at text,
                status text not null default 'Watch'
            );

            create table if not exists scan_runs (
                id integer primary key autoincrement,
                run_date text not null,
                started_at text not null,
                finished_at text,
                status text not null default 'running',
                triggered_by text not null default 'manual',
                forced integer not null default 0,
                scanned_count integer not null default 0,
                saved_count integer not null default 0,
                new_count integer not null default 0,
                updated_count integer not null default 0,
                duplicate_count integer not null default 0,
                recommended_count integer not null default 0,
                ai_recommended_count integer not null default 0,
                failures_json text not null default '[]',
                created_at text not null,
                updated_at text not null
            );

            create table if not exists scan_source_runs (
                id integer primary key autoincrement,
                scan_run_id integer not null,
                source text not null,
                started_at text not null,
                finished_at text,
                status text not null default 'running',
                scanned_count integer not null default 0,
                saved_count integer not null default 0,
                new_count integer not null default 0,
                updated_count integer not null default 0,
                duplicate_count integer not null default 0,
                failure_count integer not null default 0,
                failures_json text not null default '[]',
                created_at text not null,
                updated_at text not null,
                foreign key(scan_run_id) references scan_runs(id)
            );

            create table if not exists resume_versions (
                id integer primary key autoincrement,
                filename text not null,
                original_filename text not null,
                stored_path text not null unique,
                text_path text not null,
                content_hash text not null,
                mime_type text,
                file_size integer not null default 0,
                active integer not null default 0,
                created_at text not null
            );

            create table if not exists resume_analyses (
                id integer primary key autoincrement,
                resume_version_id integer not null,
                mode text not null default 'local',
                summary text not null default '',
                strengths_json text not null default '[]',
                directions_json text not null default '[]',
                gaps_json text not null default '[]',
                evidence_json text not null default '[]',
                created_at text not null,
                foreign key(resume_version_id) references resume_versions(id)
            );

            create table if not exists career_preferences (
                id integer primary key check(id = 1),
                selected_directions_json text not null default '[]',
                direction_weights_json text not null default '{}',
                exclude_keywords_json text not null default '[]',
                updated_at text not null
            );
            """
                )
                ensure_column(conn, "jobs", "jd_cn_text", "text")
                ensure_column(conn, "jobs", "region", "text not null default 'SG'")
                ensure_column(conn, "jobs", "city", "text")
                ensure_column(conn, "jobs", "source_region", "text")
                ensure_column(conn, "jobs", "employment_type", "text not null default 'Unknown'")
                ensure_column(conn, "jobs", "conversion_opportunity", "integer not null default 0")
                ensure_column(conn, "jobs", "salary_min", "real")
                ensure_column(conn, "jobs", "salary_max", "real")
                ensure_column(conn, "jobs", "salary_currency", "text")
                ensure_column(conn, "jobs", "salary_period", "text")
                ensure_column(conn, "jobs", "salary_text", "text")
                ensure_column(conn, "jobs", "salary_fit", "text not null default 'unknown'")
                ensure_column(conn, "jobs", "conversion_signal", "text not null default 'unknown'")
                ensure_column(conn, "jobs", "visa_sponsorship_signal", "text not null default 'unknown'")
                ensure_column(conn, "jobs", "language_signal", "text not null default 'unknown'")
                ensure_column(conn, "jobs", "pathway_score", "real not null default 0")
                ensure_column(conn, "jobs", "pathway_evidence_json", "text not null default '[]'")
                ensure_column(conn, "jobs", "application_deadline", "text")
                ensure_column(conn, "jobs", "last_followup_at", "text")
                ensure_column(conn, "jobs", "followup_count", "integer not null default 0")
                ensure_column(conn, "applications", "assist_payload_path", "text")
                ensure_column(conn, "applications", "assist_result_path", "text")
                ensure_column(conn, "applications", "assist_status", "text")
                ensure_column(conn, "applications", "assist_updated_at", "text")
                ensure_column(conn, "scan_runs", "region", "text not null default 'SG'")
                ensure_column(conn, "scan_runs", "city", "text")
                ensure_column(conn, "scan_runs", "source_region", "text")
                ensure_column(conn, "scan_runs", "new_count", "integer not null default 0")
                ensure_column(conn, "scan_runs", "updated_count", "integer not null default 0")
                ensure_column(conn, "scan_runs", "duplicate_count", "integer not null default 0")
                ensure_column(conn, "scan_source_runs", "new_count", "integer not null default 0")
                ensure_column(conn, "scan_source_runs", "updated_count", "integer not null default 0")
                ensure_column(conn, "scan_source_runs", "duplicate_count", "integer not null default 0")
                migrate_watch_companies_table(conn)
                seed_default_watch_companies(conn)
                backfill_job_metadata(conn)
                backfill_application_deadlines(conn)
        INITIALIZED_DB_PATHS.add(str(db_path))


def row_to_dict(row: sqlite3.Row) -> dict:
    out = dict(row)
    for key in [
        "eligibility_flags",
        "custom_questions_json",
        "failures_json",
        "strengths_json",
        "directions_json",
        "gaps_json",
        "evidence_json",
        "selected_directions_json",
        "direction_weights_json",
        "exclude_keywords_json",
        "city_tags_json",
        "aliases_json",
        "pathway_evidence_json",
    ]:
        if key in out:
            try:
                fallback = "{}" if key == "direction_weights_json" else "[]"
                out[key] = json.loads(out[key] or fallback)
            except json.JSONDecodeError:
                out[key] = {} if key == "direction_weights_json" else []
    return out


def clean_text(value: str) -> str:
    value = re.sub(r"<script\b.*?</script>", " ", value or "", flags=re.I | re.S)
    value = re.sub(r"<style\b.*?</style>", " ", value, flags=re.I | re.S)
    value = re.sub(r"<[^>]+>", " ", value)
    value = unescape(value)
    return re.sub(r"\s+", " ", value).strip()


DEADLINE_MONTH_PATTERN = (
    r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
)
APPLICATION_DEADLINE_PATTERN = re.compile(
    rf"\b(?:closing\s+date|application\s+deadline|applications?\s+(?:close|closing)|apply\s+by)"
    rf"\s*[:\-–]?\s*(?:on\s+)?(?P<date>"
    rf"\d{{1,2}}[\s./-]+{DEADLINE_MONTH_PATTERN}[\s,./-]+\d{{4}}|"
    rf"{DEADLINE_MONTH_PATTERN}\s+\d{{1,2}}(?:st|nd|rd|th)?[,]?\s+\d{{4}}|"
    rf"\d{{4}}-\d{{1,2}}-\d{{1,2}}|\d{{1,2}}[/-]\d{{1,2}}[/-]\d{{4}})\b",
    re.I,
)


def normalize_application_deadline(value: str, reference_date: dt.date | None = None) -> str:
    raw = re.sub(r"(?i)(\d)(?:st|nd|rd|th)\b", r"\1", clean_text(value)).strip(" ,.;")
    if not raw:
        return ""
    parsed = None
    for date_format in ["%d %B %Y", "%d %b %Y", "%B %d, %Y", "%b %d, %Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"]:
        try:
            parsed = dt.datetime.strptime(raw, date_format).date()
            break
        except ValueError:
            continue
    if not parsed:
        return ""
    reference = reference_date or dt.date.today()
    if parsed.year < 2000 or parsed.year > reference.year + 5:
        return ""
    return parsed.strftime(DATE_FMT)


def extract_application_deadline(text: str, reference_date: dt.date | None = None) -> str:
    match = APPLICATION_DEADLINE_PATTERN.search(clean_text(text))
    return normalize_application_deadline(match.group("date"), reference_date) if match else ""


def application_deadline_status(value: str | None, reference_date: dt.date | None = None) -> dict:
    try:
        deadline = dt.datetime.strptime(str(value or "")[:10], DATE_FMT).date()
    except ValueError:
        return {"code": "unknown", "label": "", "days_remaining": None}
    remaining = (deadline - (reference_date or dt.date.today())).days
    if remaining < 0:
        return {"code": "expired", "label": "截止日期已过", "days_remaining": remaining}
    if remaining == 0:
        return {"code": "today", "label": "今天截止", "days_remaining": 0}
    if remaining <= 3:
        return {"code": "urgent", "label": f"{remaining} 天后截止", "days_remaining": remaining}
    if remaining <= 7:
        return {"code": "soon", "label": f"{remaining} 天后截止", "days_remaining": remaining}
    return {"code": "scheduled", "label": f"截止 {deadline.strftime('%m-%d')}", "days_remaining": remaining}


def queue_decision(job: dict, reference_date: dt.date | None = None) -> dict:
    reference = reference_date or dt.date.today()
    deadline = application_deadline_status(job.get("application_deadline"), reference)
    fit_score = float(job.get("fit_score") or job.get("rank_score") or job.get("score") or 0)
    freshness = str(job.get("listing_freshness_status") or "")
    try:
        queued_date = dt.datetime.fromisoformat(str(job.get("updated_at") or "")).date()
        queue_age = max(0, (reference - queued_date).days)
    except ValueError:
        queue_age = 0

    if deadline["code"] == "expired":
        return {"priority": "review", "label": "先确认再投", "reason": "截止日期已过，先确认岗位是否仍开放。", "order": 2}
    if deadline["code"] in {"today", "urgent", "soon"}:
        return {"priority": "today", "label": "今天优先", "reason": f"{deadline['label']}，建议优先完成投递。", "order": 0}
    if freshness in {"aging", "verify", "unknown", "likely_closed"}:
        return {"priority": "review", "label": "先确认再投", "reason": "岗位新鲜度需确认，先检查原链接是否仍开放。", "order": 2}
    if job.get("user_tag_mutes"):
        labels = "、".join(item.get("label") or item.get("id") or "" for item in job["user_tag_mutes"][:2])
        return {"priority": "review", "label": "先确认再投", "reason": f"命中少看标签{f'：{labels}' if labels else ''}，确认后再投。", "order": 2}
    if job.get("direction_mismatch_adjustment"):
        return {"priority": "review", "label": "先确认再投", "reason": "与当前求职方向偏离，确认仍值得投入时间。", "order": 2}
    if fit_score >= 4.0:
        return {"priority": "today", "label": "今天优先", "reason": f"方向一致，综合 {fit_score:.1f}，建议今天完成。", "order": 0}
    if queue_age >= 7 and fit_score >= 3.3:
        return {"priority": "today", "label": "今天优先", "reason": f"已在队列 {queue_age} 天，今天决定投递或移出。", "order": 0}
    if fit_score >= 3.3:
        return {"priority": "next", "label": "接下来", "reason": "匹配可投，处理完今天优先项后再看。", "order": 1}
    return {"priority": "review", "label": "先确认再投", "reason": f"综合 {fit_score:.1f}，先确认岗位要求和投入价值。", "order": 2}


def queue_job_sort_key(job: dict, reference_date: dt.date | None = None) -> tuple:
    decision = queue_decision(job, reference_date)
    deadline = application_deadline_status(job.get("application_deadline"), reference_date)
    status_order = {"today": 0, "urgent": 1, "soon": 2, "scheduled": 3, "unknown": 4, "expired": 5}
    remaining = deadline["days_remaining"] if deadline["days_remaining"] is not None else 999999
    if deadline["code"] == "expired":
        remaining = abs(remaining)
    try:
        updated = dt.datetime.fromisoformat(str(job.get("updated_at") or "")).timestamp()
    except ValueError:
        updated = 0
    return (
        decision["order"],
        status_order[deadline["code"]],
        remaining,
        -float(job.get("rank_score") or job.get("score") or 0),
        -updated,
    )


def deadline_recommendation_priority(job: dict, reference_date: dt.date | None = None) -> int:
    deadline = application_deadline_status(job.get("application_deadline"), reference_date)
    return 8 - int(deadline["days_remaining"] or 0) if deadline["code"] in {"today", "urgent", "soon"} else 0


def http_get(url: str, timeout: int = 25, retries: int = 1) -> str:
    request = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-SG,en;q=0.9,zh-CN;q=0.7,zh;q=0.6",
        },
    )
    last_exc: Exception | None = None
    for attempt in range(max(1, retries + 1)):
        try:
            with urllib.request.urlopen(request, timeout=timeout) as response:
                charset = response.headers.get_content_charset() or "utf-8"
                return response.read().decode(charset, errors="ignore")
        except Exception as exc:
            last_exc = exc
            if attempt < retries:
                time.sleep(0.8 * (attempt + 1))
    raise last_exc or RuntimeError(f"Failed to fetch {url}")


def http_post_json(url: str, payload: dict, timeout: int = 25) -> dict:
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124 Safari/537.36",
            "Accept": "application/json",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8", errors="ignore"))


def absolute_url(base: str, href: str) -> str:
    href = unescape(href or "").strip()
    return urljoin(base, href)


def canonical_job_url(source: str, url: str, external_job_id: str | None = None) -> str:
    parsed = urlparse(url)
    source_lowered = (source or "").lower()
    if source_lowered == "linkedin":
        job_id = external_job_id or ""
        if not job_id:
            match = re.search(r"(\d{7,})", parsed.path)
            job_id = match.group(1) if match else ""
        if job_id:
            return f"https://www.linkedin.com/jobs/view/{job_id}"
    if source_lowered == "internsg":
        return parsed._replace(query="", fragment="").geturl()
    if source_lowered == "indeed":
        query = parse_qs(parsed.query)
        job_key = (query.get("jk") or [""])[0]
        if job_key:
            return f"https://sg.indeed.com/viewjob?jk={job_key}"
        return parsed._replace(fragment="").geturl()
    if source_lowered == "jobstreet":
        match = re.search(r"/job/(\d+)", parsed.path)
        if match:
            return f"https://sg.jobstreet.com/job/{match.group(1)}"
        return parsed._replace(query="", fragment="").geturl()
    return parsed._replace(fragment="").geturl()


def read_resume_file(path: Path) -> str:
    if not path.exists():
        return ""
    if path.suffix.lower() == ".pdf":
        try:
            import fitz

            with fitz.open(str(path)) as doc:
                return "\n\n".join(page.get_text("text") for page in doc).strip()
        except Exception:
            try:
                from pypdf import PdfReader

                reader = PdfReader(str(path))
                return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
            except Exception:
                return ""
    if path.suffix.lower() == ".docx":
        try:
            from docx import Document

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs).strip()
        except Exception:
            return ""
    return path.read_text(encoding="utf-8", errors="ignore")


def read_resume_text() -> str:
    try:
        active_resume = get_active_resume_version()
        if active_resume and Path(active_resume["text_path"]).exists():
            text = Path(active_resume["text_path"]).read_text(encoding="utf-8", errors="ignore")
            if text.strip():
                return text
    except Exception:
        pass
    try:
        profile_resume = Path(load_profile().get("resume_path") or "")
    except Exception:
        profile_resume = current_resume_path()
    candidates = [
        profile_resume,
        current_resume_path(),
        current_default_resume_pdf(),
        current_resume_upload_dir() / "active-resume.md",
    ]
    candidates.extend(APP_DIR.parent.glob("*resume.md"))
    for path in candidates:
        if path.exists():
            text = read_resume_file(path)
            if text.strip():
                return text
    return FALLBACK_PROFILE_TEXT


def default_profile() -> dict:
    return {
        "full_name": "Your Name",
        "first_name": "Your",
        "last_name": "Name",
        "email": "you@example.com",
        "phone": "",
        "location": "",
        "school": "",
        "degree": "",
        "visa_status": "",
        "work_authorisation": "I will confirm work eligibility for each role before final submission.",
        "linkedin": "",
        "portfolio": "",
        "resume_path": str(current_resume_path()),
        "cover_letter_path": "",
        "availability": "Available for suitable internship, graduate, or early-career roles, subject to schedule and employer requirements.",
        "common_answers": [
            {
                "question": "Why are you interested in this role?",
                "answer": "I am interested in this role because it connects product, service, and human-centred design work with real user and business needs. I would like to bring my research, prototyping, and AI-assisted workflow experience to the team.",
            },
            {
                "question": "Do you require work sponsorship?",
                "answer": "I will confirm the exact work authorisation and arrangement before final submission.",
            },
        ],
    }


def load_profile() -> dict:
    ensure_dirs()
    ensure_cloud_state_loaded()
    profile = default_profile()
    profile_path = current_profile_path()
    if profile_path.exists():
        try:
            stored = json.loads(profile_path.read_text(encoding="utf-8"))
            if isinstance(stored, dict):
                profile.update(stored)
        except json.JSONDecodeError:
            pass
    return profile


def save_profile(payload: dict) -> dict:
    ensure_dirs()
    profile = load_profile()
    allowed = set(default_profile().keys())
    for key, value in payload.items():
        if key in allowed:
            profile[key] = value
    if not isinstance(profile.get("common_answers"), list):
        profile["common_answers"] = []
    current_profile_path().write_text(json.dumps(profile, ensure_ascii=False, indent=2), encoding="utf-8")
    return profile


def normalize_region(region: str | None) -> str:
    code = (region or "").strip().upper()
    aliases = {
        "SINGAPORE": "SG",
        "CHINA": "CN",
        "MAINLAND": "CN",
        "CHINA MAINLAND": "CN",
        "MAINLAND CHINA": "CN",
        "HONG KONG": "HK",
        "HONGKONG": "HK",
    }
    code = aliases.get(code, code)
    return code if code in REGION_CONFIGS else "SG"


def default_region_context(region: str) -> dict:
    code = normalize_region(region)
    config = REGION_CONFIGS[code]
    defaults = {
        "SG": {
            "work_authorisation": "Singapore work eligibility to be confirmed",
            "target_directions": [],
            "job_types": ["Internship", "Graduate", "Full-time"],
            "employment_priority": "both",
        },
        "CN": {
            "work_authorisation": "China mainland work eligibility to be confirmed",
            "target_directions": [],
            "job_types": ["Internship", "Graduate", "Full-time"],
            "employment_priority": "both",
        },
        "HK": {
            "work_authorisation": "Hong Kong work eligibility to be confirmed",
            "target_directions": [],
            "job_types": ["Internship", "Graduate", "Full-time"],
            "employment_priority": "both",
        },
    }[code]
    return {
        "region": code,
        "city": config["default_city"],
        "work_authorisation": defaults["work_authorisation"],
        "target_directions": list(defaults["target_directions"]),
        "job_types": list(defaults["job_types"]),
        "employment_priority": defaults["employment_priority"],
        "career_goal": "sg_internship_to_fulltime" if code == "SG" else "experience_first",
        "sponsorship_priority": "high" if code == "SG" else "medium",
        "language_preference": "chinese_friendly" if code == "SG" else "bilingual",
        "conversion_priority": "high" if code == "SG" else "medium",
        "preferred_company_groups": ["greater_china", "ai_startup", "sg_anchor"] if code == "SG" else [],
        "preferred_job_tags": [
            "internship",
            "conversion_strong",
            "conversion_possible",
            "visa_possible",
            "chinese_friendly",
            "company_greater_china",
            "company_ai_startup",
            "ai_related",
            "product_related",
        ] if code == "SG" else [],
        "muted_job_tags": ["visa_unlikely", "conversion_none", "high_experience"] if code == "SG" else [],
        "salary_currency": REGION_CURRENCIES.get(code, ""),
        "salary_period": "monthly",
        "salary_min": None,
        "salary_preferred": None,
        "company_focus": [],
        "exclude_keywords": [],
        "updated_at": now_iso(),
    }


def default_user_context() -> dict:
    return {
        "active_region": "SG",
        "contexts": {code: default_region_context(code) for code in REGION_CONFIGS},
        "onboarding_completed": False,
        "onboarding_step": 1,
        "resume_analyzed": False,
        "updated_at": now_iso(),
    }


def merge_user_context(stored: dict) -> dict:
    merged = default_user_context()
    if not isinstance(stored, dict):
        return merged
    merged["active_region"] = normalize_region(stored.get("active_region") or "SG")
    merged["onboarding_completed"] = bool(stored.get("onboarding_completed", merged["onboarding_completed"]))
    merged["resume_analyzed"] = bool(stored.get("resume_analyzed", merged["resume_analyzed"]))
    try:
        merged["onboarding_step"] = max(1, min(3, int(stored.get("onboarding_step") or merged["onboarding_step"])))
    except (TypeError, ValueError):
        merged["onboarding_step"] = 1
    contexts = stored.get("contexts") if isinstance(stored.get("contexts"), dict) else {}
    for code in REGION_CONFIGS:
        if isinstance(contexts.get(code), dict):
            base = merged["contexts"][code]
            for key in [
                "city",
                "work_authorisation",
                "target_directions",
                "job_types",
                "employment_priority",
                "career_goal",
                "sponsorship_priority",
                "language_preference",
                "conversion_priority",
                "preferred_company_groups",
                "preferred_job_tags",
                "muted_job_tags",
                "salary_currency",
                "salary_period",
                "salary_min",
                "salary_preferred",
                "company_focus",
                "exclude_keywords",
            ]:
                if key in contexts[code]:
                    base[key] = contexts[code][key]
            base["region"] = code
            base["updated_at"] = contexts[code].get("updated_at") or base["updated_at"]
    merged["updated_at"] = stored.get("updated_at") or merged["updated_at"]
    return merged


def load_user_context() -> dict:
    ensure_dirs()
    ensure_cloud_state_loaded()
    context_path = current_user_context_path()
    if not context_path.exists():
        return default_user_context()
    try:
        return merge_user_context(json.loads(context_path.read_text(encoding="utf-8")))
    except json.JSONDecodeError:
        return default_user_context()


def save_user_context(payload: dict) -> dict:
    ensure_dirs()
    context = load_user_context()
    active_region = normalize_region(payload.get("active_region") or payload.get("region") or context["active_region"])
    context["active_region"] = active_region
    target = context["contexts"][active_region]
    updates = payload.get("context") if isinstance(payload.get("context"), dict) else payload
    for key in ["city", "work_authorisation"]:
        if key in updates:
            target[key] = str(updates.get(key) or "").strip() or target[key]
    if "employment_priority" in updates:
        priority = str(updates.get("employment_priority") or "unspecified").strip()
        target["employment_priority"] = priority if priority in EMPLOYMENT_PRIORITY_VALUES else "unspecified"
    if "career_goal" in updates:
        goal = str(updates.get("career_goal") or "sg_internship_to_fulltime").strip()
        target["career_goal"] = goal if goal in CAREER_GOAL_VALUES else "sg_internship_to_fulltime"
    for key in ["sponsorship_priority", "conversion_priority"]:
        if key in updates:
            value = str(updates.get(key) or "unspecified").strip()
            target[key] = value if value in PRIORITY_LEVEL_VALUES else "unspecified"
    if "language_preference" in updates:
        value = str(updates.get("language_preference") or "unspecified").strip()
        target["language_preference"] = value if value in LANGUAGE_PREFERENCE_VALUES else "unspecified"
    if "salary_currency" in updates:
        target["salary_currency"] = str(updates.get("salary_currency") or "").strip().upper()
    if "salary_period" in updates:
        period = str(updates.get("salary_period") or "monthly").strip()
        target["salary_period"] = period if period in SALARY_PERIODS else "monthly"
    for key in ["salary_min", "salary_preferred"]:
        if key in updates:
            raw_value = str(updates.get(key) or "").replace(",", "").strip()
            if not raw_value:
                target[key] = None
            else:
                try:
                    target[key] = max(0.0, float(raw_value))
                except ValueError:
                    target[key] = None
    for key in ["target_directions", "job_types", "company_focus", "exclude_keywords", "preferred_company_groups", "preferred_job_tags", "muted_job_tags"]:
        if key in updates:
            values = updates.get(key) or []
            if isinstance(values, str):
                values = [item.strip() for item in re.split(r"[,，\n]+", values) if item.strip()]
            cleaned = [str(item).strip() for item in values if str(item).strip()]
            if key == "preferred_company_groups":
                cleaned = [item for item in cleaned if item in COMPANY_GROUP_VALUES]
            elif key in {"preferred_job_tags", "muted_job_tags"}:
                cleaned = [item for item in cleaned if item in USER_JOB_TAG_VALUES]
            target[key] = cleaned
    if "onboarding_completed" in payload:
        context["onboarding_completed"] = bool(payload.get("onboarding_completed"))
    if "resume_analyzed" in payload:
        context["resume_analyzed"] = bool(payload.get("resume_analyzed"))
    if "onboarding_step" in payload:
        try:
            context["onboarding_step"] = max(1, min(3, int(payload.get("onboarding_step") or 1)))
        except (TypeError, ValueError):
            context["onboarding_step"] = 1
    target["updated_at"] = now_iso()
    context["updated_at"] = target["updated_at"]
    current_user_context_path().write_text(json.dumps(context, ensure_ascii=False, indent=2), encoding="utf-8")
    return context


def mark_resume_analyzed() -> None:
    context = load_user_context()
    context["resume_analyzed"] = True
    context["onboarding_step"] = max(3, int(context.get("onboarding_step") or 1))
    context["updated_at"] = now_iso()
    current_user_context_path().write_text(json.dumps(context, ensure_ascii=False, indent=2), encoding="utf-8")


def active_region_code(region: str | None = None) -> str:
    return normalize_region(region or load_user_context().get("active_region"))


def active_region_context(region: str | None = None) -> dict:
    context = load_user_context()
    code = active_region_code(region or context.get("active_region"))
    return context["contexts"].get(code) or default_region_context(code)


def regions_payload() -> dict:
    catalog_counts = {code: 0 for code in REGION_CONFIGS}
    for item in COMPANY_CATALOG:
        catalog_counts[item["region"]] += 1
    return {
        "active_region": active_region_code(),
        "regions": [
            {
                **config,
                "catalog_count": catalog_counts.get(code, 0),
            }
            for code, config in REGION_CONFIGS.items()
        ],
    }


def profile_options_payload(region: str | None = None) -> dict:
    code = active_region_code(region)
    config = REGION_CONFIGS[code]
    currency = REGION_CURRENCIES.get(code, "")
    salary_bands = {key: list(value) for key, value in SALARY_BAND_OPTIONS.get(code, {}).items()}
    if "daily" not in salary_bands:
        salary_bands["daily"] = SALARY_BAND_OPTIONS["SG"]["daily"]
    if "hourly" not in salary_bands:
        salary_bands["hourly"] = SALARY_BAND_OPTIONS["SG"]["hourly"]
    return {
        "region": code,
        "label": config["label"],
        "cities": config["cities"],
        "default_city": config["default_city"],
        "city_required": code == "CN",
        "salary_currency": currency,
        "regions": [{"value": key, "label": value["label"]} for key, value in REGION_CONFIGS.items()],
        "work_authorisation_options": WORK_AUTH_OPTIONS.get(code, WORK_AUTH_OPTIONS["SG"]),
        "employment_priority_options": EMPLOYMENT_PRIORITY_OPTIONS,
        "career_goal_options": CAREER_GOAL_OPTIONS,
        "priority_level_options": PRIORITY_LEVEL_OPTIONS,
        "language_preference_options": LANGUAGE_PREFERENCE_OPTIONS,
        "job_tag_options": USER_JOB_TAG_OPTIONS,
        "company_group_options": [
            {"value": "greater_china", "label": "大中华/中文友好"},
            {"value": "sg_anchor", "label": "新加坡本地大厂"},
            {"value": "ai_startup", "label": "AI/高潜力初创"},
            {"value": "product_design", "label": "产品/设计"},
            {"value": "fintech", "label": "Fintech"},
            {"value": "service_brand", "label": "服务体验品牌"},
        ],
        "direction_options": [
            {"value": item["id"], "label": item["label"], "category": DIRECTION_CATEGORIES.get(item["id"], "其他方向")}
            for item in CAREER_DIRECTIONS
        ],
        "direction_categories": list(dict.fromkeys(DIRECTION_CATEGORIES.values())),
        "job_type_options": JOB_TYPE_OPTIONS,
        "salary_period_options": SALARY_PERIOD_OPTIONS,
        "salary_band_options": salary_bands,
    }


def json_list(value) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except json.JSONDecodeError:
            parsed = [item.strip() for item in re.split(r"[,，\n]+", value) if item.strip()]
        if isinstance(parsed, list):
            return [str(item).strip() for item in parsed if str(item).strip()]
    return []


@lru_cache(maxsize=4096)
def clean_company_name(value: str) -> str:
    cleaned = clean_text(value or "")
    return re.sub(r"\s+featured\s*$", "", cleaned, flags=re.I).strip()


@lru_cache(maxsize=4096)
def normalize_company_phrase(value: str) -> str:
    lowered = clean_company_name(value).lower().replace("&", " and ")
    lowered = re.sub(r"\b(private limited|pte ltd|pte\.?\s*ltd\.?|limited|ltd|inc|corp|corporation|co)\b\.?", " ", lowered)
    lowered = re.sub(r"[^a-z0-9]+", " ", lowered)
    return re.sub(r"\s+", " ", lowered).strip()


def company_alias_values(company: str, item: dict | None = None) -> list[str]:
    aliases: list[str] = []
    for value in [company, *(json_list((item or {}).get("aliases"))), *(json_list((item or {}).get("aliases_json")))]:
        if value and value not in aliases:
            aliases.append(value)
    for key in [company.lower(), normalize_company_phrase(company)]:
        for value in COMPANY_ALIAS_OVERRIDES.get(key, []):
            if value and value not in aliases:
                aliases.append(value)
    for suffix in [" Singapore", " SG", " Hong Kong", " China"]:
        if company.endswith(suffix):
            base = company[: -len(suffix)].strip()
            if len(base) > 2 and base.lower() not in {"ai"} and base not in aliases:
                aliases.append(base)
    return aliases or [company]


def company_match_terms(company: str, item: dict | None = None) -> list[str]:
    terms: list[str] = []
    for alias in company_alias_values(company, item):
        normalized = normalize_company_phrase(alias)
        if not normalized:
            continue
        if len(normalized) <= 2 and normalized not in {"ai"}:
            continue
        if normalized not in terms:
            terms.append(normalized)
    return sorted(terms, key=len, reverse=True)


def company_text_has_term(text: str, term: str) -> bool:
    normalized = f" {normalize_company_phrase(text)} "
    return f" {term} " in normalized


def normalize_job_title_for_dedupe(value: str) -> str:
    lowered = (value or "").lower().replace("&", " and ")
    lowered = re.sub(
        r"\((?:\s*(?:ai|llm|rag|genai|generative ai|machine learning)\s*[/,+&-]?\s*)+\)",
        " ",
        lowered,
    )
    lowered = re.sub(r"[^a-z0-9]+", " ", lowered)
    return re.sub(r"\s+", " ", lowered).strip()


COMPANY_DEDUPE_ALIASES = {
    "cpf board": "central provident fund board",
    "cyber security agency of singapore": "cyber security agency",
    "csa": "cyber security agency",
    "housing and development board": "hdb",
}


def job_dedupe_key(job: dict) -> str:
    company = normalize_company_phrase(job.get("company") or "")
    company = COMPANY_DEDUPE_ALIASES.get(company, company)
    position = normalize_job_title_for_dedupe(job.get("position") or "")
    region = active_region_code(job.get("region") or job.get("source_region") or None)
    city = normalize_company_phrase(job.get("city") or job.get("location") or "")
    generic_companies = {
        "",
        "unknown company",
        "linkedin company",
        "jobstreet company",
        "indeed company",
        "internsg company",
        "google jobs",
        "glints nodeflair startups",
    }
    if company in generic_companies or len(position) < 4:
        return f"url:{job.get('url') or job.get('id') or ''}"
    return "|".join([region, city, company, position])


def collapse_duplicate_job_groups(jobs: list[dict]) -> list[dict]:
    collapsed: list[dict] = []
    by_key: dict[str, dict] = {}
    for job in jobs:
        key = job_dedupe_key(job)
        existing = by_key.get(key)
        if existing is None:
            item = dict(job)
            item["dedupe_key"] = key
            item["duplicate_count"] = 0
            item["alternate_links"] = []
            item["alternate_sources"] = []
            item["source_count"] = 1
            collapsed.append(item)
            by_key[key] = item
            continue
        url = job.get("url") or ""
        if url and url != existing.get("url") and not any(link.get("url") == url for link in existing["alternate_links"]):
            existing["alternate_links"].append({
                "id": job.get("id"),
                "source": job.get("source") or "其它来源",
                "url": url,
            })
        source = job.get("source") or "其它来源"
        if source != existing.get("source") and source not in existing["alternate_sources"]:
            existing["alternate_sources"].append(source)
        existing["duplicate_count"] += 1
        existing["source_count"] = 1 + len(existing["alternate_sources"])
    return collapsed


def match_job_to_company(job: dict, company: str, item: dict | None = None) -> tuple[bool, str]:
    terms = company_match_terms(company, item)
    company_text = job.get("company") or ""
    for term in terms:
        if company_text_has_term(company_text, term):
            return True, f"公司名匹配：{term}"
    generic_company_names = {
        "",
        "unknown company",
        "linkedin company",
        "jobstreet company",
        "indeed company",
        "internsg company",
        "google jobs",
    }
    if normalize_company_phrase(company_text) not in generic_company_names:
        return False, ""
    broader_text = f"{job.get('name') or ''} {job.get('position') or ''} {(job.get('jd_text') or '')[:900]}"
    for term in terms:
        if len(term) >= 6 and company_text_has_term(broader_text, term):
            return True, f"岗位文本匹配：{term}"
    return False, ""


def company_source_label(source: str) -> str:
    source_text = source or ""
    for key, label in COMPANY_SCAN_SOURCE_LABELS.items():
        if key.lower() in source_text.lower():
            return label
    return f"{source_text or '公共来源'} 匹配"


def company_source_group(source: str) -> str:
    lowered = (source or "").lower()
    if "company site" in lowered or "ats" in lowered:
        return "official"
    if "linkedin" in lowered or "jobstreet" in lowered or "google jobs" in lowered:
        return "public_match"
    return "supplemental"


def company_item_from_row(row: dict) -> dict:
    out = dict(row)
    out["city_tags"] = out.get("city_tags") or out.get("city_tags_json") or []
    out["aliases"] = company_alias_values(out.get("company") or "", out)
    catalog_item = next(
        (
            dict(item)
            for item in COMPANY_CATALOG
            if item.get("region") == out.get("region") and item.get("company", "").lower() == str(out.get("company") or "").lower()
        ),
        None,
    )
    if catalog_item:
        for key in ["tags", "recommend_reason", "language_signal"]:
            if not out.get(key) and catalog_item.get(key):
                out[key] = catalog_item[key]
    out.update(company_pathway_profile({**(catalog_item or {}), **out}))
    return out


def company_catalog_item(company: str, region: str | None = None) -> dict | None:
    code = active_region_code(region)
    lowered = (company or "").lower()
    for item in COMPANY_CATALOG:
        if item["region"] == code and item["company"].lower() == lowered:
            return dict(item)
    return None


def company_pathway_profile(item: dict | None) -> dict:
    item = item or {}
    name = str(item.get("company") or "")
    tags = [str(tag) for tag in item.get("tags") or []]
    tag_text = " ".join(tags).lower()
    company_type = str(item.get("company_type") or item.get("source") or "Company")
    type_text = company_type.lower()
    name_text = name.lower()
    greater_china = bool(
        "中文友好" in tag_text
        or "greater china" in type_text
        or "出海" in tag_text
        or any(term in name_text for term in [
            "bytedance",
            "tiktok",
            "tencent",
            "trip.com",
            "alibaba",
            "ant ",
            "alipay",
            "huawei",
            "xiaomi",
            "oppo",
            "vivo",
            "shein",
            "pop mart",
            "lark",
            "patsnap",
            "wiz",
            "moomoo",
        ])
    )
    ai_startup = bool("ai startup" in type_text or ("ai" in tag_text and ("初创" in tag_text or "startup" in type_text)))
    sg_anchor = bool("本地标杆" in tag_text or any(term in name_text for term in ["grab", "dbs", "govtech", "singtel", "changi", "ncs", "sea", "shopee"]))
    fintech = bool("fintech" in tag_text or "fintech" in type_text or "bank" in type_text or "payments" in type_text)
    product_design = bool("产品" in tag_text or "设计" in tag_text or "product" in type_text or "design" in type_text)
    service_brand = bool("服务" in tag_text or "retail" in type_text or "service" in type_text or "travel" in type_text)
    if greater_china:
        company_group = "greater_china"
    elif ai_startup:
        company_group = "ai_startup"
    elif sg_anchor:
        company_group = "sg_anchor"
    elif fintech:
        company_group = "fintech"
    elif product_design:
        company_group = "product_design"
    elif service_brand:
        company_group = "service_brand"
    else:
        company_group = "other"
    language_signal = item.get("language_signal") or ("Chinese-friendly possible" if greater_china else "English first")
    sponsorship_signal = item.get("sponsorship_signal") or ("possible" if sg_anchor or greater_china or "MNC" in company_type else "unknown")
    conversion_signal = item.get("intern_to_fulltime_signal") or ("possible" if sg_anchor or "实习" in tag_text or "graduate" in type_text else "unknown")
    startup_stage = item.get("startup_stage") or ("high-potential startup" if ai_startup or "高潜力初创" in tag_text else "")
    return {
        "company_group": item.get("company_group") or company_group,
        "china_connection_level": item.get("china_connection_level") or ("high" if greater_china else "low"),
        "language_signal": language_signal,
        "sponsorship_signal": sponsorship_signal,
        "intern_to_fulltime_signal": conversion_signal,
        "startup_stage": startup_stage,
        "official_careers_url": item.get("official_careers_url") or item.get("url") or "",
    }


@lru_cache(maxsize=2048)
def company_catalog_match_for_job(company: str, region: str | None = None) -> dict | None:
    code = active_region_code(region)
    job_stub = {"company": company or "", "name": company or "", "position": "", "jd_text": ""}
    for item in COMPANY_CATALOG:
        if item.get("region") != code:
            continue
        ok, _reason = match_job_to_company(job_stub, item.get("company") or "", item)
        if ok:
            out = dict(item)
            out.update(company_pathway_profile(item))
            return out
    return None


def company_match_rows(region: str, city: str | None = None) -> list[dict]:
    code = active_region_code(region)
    query = "select * from jobs where region=?"
    values: list[str] = [code]
    if city and code == "CN":
        query += " and (city=? or location like ? or coalesce(city, '')='')"
        values.extend([city, f"%{city}%"])
    query += " order by updated_at desc limit 5000"
    with get_db() as conn:
        rows = conn.execute(query, values).fetchall()
    return [row_to_dict(row) for row in rows]


def matched_company_jobs(item: dict, jobs: list[dict], limit: int | None = None) -> list[dict]:
    company = item.get("company") or ""
    matches: list[dict] = []
    seen: set[str] = set()
    for job in jobs:
        if "company site" in (job.get("source") or "").lower():
            position = job.get("position") or ""
            if len(position) > 120 or not is_actionable_job_title(position) or not COMPANY_SCAN_TITLE_ROLE_PATTERN.search(position):
                continue
        ok, reason = match_job_to_company(job, company, item)
        if not ok:
            continue
        key = canonical_job_url(job.get("source") or "", job.get("url") or "", job.get("external_job_id")) or normalize_company_phrase(f"{job.get('company')} {job.get('position')}")
        if key in seen:
            continue
        seen.add(key)
        out = dict(job)
        out["company_match_reason"] = reason
        out["company_match_source_label"] = company_source_label(out.get("source") or "")
        out["company_match_source_group"] = company_source_group(out.get("source") or "")
        matches.append(out)
        if limit and len(matches) >= limit:
            break
    return matches


def company_scan_note(item: dict, matched_count: int) -> str:
    status = item.get("last_scan_status") or "not_scanned"
    note = item.get("last_scan_note") or ""
    if note:
        return note
    if matched_count:
        return "已从官网/ATS或公共来源匹配到岗位。"
    if status == "empty":
        return "官网未暴露可识别岗位列表，公共来源暂无匹配。"
    if status in {"failed", "limited"}:
        return "官网本次访问受限或失败，公共来源暂无匹配。"
    return "还没有扫描到可展示岗位。"


def enrich_company_items(items: list[dict], region: str, city: str | None = None) -> list[dict]:
    jobs = company_match_rows(region, city)
    enriched = []
    for item in items:
        out = company_item_from_row(item)
        out["dismissed"] = out.get("status") == "Dropped" or bool(out.get("dismissed"))
        if out["dismissed"]:
            out["aliases"] = company_alias_values(out.get("company") or "", out)
            out["matched_jobs_count"] = 0
            out["matched_official_count"] = 0
            out["last_scan_status"] = "hidden"
            out["last_scan_note"] = "已暂时隐藏这家公司岗位；重新关注后恢复。"
            enriched.append(out)
            continue
        matches = matched_company_jobs(out, jobs)
        out["aliases"] = company_alias_values(out.get("company") or "", out)
        out["matched_jobs_count"] = len(matches)
        out["matched_official_count"] = sum(1 for job in matches if job.get("company_match_source_group") == "official")
        out["last_scan_status"] = out.get("last_scan_status") or "not_scanned"
        out["last_scan_note"] = company_scan_note(out, len(matches))
        enriched.append(out)
    return enriched


def company_jobs_payload(company: str, region: str | None = None, city: str | None = None, company_id: int | None = None, limit: int = 40) -> dict:
    code = active_region_code(region)
    city_name = (city or active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]).strip()
    item: dict | None = None
    if company_id:
        with get_db() as conn:
            row = conn.execute("select * from watch_companies where id=?", (company_id,)).fetchone()
        if row:
            item = company_item_from_row(row_to_dict(row))
            code = active_region_code(item.get("region"))
    if item is None:
        item = company_catalog_item(company, code) or {"company": company, "region": code}
        with get_db() as conn:
            watched_row = conn.execute(
                "select * from watch_companies where region=? and lower(company)=lower(?)",
                (code, item.get("company") or company),
            ).fetchone()
        if watched_row:
            item = {**item, **company_item_from_row(row_to_dict(watched_row))}
    if item.get("status") == "Dropped":
        return {
            "company": item.get("company") or company,
            "region": code,
            "city": city_name,
            "aliases": company_alias_values(item.get("company") or company, item),
            "matched_jobs_count": 0,
            "jobs": [],
            "last_scan_status": "hidden",
            "last_scan_note": "你已选择暂时不关注这家公司；重新关注后才会显示它的岗位。",
        }
    rows = company_match_rows(code, city_name)
    company_item = company_item_from_row(item)
    all_matches = matched_company_jobs(company_item, rows)
    jobs = all_matches[:limit]
    return {
        "company": item.get("company") or company,
        "region": code,
        "city": city_name,
        "aliases": company_alias_values(item.get("company") or company, item),
        "matched_jobs_count": len(all_matches),
        "jobs": jobs,
        "last_scan_status": item.get("last_scan_status") or "not_scanned",
        "last_scan_note": company_scan_note(item, len(all_matches)),
    }


def watched_company_keys(region: str | None = None) -> set[str]:
    code = active_region_code(region)
    with get_db() as conn:
        rows = conn.execute(
            "select * from watch_companies where region=? and status='Watch'",
            (code,),
        ).fetchall()
    terms: set[str] = set()
    for row in rows:
        item = company_item_from_row(row_to_dict(row))
        terms.update(company_match_terms(item["company"], item))
    return terms


def dismissed_company_keys(region: str | None = None) -> set[str]:
    code = active_region_code(region)
    with get_db() as conn:
        rows = conn.execute(
            "select * from watch_companies where region=? and status='Dropped'",
            (code,),
        ).fetchall()
    terms: set[str] = set()
    for row in rows:
        item = company_item_from_row(row_to_dict(row))
        terms.update(company_match_terms(item["company"], item))
    return terms


def company_catalog(region: str | None = None, city: str | None = None) -> list[dict]:
    code = active_region_code(region)
    city_name = (city or active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]).strip()
    with get_db() as conn:
        company_rows = conn.execute(
            "select * from watch_companies where region=? and status in ('Watch', 'Dropped')",
            (code,),
        ).fetchall()
    company_states = {row["company"].lower(): row_to_dict(row) for row in company_rows}
    items = []
    for item in COMPANY_CATALOG:
        if item["region"] != code:
            continue
        out = dict(item)
        state_item = company_states.get(item["company"].lower())
        if state_item:
            out.update(state_item)
        out["watched"] = out.get("status") == "Watch"
        out["dismissed"] = out.get("status") == "Dropped"
        out["city_match"] = city_name in (item.get("city_tags") or [])
        items.append(out)
    items = enrich_company_items(items, code, city_name)
    items.sort(key=lambda value: (value.get("watched", False), not value.get("dismissed", False), value.get("city_match", False), int(value.get("priority") or 0)), reverse=True)
    return items


def validate_http_url(url: str) -> str:
    cleaned = (url or "").strip()
    parsed = urlparse(cleaned)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError("Please enter a valid http(s) URL.")
    return cleaned


def resume_version_to_dict(row: sqlite3.Row | dict | None) -> dict | None:
    if not row:
        return None
    out = dict(row)
    out["active"] = bool(out.get("active"))
    return out


def resume_analysis_to_dict(row: sqlite3.Row | dict | None) -> dict | None:
    if not row:
        return None
    out = row_to_dict(row) if isinstance(row, sqlite3.Row) else dict(row)
    return {
        "id": out.get("id"),
        "resume_version_id": out.get("resume_version_id"),
        "mode": out.get("mode"),
        "summary": out.get("summary") or "",
        "strengths": out.get("strengths_json") or [],
        "directions": out.get("directions_json") or [],
        "gaps": out.get("gaps_json") or [],
        "evidence": out.get("evidence_json") or [],
        "created_at": out.get("created_at"),
    }


def get_active_resume_version() -> dict | None:
    try:
        with get_db() as conn:
            row = conn.execute("select * from resume_versions where active=1 order by id desc limit 1").fetchone()
            return resume_version_to_dict(row)
    except sqlite3.Error:
        return None


def latest_resume_analysis(resume_version_id: int | None = None) -> dict | None:
    with get_db() as conn:
        if resume_version_id:
            row = conn.execute(
                "select * from resume_analyses where resume_version_id=? order by id desc limit 1",
                (resume_version_id,),
            ).fetchone()
        else:
            row = conn.execute("select * from resume_analyses order by id desc limit 1").fetchone()
    return resume_analysis_to_dict(row)


def active_resume_payload() -> dict:
    version = get_active_resume_version()
    if version:
        return version
    profile_path = Path(load_profile().get("resume_path") or current_resume_path())
    return {
        "id": None,
        "filename": profile_path.name,
        "original_filename": profile_path.name,
        "stored_path": str(profile_path),
        "text_path": "",
        "content_hash": "",
        "mime_type": "",
        "file_size": profile_path.stat().st_size if profile_path.exists() else 0,
        "active": True,
        "created_at": None,
    }


def set_active_resume_version(version_id: int) -> dict:
    stamp = now_iso()
    with get_db() as conn:
        row = conn.execute("select * from resume_versions where id=?", (version_id,)).fetchone()
        if not row:
            raise KeyError(f"Resume version {version_id} not found.")
        conn.execute("update resume_versions set active=0 where active=1")
        conn.execute("update resume_versions set active=1 where id=?", (version_id,))
    version = get_active_resume_version()
    if version:
        save_profile({"resume_path": version["stored_path"]})
    return version or resume_version_to_dict(row)


def ensure_resume_version_for_path(path: Path) -> dict:
    ensure_dirs()
    path = path.expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"Cannot find resume file: {path}")
    suffix = path.suffix.lower()
    if suffix not in {".pdf", ".docx", ".md", ".txt"}:
        raise ValueError("Resume must be a PDF, DOCX, MD, or TXT file.")
    content = path.read_bytes()
    content_hash = hashlib.sha256(content).hexdigest()
    text = read_resume_file(path)
    if not text.strip():
        raise ValueError("Could not extract readable text from this resume.")
    with get_db() as conn:
        existing = conn.execute("select * from resume_versions where stored_path=?", (str(path),)).fetchone()
        if existing:
            conn.execute("update resume_versions set active=0 where active=1")
            conn.execute("update resume_versions set active=1 where id=?", (existing["id"],))
            return resume_version_to_dict(conn.execute("select * from resume_versions where id=?", (existing["id"],)).fetchone())
        text_path = current_resume_upload_dir() / f"{path.stem}-{content_hash[:10]}.txt"
        text_path.write_text(text, encoding="utf-8")
        stamp = now_iso()
        conn.execute("update resume_versions set active=0 where active=1")
        conn.execute(
            """
            insert into resume_versions(
                filename, original_filename, stored_path, text_path, content_hash,
                mime_type, file_size, active, created_at
            )
            values(?, ?, ?, ?, ?, ?, ?, 1, ?)
            """,
            (path.name, path.name, str(path), str(text_path), content_hash, "", len(content), stamp),
        )
        version_id = conn.execute("select last_insert_rowid()").fetchone()[0]
    save_profile({"resume_path": str(path)})
    with get_db() as conn:
        return resume_version_to_dict(conn.execute("select * from resume_versions where id=?", (version_id,)).fetchone())


def save_uploaded_resume(original_filename: str, content: bytes, mime_type: str = "") -> dict:
    ensure_dirs()
    if not original_filename:
        raise ValueError("Resume filename is required.")
    suffix = Path(original_filename).suffix.lower()
    if suffix not in {".pdf", ".docx", ".md", ".txt"}:
        raise ValueError("Resume must be a PDF, DOCX, MD, or TXT file.")
    if not content:
        raise ValueError("Resume file is empty.")
    content_hash = hashlib.sha256(content).hexdigest()
    safe_name = sanitize_filename(Path(original_filename).stem) or "resume"
    filename = f"{dt.datetime.now().strftime('%Y%m%d-%H%M%S')}-{content_hash[:10]}-{safe_name}{suffix}"
    stored_path = current_resume_upload_dir() / filename
    stored_path.write_bytes(content)
    text = read_resume_file(stored_path)
    if not text.strip():
        stored_path.unlink(missing_ok=True)
        raise ValueError("Could not extract readable text from this resume.")
    text_path = stored_path.with_suffix(".txt")
    text_path.write_text(text, encoding="utf-8")
    stamp = now_iso()
    with get_db() as conn:
        conn.execute("update resume_versions set active=0 where active=1")
        conn.execute(
            """
            insert into resume_versions(
                filename, original_filename, stored_path, text_path, content_hash,
                mime_type, file_size, active, created_at
            )
            values(?, ?, ?, ?, ?, ?, ?, 1, ?)
            """,
            (filename, original_filename, str(stored_path), str(text_path), content_hash, mime_type, len(content), stamp),
        )
        version_id = conn.execute("select last_insert_rowid()").fetchone()[0]
    save_profile({"resume_path": str(stored_path)})
    return {"resume": set_active_resume_version(version_id), "analysis": analyze_resume_version(version_id, mode="local")}


def split_resume_sentences(text: str) -> list[str]:
    chunks = re.split(r"[\n\r]+|(?<=[.!?])\s+", text)
    return [re.sub(r"\s+", " ", chunk).strip() for chunk in chunks if len(chunk.strip()) >= 18]


def evidence_snippets(text: str, keywords: list[str], limit: int = 3) -> list[str]:
    snippets: list[str] = []
    for sentence in split_resume_sentences(text):
        lowered = sentence.lower()
        if any(has_keyword(lowered, keyword) for keyword in keywords):
            snippets.append(sentence[:260])
        if len(snippets) >= limit:
            break
    return snippets


def score_resume_direction(text: str, direction: dict) -> dict:
    lowered = text.lower()
    keyword_hits = [kw for kw in direction["keywords"] if has_keyword(lowered, kw)]
    evidence_hits = [kw for kw in direction["evidence"] if has_keyword(lowered, kw)]
    snippets = evidence_snippets(text, direction["keywords"] + direction["evidence"], limit=2)
    score = round(min(1.0, (len(keyword_hits) * 0.14) + (len(evidence_hits) * 0.18) + (0.16 if snippets else 0)), 2)
    return {
        "id": direction["id"],
        "label": direction["label"],
        "score": score,
        "matched_keywords": sorted(set(keyword_hits + evidence_hits))[:10],
        "evidence": snippets,
        "gaps": direction["gaps"],
        "source": "resume",
    }


def build_local_resume_analysis(text: str) -> dict:
    lowered = text.lower()
    directions = []
    evidence_items = []
    for direction in CAREER_DIRECTIONS:
        scored = score_resume_direction(text, direction)
        if scored["score"] >= MIN_RESUME_DIRECTION_SCORE:
            directions.append(scored)
        for snippet in scored["evidence"]:
            evidence_items.append({"direction_id": direction["id"], "direction": direction["label"], "text": snippet})
    directions.sort(key=lambda item: (item["score"], len(item["matched_keywords"])), reverse=True)

    strength_keywords = {
        "Human-centred/service design": ["human-centred", "human-centered", "service design", "service blueprint", "journey"],
        "UX research and synthesis": ["user research", "ux research", "interview", "research synthesis", "insight"],
        "Prototyping and visualisation": ["prototype", "figma", "visual design", "touchdesigner", "rhino", "arduino"],
        "AI-assisted workflows": ["ai-assisted", "prompt", "workflow automation", "jd/capability matching", "scenario exploration"],
        "Operations and coordination": ["operations", "on-site", "coordination", "workflow", "documentation"],
    }
    strengths = []
    for label, keywords in strength_keywords.items():
        hits = [kw for kw in keywords if has_keyword(lowered, kw)]
        if hits:
            strengths.append({"label": label, "evidence_terms": hits[:6], "snippets": evidence_snippets(text, keywords, limit=2)})

    top_labels = [item["label"] for item in directions[:3] if item["score"] > 0]
    if top_labels:
        summary_text = "This resume is strongest for " + ", ".join(top_labels) + "."
    else:
        summary_text = "This resume has readable content, but the current target-role signals are still weak."
    gaps = []
    for direction in directions[:3]:
        missing = [gap for gap in direction["gaps"] if not any(word in lowered for word in gap.lower().split())]
        if missing:
            gaps.append({"direction_id": direction["id"], "direction": direction["label"], "items": missing[:3]})

    return {
        "summary": summary_text,
        "strengths": strengths[:6],
        "directions": directions,
        "gaps": gaps[:5],
        "evidence": evidence_items[:12],
    }


def call_ai_resume_analysis(text: str, local_analysis: dict) -> dict:
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not configured. No resume text was sent out.")
    prompt = {
        "task": "Analyze this Singapore early-career resume for suitable job directions. Return concise JSON only.",
        "directions": [item["label"] for item in CAREER_DIRECTIONS],
        "local_baseline": local_analysis,
        "resume_text": text[:14000],
    }
    request = urllib.request.Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(
            {
                "model": OPENAI_MODEL,
                "input": [
                    {
                        "role": "system",
                        "content": "You are a truthful resume analyst. Do not invent facts. Return JSON with summary, strengths, directions, gaps, and evidence.",
                    },
                    {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)},
                ],
            }
        ).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=45) as response:
        data = json.loads(response.read().decode("utf-8"))
    output_text = data.get("output_text") or ""
    if not output_text:
        chunks = []
        for item in data.get("output", []):
            for part in item.get("content", []):
                if part.get("text"):
                    chunks.append(part["text"])
        output_text = "\n".join(chunks)
    try:
        parsed = json.loads(output_text)
    except json.JSONDecodeError:
        parsed = {"summary": output_text.strip()[:900]}
    merged = dict(local_analysis)
    for key in ["summary", "strengths", "directions", "gaps", "evidence"]:
        if parsed.get(key):
            merged[key] = parsed[key]
    return merged


def analyze_resume_version(resume_version_id: int | None = None, mode: str = "local") -> dict:
    version = None
    if resume_version_id:
        with get_db() as conn:
            version = resume_version_to_dict(conn.execute("select * from resume_versions where id=?", (resume_version_id,)).fetchone())
    if not version:
        version = get_active_resume_version()
    if not version:
        version = ensure_resume_version_for_path(Path(load_profile().get("resume_path") or current_resume_path()))
    text_path = Path(version["text_path"])
    text = text_path.read_text(encoding="utf-8", errors="ignore") if text_path.exists() else read_resume_file(Path(version["stored_path"]))
    if not text.strip():
        raise ValueError("Active resume has no readable text.")
    mode = mode if mode in {"local", "ai"} else "local"
    analysis = build_local_resume_analysis(text)
    if mode == "ai":
        analysis = call_ai_resume_analysis(text, analysis)
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            insert into resume_analyses(
                resume_version_id, mode, summary, strengths_json,
                directions_json, gaps_json, evidence_json, created_at
            )
            values(?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                version["id"],
                mode,
                analysis.get("summary") or "",
                json.dumps(analysis.get("strengths") or [], ensure_ascii=False),
                json.dumps(analysis.get("directions") or [], ensure_ascii=False),
                json.dumps(analysis.get("gaps") or [], ensure_ascii=False),
                json.dumps(analysis.get("evidence") or [], ensure_ascii=False),
                stamp,
            ),
        )
        analysis_id = conn.execute("select last_insert_rowid()").fetchone()[0]
        row = conn.execute("select * from resume_analyses where id=?", (analysis_id,)).fetchone()
    mark_resume_analyzed()
    return resume_analysis_to_dict(row)


def get_career_preferences() -> dict:
    stamp = now_iso()
    with get_db() as conn:
        row = conn.execute("select * from career_preferences where id=1").fetchone()
        if not row:
            conn.execute(
                """
                insert into career_preferences(
                    id, selected_directions_json, direction_weights_json,
                    exclude_keywords_json, updated_at
                )
                values(1, '[]', '{}', '[]', ?)
                """,
                (stamp,),
            )
            row = conn.execute("select * from career_preferences where id=1").fetchone()
    out = row_to_dict(row)
    return {
        "selected_directions": out.get("selected_directions_json") or [],
        "direction_weights": out.get("direction_weights_json") or {},
        "exclude_keywords": out.get("exclude_keywords_json") or [],
        "updated_at": out.get("updated_at"),
    }


def save_career_preferences(payload: dict) -> dict:
    valid_ids = {item["id"] for item in CAREER_DIRECTIONS}
    selected = [item for item in payload.get("selected_directions", []) if item in valid_ids]
    weights = payload.get("direction_weights") or {}
    weights = {key: float(value) for key, value in weights.items() if key in valid_ids and isinstance(value, (int, float))}
    excludes = [str(item).strip().lower() for item in payload.get("exclude_keywords", []) if str(item).strip()]
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            insert into career_preferences(
                id, selected_directions_json, direction_weights_json,
                exclude_keywords_json, updated_at
            )
            values(1, ?, ?, ?, ?)
            on conflict(id) do update set
                selected_directions_json=excluded.selected_directions_json,
                direction_weights_json=excluded.direction_weights_json,
                exclude_keywords_json=excluded.exclude_keywords_json,
                updated_at=excluded.updated_at
            """,
            (
                json.dumps(selected, ensure_ascii=False),
                json.dumps(weights, ensure_ascii=False),
                json.dumps(excludes, ensure_ascii=False),
                stamp,
            ),
        )
    return get_career_preferences()


def career_fit() -> dict:
    active_resume = active_resume_payload()
    analysis = latest_resume_analysis(active_resume.get("id")) if active_resume.get("id") else latest_resume_analysis()
    preferences = get_career_preferences()
    suggested = analysis["directions"][:6] if analysis else []
    return {
        "active_resume": active_resume,
        "analysis": analysis,
        "suggested_directions": suggested,
        "selected_directions": preferences["selected_directions"],
        "direction_weights": preferences["direction_weights"],
        "exclude_keywords": preferences["exclude_keywords"],
        "all_directions": [
            {"id": item["id"], "label": item["label"], "keywords": item["keywords"], "category": DIRECTION_CATEGORIES.get(item["id"], "其他方向")}
            for item in CAREER_DIRECTIONS
        ],
        "resume_analyzed": bool(analysis),
        "ai_available": bool(os.environ.get("OPENAI_API_KEY")),
    }


def hard_flag_patterns(text: str) -> list[str]:
    flags: list[str] = []
    lowered = text.lower()
    checks = [
        (
            "citizen_or_pr_only",
            r"\b(singaporeans?\s+only|singapore\s+citizens?\s+only|pr\s+only|permanent\s+residents?\s+only|only\s+(singaporeans?|singapore\s+citizens?|prs?|permanent\s+residents?)|must\s+be\s+(a\s+)?(singaporean|singapore\s+citizen|pr|permanent\s+resident)|requires?\s+(singaporean|singapore\s+citizen|pr\s+status|permanent\s+resident))\b",
        ),
        ("local_only", r"\b(local candidates? only|locals? only|only singaporeans)\b"),
        ("clearance_required", r"\b(security clearance|government clearance|clearance required)\b"),
        ("experience_too_high", r"\b([3-9]|\d{2,})\+?\s*(years?|yrs?)\b"),
    ]
    for flag, pattern in checks:
        if re.search(pattern, lowered):
            flags.append(flag)
    if re.search(r"\b(work authorization|work authorisation|visa sponsorship|sponsorship)\b", lowered):
        flags.append("visa_unclear")
    if re.search(r"\b(captcha|login required|answer the following questions)\b", lowered):
        flags.append("custom_questions")
    return sorted(set(flags))


def detect_employment_type(position: str, jd_text: str = "", job_type: str = "") -> str:
    title = (position or "").lower()
    type_text = (job_type or "").lower()
    jd_lower = (jd_text or "").lower()
    text = f"{title}\n{type_text}\n{jd_lower}"
    senior_title = re.search(r"\b(principal|lead|senior|staff|manager|director|head)\b", title)
    high_experience = re.search(r"\b([3-9]|\d{2,})\+?\s*(years?|yrs?)\b", jd_lower)

    if re.search(r"\b(intern|internship)\b|实习|實習", title):
        return "Internship"
    if re.search(r"\b(graduate|graduate programme|graduate program|new grad|fresh graduate|management associate|graduate trainee|trainee)\b|校招|应届|應屆", title):
        return "Graduate"
    if re.search(r"\b(contract|contractor|temporary|temp|fixed[-\s]?term|freelance|part[-\s]?time)\b|兼职|合约|合約", title):
        return "Contract"
    if re.search(r"\b(full[-\s]?time|permanent)\b|正式|全职|全職", title) or (senior_title and high_experience):
        return "Full-time"

    explicit_jd_type = re.search(
        r"\bemployment\s+type\s*[:\-]?\s*(intern(?:ship)?|graduate|contract(?:or)?|temporary|full[-\s]?time|permanent|part[-\s]?time)\b",
        jd_lower,
    )
    if explicit_jd_type:
        explicit_value = explicit_jd_type.group(1)
        if explicit_value.startswith("intern"):
            return "Internship"
        if explicit_value == "graduate":
            return "Graduate"
        if explicit_value in {"full-time", "full time", "permanent"}:
            return "Full-time"
        return "Contract"

    type_matches = set()
    if re.search(r"\b(intern|internship)\b|实习|實習", type_text):
        type_matches.add("Internship")
    if re.search(r"\b(graduate|new grad|fresh graduate)\b|校招|应届|應屆", type_text):
        type_matches.add("Graduate")
    if re.search(r"\b(contract|contractor|temporary|temp|fixed[-\s]?term|freelance|part[-\s]?time)\b|兼职|合约|合約", type_text):
        type_matches.add("Contract")
    if re.search(r"\b(full[-\s]?time|permanent)\b|正式|全职|全職", type_text):
        type_matches.add("Full-time")
    if len(type_matches) == 1 and not senior_title:
        return next(iter(type_matches))
    if re.search(r"\b(graduate|graduate programme|graduate program|new grad|fresh graduate|management associate|graduate trainee)\b|校招|应届|應屆", text):
        return "Graduate"
    if re.search(r"\b(contract|contractor|temporary|temp|fixed[-\s]?term|freelance|part[-\s]?time)\b|兼职|合约|合約", text):
        return "Contract"
    if re.search(r"\b(full[-\s]?time|permanent)\b|正式|全职|全職", text):
        return "Full-time"
    return "Unknown"


def detect_conversion_opportunity(position: str, jd_text: str = "", job_type: str = "") -> bool:
    return detect_conversion_signal(position, jd_text, job_type)[0] in {"strong", "possible"}


def detection_evidence(text: str, patterns: list[str], label: str, limit: int = 2) -> list[str]:
    evidence: list[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.I):
            start, end = match.span()
            snippet = clean_text(text[max(0, start - 42): min(len(text), end + 58)])
            if snippet and snippet not in evidence:
                evidence.append(f"{label}: {snippet}")
            if len(evidence) >= limit:
                return evidence
    return evidence


def detect_conversion_signal(position: str, jd_text: str = "", job_type: str = "") -> tuple[str, list[str]]:
    text = f"{position}\n{job_type}\n{jd_text}".lower()
    if re.search(r"\b(no|not|without)\s+(guarantee|guaranteed|promise|possibility)?\s*(of\s+)?(conversion|return offer|full[-\s]?time offer)", text):
        return "none", detection_evidence(text, [r"\b(no|not|without)\s+(guarantee|guaranteed|promise|possibility)?\s*(of\s+)?(conversion|return offer|full[-\s]?time offer)[^.。\n]*"], "转正风险")
    strong_patterns = [
        r"\bconvert(?:ed|ible|sion)?\s+(?:to|into)?\s*(?:a\s+)?full[-\s]?time\b",
        r"\bfull[-\s]?time\s+conversion\b",
        r"\breturn\s+offer\b",
        r"\bfull[-\s]?time\s+offer\b",
        r"\bpermanent\s+conversion\b",
        r"(转正|轉正|留用|留任)",
    ]
    possible_patterns = [
        r"\bgraduate\s+(programme|program|pipeline|scheme)\b",
        r"\bmanagement\s+associate\b",
        r"\btrainee\s+(programme|program)\b",
        r"\bpotential\s+(?:for\s+)?(?:full[-\s]?time|permanent)\b",
        r"\bmay\s+lead\s+to\s+(?:a\s+)?full[-\s]?time\b",
        r"(校招|管培|毕业生项目|畢業生項目)",
    ]
    strong = detection_evidence(text, strong_patterns, "可转正")
    if strong:
        return "strong", strong
    possible = detection_evidence(text, possible_patterns, "可能转正")
    if possible:
        return "possible", possible
    return "unknown", []


def detect_visa_sponsorship_signal(position: str, jd_text: str = "", job_type: str = "", region: str | None = None) -> tuple[str, list[str]]:
    text = f"{position}\n{job_type}\n{jd_text}".lower()
    negative_patterns = [
        r"\b(no|not|without)\s+(visa|work\s+pass|work\s+permit|employment\s+pass|s\s+pass|sponsorship)[^.。\n]*",
        r"\b(no|not|without)\s+(sponsorship|sponsor)[^.。\n]*",
        r"\b(visa|work\s+pass|work\s+permit|employment\s+pass|s\s+pass)\s+sponsorship\s+(?:is\s+)?(?:not|unavailable)[^.。\n]*",
        r"\b(?:do(?:es)?\s+not|will\s+not|cannot|can't|won't|unable\s+to)\s+(?:provide|offer|support|sponsor)[^.。\n]*(?:visa|work\s+pass|work\s+permit|employment\s+pass|s\s+pass|sponsorship)[^.。\n]*",
        r"\b(?:must|need(?:s)?\s+to)\s+(?:already\s+)?(?:have|hold)[^.。\n]*(?:right\s+to\s+work|valid\s+work\s+authori[sz]ation)[^.。\n]*",
        r"\b(singaporeans?\s+only|singapore\s+citizens?\s+only|pr\s+only|permanent\s+residents?\s+only)[^.。\n]*",
        r"\bmust\s+be\s+(a\s+)?(singaporean|singapore\s+citizen|pr|permanent\s+resident)[^.。\n]*",
    ]
    positive_patterns = [
        r"\b(visa|work\s+pass|work\s+permit|employment\s+pass|ep|s\s+pass)\s+(sponsor|sponsorship|support|provided|available)[^.。\n]*",
        r"\b(sponsor|sponsorship|support)\s+(visa|work\s+pass|work\s+permit|employment\s+pass|ep|s\s+pass)[^.。\n]*",
        r"\bopen\s+to\s+(international|foreign)\s+(candidates|students|applicants)[^.。\n]*",
        r"\bwork\s+authori[sz]ation\s+(support|sponsorship)[^.。\n]*",
        r"(工签|工作准证|就业准证|EP|S Pass).{0,40}(支持|担保|可办|sponsor)",
    ]
    negative = detection_evidence(text, negative_patterns, "工签风险")
    if negative:
        return "unlikely", negative
    positive = detection_evidence(text, positive_patterns, "工签可能")
    if positive:
        return "possible", positive
    if re.search(r"\b(work authorization|work authorisation|visa sponsorship|sponsorship|employment pass|s pass)\b", text):
        return "unclear", detection_evidence(text, [r"\b(work authorization|work authorisation|visa sponsorship|sponsorship|employment pass|s pass)[^.。\n]*"], "工签待确认", 1)
    return "unknown", []


def detect_language_signal(position: str, jd_text: str = "", job_type: str = "", company: str = "") -> tuple[str, list[str]]:
    text = f"{company}\n{position}\n{job_type}\n{jd_text}".lower()
    chinese_patterns = [
        r"\b(mandarin|chinese|putonghua|bilingual\s+.*?(english|mandarin|chinese)|china\s+market|greater\s+china|apac\s+china)[^.。\n]*",
        r"(中文|普通话|普通話|华语|華語|中国市场|中國市場|大中华|大中華|双语|雙語).{0,60}",
    ]
    english_patterns = [
        r"\b(excellent|strong|fluent)\s+english\s+(communication|skills)[^.。\n]*",
        r"\benglish[-\s]?first\b",
    ]
    chinese = detection_evidence(text, chinese_patterns, "中文友好可能")
    if chinese:
        return "chinese_friendly_possible", chinese
    english = detection_evidence(text, english_patterns, "英文为主", 1)
    if english:
        return "english_first", english
    return "unknown", []


def pathway_score_from_signals(
    employment_type: str,
    conversion_signal: str,
    visa_sponsorship_signal: str,
    language_signal: str,
    company_profile: dict | None = None,
) -> float:
    score = 1.8
    if employment_type == "Internship":
        score += 0.9
    elif employment_type == "Graduate":
        score += 0.65
    elif employment_type == "Full-time":
        score += 0.25
    elif employment_type == "Contract":
        score -= 0.15
    score += {"strong": 0.8, "possible": 0.45, "unknown": 0.0, "none": -0.5}.get(conversion_signal, 0.0)
    score += {"possible": 0.65, "unclear": 0.1, "unknown": 0.0, "unlikely": -0.85}.get(visa_sponsorship_signal, 0.0)
    score += {"chinese_friendly_possible": 0.35, "english_first": 0.0, "unknown": 0.0}.get(language_signal, 0.0)
    profile = company_profile or {}
    if profile.get("china_connection_level") == "high":
        score += 0.25
    if profile.get("company_group") in {"sg_anchor", "ai_startup", "greater_china"}:
        score += 0.18
    if profile.get("sponsorship_signal") == "possible" and visa_sponsorship_signal != "unlikely":
        score += 0.2
    if profile.get("intern_to_fulltime_signal") == "possible" and conversion_signal == "unknown":
        score += 0.16
    return round(max(0.0, min(5.0, score)), 2)


def normalize_salary_currency(value: str) -> str:
    token = (value or "").upper().replace(" ", "")
    if token in {"S$", "SG$", "SGD", "$"}:
        return "SGD"
    if token in {"HK$", "HKD"}:
        return "HKD"
    if token in {"RMB", "CNY", "¥"}:
        return "CNY"
    if token in {"US$", "USD"}:
        return "USD"
    return token or "SGD"


def normalize_salary_period(text: str, amount: float | None = None) -> str:
    lowered = (text or "").lower()
    if re.search(r"\b(per\s+annum|annually|annual|yearly|per\s+year|p\.?\s*a\.?|/year|/yr)\b", lowered):
        return "yearly"
    if re.search(r"\b(monthly|per\s+month|a\s+month|/month|/mo|pm)\b", lowered):
        return "monthly"
    if re.search(r"\b(daily|per\s+day|a\s+day|/day|pd)\b", lowered):
        return "daily"
    if re.search(r"\b(hourly|per\s+hour|an\s+hour|/hour|/hr|ph)\b", lowered):
        return "hourly"
    if amount and amount >= 30000:
        return "yearly"
    if amount and 500 <= amount <= 20000:
        return "monthly"
    if amount and amount < 80:
        return "hourly"
    return "unknown"


def parse_salary_number(value: str, suffix: str = "") -> float:
    amount = float((value or "0").replace(",", ""))
    if suffix.lower() == "k":
        amount *= 1000
    return amount


def salary_match_context(text: str, start: int, end: int) -> str:
    sentence_start = max(text.rfind(mark, 0, start) for mark in [".", "!", "?", ";", "\n"]) + 1
    sentence_ends = [text.find(mark, end) for mark in [".", "!", "?", ";", "\n"]]
    sentence_ends = [value for value in sentence_ends if value >= 0]
    sentence_end = min(sentence_ends) if sentence_ends else len(text)
    return clean_text(text[max(sentence_start, start - 80): min(sentence_end, end + 80)])


def is_salary_context(context: str) -> bool:
    lowered = context.lower()
    business_metric = re.search(
        r"\b(raised|funding|funded|financing|investors?|investment|invested|valuation|valued|"
        r"revenue|sales|assets?|portfolio\s+value|market\s+cap|worth|fund\s+of|budget|spend\s+on)\b|"
        r"\bclients?\b.{0,50}\bpay(?:ing|s|ed)?\b|"
        r"\b(million|billion|mn|bn)\b",
        lowered,
    )
    if business_metric:
        return False
    return bool(re.search(
        r"\b(salary|pay|compensation|stipend|remuneration|allowance|wage|earnings?)\b|"
        r"\b(monthly|hourly|yearly|annually|daily|per\s+(?:month|year|annum|hour|day)|"
        r"a\s+(?:month|year|hour|day))\b|/\s*(?:mo(?:nth)?|yr|year|hr|hour|day)\b|"
        r"薪资|薪酬|月薪|年薪|时薪|津贴",
        lowered,
    ))


def parse_salary_info(position: str, jd_text: str = "", job_type: str = "", region: str | None = None) -> dict:
    text = clean_text(f"{position}\n{job_type}\n{jd_text}")
    period_pattern = r"(?:/\s*(?:mo(?:nth)?|yr|year|hr|hour|day)|per\s+(?:month|year|annum|hour|day)|a\s+(?:month|year|hour|day))"
    pattern = re.compile(
        r"(?P<currency>SGD|S\$|SG\$|HKD|HK\$|RMB|CNY|USD|US\$|\$|¥)\s*"
        r"(?P<first>\d[\d,]*(?:\.\d+)?)\s*(?P<first_suffix>[kK])?"
        rf"(?P<first_period>\s*{period_pattern})?"
        r"(?:\s*(?:-|–|—|~|to|至|到)\s*"
        r"(?:(?:SGD|S\$|SG\$|HKD|HK\$|RMB|CNY|USD|US\$|\$|¥)\s*)?"
        r"(?P<second>\d[\d,]*(?:\.\d+)?)\s*(?P<second_suffix>[kK])?"
        rf"(?P<second_period>\s*{period_pattern})?)?",
        flags=re.I,
    )
    for match in pattern.finditer(text):
        start, end = match.span()
        window = salary_match_context(text, start, end)
        if not is_salary_context(window):
            continue
        first_suffix = match.group("first_suffix") or ""
        second_suffix = match.group("second_suffix") or ""
        if match.group("second") and not first_suffix and second_suffix:
            first_suffix = second_suffix
        if match.group("second") and first_suffix and not second_suffix:
            second_suffix = first_suffix
        first = parse_salary_number(match.group("first"), first_suffix)
        second = parse_salary_number(match.group("second"), second_suffix) if match.group("second") else first
        salary_min = min(first, second)
        salary_max = max(first, second)
        if salary_max <= 0:
            continue
        currency = normalize_salary_currency(match.group("currency"))
        period = normalize_salary_period(window, salary_max)
        if period == "monthly" and (salary_max < 300 or (salary_min < 100 and salary_max / max(salary_min, 1) > 20)):
            continue
        if period == "yearly" and salary_max < 10000:
            continue
        if period == "hourly" and salary_max > 500:
            continue
        return {
            "salary_min": salary_min,
            "salary_max": salary_max,
            "salary_currency": currency,
            "salary_period": period,
            "salary_text": clean_text(window),
            "salary_fit": "unknown",
        }
    return {
        "salary_min": None,
        "salary_max": None,
        "salary_currency": "",
        "salary_period": "unknown",
        "salary_text": "",
        "salary_fit": "unknown",
    }


def salary_to_monthly(amount: float | int | None, period: str | None) -> float | None:
    if amount is None:
        return None
    value = float(amount)
    if value <= 0:
        return None
    period = period or "unknown"
    if period == "monthly":
        return value
    if period == "yearly":
        return value / 12
    if period == "daily":
        return value * 22
    if period == "hourly":
        return value * 160
    return value


def job_metadata(position: str, jd_text: str = "", job_type: str = "", region: str | None = None, company: str = "", source: str = "") -> dict:
    salary = parse_salary_info(position, jd_text, job_type, region)
    employment_type = detect_employment_type(position, jd_text, job_type)
    conversion_signal, conversion_evidence = detect_conversion_signal(position, jd_text, job_type)
    visa_signal, visa_evidence = detect_visa_sponsorship_signal(position, jd_text, job_type, region)
    language_signal, language_evidence = detect_language_signal(position, jd_text, job_type, company)
    company_profile = company_catalog_match_for_job(company, region) if company else None
    if language_signal == "unknown" and company_profile:
        company_language = str(company_profile.get("language_signal") or "").lower()
        if "chinese" in company_language:
            language_signal = "chinese_friendly_possible"
            language_evidence.append(f"公司信号: {company_profile.get('language_signal')}")
        elif "english" in company_language:
            language_signal = "english_first"
    evidence = [*conversion_evidence, *visa_evidence, *language_evidence]
    if company_profile:
        evidence.append(f"公司分组: {company_profile.get('company_group')} / {company_profile.get('sponsorship_signal')}")
    pathway_score = pathway_score_from_signals(employment_type, conversion_signal, visa_signal, language_signal, company_profile)
    return {
        "employment_type": employment_type,
        "conversion_opportunity": 1 if conversion_signal in {"strong", "possible"} else 0,
        "conversion_signal": conversion_signal,
        "visa_sponsorship_signal": visa_signal,
        "language_signal": language_signal,
        "pathway_score": pathway_score,
        "pathway_evidence_json": evidence[:8],
        **salary,
    }


def backfill_job_metadata(conn: sqlite3.Connection) -> None:
    rows = conn.execute(
        """
        select id, company, position, source, jd_text, job_type, region, employment_type, salary_fit,
               conversion_signal, visa_sponsorship_signal, language_signal, pathway_evidence_json
        from jobs
        where employment_type is null
           or salary_period is null
           or salary_fit is null
           or salary_fit = ''
           or conversion_signal is null
           or conversion_signal = ''
           or visa_sponsorship_signal is null
           or visa_sponsorship_signal = ''
           or language_signal is null
           or language_signal = ''
           or pathway_evidence_json is null
           or pathway_evidence_json = ''
           or (
               salary_min is not null
               and (
                   lower(salary_text) like '% billion%'
                   or lower(salary_text) like '% million%'
                   or lower(salary_text) like '%raised%'
                   or lower(salary_text) like '%revenue%'
                   or lower(salary_text) like '%annual sales%'
                   or lower(salary_text) like '%investment%'
                   or lower(salary_text) like '%assets%'
                   or lower(salary_text) like '%portfolio value%'
                   or lower(salary_text) like '%fund of%'
                   or lower(salary_text) like '%spend on%'
                   or (lower(salary_text) like '%client%' and lower(salary_text) like '%paying%')
                   or (salary_min < 1000 and salary_max >= 30000)
                   or (salary_period = 'monthly' and salary_max < 300)
                   or (salary_period = 'monthly' and salary_min < 100 and salary_max / max(salary_min, 1) > 20)
                   or (salary_period = 'yearly' and salary_max < 10000)
                   or (salary_period = 'hourly' and salary_max > 500)
                   or (salary_min < 100 and lower(salary_text) like '%a month%')
               )
           )
           or (
               visa_sponsorship_signal != 'unlikely'
               and (
                   lower(jd_text) like '%visa sponsorship not available%'
                   or lower(jd_text) like '%visa sponsorship is not available%'
                   or lower(jd_text) like '%do not provide visa sponsorship%'
                   or lower(jd_text) like '%does not provide visa sponsorship%'
                   or lower(jd_text) like '%will not sponsor%'
                   or lower(jd_text) like '%cannot sponsor%'
                   or (lower(jd_text) like '%must already have%' and lower(jd_text) like '%right to work%')
               )
           )
        limit 1000
        """
    ).fetchall()
    for row in rows:
        metadata = job_metadata(row["position"] or "", row["jd_text"] or "", row["job_type"] or "", row["region"], row["company"] or "", row["source"] or "")
        conn.execute(
            """
            update jobs set
                employment_type=?,
                conversion_opportunity=?,
                salary_min=?,
                salary_max=?,
                salary_currency=?,
                salary_period=?,
                salary_text=?,
                salary_fit=?,
                conversion_signal=?,
                visa_sponsorship_signal=?,
                language_signal=?,
                pathway_score=?,
                pathway_evidence_json=?,
                updated_at=coalesce(updated_at, ?)
            where id=?
            """,
            (
                metadata["employment_type"],
                metadata["conversion_opportunity"],
                metadata["salary_min"],
                metadata["salary_max"],
                metadata["salary_currency"],
                metadata["salary_period"],
                metadata["salary_text"],
                metadata["salary_fit"],
                metadata["conversion_signal"],
                metadata["visa_sponsorship_signal"],
                metadata["language_signal"],
                metadata["pathway_score"],
                json.dumps(metadata["pathway_evidence_json"], ensure_ascii=False),
                now_iso(),
                row["id"],
            ),
        )


def backfill_application_deadlines(conn: sqlite3.Connection) -> None:
    rows = conn.execute(
        "select id, jd_text from jobs where coalesce(application_deadline, '') = '' and jd_text != ''"
    ).fetchall()
    for row in rows:
        deadline = extract_application_deadline(row["jd_text"] or "")
        if deadline:
            conn.execute("update jobs set application_deadline=? where id=?", (deadline, row["id"]))


def keyword_score(text: str, keywords: list[str], full_points: int) -> float:
    lowered = text.lower()
    hits = sum(1 for kw in keywords if has_keyword(lowered, kw))
    return min(1.0, hits / max(1, full_points))


def evidence_overlap_score(jd_text: str, resume_text: str, keywords: list[str], full_points: int) -> float:
    jd_lowered = jd_text.lower()
    resume_lowered = resume_text.lower()
    hits = sum(1 for kw in keywords if has_keyword(jd_lowered, kw) and has_keyword(resume_lowered, kw))
    return min(1.0, hits / max(1, full_points))


def contains_bounded_ascii_word(text: str, keyword: str) -> bool:
    start = 0
    while True:
        position = text.find(keyword, start)
        if position < 0:
            return False
        end = position + len(keyword)
        before_is_word = position > 0 and (text[position - 1].isalnum() or text[position - 1] == "_")
        after_is_word = end < len(text) and (text[end].isalnum() or text[end] == "_")
        if not before_is_word and not after_is_word:
            return True
        start = position + 1


def has_keyword(text: str, keyword: str) -> bool:
    keyword = keyword.lower().strip()
    if not keyword:
        return False
    if keyword.isascii() and keyword.isalnum():
        return contains_bounded_ascii_word(text, keyword)
    if re.fullmatch(r"[a-z0-9+#-]+", keyword):
        return re.search(rf"\b{re.escape(keyword)}\b", text) is not None
    return keyword in text


def score_job(company: str, position: str, jd_text: str, source: str) -> tuple[float, list[str], str]:
    resume = read_resume_text().lower()
    combined = f"{company}\n{position}\n{jd_text}\n{source}"
    lowered = combined.lower()

    role_keywords = [
        "ux",
        "user research",
        "service design",
        "product design",
        "product manager",
        "product operations",
        "experience design",
        "design intern",
        "content",
        "operations",
        "healthcare",
        "prototype",
        "figma",
        "human-centred",
        "human-centered",
    ]
    seniority_positive = ["intern", "internship", "graduate", "entry level", "entry-level", "junior", "associate", "trainee"]
    seniority_negative = ["senior", "lead", "manager", "principal", "director", "head of"]
    evidence_keywords = [
        "service design",
        "ux research",
        "user research",
        "figma",
        "prototype",
        "prototyping",
        "healthcare",
        "product design",
        "experience design",
        "touchdesigner",
        "arduino",
        "3d printing",
        "laser cutting",
        "user journey",
        "service blueprint",
        "visual design",
    ]
    strategic_companies = ["bytedance", "tiktok", "shopee", "lazada", "sea", "grab", "govtech", "dbs", "pinduoduo", "pdd", "tencent"]

    flags = hard_flag_patterns(combined)
    role_fit = keyword_score(lowered, role_keywords, 5)
    title_lowered = position.lower()
    if any(has_keyword(title_lowered, word) for word in ["design", "product", "ux", "user", "experience"]):
        role_fit = max(role_fit, 0.75)

    has_seniority_negative = any(has_keyword(title_lowered, word) for word in seniority_negative)
    if has_seniority_negative:
        seniority_fit = 0.2
    else:
        employment_type = detect_employment_type(position, jd_text)
        if employment_type in {"Internship", "Graduate"}:
            seniority_fit = 1.0
        elif employment_type == "Full-time":
            seniority_fit = 0.7
        elif employment_type == "Contract":
            seniority_fit = 0.6
        elif any(word in lowered for word in seniority_positive):
            seniority_fit = 1.0
        else:
            seniority_fit = 0.55
        if any(word in lowered for word in seniority_positive) and seniority_fit < 1.0:
            seniority_fit = 1.0

    if any(flag in flags for flag in ["citizen_or_pr_only", "local_only", "clearance_required"]):
        eligibility_fit = 0.0
    elif "visa_unclear" in flags:
        eligibility_fit = 0.5
    else:
        eligibility_fit = 1.0

    evidence_fit = evidence_overlap_score(lowered, resume, evidence_keywords, 4)

    if any(company_key in company.lower() for company_key in strategic_companies):
        strategic_value = 1.0
    elif any(word in lowered for word in ["healthcare", "public service", "graduate programme", "innovation", "product"]):
        strategic_value = 0.75
    else:
        strategic_value = 0.55

    score = round(role_fit + seniority_fit + eligibility_fit + evidence_fit + strategic_value, 1)
    notes = [
        f"岗位匹配 {role_fit:.1f}",
        f"级别/类型 {seniority_fit:.1f}",
        f"身份限制 {eligibility_fit:.1f}",
        f"简历证据 {evidence_fit:.1f}",
        f"战略价值 {strategic_value:.1f}",
    ]
    if flags:
        notes.append("风险标记: " + ", ".join(flags))
    return score, flags, " | ".join(notes)


def parse_linkedin_jobs_from_html(html: str, query: str, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    for card in re.split(r"<li\b", html, flags=re.I):
        if len(jobs) >= limit:
            break
        id_match = re.search(r"urn:li:jobPosting:(\d+)", card)
        if not id_match:
            continue
        external_id = id_match.group(1)
        if external_id in seen:
            continue
        seen.add(external_id)
        href_match = re.search(r'href="([^"]*?/jobs/view/[^"]+)"', card, flags=re.I)
        title_match = re.search(r'<span class="sr-only">\s*(.*?)\s*</span>', card, flags=re.I | re.S)
        company_match = re.search(
            r'base-search-card__subtitle[^>]*>\s*(?:<a[^>]*>)?\s*(.*?)\s*(?:</a>)?\s*</',
            card,
            flags=re.I | re.S,
        )
        location_match = re.search(r'job-search-card__location[^>]*>\s*(.*?)\s*</span>', card, flags=re.I | re.S)
        position = clean_text(title_match.group(1)) if title_match else "LinkedIn Role"
        company = clean_text(company_match.group(1)) if company_match else "LinkedIn Company"
        job_url = absolute_url("https://www.linkedin.com", href_match.group(1)) if href_match else f"https://www.linkedin.com/jobs/view/{external_id}"
        location = clean_text(location_match.group(1)) if location_match else "Singapore"
        jobs.append(
            {
                "company": company,
                "position": position,
                "source": "LinkedIn",
                "url": job_url,
                "location": location,
                "job_type": "Internship / Full-time",
                "jd_text": f"{position}\n{company}\n{location}\nSource query: {query}",
                "external_job_id": external_id,
            }
        )
    return jobs


def parse_linkedin_public_detail_html(html: str) -> str:
    match = re.search(
        r'<div[^>]*class="[^"]*\bshow-more-less-html__markup\b[^"]*"[^>]*>(.*?)</div>',
        html or "",
        flags=re.I | re.S,
    )
    detail = clean_text(match.group(1)) if match else ""
    return detail if len(detail) >= 120 else ""


def enrich_linkedin_jobs_from_public_pages(jobs: list[dict], limit: int = 12) -> int:
    candidates = [job for job in jobs if len(clean_text(job.get("jd_text") or "")) < 500][:limit]
    if not candidates:
        return 0

    def fetch_detail(job: dict) -> tuple[dict, str]:
        html = http_get(job.get("url") or "", timeout=12, retries=0)
        return job, parse_linkedin_public_detail_html(html)

    enriched = 0
    with ThreadPoolExecutor(max_workers=min(4, len(candidates))) as executor:
        futures = [executor.submit(fetch_detail, job) for job in candidates]
        for future in as_completed(futures):
            try:
                job, detail = future.result()
            except Exception:
                continue
            if detail:
                job["jd_text"] = detail
                enriched += 1
    return enriched


def parse_internsg_jobs_from_html(html: str, query: str, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    rows = re.findall(
        r'<div class="ast-row list-(?:even|odd|featured)[^"]*">(.*?)</div>\s*(?=<div class="ast-row list-|</div>\s*</div>)',
        html,
        flags=re.I | re.S,
    )
    if not rows:
        rows = re.split(r'<div class="ast-row list-(?:even|odd|featured)[^"]*">', html, flags=re.I)[1:]
    for row in rows:
        if len(jobs) >= limit:
            break
        link_match = re.search(r'<a href="([^"]*?/job/[^"]+)">(.*?)</a>', row, flags=re.I | re.S)
        if not link_match:
            continue
        job_url = absolute_url("https://www.internsg.com", link_match.group(1))
        if job_url in seen:
            continue
        seen.add(job_url)
        cols = re.findall(r'<div class="ast-col-lg-\d+[^"]*">(.*?)</div>', row, flags=re.I | re.S)
        company = clean_text(cols[0]) if cols else "InternSG Company"
        company = re.sub(r"\b[\w.-]+\.[a-z]{2,}\b.*$", "", company).strip() or "InternSG Company"
        company = clean_company_name(company) or "InternSG Company"
        position = clean_text(link_match.group(2)) or "InternSG Role"
        location = clean_text(cols[2]) if len(cols) >= 3 else "Singapore"
        period = clean_text(cols[3]) if len(cols) >= 4 else ""
        jobs.append(
            {
                "company": company,
                "position": position,
                "source": "InternSG",
                "url": job_url,
                "location": location,
                "job_type": period,
                "jd_text": f"{position}\n{company}\n{location}\n{period}\nSource query: {query}",
            }
        )
    return jobs


def parse_internsg_detail_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return clean_text(html)[:12000]
    soup = BeautifulSoup(html or "", "html.parser")
    container = soup.select_one(".isg-detail-container")
    if not container:
        return clean_text(html)[:12000]
    return clean_text(container.get_text(" ", strip=True))[:12000]


def parse_indeed_jobs_from_html(html: str, query: str, limit: int) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        BeautifulSoup = None

    jobs: list[dict] = []
    seen: set[str] = set()
    seen_signatures: set[str] = set()
    seen_positions: set[str] = set()
    if BeautifulSoup:
        soup = BeautifulSoup(html, "html.parser")
        links = soup.select('a[data-jk][href], a[href*="/rc/clk"], a[href*="/viewjob"]')
        for link in links:
            if len(jobs) >= limit:
                break
            href = link.get("href") or ""
            parsed_href = urlparse(absolute_url("https://sg.indeed.com", href))
            job_key = link.get("data-jk") or (parse_qs(parsed_href.query).get("jk") or [""])[0]
            if not job_key:
                continue
            job_url = canonical_job_url("Indeed", parsed_href.geturl(), job_key)
            if job_url in seen:
                continue
            position = soup_text(link) or clean_text(link.get("title") or "")
            if not position:
                title_node = link.find(attrs={"title": True}) if hasattr(link, "find") else None
                position = clean_text(title_node.get("title") or soup_text(title_node)) if title_node else ""
            if not position:
                continue
            card = link.find_parent(attrs={"class": re.compile(r"(job_seen_beacon|result|cardOutline)", re.I)}) if hasattr(link, "find_parent") else None
            if not card:
                card = link.find_parent(["li", "td", "div"]) if hasattr(link, "find_parent") else link
            company = soup_text(card.select_one('[data-testid="company-name"]')) if hasattr(card, "select_one") else ""
            if not company:
                company_node = card.select_one('[data-testid="company-name"], [class*="company"]') if hasattr(card, "select_one") else None
                company = soup_text(company_node)
            if not company:
                company = "Indeed Company"
            location = soup_text(card.select_one('[data-testid="text-location"]')) if hasattr(card, "select_one") else ""
            signature = f"{company.lower()}|{position.lower()}"
            position_signature = position.lower()
            if signature in seen_signatures or (company == "Indeed Company" and position_signature in seen_positions):
                continue
            seen_signatures.add(signature)
            seen_positions.add(position_signature)
            seen.add(job_url)
            jobs.append(
                {
                    "company": company,
                    "position": position,
                    "source": "Indeed",
                    "url": job_url,
                    "location": location or "Singapore",
                    "job_type": "Internship / Full-time",
                    "jd_text": soup_text(card)[:8000] or f"{position}\n{company}\nSource query: {query}",
                    "external_job_id": job_key,
                }
            )
        if jobs:
            return jobs[:limit]

    for card in re.split(r"<div[^>]+class=\"[^\"]*job_seen_beacon", html, flags=re.I):
        if len(jobs) >= limit:
            break
        link_match = re.search(r'href="([^"]*?/(?:viewjob|rc/clk)\?[^"]+)"', card, flags=re.I)
        title_match = re.search(r'title="([^"]+)"', card, flags=re.I)
        company_match = re.search(r'data-testid="company-name"[^>]*>(.*?)</', card, flags=re.I | re.S)
        if not link_match or not title_match:
            continue
        raw_url = absolute_url("https://sg.indeed.com", link_match.group(1))
        job_url = canonical_job_url("Indeed", raw_url)
        if job_url in seen:
            continue
        seen.add(job_url)
        position = clean_text(title_match.group(1))
        company = clean_text(company_match.group(1)) if company_match else "Indeed Company"
        jobs.append(
            {
                "company": company,
                "position": position,
                "source": "Indeed",
                "url": job_url,
                "location": "Singapore",
                "job_type": "Internship / Full-time",
                "jd_text": clean_text(card)[:8000] or f"{position}\n{company}\nSource query: {query}",
            }
        )
    return jobs


def indeed_search_url(host: str, query: str, location: str) -> str:
    return f"https://{host}/jobs?q={quote_plus(query)}&l={quote_plus(location)}&fromage=7&sort=date"


def fetch_indeed_jobs_with_browser(
    limit: int,
    queries: list[str],
    region: str | None = None,
    time_budget_seconds: float = 24,
) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    host = REGION_CONFIGS[code].get("indeed_host")
    if not host:
        return [], [f"Indeed is not configured for {REGION_CONFIGS[code]['label']}."]
    try:
        from playwright.sync_api import sync_playwright
    except Exception:
        return [], ["Indeed 浏览器兜底不可用：Playwright 未安装。"]

    location = REGION_CONFIGS[code]["search_location"]
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    started = time.monotonic()
    browser = None
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page(
                user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 Chrome/124 Safari/537.36",
                locale="en-SG",
                viewport={"width": 1365, "height": 900},
            )
            for query in queries:
                if len(jobs) >= limit:
                    break
                if time.monotonic() - started > time_budget_seconds:
                    failures.append("Indeed 浏览器兜底：来源级时间预算已用完。")
                    break
                url = indeed_search_url(host, query, location)
                try:
                    response = page.goto(url, wait_until="domcontentloaded", timeout=18000)
                    page.wait_for_timeout(1800)
                    body_text = page.locator("body").inner_text(timeout=5000).lower()
                    if "captcha" in body_text or "verify you are human" in body_text:
                        failures.append("Indeed 浏览器兜底受限：页面要求验证码或人工验证。")
                        break
                    parsed_jobs = parse_indeed_jobs_from_html(page.content(), query, limit - len(jobs))
                    if not parsed_jobs and response and response.status >= 400:
                        failures.append(f"Indeed browser {query}: HTTP {response.status}")
                except Exception as exc:
                    failures.append(f"Indeed browser {query}: {exc}")
                    continue
                for job in parsed_jobs:
                    job_url = job.get("url", "")
                    if not job_url or job_url in seen:
                        continue
                    seen.add(job_url)
                    job["region"] = code
                    job["city"] = city
                    job["source_region"] = code
                    if code != "SG":
                        job["location"] = job.get("location") or location
                    jobs.append(job)
                    if len(jobs) >= limit:
                        break
    except Exception as exc:
        failures.append(f"Indeed 浏览器兜底启动失败：{exc}")
    finally:
        if browser:
            try:
                browser.close()
            except Exception:
                pass
    return jobs[:limit], failures


def parse_serpapi_indeed_jobs(payload: dict, query: str, limit: int, region: str | None = None) -> list[dict]:
    code = active_region_code(region)
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    jobs: list[dict] = []
    seen: set[str] = set()
    for item in payload.get("jobs_results") or []:
        if len(jobs) >= limit:
            break
        title = clean_text(item.get("title") or "")
        company = clean_text(item.get("company_name") or "")
        if not title or not company:
            continue
        apply_options = item.get("apply_options") or []
        indeed_link = ""
        for option in apply_options:
            link = option.get("link") or ""
            title_text = f"{option.get('title') or ''} {link}".lower()
            if "indeed" in title_text:
                indeed_link = link
                break
        if not indeed_link:
            continue
        canonical = canonical_job_url("Indeed", indeed_link)
        dedupe_key = canonical or f"{company}|{title}|{item.get('location')}"
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        description = clean_text(item.get("description") or "")
        extensions = item.get("detected_extensions") or {}
        job_type = " ".join(str(value) for value in extensions.values() if value)
        jobs.append(
            {
                "company": company,
                "position": title,
                "source": "Indeed",
                "url": canonical or indeed_link,
                "location": clean_text(item.get("location") or city),
                "region": code,
                "city": city,
                "source_region": code,
                "job_type": clean_text(job_type) or "Google Jobs / Indeed",
                "jd_text": (description or f"{title}\n{company}\nSource query: {query}\nIndeed apply option via Google Jobs")[:12000],
            }
        )
    return jobs[:limit]


def jobstreet_slug(query: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (query or "").lower()).strip("-")
    return slug or "internship"


def jobstreet_search_urls(query: str) -> list[str]:
    slug = jobstreet_slug(query)
    return [
        f"https://sg.jobstreet.com/{slug}-jobs/in-Singapore",
        f"https://sg.jobstreet.com/jobs?keywords={quote_plus(query)}&where=Singapore",
    ]


def jobstreet_api_search_url(query: str, page: int = 1) -> str:
    return (
        "https://sg.jobstreet.com/api/jobsearch/v5/search?"
        f"siteKey=SG-Main&keywords={quote_plus(query)}&where=Singapore&page={page}"
    )


def soup_text(element) -> str:
    return clean_text(element.get_text(" ", strip=True)) if element else ""


def parse_jobstreet_api_jobs(payload: dict, query: str, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    for item in payload.get("data") or []:
        if len(jobs) >= limit:
            break
        job_id = str(item.get("id") or "").strip()
        title = clean_text(item.get("title") or "")
        company = clean_text(item.get("companyName") or item.get("advertiser", {}).get("description") or "")
        if not job_id or not title or not is_actionable_company_name(company, title):
            continue
        job_url = canonical_job_url("JobStreet", f"https://sg.jobstreet.com/job/{job_id}")
        if job_url in seen:
            continue
        seen.add(job_url)
        locations = item.get("locations") or []
        location = ", ".join(clean_text(location.get("label") or "") for location in locations if location.get("label"))
        work_types = item.get("workTypes") or []
        salary = clean_text(item.get("salaryLabel") or "")
        teaser = clean_text(item.get("teaser") or "")
        bullets = [clean_text(value) for value in item.get("bulletPoints") or [] if clean_text(value)]
        classifications = []
        for value in item.get("classifications") or []:
            classification = value.get("classification") or {}
            subclassification = value.get("subclassification") or {}
            label = " / ".join(
                clean_text(part.get("description") or "")
                for part in [classification, subclassification]
                if part.get("description")
            )
            if label:
                classifications.append(label)
        jd_parts = [
            title,
            company,
            location or "Singapore",
            " ".join(work_types),
            salary,
            teaser,
            *bullets,
            *classifications,
            f"Source query: {query}",
        ]
        jobs.append(
            {
                "company": company or "JobStreet Company",
                "position": title,
                "source": "JobStreet",
                "url": job_url,
                "location": location or "Singapore",
                "job_type": ", ".join(work_types) or "Internship / Full-time",
                "jd_text": "\n".join(part for part in jd_parts if part)[:12000],
                "external_job_id": job_id,
            }
        )
    return jobs[:limit]


def parse_cultjobs_listing_urls(html: str, limit: int) -> list[str]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return []
    soup = BeautifulSoup(html or "", "html.parser")
    urls: list[str] = []
    for article in soup.select("article.job_listing_type-internship.job_listing_location-singapore"):
        link = article.select_one("h2.job-title a[href*='/job/']")
        url = canonical_job_url("Cultjobs", link.get("href") if link else "")
        if not url or url in urls:
            continue
        urls.append(url)
        if len(urls) >= limit:
            break
    return urls


def parse_cultjobs_detail(html: str, url: str) -> dict | None:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return None
    soup = BeautifulSoup(html or "", "html.parser")
    title = soup_text(soup.select_one("h1.job-detail-title"))
    company_link = soup.select_one(".job-metas-detail a[href*='/employer/']")
    company = soup_text(company_link)
    description = soup_text(soup.select_one(".job-detail-description"))
    location = soup_text(soup.select_one(".job-metas-detail .job-location")) or "Singapore"
    job_type = soup_text(soup.select_one(".job-metas-detail-bottom .type-job")) or "Internship"
    salary = soup_text(soup.select_one(".job-metas-detail .job-salary"))
    expiration = ""
    for item in soup.select(".job-detail-detail li"):
        if "expiration date" in soup_text(item).lower():
            expiration = soup_text(item.select_one(".value"))
            break
    if expiration:
        try:
            expiry_date = dt.datetime.strptime(expiration, "%B %d, %Y").date()
            if expiry_date < dt.date.today():
                return None
        except ValueError:
            pass
    if (
        not title
        or not COMPANY_SCAN_TITLE_ROLE_PATTERN.search(title)
        or not is_actionable_company_name(company, title)
        or "singapore" not in location.lower()
    ):
        return None
    salary_info = parse_salary_info(title, salary, job_type, "SG")
    if salary_info.get("salary_max") and salary:
        salary_info["salary_text"] = re.sub(r"([$])\s+(?=\d)", r"\1", salary)
    return {
        "company": company,
        "position": title,
        "source": "Cultjobs",
        "url": canonical_job_url("Cultjobs", url),
        "location": location,
        "job_type": job_type,
        "jd_text": "\n".join(part for part in [title, company, location, job_type, salary, description] if part)[:12000],
        "region": "SG",
        "city": "Singapore",
        "source_region": "SG",
        **salary_info,
    }


def parse_jobstreet_jobs_from_html(html: str, query: str, limit: int) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        BeautifulSoup = None

    jobs: list[dict] = []
    seen: set[str] = set()
    if BeautifulSoup:
        soup = BeautifulSoup(html, "html.parser")
        cards = soup.select('[data-automation="normalJob"], [data-automation="jobCard"], article')
        if not cards:
            title_links = [
                link for link in soup.find_all("a", href=True)
                if "/job/" in link.get("href", "") or link.get("data-automation") == "jobTitle"
            ]
            cards = []
            for link in title_links:
                card = link.find_parent(["article", "section", "div"]) or link
                if card not in cards:
                    cards.append(card)

        for card in cards:
            if len(jobs) >= limit:
                break
            title_link = card.select_one('[data-automation="jobTitle"]') if hasattr(card, "select_one") else None
            if not title_link:
                title_link = card.find("a", href=re.compile(r"/job/")) if hasattr(card, "find") else None
            if not title_link:
                continue
            href = title_link.get("href", "")
            if "/job/" not in href and not href.startswith("http"):
                continue
            job_url = absolute_url("https://sg.jobstreet.com", href)
            canonical_url = canonical_job_url("JobStreet", job_url)
            if canonical_url in seen:
                continue
            seen.add(canonical_url)
            position = soup_text(title_link) or "JobStreet Role"
            company = soup_text(card.select_one('[data-automation="jobCompany"]')) if hasattr(card, "select_one") else ""
            if not company:
                company_link = card.find("a", href=re.compile(r"/companies/")) if hasattr(card, "find") else None
                company = soup_text(company_link)
            if not company:
                card_text = soup_text(card)
                match = re.search(r"\bat\s+(.+?)(?:\s+This is|\s+Be an|\s+\$|\s+Singapore|\s+Central|\s+East|\s+West|\s+North|$)", card_text, flags=re.I)
                company = clean_text(match.group(1)) if match else "JobStreet Company"
            location = soup_text(card.select_one('[data-automation="jobLocation"]')) if hasattr(card, "select_one") else ""
            work_type = soup_text(card.select_one('[data-automation="jobWorkType"]')) if hasattr(card, "select_one") else ""
            summary = soup_text(card.select_one('[data-automation="jobShortDescription"]')) if hasattr(card, "select_one") else ""
            text = soup_text(card)
            jobs.append(
                {
                    "company": company or "JobStreet Company",
                    "position": position,
                    "source": "JobStreet",
                    "url": canonical_url,
                    "location": location or "Singapore",
                    "job_type": work_type or "Internship / Full-time",
                    "jd_text": (summary or text or f"{position}\n{company}\nSource query: {query}")[:12000],
                }
            )

    if jobs:
        return jobs[:limit]

    for block in re.split(r"<h[1-4]|<article|data-automation=\"normalJob\"", html, flags=re.I):
        if len(jobs) >= limit:
            break
        link_match = re.search(r'href="([^"]*/job/\d+[^"]*)"', block, flags=re.I)
        title_match = re.search(r">(.*?)</a>", block, flags=re.I | re.S)
        if not link_match or not title_match:
            continue
        job_url = canonical_job_url("JobStreet", absolute_url("https://sg.jobstreet.com", link_match.group(1)))
        if job_url in seen:
            continue
        seen.add(job_url)
        position = clean_text(title_match.group(1)) or "JobStreet Role"
        text = clean_text(block)
        jobs.append(
            {
                "company": "JobStreet Company",
                "position": position,
                "source": "JobStreet",
                "url": job_url,
                "location": "Singapore",
                "job_type": "Internship / Full-time",
                "jd_text": text[:12000] or f"{position}\nSource query: {query}",
            }
        )
    return jobs[:limit]


def region_queries(region: str | None = None, ai_only: bool = False) -> list[str]:
    code = active_region_code(region)
    base = AI_TARGET_QUERIES if ai_only else TARGET_QUERIES
    if code == "SG":
        return list(base)
    if code == "CN":
        city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
        return [
            f"{city} AI product intern",
            f"{city} UX design intern",
            f"{city} product operations intern",
            f"{city} user research intern",
            f"{city} service design intern",
            "AI product manager intern China",
        ][:6 if not ai_only else 4]
    if code == "HK":
        return [
            "Hong Kong AI product intern",
            "Hong Kong UX design intern",
            "Hong Kong product operations intern",
            "Hong Kong user research intern",
            "Hong Kong graduate product design",
        ][:5 if not ai_only else 4]
    return list(base)


def fetch_google_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    api_key = os.environ.get("SERPAPI_KEY", "").strip()
    if not api_key:
        return [], []
    code = active_region_code(region)
    location = REGION_CONFIGS[code]["search_location"]
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    for query in region_queries(code)[:4]:
        if len(jobs) >= limit:
            break
        params = urlencode(
            {
                "engine": "google_jobs",
                "q": query,
                "location": location,
                "hl": "en",
                "api_key": api_key,
            }
        )
        request = urllib.request.Request(
            f"https://serpapi.com/search.json?{params}",
            headers={"User-Agent": "Job Assistant local app"},
        )
        try:
            with urllib.request.urlopen(request, timeout=12) as response:
                payload = json.loads(response.read().decode("utf-8", errors="ignore"))
        except Exception as exc:
            failures.append(f"Google Jobs {query}: {exc}")
            continue
        for item in payload.get("jobs_results") or []:
            if len(jobs) >= limit:
                break
            title = clean_text(item.get("title") or "")
            company = clean_text(item.get("company_name") or "")
            if not title or not company:
                continue
            apply_link = ""
            related_links = item.get("related_links") or []
            detected_extensions = item.get("detected_extensions") or {}
            apply_options = item.get("apply_options") or []
            if apply_options:
                apply_link = apply_options[0].get("link") or ""
            if not apply_link and related_links:
                apply_link = related_links[0].get("link") or ""
            if not apply_link:
                apply_link = item.get("share_link") or item.get("job_id") or ""
            dedupe_key = apply_link or f"{company}|{title}|{item.get('location')}"
            if dedupe_key in seen:
                continue
            seen.add(dedupe_key)
            description = clean_text(item.get("description") or "")
            extensions = " ".join(str(value) for value in detected_extensions.values() if value)
            jobs.append(
                {
                    "company": company,
                    "position": title,
                    "source": "Google Jobs",
                    "url": apply_link if apply_link.startswith("http") else f"https://www.google.com/search?q={quote_plus(company + ' ' + title + ' job')}",
                    "location": clean_text(item.get("location") or location),
                    "region": code,
                    "city": city,
                    "source_region": code,
                    "job_type": clean_text(extensions) or "Google Jobs",
                    "jd_text": (description or f"{title}\n{company}\nSource query: {query}")[:12000],
                }
            )
    return jobs, failures


def fetch_indeed_jobs_via_google_jobs(
    limit: int,
    queries: list[str],
    region: str | None = None,
) -> tuple[list[dict], list[str]]:
    api_key = os.environ.get("SERPAPI_KEY", "").strip()
    if not api_key:
        return [], ["Indeed 受限：直连被站点拒绝；配置 SERPAPI_KEY 后可用 Google Jobs 补充 Indeed apply option。"]
    code = active_region_code(region)
    location = REGION_CONFIGS[code]["search_location"]
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    for query in queries:
        if len(jobs) >= limit:
            break
        params = urlencode(
            {
                "engine": "google_jobs",
                "q": query,
                "location": location,
                "hl": "en",
                "api_key": api_key,
            }
        )
        request = urllib.request.Request(
            f"https://serpapi.com/search.json?{params}",
            headers={"User-Agent": "Job Assistant local app"},
        )
        try:
            with urllib.request.urlopen(request, timeout=12) as response:
                payload = json.loads(response.read().decode("utf-8", errors="ignore"))
        except Exception as exc:
            failures.append(f"Indeed Google Jobs fallback {query}: {exc}")
            continue
        for job in parse_serpapi_indeed_jobs(payload, query, limit - len(jobs), code):
            url = canonical_job_url("Indeed", job.get("url", ""))
            if not url or url in seen:
                continue
            seen.add(url)
            job["url"] = url
            jobs.append(job)
            if len(jobs) >= limit:
                break
    if jobs:
        failures.append("Indeed 直连受限，已用 Google Jobs 中的 Indeed apply option 补充。")
    return jobs[:limit], failures


def fetch_linkedin_jobs(limit: int, queries: list[str] | None = None, region: str | None = None) -> tuple[list[dict], list[str]]:
    jobs: list[dict] = []
    failures: list[str] = []
    detail_failures: list[str] = []
    seen: set[str] = set()
    code = active_region_code(region)
    location = REGION_CONFIGS[code]["search_location"]
    detail_limited = False
    for query in (queries or region_queries(code)):
        if len(jobs) >= limit:
            break
        url = (
            "https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"
            f"?keywords={quote_plus(query)}&location={quote_plus(location)}&f_TPR=r604800&sortBy=DD&start=0"
        )
        try:
            html = http_get(url)
        except Exception as exc:
            failures.append(f"LinkedIn {query}: {exc}")
            continue
        for job in parse_linkedin_jobs_from_html(html, query, limit - len(jobs)):
            if len(jobs) >= limit:
                break
            external_id = job.get("external_job_id") or ""
            if external_id in seen:
                continue
            seen.add(external_id)
            if not detail_limited:
                try:
                    detail_html = http_get(f"https://www.linkedin.com/jobs-guest/jobs/api/jobPosting/{external_id}", timeout=18)
                    detail_text = clean_text(detail_html)
                    if detail_text:
                        job["jd_text"] = detail_text
                except Exception as exc:
                    detail_failures.append(f"LinkedIn detail {external_id}: {exc}")
                    if any(token in str(exc).lower() for token in ["429", "too many requests", "403", "forbidden"]):
                        detail_limited = True
            job["region"] = code
            job["city"] = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
            job["source_region"] = code
            if code != "SG":
                job["location"] = job.get("location") or location
            jobs.append(job)
    if detail_failures:
        sparse_job_count = sum(len(clean_text(job.get("jd_text") or "")) < 500 for job in jobs)
        enriched_job_count = enrich_linkedin_jobs_from_public_pages(jobs)
        if enriched_job_count < sparse_job_count:
            failures.append(detail_failures[0])
    return jobs, failures


def fetch_internsg_jobs(limit: int, queries: list[str] | None = None) -> tuple[list[dict], list[str]]:
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    for query in (queries or TARGET_QUERIES):
        if len(jobs) >= limit:
            break
        url = f"https://www.internsg.com/jobs/?f_p={quote_plus(query)}&f_i=&f_c=&f_s="
        try:
            html = http_get(url)
        except Exception as exc:
            failures.append(f"InternSG {query}: {exc}")
            continue
        for job in parse_internsg_jobs_from_html(html, query, limit - len(jobs)):
            if len(jobs) >= limit:
                break
            job_url = job.get("url", "")
            if job_url in seen:
                continue
            seen.add(job_url)
            jobs.append(job)

    def fetch_detail(index_job: tuple[int, dict]) -> tuple[int, str, str]:
        index, job = index_job
        try:
            detail_html = http_get(job.get("url", ""), timeout=18)
            return index, parse_internsg_detail_text(detail_html)[:12000], ""
        except Exception as exc:
            return index, "", f"InternSG detail {job.get('position')}: {exc}"

    details: dict[int, tuple[str, str]] = {}
    with ThreadPoolExecutor(max_workers=min(6, max(1, len(jobs)))) as executor:
        futures = [executor.submit(fetch_detail, item) for item in enumerate(jobs)]
        for future in as_completed(futures):
            index, detail_text, failure = future.result()
            details[index] = (detail_text, failure)
    for index, job in enumerate(jobs):
        detail_text, failure = details.get(index, ("", ""))
        if detail_text:
            job["jd_text"] = detail_text
        if failure:
            failures.append(failure)
    return jobs, failures


def fetch_indeed_jobs(
    limit: int,
    queries: list[str] | None = None,
    region: str | None = None,
    time_budget_seconds: float = 18,
    failure_limit: int = 2,
) -> tuple[list[dict], list[str]]:
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    code = active_region_code(region)
    host = REGION_CONFIGS[code].get("indeed_host")
    if not host:
        return [], [f"Indeed is not configured for {REGION_CONFIGS[code]['label']}."]
    location = REGION_CONFIGS[code]["search_location"]
    query_list = queries if queries is not None else region_queries(code)[:4]
    started = time.monotonic()
    for query in query_list:
        if len(jobs) >= limit:
            break
        if time.monotonic() - started > time_budget_seconds:
            failures.append("Indeed 受限：来源级时间预算已用完，已跳过剩余查询。")
            break
        if len(failures) >= failure_limit:
            failures.append("Indeed 受限：连续失败较多，已跳过剩余查询。")
            break
        url = indeed_search_url(host, query, location)
        try:
            html = http_get(url, timeout=8, retries=0)
        except Exception as exc:
            failures.append(f"Indeed {query}: {exc}")
            continue
        for job in parse_indeed_jobs_from_html(html, query, limit - len(jobs)):
            if len(jobs) >= limit:
                break
            job_url = job.get("url", "")
            if job_url in seen:
                continue
            seen.add(job_url)
            job["region"] = code
            job["city"] = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
            job["source_region"] = code
            if code != "SG":
                job["location"] = job.get("location") or location
            jobs.append(job)
    if len(jobs) < limit:
        browser_jobs, browser_failures = fetch_indeed_jobs_with_browser(
            limit - len(jobs),
            query_list,
            code,
            max(8, time_budget_seconds),
        )
        if browser_jobs:
            failures = []
        else:
            failures.extend(browser_failures)
        for job in browser_jobs:
            job_url = job.get("url", "")
            if job_url in seen:
                continue
            seen.add(job_url)
            jobs.append(job)
            if len(jobs) >= limit:
                break
    if len(jobs) < limit and (not jobs or os.environ.get("SERPAPI_KEY")):
        api_jobs, api_failures = fetch_indeed_jobs_via_google_jobs(limit - len(jobs), query_list, code)
        failures.extend(api_failures)
        for job in api_jobs:
            job_url = job.get("url", "")
            if job_url in seen:
                continue
            seen.add(job_url)
            jobs.append(job)
            if len(jobs) >= limit:
                break
    return jobs, failures


def fetch_jobstreet_jobs(
    limit: int,
    queries: list[str] | None = None,
    region: str | None = None,
    time_budget_seconds: float = 24,
    failure_limit: int = 3,
    use_html_fallback: bool = True,
) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], [f"JobStreet connector is Singapore-only in this version."]
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    started = time.monotonic()
    query_list = queries or TARGET_QUERIES
    query_cap = max(1, (limit + len(query_list) - 1) // max(1, len(query_list)))
    for query in query_list:
        if len(jobs) >= limit:
            break
        if time.monotonic() - started > time_budget_seconds:
            failures.append("JobStreet 受限：来源级时间预算已用完，已跳过剩余查询。")
            break
        if len(failures) >= failure_limit:
            failures.append("JobStreet 受限：连续失败较多，已跳过剩余查询。")
            break
        query_jobs: list[dict] = []
        api_url = jobstreet_api_search_url(query)
        try:
            payload = json.loads(http_get(api_url, timeout=8, retries=0))
            query_jobs.extend(parse_jobstreet_api_jobs(payload, query, query_cap))
        except Exception as exc:
            failures.append(f"JobStreet API {query}: {exc}")
        for url in jobstreet_search_urls(query) if use_html_fallback else []:
            if len(query_jobs) >= query_cap:
                break
            try:
                html = http_get(url, timeout=8, retries=0)
            except Exception as exc:
                failures.append(f"JobStreet {query}: {exc}")
                continue
            parsed_jobs = parse_jobstreet_jobs_from_html(html, query, limit - len(jobs))
            if parsed_jobs:
                query_jobs.extend(parsed_jobs)
                break
        for job in query_jobs[:query_cap]:
            if len(jobs) >= limit:
                break
            url = canonical_job_url("JobStreet", job.get("url", ""))
            if not url or url in seen:
                continue
            seen.add(url)
            job["url"] = url
            if not job.get("external_job_id"):
                try:
                    detail_html = http_get(url, timeout=6, retries=0)
                    detail_text = clean_text(detail_html)
                    if detail_text:
                        job["jd_text"] = detail_text[:12000]
                except Exception as exc:
                    failures.append(f"JobStreet detail {job.get('position')}: {exc}")
            job["region"] = code
            job["city"] = "Singapore"
            job["source_region"] = code
            jobs.append(job)
    return jobs, failures


def watched_company_scan_items(region: str, limit: int = 24) -> list[dict]:
    with get_db() as conn:
        rows = conn.execute(
            """
            select * from watch_companies
            where region=? and status='Watch'
            order by priority desc, company
            limit ?
            """,
            (active_region_code(region), limit),
        ).fetchall()
    return [company_item_from_row(row_to_dict(row)) for row in rows]


def fetch_watched_company_public_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], []
    companies = watched_company_scan_items(code)
    if not companies:
        return [], []
    queries = [item.get("company") or "" for item in companies if item.get("company")]
    raw_jobs, failures = fetch_jobstreet_jobs(
        max(limit * 2, len(queries) * 3),
        queries,
        code,
        time_budget_seconds=18,
        failure_limit=3,
        use_html_fallback=False,
    )
    matches_by_company: list[list[dict]] = [[] for _item in companies]
    for job in raw_jobs:
        for index, item in enumerate(companies):
            if match_job_to_company(job, item.get("company") or "", item)[0]:
                matches_by_company[index].append(job)
                break
    matched: list[dict] = []
    round_index = 0
    while len(matched) < limit and any(round_index < len(items) for items in matches_by_company):
        for items in matches_by_company:
            if round_index < len(items):
                matched.append(items[round_index])
                if len(matched) >= limit:
                    break
        round_index += 1
    return matched, failures


def merge_fetch_results(fetch_calls: list[tuple[int, object]], limit: int) -> tuple[list[dict], list[str]]:
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    for call_limit, fetcher in fetch_calls:
        if len(jobs) >= limit:
            break
        try:
            raw_jobs, raw_failures = fetcher(min(call_limit, max(1, limit - len(jobs))))
        except Exception as exc:
            raw_jobs, raw_failures = [], [str(exc)]
        failures.extend(str(item) for item in raw_failures)
        for job in raw_jobs:
            source = job.get("source") or ""
            url = canonical_job_url(source, job.get("url", ""), job.get("external_job_id"))
            if not url or url in seen:
                continue
            seen.add(url)
            job["url"] = url
            jobs.append(job)
            if len(jobs) >= limit:
                break
    return jobs, failures


def summarize_scan_source_failures(source_name: str, failures: list[str], saved_count: int) -> list[str]:
    unique = list(dict.fromkeys(str(item) for item in failures if str(item).strip()))
    if "LinkedIn" not in source_name:
        return unique
    rate_limited = [
        item for item in unique
        if any(token in item.lower() for token in ["429", "too many requests", "403", "forbidden"])
    ]
    if not rate_limited:
        return unique
    other = [item for item in unique if item not in rate_limited]
    return [
        *other,
        f"LinkedIn 限流：详情或部分关键词请求受限；已保留 {saved_count} 条列表结果，本轮已停止重复请求。",
    ]


def is_nonblocking_scan_warning(source_name: str, error: str, saved_count: int) -> bool:
    lowered = (error or "").lower()
    return (
        "linkedin" in (source_name or "").lower()
        and saved_count >= SOURCE_LIMITS["LinkedIn"]
        and any(token in lowered for token in ["429", "too many requests", "403", "forbidden", "限流"])
    )


def fetch_linkedin_jobs_with_ai(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    return merge_fetch_results(
        [
            (SOURCE_LIMITS["LinkedIn"], lambda call_limit: fetch_linkedin_jobs(call_limit, region=code)),
            (SOURCE_LIMITS["LinkedIn AI"], lambda call_limit: fetch_linkedin_jobs(call_limit, region_queries(code, ai_only=True), code)),
        ],
        limit,
    )


def fetch_internsg_jobs_with_ai(limit: int) -> tuple[list[dict], list[str]]:
    return merge_fetch_results(
        [
            (SOURCE_LIMITS["InternSG"], lambda call_limit: fetch_internsg_jobs(call_limit)),
            (SOURCE_LIMITS["InternSG AI"], lambda call_limit: fetch_internsg_jobs(call_limit, AI_TARGET_QUERIES)),
        ],
        limit,
    )


def fetch_indeed_jobs_with_ai(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    return merge_fetch_results(
        [
            (SOURCE_LIMITS["Indeed"], lambda call_limit: fetch_indeed_jobs(call_limit, region=code)),
            (SOURCE_LIMITS["Indeed AI"], lambda call_limit: fetch_indeed_jobs(call_limit, region_queries(code, ai_only=True), code, 10, 1)),
        ],
        limit,
    )


def fetch_jobstreet_jobs_with_ai(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    return merge_fetch_results(
        [
            (SOURCE_LIMITS["JobStreet"], lambda call_limit: fetch_jobstreet_jobs(call_limit, region=code)),
            (SOURCE_LIMITS["JobStreet AI"], lambda call_limit: fetch_jobstreet_jobs(call_limit, region_queries(code, ai_only=True), code, 12, 1)),
        ],
        limit,
    )


def company_job_record(
    company: str,
    title: str,
    url: str,
    region: str,
    city: str,
    focus: str,
    source_url: str,
    description: str = "",
    source: str = "Company Site",
    location: str | None = None,
) -> dict | None:
    raw_position = clean_text(title)
    position = raw_position[:180]
    normalized_company = normalize_company_phrase(company)
    if (
        not is_actionable_job_title(position)
        or len(raw_position) > 120
        or not COMPANY_SCAN_TITLE_ROLE_PATTERN.search(position)
        or (normalized_company in COMPANIES_REQUIRE_EXPLICIT_LOCATION and not location)
        or (location and not location_matches_region(location, region, city))
    ):
        return None
    job_url = absolute_url(source_url, url)
    if not job_url.startswith("http"):
        return None
    jd_text = clean_text(description)[:9000]
    return {
        "company": company,
        "position": position,
        "source": source,
        "url": job_url,
        "location": location or city,
        "region": region,
        "city": city,
        "source_region": region,
        "job_type": "Company career page",
        "jd_text": f"{company} official career match\nRole: {position}\nFocus: {focus}\nSource: {source_url}\nURL: {job_url}\n\n{jd_text}".strip(),
    }


def append_company_job(jobs: list[dict], seen: set[str], job: dict | None, limit: int) -> None:
    if not job or len(jobs) >= limit:
        return
    key = canonical_job_url(job.get("source") or "", job.get("url") or "", job.get("external_job_id"))
    if key in seen:
        return
    seen.add(key)
    jobs.append(job)


def iter_json_objects(value):
    if isinstance(value, dict):
        yield value
        for child in value.values():
            yield from iter_json_objects(child)
    elif isinstance(value, list):
        for child in value:
            yield from iter_json_objects(child)


def parse_company_jsonld_jobs(html: str, base_url: str, company: str, focus: str, region: str, city: str, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    for script in re.findall(r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>', html or "", flags=re.I | re.S):
        try:
            payload = json.loads(unescape(script).strip())
        except Exception:
            continue
        for item in iter_json_objects(payload):
            item_type = item.get("@type") or item.get("type")
            types = item_type if isinstance(item_type, list) else [item_type]
            if not any(str(value).lower() == "jobposting" for value in types):
                continue
            title = item.get("title") or item.get("name") or ""
            location_value = ""
            job_location = item.get("jobLocation")
            if isinstance(job_location, dict):
                address = job_location.get("address") or {}
                if isinstance(address, dict):
                    location_value = clean_text(" ".join(str(address.get(key) or "") for key in ["addressLocality", "addressRegion", "addressCountry"])) or city
            description = item.get("description") or item.get("responsibilities") or ""
            append_company_job(
                jobs,
                seen,
                company_job_record(company, title, item.get("url") or base_url, region, city, focus, base_url, description, "Company Site / ATS", location_value),
                limit,
            )
            if len(jobs) >= limit:
                return jobs
    return jobs


def parse_company_anchor_jobs(html: str, page_url: str, company: str, focus: str, region: str, city: str, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    for href, label in re.findall(r'<a[^>]+href=["\']([^"\']+)["\'][^>]*>(.*?)</a>', html or "", flags=re.I | re.S):
        if len(jobs) >= limit:
            break
        text = clean_text(label)
        if len(text) < 5:
            continue
        append_company_job(
            jobs,
            seen,
            company_job_record(company, text, href, region, city, focus, page_url, text, "Company Site"),
            limit,
        )
    return jobs


def extract_embedded_ats_links(html: str, base_url: str) -> list[str]:
    links: list[str] = []
    for raw in re.findall(r'https?://[^"\'\s<>()]+', html or "", flags=re.I):
        cleaned = unescape(raw).replace("\\/", "/").rstrip("\\).,;")
        if COMPANY_CAREER_LINK_PATTERN.search(cleaned):
            links.append(cleaned)
    for href, label in re.findall(r'<a[^>]+href=["\']([^"\']+)["\'][^>]*>(.*?)</a>', html or "", flags=re.I | re.S):
        text = f"{href} {clean_text(label)}"
        if COMPANY_CAREER_LINK_PATTERN.search(text):
            links.append(absolute_url(base_url, href))
    parsed = urlparse(base_url)
    if parsed.scheme and parsed.netloc:
        for path in ["/careers", "/career", "/jobs", "/openings", "/join-us", "/work-with-us"]:
            links.append(f"{parsed.scheme}://{parsed.netloc}{path}")
    out: list[str] = []
    seen: set[str] = set()
    base_host = parsed.netloc.lower().removeprefix("www.")
    ats_hosts = (
        "greenhouse.io",
        "lever.co",
        "ashbyhq.com",
        "workdayjobs.com",
        "myworkdayjobs.com",
        "smartrecruiters.com",
        "apply.workable.com",
        "bamboohr.com",
        "careers-page.com",
        "surgeahead.com",
        "lifeatcanva.com",
        "wise.jobs",
    )
    for link in links:
        if any(token in link.lower() for token in [");", "background:", "background-", "{", "}"]):
            continue
        parsed_link = urlparse(link)
        host = parsed_link.netloc.lower().removeprefix("www.")
        if not parsed_link.scheme.startswith("http"):
            continue
        if parsed_link.path.lower().startswith(("/wp-json/", "/wp-content/", "/wp-includes/")):
            continue
        if re.search(r"\.(?:avif|gif|jpe?g|png|svg|webp|css|js|woff2?|ttf|ico)$", parsed_link.path, flags=re.I):
            continue
        if host != base_host and not any(ats in host for ats in ats_hosts):
            continue
        key = parsed_link._replace(fragment="").geturl()
        if key not in seen:
            seen.add(key)
            out.append(key)
    return out


def greenhouse_board_token(url: str) -> str:
    parsed = urlparse(url)
    if "greenhouse.io" not in parsed.netloc:
        return ""
    parts = [part for part in parsed.path.split("/") if part]
    if not parts:
        return ""
    if parts[0] in {"embed", "jobs"} and len(parts) > 1:
        return parts[1]
    return parts[0]


def lever_company_token(url: str) -> str:
    parsed = urlparse(url)
    if "lever.co" not in parsed.netloc:
        return ""
    parts = [part for part in parsed.path.split("/") if part]
    return parts[0] if parts else ""


def ashby_board_token(url: str) -> str:
    parsed = urlparse(url)
    if "ashbyhq.com" not in parsed.netloc:
        return ""
    parts = [part for part in parsed.path.split("/") if part]
    return parts[0] if parts else ""


def smartrecruiters_company_token(url: str) -> str:
    parsed = urlparse(url)
    if "smartrecruiters.com" not in parsed.netloc:
        return ""
    parts = [part for part in parsed.path.split("/") if part]
    return parts[0] if len(parts) >= 1 and parts[0].lower() != "job" else ""


def smartrecruiters_public_job_url(item: dict, company_token: str) -> str:
    if item.get("postingUrl") or item.get("applyUrl"):
        return item.get("postingUrl") or item.get("applyUrl")
    posting_id = clean_text(item.get("id") or "")
    if not posting_id:
        return item.get("ref") or ""
    company_data = item.get("company") if isinstance(item.get("company"), dict) else {}
    identifier = clean_text(company_data.get("identifier") or company_token)
    slug = re.sub(r"[^a-z0-9]+", "-", clean_text(item.get("name") or "").lower()).strip("-") or "job"
    return f"https://jobs.smartrecruiters.com/{identifier}/{posting_id}-{slug}"


def workable_account_token(url: str) -> str:
    parsed = urlparse(url)
    if "apply.workable.com" not in parsed.netloc:
        return ""
    parts = [part for part in parsed.path.split("/") if part]
    if not parts or parts[0].lower() == "j":
        return ""
    return parts[0]


def workday_board_parts(url: str) -> tuple[str, str, str] | None:
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    if not any(domain in host for domain in ["myworkdayjobs.com", "workdayjobs.com"]):
        return None
    tenant = host.split(".", 1)[0]
    parts = [part for part in parsed.path.split("/") if part]
    if parts and re.fullmatch(r"[a-z]{2}-[a-z]{2}", parts[0], flags=re.I):
        parts = parts[1:]
    if not tenant or not parts:
        return None
    origin = f"{parsed.scheme or 'https'}://{parsed.netloc}"
    return origin, tenant, parts[0]


def location_matches_region(location_text: str, region: str, city: str) -> bool:
    text = clean_text(location_text).lower()
    if not text:
        return True
    if "remote" in text:
        return True
    sg_terms = ["singapore", " sg ", "sgp"]
    hk_terms = ["hong kong", "hong kong sar", " hk "]
    cn_cities = ["beijing", "shanghai", "shenzhen", "hangzhou", "guangzhou", "chengdu", "nanjing", "suzhou", "wuhan", "xiamen"]
    cn_terms = ["china", "mainland", *cn_cities]
    other_terms = [
        "united states", "usa", "u.s.", "new york", "california", "san francisco",
        "india", "bangalore", "bengaluru", "hyderabad", "united kingdom", "london",
        "europe", "germany", "berlin", "france", "paris", "australia", "sydney",
        "melbourne", "canada", "toronto", "japan", "tokyo", "korea", "seoul",
        "taiwan", "taipei", "vietnam", "ho chi minh", "indonesia", "jakarta",
        "malaysia", "kuala lumpur", "thailand", "bangkok", "philippines", "manila",
    ]

    padded = f" {text} "
    if region == "SG":
        if any(term in padded for term in sg_terms):
            return True
        return not any(term in padded for term in [*hk_terms, *cn_terms, *other_terms])
    if region == "HK":
        if any(term in padded for term in hk_terms):
            return True
        return not any(term in padded for term in [*sg_terms, *cn_terms, *other_terms])
    if region == "CN":
        selected_city = clean_text(city).lower()
        if selected_city and selected_city in text:
            return True
        if any(term in padded for term in [*sg_terms, *hk_terms, *other_terms]):
            return False
        if any(term in text for term in cn_cities):
            return False
        return not any(term in text for term in ["china", "mainland"])
    return True


def explicit_ats_location_matches_region(location_text: str, region: str, city: str) -> bool:
    text = clean_text(location_text).lower()
    if not text:
        return False
    padded = f" {text} "
    if region == "SG":
        return "singapore" in text or bool(re.search(r"\bsg\b", padded))
    if region == "HK":
        return "hong kong" in text or bool(re.search(r"\bhk\b", padded))
    if region == "CN":
        selected_city = clean_text(city).lower()
        return bool(selected_city and selected_city in text) or (not selected_city and ("china" in text or "mainland" in text))
    return location_matches_region(location_text, region, city)


def fetch_company_ats_jobs(url: str, company: str, focus: str, region: str, city: str, limit: int) -> tuple[list[dict], list[str]]:
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    greenhouse = greenhouse_board_token(url)
    lever = lever_company_token(url)
    ashby = ashby_board_token(url)
    smartrecruiters = smartrecruiters_company_token(url)
    workable = workable_account_token(url)
    workday = workday_board_parts(url)
    try:
        if greenhouse:
            payload = json.loads(http_get(f"https://boards-api.greenhouse.io/v1/boards/{greenhouse}/jobs?content=true", timeout=COMPANY_SCAN_TIMEOUT, retries=0))
            for item in payload.get("jobs") or []:
                location_text = (item.get("location") or {}).get("name") if isinstance(item.get("location"), dict) else ""
                if not explicit_ats_location_matches_region(location_text, region, city):
                    continue
                append_company_job(
                    jobs,
                    seen,
                    company_job_record(
                        company,
                        item.get("title") or "",
                        item.get("absolute_url") or url,
                        region,
                        city,
                        focus,
                        url,
                        clean_text(item.get("content") or ""),
                        "Company Site / ATS",
                        location_text,
                    ),
                    limit,
                )
        elif lever:
            payload = json.loads(http_get(f"https://api.lever.co/v0/postings/{lever}?mode=json", timeout=COMPANY_SCAN_TIMEOUT, retries=0))
            for item in payload if isinstance(payload, list) else []:
                categories = item.get("categories") or {}
                location_text = categories.get("location") if isinstance(categories, dict) else ""
                if not explicit_ats_location_matches_region(location_text, region, city):
                    continue
                append_company_job(
                    jobs,
                    seen,
                    company_job_record(
                        company,
                        item.get("text") or "",
                        item.get("hostedUrl") or item.get("applyUrl") or url,
                        region,
                        city,
                        focus,
                        url,
                        item.get("descriptionPlain") or item.get("description") or "",
                        "Company Site / ATS",
                        location_text,
                    ),
                    limit,
                )
        elif ashby:
            payload = json.loads(http_get(f"https://api.ashbyhq.com/posting-api/job-board/{ashby}", timeout=COMPANY_SCAN_TIMEOUT, retries=0))
            for item in payload.get("jobs") or []:
                location_text = item.get("locationName") or item.get("location") or ""
                if not explicit_ats_location_matches_region(location_text, region, city):
                    continue
                append_company_job(
                    jobs,
                    seen,
                    company_job_record(
                        company,
                        item.get("title") or "",
                        item.get("jobUrl") or item.get("applyUrl") or url,
                        region,
                        city,
                        focus,
                        url,
                        item.get("descriptionHtml") or item.get("descriptionPlain") or "",
                        "Company Site / ATS",
                        location_text,
                    ),
                    limit,
                )
        elif smartrecruiters:
            offset = 0
            country_code = {"SG": "sg", "HK": "hk", "CN": "cn"}.get(region)
            listing_items: list[dict] = []
            for _page in range(5):
                query = {"limit": 100, "offset": offset}
                if country_code:
                    query["country"] = country_code
                payload = json.loads(http_get(
                    f"https://api.smartrecruiters.com/v1/companies/{smartrecruiters}/postings?{urlencode(query)}",
                    timeout=COMPANY_SCAN_TIMEOUT,
                    retries=0,
                ))
                postings = payload.get("content") or []
                for item in postings:
                    location = item.get("location") if isinstance(item.get("location"), dict) else {}
                    location_text = re.sub(
                        r",\s*,",
                        ",",
                        clean_text(location.get("fullLocation") or location.get("city") or ""),
                    )
                    if not explicit_ats_location_matches_region(location_text, region, city):
                        continue
                    listing_items.append(item)
                    if len(listing_items) >= limit:
                        break
                total = int(payload.get("totalFound") or len(postings))
                offset += int(payload.get("limit") or 100)
                if len(listing_items) >= limit or not postings or offset >= total:
                    break

            def fetch_smartrecruiters_detail(item: dict) -> tuple[dict, dict]:
                posting_id = clean_text(item.get("id") or "")
                official_detail_url = (
                    f"https://api.smartrecruiters.com/v1/companies/{smartrecruiters}/postings/{posting_id}"
                    if posting_id else ""
                )
                detail_url = clean_text(item.get("ref") or "")
                if urlparse(detail_url).netloc.lower() != "api.smartrecruiters.com":
                    detail_url = official_detail_url
                if not detail_url:
                    return item, {}
                try:
                    detail = json.loads(http_get(detail_url, timeout=COMPANY_SCAN_TIMEOUT, retries=0))
                except Exception:
                    detail = {}
                return item, detail if isinstance(detail, dict) else {}

            with ThreadPoolExecutor(max_workers=min(6, len(listing_items) or 1)) as detail_executor:
                detail_rows = list(detail_executor.map(fetch_smartrecruiters_detail, listing_items))
            for item, detail in detail_rows:
                location = item.get("location") if isinstance(item.get("location"), dict) else {}
                employment = item.get("typeOfEmployment") if isinstance(item.get("typeOfEmployment"), dict) else {}
                experience = item.get("experienceLevel") if isinstance(item.get("experienceLevel"), dict) else {}
                department = item.get("department") if isinstance(item.get("department"), dict) else {}
                function = item.get("function") if isinstance(item.get("function"), dict) else {}
                location_text = re.sub(
                    r",\s*,",
                    ",",
                    clean_text(location.get("fullLocation") or location.get("city") or ""),
                )
                metadata = "\n".join(
                    value for value in [
                        f"Employment: {clean_text(employment.get('label') or '')}" if employment.get("label") else "",
                        f"Experience: {clean_text(experience.get('label') or '')}" if experience.get("label") else "",
                        f"Department: {clean_text(department.get('label') or '')}" if department.get("label") else "",
                        f"Function: {clean_text(function.get('label') or '')}" if function.get("label") else "",
                    ]
                    if value
                )
                job_ad = detail.get("jobAd") if isinstance(detail.get("jobAd"), dict) else {}
                sections = job_ad.get("sections") if isinstance(job_ad.get("sections"), dict) else {}
                section_text = "\n".join(
                    clean_text(section.get("text") or "")
                    for key in ["companyDescription", "jobDescription", "qualifications", "additionalInformation"]
                    for section in [sections.get(key)]
                    if isinstance(section, dict) and clean_text(section.get("text") or "")
                )
                posting_item = {**item, **{key: detail.get(key) for key in ["postingUrl", "applyUrl"] if detail.get(key)}}
                append_company_job(
                    jobs,
                    seen,
                    company_job_record(
                        company,
                        item.get("name") or "",
                        smartrecruiters_public_job_url(posting_item, smartrecruiters) or url,
                        region,
                        city,
                        focus,
                        url,
                        "\n".join(value for value in [metadata, section_text] if value) or item.get("name") or "",
                        "Company Site / ATS",
                        location_text,
                    ),
                    limit,
                )
        elif workable:
            payload = json.loads(http_get(f"https://apply.workable.com/api/v1/widget/accounts/{workable}?details=true", timeout=COMPANY_SCAN_TIMEOUT, retries=0))
            for item in payload.get("jobs") or []:
                locations = item.get("locations") if isinstance(item.get("locations"), list) else []
                location_text = " ".join(
                    clean_text(" ".join(str(location.get(key) or "") for key in ["city", "region", "country", "countryCode"]))
                    for location in locations
                    if isinstance(location, dict)
                )
                location_text = location_text or clean_text(" ".join(str(item.get(key) or "") for key in ["city", "state", "country"]))
                if not explicit_ats_location_matches_region(location_text, region, city):
                    continue
                append_company_job(
                    jobs,
                    seen,
                    company_job_record(
                        company,
                        item.get("title") or "",
                        item.get("url") or item.get("shortlink") or item.get("application_url") or url,
                        region,
                        city,
                        focus,
                        url,
                        item.get("description") or "",
                        "Company Site / ATS",
                        location_text or city,
                    ),
                    limit,
                )
        elif workday:
            origin, tenant, site = workday
            api_root = f"{origin}/wday/cxs/{tenant}/{site}"
            seen_paths: set[str] = set()
            listing_items: list[tuple[dict, str, str, str]] = []
            for search_term in ["intern", "graduate", ""]:
                if len(listing_items) >= limit:
                    break
                offset = 0
                for _page in range(5):
                    payload = http_post_json(
                        f"{api_root}/jobs",
                        {
                            "appliedFacets": {},
                            "limit": 20,
                            "offset": offset,
                            "searchText": search_term,
                        },
                        timeout=COMPANY_SCAN_TIMEOUT,
                    )
                    postings = payload.get("jobPostings") if isinstance(payload, dict) else []
                    if not isinstance(postings, list) or not postings:
                        break
                    for item in postings:
                        title = clean_text(item.get("title") or "")
                        external_path = item.get("externalPath") or ""
                        location_text = clean_text(item.get("locationsText") or "")
                        if not external_path or external_path in seen_paths:
                            continue
                        if not explicit_ats_location_matches_region(location_text, region, city):
                            continue
                        employment = detect_employment_type(title, "", "")
                        if search_term == "intern" and employment != "Internship":
                            continue
                        if search_term == "graduate" and employment != "Graduate":
                            continue
                        seen_paths.add(external_path)
                        listing_items.append((item, external_path, title, location_text))
                        if len(listing_items) >= limit:
                            break
                    if len(listing_items) >= limit:
                        break
                    offset += len(postings)
                    if offset >= int(payload.get("total") or 0):
                        break

            def fetch_workday_detail(entry: tuple[dict, str, str, str]) -> tuple[dict, str, str, str]:
                item, external_path, title, location_text = entry
                try:
                    detail = json.loads(http_get(f"{api_root}{external_path}", timeout=COMPANY_SCAN_TIMEOUT, retries=0))
                except Exception:
                    detail = {}
                info = detail.get("jobPostingInfo") if isinstance(detail, dict) else {}
                return (info if isinstance(info, dict) else {}, external_path, title, location_text)

            with ThreadPoolExecutor(max_workers=min(6, len(listing_items) or 1)) as detail_executor:
                detail_rows = list(detail_executor.map(fetch_workday_detail, listing_items))
            for (item, _external_path, _title, _location_text), (info, external_path, title, location_text) in zip(listing_items, detail_rows):
                detail_location = clean_text(info.get("location") or location_text)
                metadata = "\n".join(
                    value for value in [
                        clean_text(info.get("jobDescription") or ""),
                        f"Employment: {clean_text(info.get('timeType') or '')}" if info.get("timeType") else "",
                        f"Posted: {clean_text(item.get('postedOn') or '')}" if item.get("postedOn") else "",
                        f"Requisition: {clean_text(info.get('jobReqId') or '')}" if info.get("jobReqId") else "",
                    ]
                    if value
                )
                public_url = info.get("externalUrl") or f"{origin}/{site}{external_path}"
                append_company_job(
                    jobs,
                    seen,
                    company_job_record(
                        company,
                        info.get("title") or title,
                        public_url,
                        region,
                        city,
                        focus,
                        url,
                        metadata,
                        "Company Site / ATS",
                        detail_location,
                    ),
                    limit,
                )
    except Exception as exc:
        failures.append(f"{company} ATS {url}: {exc}")
    return jobs[:limit], failures


def update_company_scan_result(company: str, region: str, status: str, note: str, jobs_found: int) -> None:
    with get_db() as conn:
        conn.execute(
            """
            update watch_companies
            set last_checked_at=?, last_scan_status=?, last_scan_note=?, last_jobs_found=?
            where region=? and company=?
            """,
            (now_iso(), status, note[:500], jobs_found, region, company),
        )


def fetch_company_site_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    jobs: list[dict] = []
    failures: list[str] = []
    code = active_region_code(region)
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    with get_db() as conn:
        companies = [row_to_dict(row) for row in conn.execute(
            """
            select company, url, focus, region, last_checked_at, last_scan_status
            from watch_companies
            where region=? and status='Watch'
            order by priority desc, company
            """,
            (code,),
        ).fetchall()]
    scan_now = dt.datetime.now()
    active_companies: list[dict] = []
    for company in companies:
        checked_at = company.get("last_checked_at") or ""
        try:
            checked_time = dt.datetime.fromisoformat(checked_at)
        except ValueError:
            checked_time = None
        on_cooldown = (
            company.get("last_scan_status") in {"failed", "limited"}
            and checked_time is not None
            and (scan_now - checked_time).total_seconds() < 6 * 60 * 60
        )
        if not on_cooldown:
            active_companies.append(company)
    companies = active_companies
    per_company = max(2, min(COMPANY_SCAN_PER_COMPANY_CAP, (limit // max(1, len(companies))) + 1))

    def scan_company(index_row: tuple[int, dict]) -> tuple[int, list[dict], list[str], str, str]:
        index, row = index_row
        company = row["company"]
        base_url = row["url"]
        company_jobs: list[dict] = []
        seen_urls: set[str] = set()
        company_failures: list[str] = []
        try:
            html = http_get(base_url, timeout=COMPANY_SCAN_TIMEOUT, retries=0)
        except Exception as exc:
            return index, [], [f"{company}: {exc}"], "failed", "官网访问失败，已保留公共来源匹配。"
        pages = [base_url, *extract_embedded_ats_links(html, base_url)]
        child_page_misses = 0
        for page_url in pages[:COMPANY_SCAN_PAGE_CAP]:
            if len(company_jobs) >= per_company:
                break
            ats_jobs, ats_failures = fetch_company_ats_jobs(page_url, company, row["focus"], code, city, per_company - len(company_jobs))
            company_failures.extend(ats_failures)
            for job in ats_jobs:
                append_company_job(company_jobs, seen_urls, job, per_company)
            if len(company_jobs) >= per_company:
                break
            page_html = html if page_url == base_url else ""
            if page_url != base_url:
                try:
                    page_html = http_get(page_url, timeout=COMPANY_SCAN_TIMEOUT, retries=0)
                except Exception:
                    child_page_misses += 1
                    if child_page_misses >= 2:
                        break
                    continue
            for job in parse_company_jsonld_jobs(page_html, page_url, company, row["focus"], code, city, per_company - len(company_jobs)):
                append_company_job(company_jobs, seen_urls, job, per_company)
            for job in parse_company_anchor_jobs(page_html, page_url, company, row["focus"], code, city, per_company - len(company_jobs)):
                append_company_job(company_jobs, seen_urls, job, per_company)
        if company_jobs:
            return index, company_jobs, [], "success", f"官网/ATS 本次找到 {len(company_jobs)} 个可识别岗位。"
        if company_failures:
            return index, [], company_failures[:3], "limited", "官网部分页面访问受限，公共来源匹配会继续补充。"
        return index, [], [], "empty", "官网未暴露可识别岗位列表，公共来源匹配会继续补充。"

    results: dict[int, tuple[list[dict], list[str], str, str]] = {}
    with ThreadPoolExecutor(max_workers=min(8, max(1, len(companies)))) as executor:
        futures = [executor.submit(scan_company, item) for item in enumerate(companies)]
        for future in as_completed(futures):
            index, company_jobs, company_failures, status, note = future.result()
            results[index] = (company_jobs, company_failures, status, note)

    for index, row in enumerate(companies):
        company_jobs, company_failures, status, note = results.get(index, ([], [f"{row['company']}: 扫描未返回。"], "failed", "官网扫描未返回，已保留公共来源匹配。"))
        jobs.extend(company_jobs[: max(0, limit - len(jobs))])
        failures.extend(company_failures)
        update_company_scan_result(row["company"], code, status, note, len(company_jobs))
    return jobs, failures


def classify_public_job_link(url: str, title: str, context_text: str, region: str) -> str | None:
    parsed = urlparse(url)
    host = parsed.netloc.lower().removeprefix("www.")
    path = parsed.path.lower().rstrip("/")
    title_lower = clean_text(title).lower()
    context_lower = clean_text(context_text).lower()
    if not title_lower or any(title_lower.startswith(prefix) for prefix in ["view all", "see all", "browse", "explore", "search"]):
        return None
    if host.endswith("mycareersfuture.gov.sg"):
        return "MyCareersFuture" if path.startswith("/job/") else None
    if host.endswith("glints.com"):
        if "/opportunities/jobs/" in path and not path.endswith("/explore"):
            return "Glints"
        return None
    if host.endswith("nodeflair.com"):
        return "NodeFlair" if re.search(r"/jobs/[^/?]*\d+", path) else None
    if host.endswith("wellfound.com"):
        if not re.fullmatch(r"/jobs/\d+-[^/]+", path):
            return None
        return "Wellfound" if region == "SG" and "singapore" in context_lower else None
    return None


def public_card_company(container, title: str) -> str:
    company_node = container.select_one('[data-company-name], [data-testid*="company"], [class*="company"], [class*="Company"]')
    if company_node:
        company = clean_text(company_node.get_text(" ", strip=True))
        if company and company.lower() != title.lower():
            return company[:140]
    title_lower = clean_text(title).lower()
    blocked = {
        title_lower,
        "singapore",
        "remote",
        "hybrid",
        "apply",
        "apply now",
        "view job",
        "save",
    }
    for value in container.stripped_strings:
        candidate = clean_text(str(value))
        lowered = candidate.lower()
        if not candidate or lowered in blocked or lowered.startswith("sgd "):
            continue
        if re.fullmatch(r"(?:singapore|remote|hybrid)(?:,?\s+singapore)?", lowered):
            continue
        if 2 <= len(candidate) <= 140:
            return candidate
    return ""


def parse_public_search_jobs_from_html(
    html: str,
    page_url: str,
    source: str,
    region: str,
    city: str,
    limit: int,
) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return []
    soup = BeautifulSoup(html or "", "html.parser")
    jobs: list[dict] = []
    seen: set[str] = set()
    fallback_sources = {"Mainland Public Search", "JobsDB"}
    role_pattern = re.compile(
        r"\b(intern|internship|graduate|associate|junior|ux|user research|design|designer|product|operations|analyst|content|marketing|ai)\b",
        flags=re.I,
    )
    for anchor in soup.find_all("a", href=True):
        if len(jobs) >= limit:
            break
        title = clean_text(anchor.get_text(" ", strip=True))
        if len(title) < 8 or not role_pattern.search(title):
            continue
        job_url = absolute_url(page_url, anchor.get("href") or "")
        if not job_url.startswith("http") or job_url in seen:
            continue
        container = anchor.find_parent(["article", "li"]) or anchor.parent
        context_text = clean_text(container.get_text(" ", strip=True)) if container else title
        channel = classify_public_job_link(job_url, title, context_text, region)
        if channel is None and source not in fallback_sources:
            continue
        channel = channel or source
        company = public_card_company(container, title) if container else ""
        if not company and source not in fallback_sources:
            continue
        seen.add(job_url)
        jobs.append(
            {
                "company": company or source,
                "position": title[:180],
                "source": channel,
                "url": job_url,
                "location": city,
                "region": region,
                "city": city,
                "source_region": region,
                "job_type": "Public search",
                "jd_text": f"{channel} public search match\nTitle: {title}\nCompany: {company or source}\nLocation: {city}\nURL: {job_url}\nContext: {context_text[:1200]}",
            }
        )
    return jobs


def public_search_unavailable_reason(html: str) -> str:
    text = clean_text(html).lower()
    if "temporarily unable to search for jobs" in text:
        return "limited: site reports that job search is temporarily unavailable"
    if "service temporarily unavailable" in text or "temporarily unavailable" in text:
        return "limited: site is temporarily unavailable"
    return ""


def generic_public_search_jobs(
    source: str,
    base_urls: list[str],
    limit: int,
    region: str,
    time_budget_seconds: float = 35,
    host_failure_limit: int = 2,
    unparseable_page_is_limited: bool = False,
) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    jobs: list[dict] = []
    failures: list[str] = []
    seen: set[str] = set()
    started = time.monotonic()
    host_failures: dict[str, int] = {}
    blocked_hosts: set[str] = set()
    source_hosts = {urlparse(template).netloc.lower() for template in base_urls}

    def record_failure(host: str, query: str, detail: object) -> None:
        failures.append(f"{source} {host} {query}: {detail}")
        host_failures[host] = host_failures.get(host, 0) + 1
        if host_failures[host] >= host_failure_limit:
            blocked_hosts.add(host)

    for query in region_queries(code):
        if len(jobs) >= limit or time.monotonic() - started >= time_budget_seconds or blocked_hosts == source_hosts:
            break
        for template in base_urls:
            if len(jobs) >= limit or time.monotonic() - started >= time_budget_seconds:
                break
            url = template.format(query=quote_plus(query), city=quote_plus(city))
            host = urlparse(url).netloc.lower()
            if host in blocked_hosts:
                continue
            try:
                html = http_get(url, timeout=10, retries=0)
            except Exception as exc:
                record_failure(host, query, exc)
                continue
            unavailable_reason = public_search_unavailable_reason(html)
            if unavailable_reason:
                record_failure(host, query, unavailable_reason)
                continue
            parsed_jobs = parse_public_search_jobs_from_html(html, url, source, code, city, limit - len(jobs))
            if unparseable_page_is_limited and not parsed_jobs:
                record_failure(host, query, "limited: public response contains no parseable job cards")
                continue
            for job in parsed_jobs:
                if job["url"] in seen:
                    continue
                seen.add(job["url"])
                job["jd_text"] += f"\nQuery: {query}"
                jobs.append(job)
            if jobs:
                break
    return jobs[:limit], failures


def fetch_mainland_public_jobs(limit: int) -> tuple[list[dict], list[str]]:
    return generic_public_search_jobs(
        "Mainland Public Search",
        [
            "https://www.zhipin.com/web/geek/job?query={query}",
            "https://www.lagou.com/wn/jobs?kd={query}",
        ],
        limit,
        "CN",
    )


def fetch_jobsdb_hk_jobs(limit: int) -> tuple[list[dict], list[str]]:
    return generic_public_search_jobs(
        "JobsDB",
        [
            "https://hk.jobsdb.com/{query}-jobs",
            "https://hk.jobsdb.com/jobs-in-{city}?keywords={query}",
        ],
        limit,
        "HK",
    )


def parse_mycareersfuture_api_jobs(payload: dict, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    for item in payload.get("results") or []:
        if len(jobs) >= limit or not isinstance(item, dict):
            break
        job_id = clean_text(item.get("uuid") or "")
        title = clean_text(item.get("title") or "")
        status_data = item.get("status")
        status = clean_text(
            status_data.get("jobStatus") if isinstance(status_data, dict) else status_data or ""
        ).lower()
        metadata = item.get("metadata") if isinstance(item.get("metadata"), dict) else {}
        expiry_date = clean_text(metadata.get("expiryDate") or "")[:10]
        address = item.get("address") if isinstance(item.get("address"), dict) else {}
        employment_types = [
            clean_text(value.get("employmentType") or "")
            for value in (item.get("employmentTypes") or [])
            if isinstance(value, dict) and value.get("employmentType")
        ]
        if (
            not job_id
            or not is_actionable_job_title(title)
            or status in {"closed", "expired", "deleted"}
            or (expiry_date and expiry_date < today())
            or bool(address.get("isOverseas"))
            or (employment_types and not any("internship" in value.lower() for value in employment_types))
        ):
            continue
        if not employment_types and detect_employment_type(title, clean_text(item.get("description") or ""), "") not in {"Internship", "Graduate"}:
            continue
        company_data = item.get("hiringCompany") if isinstance(item.get("hiringCompany"), dict) else None
        company_data = company_data or (item.get("postedCompany") if isinstance(item.get("postedCompany"), dict) else {})
        company = clean_text(company_data.get("name") or "")
        if not is_actionable_company_name(company, title):
            continue
        job_url = clean_text(metadata.get("jobDetailsUrl") or "") or f"https://www.mycareersfuture.gov.sg/job/{job_id}"
        if job_url in seen:
            continue
        seen.add(job_url)
        districts = [
            clean_text(value.get("location") or "")
            for value in (address.get("districts") or [])
            if isinstance(value, dict) and value.get("location")
        ]
        location = "Singapore" + (f" · {districts[0]}" if districts else "")
        salary = item.get("salary") if isinstance(item.get("salary"), dict) else {}
        salary_type = salary.get("type") if isinstance(salary.get("type"), dict) else {}

        def salary_number(value) -> str:
            try:
                number = float(value)
                return str(int(number)) if number.is_integer() else f"{number:g}"
            except (TypeError, ValueError):
                return ""

        minimum_salary = salary_number(salary.get("minimum"))
        maximum_salary = salary_number(salary.get("maximum"))
        salary_parts = [value for value in [minimum_salary, maximum_salary] if value]
        salary_text = ""
        if salary_parts:
            salary_range = " - ".join(dict.fromkeys(salary_parts))
            salary_text = f"SGD {salary_range} {clean_text(salary_type.get('salaryType') or '')}".strip()
        position_levels = [
            clean_text(value.get("position") or "")
            for value in (item.get("positionLevels") or [])
            if isinstance(value, dict) and value.get("position")
        ]
        skills = [
            clean_text(value.get("skill") or "")
            for value in (item.get("skills") or [])
            if isinstance(value, dict) and value.get("skill")
        ]
        description = clean_text(item.get("description") or "")
        jd_metadata = "\n".join(
            value for value in [
                f"Employment: {', '.join(employment_types)}" if employment_types else "",
                f"Position level: {', '.join(position_levels)}" if position_levels else "",
                f"Salary: {salary_text}" if salary_text else "",
                f"Minimum experience: {item.get('minimumYearsExperience')} years" if item.get("minimumYearsExperience") is not None else "",
                f"Skills: {', '.join(skills[:20])}" if skills else "",
                f"Posted: {clean_text(metadata.get('newPostingDate') or metadata.get('originalPostingDate') or '')}",
                f"Expires: {expiry_date}" if expiry_date else "",
            ]
            if value
        )
        jobs.append({
            "company": company,
            "position": title[:180],
            "source": "MyCareersFuture",
            "url": job_url,
            "external_job_id": job_id,
            "location": location,
            "region": "SG",
            "city": "Singapore",
            "source_region": "SG",
            "job_type": employment_types[0] if employment_types else "Internship",
            "jd_text": f"{description}\n\n{jd_metadata}".strip(),
        })
    return jobs


def fetch_mycareersfuture_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], ["MyCareersFuture connector is Singapore-only in this version."]
    queries = [
        "ai product intern",
        "product design intern",
        "user research intern",
        "marketing intern",
        "internship",
    ]

    def fetch_query(index_query: tuple[int, str]) -> tuple[int, list[dict], str | None]:
        index, query = index_query
        try:
            payload = http_post_json(
                "https://api.mycareersfuture.gov.sg/v2/search?limit=40&page=0",
                {
                    "search": query,
                    "sortBy": ["new_posting_date"],
                    "sessionId": str(uuid.uuid4()),
                    "employmentTypes": ["Internship/Attachment"],
                },
                timeout=12,
            )
            return index, parse_mycareersfuture_api_jobs(payload, 40), None
        except Exception as exc:
            return index, [], f"MyCareersFuture {query}: limited: public API request failed ({exc})"

    results: dict[int, list[dict]] = {}
    failures: list[str] = []
    with ThreadPoolExecutor(max_workers=len(queries)) as executor:
        futures = [executor.submit(fetch_query, item) for item in enumerate(queries)]
        for future in as_completed(futures):
            index, query_jobs, failure = future.result()
            results[index] = query_jobs
            if failure:
                failures.append(failure)

    jobs: list[dict] = []
    seen: set[str] = set()
    round_index = 0
    while len(jobs) < limit and any(round_index < len(results.get(index, [])) for index in range(len(queries))):
        for index in range(len(queries)):
            query_jobs = results.get(index, [])
            if round_index >= len(query_jobs):
                continue
            job = query_jobs[round_index]
            key = canonical_job_url(job.get("source") or "", job.get("url") or "", job.get("external_job_id"))
            if key and key not in seen:
                seen.add(key)
                jobs.append(job)
                if len(jobs) >= limit:
                    break
        round_index += 1
    return jobs, failures


CAREERS_GOV_QUERIES = [
    "product intern",
    "design intern",
    "user research intern",
    "AI intern",
    "marketing intern",
    "internship",
]
CAREERS_GOV_ALGOLIA_URL = "https://3ow7d8b4iz-dsn.algolia.net/1/indexes/job_index/query"
CAREERS_GOV_ALGOLIA_APP_ID = "3OW7D8B4IZ"
CAREERS_GOV_ALGOLIA_SEARCH_KEY = "32fa71d8b0bc06be1e6395bf8c430107"


def careers_gov_job_url(object_id: str, job_source: str = "") -> str:
    raw = clean_text(object_id)
    if ":" in raw:
        prefix, identifier = raw.split(":", 1)
    else:
        prefix, identifier = job_source, raw
    prefix = clean_text(job_source or prefix).lower()
    identifier = identifier.strip("/")
    if not prefix or not identifier:
        return ""
    return f"https://jobs.careers.gov.sg/jobs/{quote(prefix, safe='')}/{quote(identifier, safe='/')}"


def parse_careers_gov_algolia_jobs(payload: dict, query: str, limit: int) -> list[dict]:
    jobs: list[dict] = []
    seen: set[str] = set()
    for item in payload.get("hits") or []:
        if len(jobs) >= limit or not isinstance(item, dict):
            break
        position = clean_text(item.get("title") or "")
        company = clean_text(item.get("agency") or item.get("agencyAbbr") or "")
        employment_label = clean_text(item.get("employmentType") or "")
        employment_lower = employment_label.lower()
        if "intern" in employment_lower:
            employment_type = "Internship"
        elif "trainee" in employment_lower or "graduate" in employment_lower:
            employment_type = "Graduate"
        else:
            continue
        if not is_actionable_job_title(position) or not is_actionable_company_name(company, position):
            continue
        object_id = clean_text(item.get("objectID") or "")
        job_source = clean_text(item.get("jobSource") or object_id.partition(":")[0])
        url = careers_gov_job_url(object_id, job_source)
        if not url or url in seen:
            continue
        seen.add(url)
        description = clean_text(item.get("description") or "")
        department = clean_text(item.get("department") or "")
        updated = ""
        try:
            timestamp = float(item.get("activityTimestamp") or 0) / 1000
            if timestamp > 0:
                updated = dt.datetime.fromtimestamp(timestamp, tz=dt.timezone.utc).strftime(DATE_FMT)
        except (TypeError, ValueError, OverflowError, OSError):
            pass
        metadata = "\n".join(value for value in [
            f"Department: {department}" if department else "",
            f"Employment: {employment_label}" if employment_label else "",
            f"Updated: {updated}" if updated else "",
            f"Search query: {query}" if query else "",
        ] if value)
        jobs.append({
            "company": company,
            "position": position[:180],
            "source": "Careers@Gov",
            "url": url,
            "external_job_id": object_id.partition(":")[2] or object_id,
            "location": "Singapore",
            "region": "SG",
            "city": "Singapore",
            "source_region": "SG",
            "job_type": employment_label or employment_type,
            "employment_type": employment_type,
            "jd_text": f"{description}\n\n{metadata}".strip(),
        })
    return jobs


def search_careers_gov_jobs(query: str, limit: int) -> list[dict]:
    request = urllib.request.Request(
        CAREERS_GOV_ALGOLIA_URL,
        data=json.dumps({
            "query": query,
            "hitsPerPage": max(1, min(100, limit)),
            "attributesToHighlight": [],
            "queryLanguages": ["en"],
        }).encode("utf-8"),
        headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/124 Safari/537.36",
            "Accept": "application/json",
            "Content-Type": "application/json",
            "Origin": "https://jobs.careers.gov.sg",
            "Referer": "https://jobs.careers.gov.sg/",
            "X-Algolia-Application-Id": CAREERS_GOV_ALGOLIA_APP_ID,
            "X-Algolia-API-Key": CAREERS_GOV_ALGOLIA_SEARCH_KEY,
        },
        method="POST",
    )
    with urllib.request.urlopen(request, timeout=12) as response:
        payload = json.loads(response.read().decode("utf-8", errors="ignore"))
    return parse_careers_gov_algolia_jobs(payload, query, limit)


def parse_careers_gov_detail_html(html: str, listing: dict) -> dict:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return dict(listing)
    soup = BeautifulSoup(html or "", "html.parser")
    job = dict(listing)
    title_node = soup.find("h1")
    if title_node:
        job["position"] = clean_text(title_node.get_text(" ", strip=True))[:180] or job.get("position")
    sections: list[str] = []
    allowed_headings = {
        "what the role is",
        "what you will be working on",
        "what we are looking for",
    }
    for heading in soup.find_all("h2"):
        label = clean_text(heading.get_text(" ", strip=True))
        if label.lower() not in allowed_headings:
            continue
        container = heading.parent
        content = clean_text(container.get_text(" ", strip=True)) if container else ""
        if content.lower().startswith(label.lower()):
            content = clean_text(content[len(label):])
        if content:
            sections.append(f"{label}\n{content}")
    page_text = clean_text(soup.get_text(" ", strip=True))
    closing_match = re.search(r"\b(?:closing on|job is closing on)\s+(\d{1,2}\s+[A-Za-z]{3,9}\s+\d{4})", page_text, flags=re.I)
    closing_text = f"Closing date: {closing_match.group(1)}" if closing_match else ""
    if sections:
        job["jd_text"] = "\n\n".join([*sections, *([closing_text] if closing_text else [])])[:12000]
    elif closing_text:
        job["jd_text"] = f"{job.get('jd_text') or ''}\n\n{closing_text}".strip()
    job["job_type"] = detect_employment_type(job.get("position") or "", job.get("jd_text") or "", job.get("job_type") or "")
    return job


def fetch_careers_gov_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], []
    query_results: dict[int, list[dict]] = {}
    failures: list[str] = []

    def fetch_query(index_query: tuple[int, str]) -> tuple[int, list[dict], str | None]:
        index, query = index_query
        try:
            return index, search_careers_gov_jobs(query, 60), None
        except Exception as exc:
            return index, [], f"Careers@Gov {query}: {exc}"

    with ThreadPoolExecutor(max_workers=len(CAREERS_GOV_QUERIES)) as executor:
        futures = [executor.submit(fetch_query, item) for item in enumerate(CAREERS_GOV_QUERIES)]
        for future in as_completed(futures):
            index, jobs, failure = future.result()
            query_results[index] = jobs
            if failure and len(failures) < 3:
                failures.append(failure)

    selected: list[dict] = []
    seen: set[str] = set()
    round_index = 0
    while len(selected) < limit and any(round_index < len(query_results.get(index, [])) for index in range(len(CAREERS_GOV_QUERIES))):
        for index in range(len(CAREERS_GOV_QUERIES)):
            jobs = query_results.get(index, [])
            if round_index >= len(jobs):
                continue
            job = jobs[round_index]
            url = job.get("url") or ""
            if url and url not in seen:
                seen.add(url)
                selected.append(job)
                if len(selected) >= limit:
                    break
        round_index += 1

    enriched: dict[int, dict] = {}

    def fetch_detail(index_job: tuple[int, dict]) -> tuple[int, dict, str | None]:
        index, job = index_job
        try:
            html = http_get(job.get("url") or "", timeout=8, retries=0)
            return index, parse_careers_gov_detail_html(html, job), None
        except Exception as exc:
            return index, job, f"Careers@Gov detail {job.get('position')}: {exc}"

    with ThreadPoolExecutor(max_workers=min(8, max(1, len(selected)))) as executor:
        futures = [executor.submit(fetch_detail, item) for item in enumerate(selected)]
        for future in as_completed(futures):
            index, job, failure = future.result()
            enriched[index] = job
            if failure and len(failures) < 3:
                failures.append(failure)
    return [enriched.get(index, job) for index, job in enumerate(selected)], failures


INTERNSHIP_SG_QUERIES = [
    "design intern",
    "product design",
    "AI intern",
    "content intern",
    "operations intern",
    "marketing intern",
]


def parse_internship_sg_search_html(html: str, query: str, limit: int) -> list[dict]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return []
    soup = BeautifulSoup(html or "", "html.parser")
    jobs: list[dict] = []
    seen: set[str] = set()
    for card in soup.select('a[href^="/internships/"]'):
        position_node = card.select_one("h3")
        company_node = card.select_one(".lbl")
        if not position_node or not company_node:
            continue
        position = clean_text(position_node.get_text(" ", strip=True))
        company = clean_text(company_node.get_text(" ", strip=True))
        if not is_actionable_job_title(position) or not is_actionable_company_name(company, position):
            continue
        employment_type = detect_employment_type(position)
        if employment_type not in {"Internship", "Graduate"}:
            continue
        url = absolute_url("https://internship.sg", card.get("href") or "")
        if not url or url in seen:
            continue
        seen.add(url)
        metadata = [
            clean_text(node.get_text(" ", strip=True))
            for node in card.select(".mono span")
        ]
        metadata = [value for value in metadata if value and value != "·"]
        location = metadata[0] if metadata else "Singapore"
        external_match = re.search(r"-(\d{10,}(?:-\d+)?)$", urlparse(url).path.rstrip("/"))
        jobs.append({
            "company": company,
            "position": position[:180],
            "source": "Internship.sg",
            "url": url,
            "external_job_id": external_match.group(1) if external_match else None,
            "location": location,
            "region": "SG",
            "city": "Singapore",
            "source_region": "SG",
            "job_type": employment_type,
            "jd_text": f"{position}\n{company}\n{' · '.join(metadata)}\nSource query: {query}".strip(),
        })
        if len(jobs) >= limit:
            break
    return jobs


def parse_internship_sg_detail_html(html: str, listing: dict) -> dict:
    try:
        from bs4 import BeautifulSoup
    except ImportError:  # pragma: no cover - requirements include BeautifulSoup.
        return dict(listing)
    soup = BeautifulSoup(html or "", "html.parser")
    job = dict(listing)
    title_node = soup.select_one("h1")
    company_node = soup.select_one('a[href^="/companies/"]')
    if title_node:
        job["position"] = clean_text(title_node.get_text(" ", strip=True))[:180] or job.get("position")
    if company_node:
        job["company"] = clean_text(company_node.get_text(" ", strip=True)) or job.get("company")

    details: dict[str, list[str]] = {}
    for group in soup.select("dl > div"):
        term = group.select_one("dt")
        if not term:
            continue
        key = clean_text(term.get_text(" ", strip=True)).lower()
        values = [clean_text(node.get_text(" ", strip=True)) for node in group.select("dd")]
        details[key] = [value for value in values if value]
    location_parts = details.get("location") or []
    if location_parts:
        job["location"] = " · ".join(dict.fromkeys(location_parts))
    allowance = (details.get("allowance") or [""])[0]
    if allowance and "not disclosed" not in allowance.lower() and "not stated" not in allowance.lower():
        job["salary_text"] = allowance
    duration = (details.get("duration") or [""])[0]
    job["job_type"] = " · ".join(value for value in [detect_employment_type(job.get("position") or ""), duration] if value)

    descriptions = [clean_text(node.get_text(" ", strip=True)) for node in soup.select(".prose-editorial")]
    description = max(descriptions, key=len, default="")
    original_url = ""
    for link in soup.select("a[href]"):
        text = clean_text(link.get_text(" ", strip=True)).lower()
        href = absolute_url("https://internship.sg", link.get("href") or "")
        if "original listing" in text and urlparse(href).netloc not in {"", "internship.sg", "www.internship.sg"}:
            original_url = href
            break
    evidence = f"Original listing: {original_url}" if original_url else ""
    if description:
        job["jd_text"] = "\n\n".join(value for value in [description[:12000], evidence] if value)
    elif evidence:
        job["jd_text"] = f"{job.get('jd_text') or ''}\n\n{evidence}".strip()
    return job


def fetch_internship_sg_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], []

    def fetch_search(index_query: tuple[int, str]) -> tuple[int, list[dict], str | None]:
        index, query = index_query
        url = f"https://internship.sg/search?q={quote_plus(query).replace('+', '%20')}"
        try:
            html = http_get(url, timeout=10, retries=0)
            return index, parse_internship_sg_search_html(html, query, 30), None
        except Exception as exc:
            return index, [], f"Internship.sg {query}: {exc}"

    search_results: dict[int, list[dict]] = {}
    failures: list[str] = []
    with ThreadPoolExecutor(max_workers=len(INTERNSHIP_SG_QUERIES)) as executor:
        futures = [executor.submit(fetch_search, item) for item in enumerate(INTERNSHIP_SG_QUERIES)]
        for future in as_completed(futures):
            index, query_jobs, failure = future.result()
            search_results[index] = query_jobs
            if failure and len(failures) < 3:
                failures.append(failure)

    selected: list[dict] = []
    seen: set[str] = set()
    round_index = 0
    while len(selected) < limit and any(round_index < len(search_results.get(index, [])) for index in range(len(INTERNSHIP_SG_QUERIES))):
        for index in range(len(INTERNSHIP_SG_QUERIES)):
            query_jobs = search_results.get(index, [])
            if round_index >= len(query_jobs):
                continue
            job = query_jobs[round_index]
            url = job.get("url") or ""
            if url and url not in seen:
                seen.add(url)
                selected.append(job)
                if len(selected) >= limit:
                    break
        round_index += 1

    enriched: dict[int, dict] = {}

    def fetch_detail(index_job: tuple[int, dict]) -> tuple[int, dict, str | None]:
        index, job = index_job
        try:
            html = http_get(job.get("url") or "", timeout=8, retries=0)
            return index, parse_internship_sg_detail_html(html, job), None
        except Exception as exc:
            return index, job, f"Internship.sg detail {job.get('position')}: {exc}"

    with ThreadPoolExecutor(max_workers=min(6, max(1, len(selected)))) as executor:
        futures = [executor.submit(fetch_detail, item) for item in enumerate(selected)]
        for future in as_completed(futures):
            index, job, failure = future.result()
            enriched[index] = job
            if failure and len(failures) < 3:
                failures.append(failure)
    return [enriched.get(index, job) for index, job in enumerate(selected)], failures


def fetch_cultjobs_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], []
    failures: list[str] = []
    try:
        listing_html = http_get("https://cultjobs.com/job-type/internship/", timeout=12, retries=0)
    except Exception as exc:
        return [], [f"Cultjobs listing: {exc}"]
    urls = parse_cultjobs_listing_urls(listing_html, limit + 6)
    parsed_by_index: dict[int, dict] = {}

    def fetch_detail(index_url: tuple[int, str]) -> tuple[int, dict | None, str | None]:
        index, url = index_url
        try:
            detail_html = http_get(url, timeout=6, retries=0)
        except Exception as exc:
            return index, None, f"Cultjobs detail {url}: {exc}"
        return index, parse_cultjobs_detail(detail_html, url), None

    with ThreadPoolExecutor(max_workers=min(6, max(1, len(urls)))) as executor:
        futures = [executor.submit(fetch_detail, item) for item in enumerate(urls)]
        for future in as_completed(futures):
            index, job, failure = future.result()
            if failure and len(failures) < 3:
                failures.append(failure)
            if job:
                parsed_by_index[index] = job
    jobs = [parsed_by_index[index] for index in sorted(parsed_by_index)][:limit]
    if not jobs and not failures:
        failures.append("Cultjobs limited: no current Singapore internships found.")
    return jobs, failures


def fetch_sg_startup_channel_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], []
    raw_jobs, failures = fetch_jobstreet_jobs(
        max(limit * 2, limit),
        STARTUP_OPPORTUNITY_QUERIES,
        code,
        time_budget_seconds=10,
        failure_limit=2,
        use_html_fallback=False,
    )
    jobs: list[dict] = []
    for job in raw_jobs:
        employment = detect_employment_type(job.get("position") or "", job.get("jd_text") or "", job.get("job_type") or "")
        if employment not in {"Internship", "Graduate"}:
            continue
        identity_text = clean_ai_detection_text(f"{job.get('company') or ''} {job.get('position') or ''}")
        if not any(has_keyword(identity_text, signal) for signal in STARTUP_OPPORTUNITY_SIGNALS):
            continue
        job["source"] = "JobStreet · 创业/AI"
        jobs.append(job)
        if len(jobs) >= limit:
            break
    return jobs, failures


def fetch_sg_ai_startup_ats_jobs(limit: int, region: str | None = None) -> tuple[list[dict], list[str]]:
    code = active_region_code(region)
    if code != "SG":
        return [], []
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    per_board_limit = max(4, min(8, limit))
    raw_board_limit = max(40, per_board_limit)

    def fetch_board(index_board: tuple[int, tuple[str, str, str]]) -> tuple[int, list[dict], list[str]]:
        index, (company, url, focus) = index_board
        try:
            board_jobs, board_failures = fetch_company_ats_jobs(url, company, focus, code, city, raw_board_limit)
        except Exception as exc:
            return index, [], [f"{company} ATS: {exc}"]
        entry_jobs: list[dict] = []
        for job in board_jobs:
            location = clean_text(job.get("location") or "").lower()
            if "singapore" not in location:
                continue
            employment = detect_employment_type(
                job.get("position") or "",
                job.get("jd_text") or "",
                job.get("job_type") or "",
            )
            if employment not in {"Internship", "Graduate"}:
                continue
            if employment == "Graduate" and detect_employment_type(
                job.get("position") or "",
                "",
                job.get("job_type") or "",
            ) != "Graduate":
                continue
            job["employment_type"] = employment
            job["job_type"] = employment
            job["source"] = "ATS · 科技初创"
            entry_jobs.append(job)
            if len(entry_jobs) >= per_board_limit:
                break
        return index, entry_jobs, board_failures

    results: dict[int, tuple[list[dict], list[str]]] = {}
    with ThreadPoolExecutor(max_workers=min(6, len(SG_AI_STARTUP_ATS_BOARDS))) as executor:
        futures = [executor.submit(fetch_board, item) for item in enumerate(SG_AI_STARTUP_ATS_BOARDS)]
        for future in as_completed(futures):
            index, jobs, failures = future.result()
            results[index] = (jobs, failures)

    jobs_by_company: list[list[dict]] = []
    failures: list[str] = []
    for index in range(len(SG_AI_STARTUP_ATS_BOARDS)):
        board_jobs, board_failures = results.get(index, ([], [f"{SG_AI_STARTUP_ATS_BOARDS[index][0]} ATS: no response"]))
        jobs_by_company.append(board_jobs)
        failures.extend(board_failures[:1])

    jobs: list[dict] = []
    seen: set[str] = set()
    round_index = 0
    while len(jobs) < limit and any(round_index < len(items) for items in jobs_by_company):
        for company_jobs in jobs_by_company:
            if round_index >= len(company_jobs):
                continue
            job = company_jobs[round_index]
            key = canonical_job_url(job.get("source") or "", job.get("url") or "", job.get("external_job_id"))
            if key and key not in seen:
                seen.add(key)
                jobs.append(job)
                if len(jobs) >= limit:
                    break
        round_index += 1
    return jobs, failures


def scan_source_definitions(region: str | None = None) -> list[tuple[str, object, int]]:
    code = active_region_code(region)
    if code == "SG":
        sources = [
            ("LinkedIn（含 AI 关键词）", lambda limit: fetch_linkedin_jobs_with_ai(limit, code), SOURCE_LIMITS["LinkedIn"] + SOURCE_LIMITS["LinkedIn AI"]),
            ("InternSG（含 AI 关键词）", fetch_internsg_jobs_with_ai, SOURCE_LIMITS["InternSG"] + SOURCE_LIMITS["InternSG AI"]),
            ("Cultjobs", lambda limit: fetch_cultjobs_jobs(limit, code), SOURCE_LIMITS["Cultjobs"]),
            ("MyCareersFuture", lambda limit: fetch_mycareersfuture_jobs(limit, code), SOURCE_LIMITS["MyCareersFuture"]),
            ("Careers@Gov", lambda limit: fetch_careers_gov_jobs(limit, code), SOURCE_LIMITS["Careers@Gov"]),
            ("Internship.sg", lambda limit: fetch_internship_sg_jobs(limit, code), SOURCE_LIMITS["Internship.sg"]),
            ("新加坡科技与 AI ATS", lambda limit: fetch_sg_ai_startup_ats_jobs(limit, code), SOURCE_LIMITS["AI Startup ATS"]),
            ("Indeed", lambda limit: fetch_indeed_jobs_with_ai(limit, code), SOURCE_LIMITS["Indeed"] + SOURCE_LIMITS["Indeed AI"]),
            ("JobStreet", lambda limit: fetch_jobstreet_jobs_with_ai(limit, code), SOURCE_LIMITS["JobStreet"] + SOURCE_LIMITS["JobStreet AI"]),
            ("关注公司公开来源", lambda limit: fetch_watched_company_public_jobs(limit, code), SOURCE_LIMITS["Watched Companies"]),
            ("创业与 AI 机会", lambda limit: fetch_sg_startup_channel_jobs(limit, code), SOURCE_LIMITS["Startup Channels"]),
            ("公司官网", lambda limit: fetch_company_site_jobs(limit, code), SOURCE_LIMITS["Company Site"]),
        ]
        if os.environ.get("SERPAPI_KEY"):
            sources.insert(2, ("Google Jobs", lambda limit: fetch_google_jobs(limit, code), SOURCE_LIMITS["Google Jobs"]))
        return sources
    if code == "HK":
        return [
            ("LinkedIn（含 AI 关键词）", lambda limit: fetch_linkedin_jobs_with_ai(limit, code), SOURCE_LIMITS["LinkedIn"] + SOURCE_LIMITS["LinkedIn AI"]),
            ("JobsDB", fetch_jobsdb_hk_jobs, 24),
            ("公司官网", lambda limit: fetch_company_site_jobs(limit, code), SOURCE_LIMITS["Company Site"]),
        ]
    return [
        ("LinkedIn", lambda limit: fetch_linkedin_jobs(limit, region=code), SOURCE_LIMITS["LinkedIn"]),
        ("Mainland Public Search", fetch_mainland_public_jobs, 24),
        ("公司官网", lambda limit: fetch_company_site_jobs(limit, code), SOURCE_LIMITS["Company Site"]),
    ]


def expected_scan_sources(region: str | None = None) -> list[str]:
    return [source_name for source_name, _fetcher, _limit in scan_source_definitions(region)]


def scan_source_mode(source: str) -> str:
    return SCAN_SOURCE_MODES.get(source, "supplemental" if is_limited_scan_source(source) else "primary")


def expected_scan_source_details(region: str | None = None) -> list[dict]:
    return [
        {"source": source, "mode": scan_source_mode(source)}
        for source in expected_scan_sources(region)
    ]


def is_limited_scan_source(source: str) -> bool:
    return any(name.lower() in (source or "").lower() for name in LIMITED_SCAN_SOURCES)


def has_limited_failure(failures: list[dict | str]) -> bool:
    for failure in failures:
        text = failure.get("error", "") if isinstance(failure, dict) else str(failure)
        if "受限" in text or "limited" in text.lower():
            return True
    return False


def create_scan_run(triggered_by: str = "manual", forced: bool = False, region: str | None = None) -> int:
    stamp = now_iso()
    code = active_region_code(region)
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    with get_db() as conn:
        conn.execute(
            """
            insert into scan_runs(
                run_date, started_at, status, triggered_by, forced,
                region, city, source_region, created_at, updated_at
            )
            values(?, ?, 'running', ?, ?, ?, ?, ?, ?, ?)
            """,
            (today(), stamp, triggered_by, 1 if forced else 0, code, city, code, stamp, stamp),
        )
        return conn.execute("select last_insert_rowid()").fetchone()[0]


def create_scan_source_run(scan_run_id: int, source: str) -> int:
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            insert into scan_source_runs(scan_run_id, source, started_at, status, created_at, updated_at)
            values(?, ?, ?, 'running', ?, ?)
            """,
            (scan_run_id, source, stamp, stamp, stamp),
        )
        return conn.execute("select last_insert_rowid()").fetchone()[0]


def finish_scan_source_run(
    source_run_id: int,
    status: str,
    scanned: int,
    saved: int,
    failures: list[dict | str],
    new_count: int = 0,
    updated_count: int = 0,
    duplicate_count: int = 0,
) -> None:
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            update scan_source_runs set
                finished_at=?,
                status=?,
                scanned_count=?,
                saved_count=?,
                new_count=?,
                updated_count=?,
                duplicate_count=?,
                failure_count=?,
                failures_json=?,
                updated_at=?
            where id=?
            """,
            (
                stamp,
                status,
                scanned,
                saved,
                new_count,
                updated_count,
                duplicate_count,
                len(failures),
                json.dumps(failures, ensure_ascii=False),
                stamp,
                source_run_id,
            ),
        )


def finish_scan_run(
    scan_run_id: int,
    status: str,
    scanned: int,
    saved: int,
    recommended: int,
    ai_recommended: int,
    failures: list[dict],
    new_count: int = 0,
    updated_count: int = 0,
    duplicate_count: int = 0,
) -> None:
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            update scan_runs set
                finished_at=?,
                status=?,
                scanned_count=?,
                saved_count=?,
                new_count=?,
                updated_count=?,
                duplicate_count=?,
                recommended_count=?,
                ai_recommended_count=?,
                failures_json=?,
                updated_at=?
            where id=?
            """,
            (
                stamp,
                status,
                scanned,
                saved,
                new_count,
                updated_count,
                duplicate_count,
                recommended,
                ai_recommended,
                json.dumps(failures, ensure_ascii=False),
                stamp,
                scan_run_id,
            ),
        )


def scan_thread_alive(scan_run_id: int) -> bool:
    with SCAN_THREADS_LOCK:
        thread = SCAN_THREADS.get(scan_thread_key(scan_run_id))
    return bool(thread and thread.is_alive())


def mark_scan_run_interrupted(scan_run_id: int) -> None:
    stamp = now_iso()
    interrupted = {"source": "scan", "error": "Previous scan was interrupted before completion. Start a new scan."}
    with get_db() as conn:
        run = conn.execute("select failures_json from scan_runs where id=? and status='running'", (scan_run_id,)).fetchone()
        if not run:
            return
        try:
            failures = json.loads(run["failures_json"] or "[]")
        except json.JSONDecodeError:
            failures = []
        failures.append(interrupted)
        totals = conn.execute(
            """
            select
                coalesce(sum(scanned_count), 0) as scanned,
                coalesce(sum(saved_count), 0) as saved,
                coalesce(sum(new_count), 0) as new_count,
                coalesce(sum(updated_count), 0) as updated_count,
                coalesce(sum(duplicate_count), 0) as duplicate_count,
                coalesce(sum(failure_count), 0) as failure_count
            from scan_source_runs
            where scan_run_id=?
            """,
            (scan_run_id,),
        ).fetchone()
        conn.execute(
            """
            update scan_source_runs
            set finished_at=?, status='interrupted', updated_at=?
            where scan_run_id=? and status='running'
            """,
            (stamp, stamp, scan_run_id),
        )
        conn.execute(
            """
            update scan_runs
            set finished_at=?,
                status='interrupted',
                scanned_count=?,
                saved_count=?,
                new_count=?,
                updated_count=?,
                duplicate_count=?,
                failures_json=?,
                updated_at=?
            where id=? and status='running'
            """,
            (
                stamp,
                int(totals["scanned"] or 0),
                int(totals["saved"] or 0),
                int(totals["new_count"] or 0),
                int(totals["updated_count"] or 0),
                int(totals["duplicate_count"] or 0),
                json.dumps(failures, ensure_ascii=False),
                stamp,
                scan_run_id,
            ),
        )


def get_scan_run(scan_run_id: int) -> dict:
    with get_db() as conn:
        run = conn.execute("select status from scan_runs where id=?", (scan_run_id,)).fetchone()
    if run and run["status"] == "running" and not scan_thread_alive(scan_run_id):
        mark_scan_run_interrupted(scan_run_id)
    with get_db() as conn:
        run = conn.execute("select * from scan_runs where id=?", (scan_run_id,)).fetchone()
        if not run:
            raise KeyError(f"Scan run {scan_run_id} not found.")
        out = row_to_dict(run)
        rows = conn.execute("select * from scan_source_runs where scan_run_id=? order by id", (scan_run_id,)).fetchall()
        out["sources"] = [{**row_to_dict(row), "mode": scan_source_mode(row["source"])} for row in rows]
        return out


def latest_scan_run(run_date: str | None = None, region: str | None = None) -> dict | None:
    code = active_region_code(region) if region is not None else None
    with get_db() as conn:
        if run_date and code:
            row = conn.execute(
                "select * from scan_runs where run_date=? and region=? order by started_at desc limit 1",
                (run_date, code),
            ).fetchone()
        elif run_date:
            row = conn.execute(
                "select * from scan_runs where run_date=? order by started_at desc limit 1",
                (run_date,),
            ).fetchone()
        elif code:
            row = conn.execute(
                "select * from scan_runs where region=? order by started_at desc limit 1",
                (code,),
            ).fetchone()
        else:
            row = conn.execute("select * from scan_runs order by started_at desc limit 1").fetchone()
        return get_scan_run(row["id"]) if row else None


def latest_successful_scan(run_date: str | None = None, region: str | None = None) -> dict | None:
    code = active_region_code(region) if region is not None else None
    with get_db() as conn:
        if code:
            row = conn.execute(
                """
                select * from scan_runs
                where run_date=? and region=? and status in ('success', 'partial', 'limited')
                order by started_at desc
                limit 1
                """,
                (run_date or today(), code),
            ).fetchone()
        else:
            row = conn.execute(
                """
                select * from scan_runs
                where run_date=? and status in ('success', 'partial', 'limited')
                order by started_at desc
                limit 1
                """,
                (run_date or today(),),
            ).fetchone()
        return get_scan_run(row["id"]) if row else None


def scan_sources(triggered_by: str = "manual", forced: bool = True, scan_run_id: int | None = None, region: str | None = None) -> dict:
    code = active_region_code(region)
    context = active_region_context(code)
    city = context.get("city") or REGION_CONFIGS[code]["default_city"]
    scan_run_id = scan_run_id or create_scan_run(triggered_by, forced, code)
    sources = scan_source_definitions(code)
    saved: list[dict] = []
    failures: list[dict] = []
    source_counts: dict[str, int] = {}
    seen_urls: set[str] = set()
    seen_dedupe_keys: set[str] = set()
    with get_db() as conn:
        existing_urls = {row["url"] for row in conn.execute("select url from jobs where url<>''").fetchall()}
    new_count = 0
    updated_count = 0
    duplicate_count = 0
    scan_user_id = request_user_id()
    source_run_ids = {name: create_scan_source_run(scan_run_id, name) for name, _fetcher, _limit in sources}
    fetch_results: dict[str, tuple[list[dict], list[str]]] = {}

    def fetch_source(fetcher, limit: int) -> tuple[list[dict], list[str]]:
        with request_user_context(scan_user_id):
            return fetcher(limit)

    with ThreadPoolExecutor(max_workers=min(12, max(1, len(sources)))) as executor:
        futures = {
            executor.submit(fetch_source, fetcher, limit): source_name
            for source_name, fetcher, limit in sources
        }
        for future in as_completed(futures):
            source_name = futures[future]
            try:
                raw_jobs, fetch_failures = future.result()
            except Exception as exc:
                raw_jobs, fetch_failures = [], [str(exc)]
            fetch_results[source_name] = (raw_jobs, list(fetch_failures))

    for source_name, fetcher, limit in sources:
        source_run_id = source_run_ids[source_name]
        source_saved = 0
        source_new = 0
        source_updated = 0
        source_duplicates = 0
        source_failure_items: list[dict] = []
        raw_jobs, source_errors = fetch_results.get(source_name, ([], ["来源抓取未返回结果。"]))
        source_errors = summarize_scan_source_failures(source_name, source_errors, len(raw_jobs))
        source_errors = [
            error for error in source_errors
            if not is_nonblocking_scan_warning(source_name, error, len(raw_jobs))
        ]
        for failure in source_errors:
            item = {"source": source_name, "error": str(failure)}
            failures.append(item)
            source_failure_items.append(item)
        source_counts[source_name] = len(raw_jobs)
        for raw_job in raw_jobs:
            source_for_url = raw_job.get("source") or source_name
            url = canonical_job_url(source_for_url, raw_job.get("url", ""), raw_job.get("external_job_id"))
            raw_job["url"] = url
            raw_job.setdefault("region", code)
            raw_job.setdefault("city", city)
            raw_job.setdefault("source_region", code)
            if not raw_job.get("location"):
                raw_job["location"] = city
            if not url:
                continue
            if url in seen_urls:
                source_duplicates += 1
                duplicate_count += 1
                continue
            seen_urls.add(url)
            dedupe_key = job_dedupe_key(raw_job)
            if dedupe_key in seen_dedupe_keys:
                source_duplicates += 1
                duplicate_count += 1
            else:
                seen_dedupe_keys.add(dedupe_key)
            try:
                saved_job = upsert_job(raw_job)
                saved.append(saved_job)
                source_saved += 1
                if url in existing_urls:
                    source_updated += 1
                    updated_count += 1
                else:
                    source_new += 1
                    new_count += 1
            except Exception as exc:
                item = {"source": source_name, "url": url, "error": str(exc)}
                failures.append(item)
                source_failure_items.append(item)
        if not source_failure_items:
            source_status = "success"
        elif raw_jobs:
            source_status = "partial"
        elif is_limited_scan_source(source_name) or has_limited_failure(source_failure_items):
            source_status = "limited"
        else:
            source_status = "failed"
        finish_scan_source_run(
            source_run_id,
            source_status,
            len(raw_jobs),
            source_saved,
            source_failure_items,
            source_new,
            source_updated,
            source_duplicates,
        )

    recommended = [
        job for job in saved
        if job["score"] >= 3.0
        and job["status"] not in (NOTION_APPLICATION_STATUSES | {"Dropped", "Closed"})
        and not {"citizen_or_pr_only", "local_only", "clearance_required"}.intersection(set(job.get("eligibility_flags") or []))
    ]
    recommended.sort(key=lambda item: item["score"], reverse=True)
    ai_recommended = len(list_ai_jobs({"limit": ["20"], "region": [code]}))
    if not failures:
        run_status = "success"
    elif sum(source_counts.values()):
        run_status = "partial"
    elif has_limited_failure(failures):
        run_status = "limited"
    else:
        run_status = "failed"
    finish_scan_run(
        scan_run_id,
        run_status,
        sum(source_counts.values()),
        len(saved),
        len(recommended),
        ai_recommended,
        failures[:80],
        new_count,
        updated_count,
        duplicate_count,
    )
    generate_report(code)
    result = {
        "run_id": scan_run_id,
        "status": run_status,
        "region": code,
        "city": city,
        "scanned": sum(source_counts.values()),
        "saved": len(saved),
        "new": new_count,
        "updated": updated_count,
        "duplicates": duplicate_count,
        "recommended": len(recommended),
        "ai_recommended": ai_recommended,
        "top_jobs": recommended[:20],
        "source_counts": source_counts,
        "failures": failures[:40],
        "date": today(),
    }
    result["scan_run"] = get_scan_run(scan_run_id)
    return result


def running_scan_run(region: str | None = None) -> dict | None:
    code = active_region_code(region) if region is not None else None
    with get_db() as conn:
        if code:
            row = conn.execute(
                """
                select * from scan_runs
                where status='running' and region=?
                order by started_at desc
                limit 1
                """,
                (code,),
            ).fetchone()
        else:
            row = conn.execute(
                """
                select * from scan_runs
                where status='running'
                order by started_at desc
                limit 1
                """
            ).fetchone()
    if not row:
        return None
    run = get_scan_run(row["id"])
    return run if run["status"] == "running" else None


def scan_status_payload(scan_run_id: int | None = None, region: str | None = None) -> dict:
    code = active_region_code(region)
    run = get_scan_run(scan_run_id) if scan_run_id else latest_scan_run(region=code)
    return {
        "run": run,
        "region": code,
        "expected_sources": expected_scan_sources(code),
        "expected_source_details": expected_scan_source_details(code),
        "running": bool(run and run.get("status") == "running"),
    }


def _scan_async_worker(scan_run_id: int, triggered_by: str, forced: bool, region: str, user_id: str) -> None:
    with request_user_context(user_id):
        try:
            scan_sources(triggered_by=triggered_by, forced=forced, scan_run_id=scan_run_id, region=region)
        except Exception as exc:
            finish_scan_run(scan_run_id, "failed", 0, 0, 0, 0, [{"source": "scan", "error": str(exc)}])
        finally:
            safe_sync_cloud_state("async_scan")
            with SCAN_THREADS_LOCK:
                SCAN_THREADS.pop(scan_thread_key(scan_run_id, user_id), None)


def start_scan_async(triggered_by: str = "manual", forced: bool = True, region: str | None = None) -> dict:
    code = active_region_code(region)
    existing = running_scan_run(code)
    if existing:
        return {"started": False, "reason": "scan_already_running", **scan_status_payload(existing["id"], code)}
    scan_run_id = create_scan_run(triggered_by, forced, code)
    user_id = request_user_id()
    thread = threading.Thread(target=_scan_async_worker, args=(scan_run_id, triggered_by, forced, code, user_id), daemon=True)
    with SCAN_THREADS_LOCK:
        SCAN_THREADS[scan_thread_key(scan_run_id, user_id)] = thread
    thread.start()
    return {"started": True, **scan_status_payload(scan_run_id, code)}


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "-", value).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned[:120] or "job"


def ascii_pdf_text(value: str) -> str:
    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u00d7": "x",
        "\uff5c": "|",
        "\uff08": "(",
        "\uff09": ")",
    }
    for source, target in replacements.items():
        value = value.replace(source, target)
    return value.encode("latin-1", errors="replace").decode("latin-1")


def job_focus(job: dict) -> str:
    text = f"{job.get('company', '')} {job.get('position', '')} {job.get('jd_text', '')}".lower()
    if any(term in text for term in ["product management", "product manager", "product operations", "product analyst"]):
        return "product"
    if any(term in text for term in ["content", "marketing", "media", "communication", "brand", "social"]):
        return "content"
    if any(term in text for term in ["health", "healthcare", "public health", "medical"]):
        return "healthcare"
    if any(term in text for term in ["ux", "ui", "user experience", "product designer", "user research"]):
        return "ux"
    return "service"


def tailored_profile(focus: str, job: dict) -> str:
    profiles = {
        "ux": "Human-centred design master's student at SUTD with a product design foundation and hands-on experience across UX research, service design, journey mapping, interactive prototyping, and public-facing experience environments.",
        "product": "Human-centred design master's student at SUTD with experience translating ambiguous user and service problems into journeys, service blueprints, product-service concepts, prototypes, and implementation-ready design decisions.",
        "content": "Service and experience designer with visual communication experience across cultural events, social media-ready assets, storyboards, invitations, staff passes, and public-facing experience operations.",
        "healthcare": "Human-centred design master's student focused on healthcare service design, public health behaviour change, user research, service blueprinting, and product-service systems for everyday health contexts.",
        "service": "Service and experience designer with a product design foundation, strong user research orientation, and hands-on experience turning complex experience problems into service blueprints, prototypes, and validated design directions.",
    }
    return f"{profiles.get(focus, profiles['service'])} {AI_PROFILE_SIGNAL}"


def tailored_projects(focus: str) -> list[tuple[str, list[str]]]:
    projects = {
        "Smart Sugar Tray | Singapore Public Health Smart Sugar Control Plate System": [
            "Designed an NFC-enabled dining service that gives real-time sugar feedback in Singapore public dining contexts.",
            "Integrated product, service flow, UI touchpoints, and AI-assisted scenario exploration for meal-context feedback and public health communication.",
        ],
        '"Penglai" Fairy Island | Traditional Chinese Medicine Experience Service Design': [
            "Reframed Traditional Chinese Medicine into a youth-friendly service system focused on trust, clarity, and sustained participation.",
            "Developed service blueprints and user journeys connecting education, consultation, and long-term health engagement touchpoints.",
        ],
        "Music Tipsy Corner | Music x Tipsy Interactive Experience Service Design": [
            "Designed a music-driven cocktail service that translates emotional states into interaction flows and hospitality touchpoints.",
            "Built TouchDesigner prototypes to make abstract music and emotion inputs testable through a tangible customer journey.",
        ],
        "Wearable Targeted Therapy Garment | Healthcare Product Design": [
            "Explored TCM ion-based treatment, ergonomics, material choices, comfort, daily usability, and real-world feasibility.",
        ],
    }
    order = {
        "healthcare": [
            "Smart Sugar Tray | Singapore Public Health Smart Sugar Control Plate System",
            '"Penglai" Fairy Island | Traditional Chinese Medicine Experience Service Design',
            "Wearable Targeted Therapy Garment | Healthcare Product Design",
        ],
        "content": [
            "Music Tipsy Corner | Music x Tipsy Interactive Experience Service Design",
            "Smart Sugar Tray | Singapore Public Health Smart Sugar Control Plate System",
            '"Penglai" Fairy Island | Traditional Chinese Medicine Experience Service Design',
        ],
        "product": [
            "Smart Sugar Tray | Singapore Public Health Smart Sugar Control Plate System",
            '"Penglai" Fairy Island | Traditional Chinese Medicine Experience Service Design',
            "Music Tipsy Corner | Music x Tipsy Interactive Experience Service Design",
        ],
        "ux": [
            "Smart Sugar Tray | Singapore Public Health Smart Sugar Control Plate System",
            '"Penglai" Fairy Island | Traditional Chinese Medicine Experience Service Design',
            "Music Tipsy Corner | Music x Tipsy Interactive Experience Service Design",
        ],
    }
    selected = order.get(focus, order["ux"])
    return [(name, projects[name]) for name in selected]


def tailored_experience(focus: str) -> list[tuple[str, str, str, list[str]]]:
    experiences = [
        (
            "Singapore International Musicians Association",
            "Creative Content & Media Communication Intern",
            "2025.10 - 2026.01",
            [
                "Created visual and video communication assets for competitions and cultural events, translating event narratives into storyboards, invitations, staff passes, on-site materials, and social media-ready visuals.",
            ],
        ),
        (
            "FabLab O China, Digital Fabrication Workshop",
            "Course Instructor / Digital Fabrication & Design Education Support",
            "2021.09 - 2025.01",
            [
                "Designed and refined FABO curriculum modules using 3D printing and laser cutting for STE(D)M learning.",
                "Taught 6 classes with 48 students in robotics and wearable design; led public Make Faire x FABO activities with 100+ participants.",
            ],
        ),
        (
            "ALIGHT ROOM Lafayette Cultural Art Center",
            "On-site Operation & Execution Support",
            "2025.03 - 2025.07",
            [
                "Supported on-site execution for immersive exhibitions and pop-up events; observed visitor behaviour and assisted spatial flow adjustments to improve wayfinding and experience continuity.",
            ],
        ),
    ]
    if focus == "content":
        return [experiences[0], experiences[2], experiences[1]]
    return experiences


def tailored_skills(focus: str) -> list[str]:
    ai_skill = "AI-enabled workflows: AI-assisted research synthesis, prompt-based ideation, scenario generation, UX writing iteration, JD/capability mapping, workflow automation"
    base = [
        ai_skill,
        "Service and UX methods: user research, journey mapping, service blueprinting, prototype validation",
        "Design and prototyping: Figma, Photoshop, Illustrator, TouchDesigner, Arduino, rapid mock-ups",
        "Product and making: Rhino, KeyShot, Cinema 4D, 3D printing, laser cutting, ergonomics",
    ]
    if focus == "content":
        return [
            ai_skill,
            "Visual communication: storytelling, storyboards, social media-ready visuals, event materials, UX writing",
            "Service and UX methods: user research, journey mapping, service blueprinting, prototype validation",
            "Design and prototyping: Figma, Photoshop, Illustrator, TouchDesigner, Arduino, rapid mock-ups",
        ]
    if focus == "healthcare":
        return [
            ai_skill,
            "Healthcare service design: service blueprinting, public health communication, behaviour-change touchpoints, user research",
            "Service and UX methods: journey mapping, service blueprinting, prototype validation, UX writing",
            "Design and prototyping: Figma, Photoshop, Illustrator, TouchDesigner, Arduino, rapid mock-ups",
        ]
    return base


def markdown_bullets(items: list[str]) -> str:
    return "\n".join(f"- {item}" for item in items)


def profile_display_name(profile: dict | None = None) -> str:
    profile = profile or load_profile()
    return (profile.get("full_name") or "Your Name").strip()


def profile_contact_line(profile: dict | None = None) -> str:
    profile = profile or load_profile()
    parts = [
        profile.get("location"),
        profile.get("email"),
        profile.get("phone"),
        profile.get("linkedin"),
        profile.get("portfolio"),
    ]
    return " | ".join(str(part).strip() for part in parts if str(part or "").strip())


def profile_education_block(profile: dict | None = None) -> str:
    profile = profile or load_profile()
    school = (profile.get("school") or "Your School").strip()
    degree = (profile.get("degree") or "Your Degree / Program").strip()
    return f"**{school}**\n{degree}"


def build_tailored_resume_markdown(job: dict) -> str:
    profile = load_profile()
    focus = job_focus(job)
    projects = tailored_projects(focus)
    experience = tailored_experience(focus)
    skills = tailored_skills(focus)
    project_lines = []
    for name, bullets in projects:
        project_lines.append(f"### {name}")
        project_lines.append(markdown_bullets(bullets))
        project_lines.append("")
    experience_lines = []
    for org, title, dates, bullets in experience:
        experience_lines.append(f"### {org}")
        experience_lines.append(f"**{title}** | {dates}")
        experience_lines.append(markdown_bullets(bullets))
        experience_lines.append("")
    return f"""# {profile_display_name(profile)}

{AI_RESUME_HEADLINE}
{profile_contact_line(profile)}

## Profile

{tailored_profile(focus, job)}

## Education

{profile_education_block(profile)}

## Selected Projects

{chr(10).join(project_lines)}
## Experience

{chr(10).join(experience_lines)}
## Skills

{markdown_bullets(skills)}
"""


def build_cover_letter_markdown(job: dict) -> str:
    profile = load_profile()
    focus = job_focus(job)
    focus_sentence = {
        "ux": "The role's UX and product design focus aligns with my experience in user research, journey mapping, service blueprints, UI touchpoints, and prototype-led validation.",
        "product": "The role's product focus aligns with my experience translating user needs and service problems into structured journeys, prototypes, and implementation-ready design decisions.",
        "content": "The role's communication focus aligns with my experience creating event narratives, storyboards, visual assets, on-site materials, and social media-ready content.",
        "healthcare": "The role's healthcare or service focus aligns with my work on Smart Sugar Tray, Penglai Fairy Island, and other health-oriented product-service systems.",
        "service": "The role aligns with my service design, experience design, and human-centred design background.",
    }.get(focus, "The role aligns with my service design, experience design, and human-centred design background.")
    return f"""# Cover Letter

Dear Hiring Team,

I am writing to apply for the {job['position']} role at {job['company']}. My current profile is: {profile_education_block(profile).replace(chr(10), " | ")}.

{focus_sentence}

My selected work includes product, UX, service design, research, and prototyping projects. Across these projects, I have worked with user research, journey mapping, UI touchpoints, visual communication, and rapid prototyping.

I would be glad to bring this mix of service thinking, product design, prototyping, and public-facing experience sensitivity to your team.

Best regards,
{profile_display_name(profile)}

---

Target URL: {job['url']}
Score: {job['score']}/5.0
"""


def html_escape(value: str) -> str:
    from html import escape

    return escape(value or "", quote=True)


def project_tags(project_name: str) -> str:
    tags = {
        "Smart Sugar Tray | Singapore Public Health Smart Sugar Control Plate System": "Service Design, User Research, Healthcare Experience, Behaviour Change",
        '"Penglai" Fairy Island | Traditional Chinese Medicine Experience Service Design': "Service Blueprinting, User Research, Healthcare Service Design",
        "Music Tipsy Corner | Music x Tipsy Interactive Experience Service Design": "Experience Design, Emotional Design, Interactive Prototyping",
        "Wearable Targeted Therapy Garment | Healthcare Product Design": "Healthcare Product, Ergonomics, Product-Service Thinking",
    }
    return tags.get(project_name, "Service Design, Experience Design, Human-Centred Design")


def render_ul_html(items: list[str]) -> str:
    return "<ul>" + "".join(f"<li>{html_escape(item)}</li>" for item in items) + "</ul>"


def render_resume_html(job: dict, output_path: Path, photo_filename: str) -> None:
    profile = load_profile()
    focus = job_focus(job)
    projects = tailored_projects(focus)
    experience = tailored_experience(focus)
    skills = tailored_skills(focus)
    project_html = []
    for name, bullets in projects:
        if name == "Wearable Targeted Therapy Garment | Healthcare Product Design":
            project_html.append(
                f"""
        <div class="item project">
          <h3>Additional Healthcare Product Project</h3>
          <p>{html_escape(bullets[0])}</p>
        </div>"""
            )
            continue
        project_html.append(
            f"""
        <div class="item project">
          <h3>{html_escape(name)}</h3>
          <div class="tags">{html_escape(project_tags(name))}</div>
          {render_ul_html(bullets)}
        </div>"""
        )

    experience_html = []
    for org, title, dates, bullets in experience:
        experience_html.append(
            f"""
        <div class="item">
          <div class="topline">
            <h3>{html_escape(org)}</h3>
            <div class="date">{html_escape(dates)}</div>
          </div>
          <div class="sub">{html_escape(title)}</div>
          {render_ul_html(bullets)}
        </div>"""
        )

    skill_html = "".join(
        f'<div class="skill-row"><strong>{html_escape(row.split(":", 1)[0])}:</strong>{html_escape(row.split(":", 1)[1]) if ":" in row else ""}</div>'
        for row in skills
    )

    output_path.write_text(
        f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html_escape(profile_display_name(profile))} - Tailored Resume</title>
  <style>
    @page {{
      size: A4;
      margin: 10mm 14mm;
    }}

    :root {{
      --ink: #171717;
      --muted: #686868;
      --rule: #d9d9d9;
      --accent: #2f6b5f;
    }}

    * {{ box-sizing: border-box; }}

    body {{
      margin: 0;
      background: #f6f6f3;
      color: var(--ink);
      font-family: "Inter", "Aptos", "Segoe UI", Arial, sans-serif;
      font-size: 10.6px;
      line-height: 1.38;
      letter-spacing: 0;
    }}

    main {{
      width: 210mm;
      min-height: 297mm;
      margin: 0 auto;
      background: #fff;
      padding: 16mm 17mm;
    }}

    header {{
      display: grid;
      grid-template-columns: 1fr auto;
      gap: 14px;
      align-items: start;
      border-bottom: 1px solid var(--rule);
      padding-bottom: 12px;
      margin-bottom: 12px;
    }}

    h1 {{
      margin: 0;
      font-size: 28px;
      line-height: 1;
      font-weight: 650;
      letter-spacing: 0;
    }}

    .headline {{
      margin-top: 7px;
      color: var(--accent);
      font-size: 13px;
      font-weight: 650;
    }}

    .contact {{
      color: var(--muted);
      text-align: left;
      line-height: 1.48;
      margin-top: 10px;
    }}

    .photo {{
      width: 24mm;
      height: 30mm;
      border: 1px solid var(--rule);
      border-radius: 4px;
      object-fit: cover;
      object-position: center top;
      display: block;
    }}

    section {{
      display: grid;
      grid-template-columns: 34mm 1fr;
      gap: 12px;
      padding: 7px 0;
      border-bottom: 1px solid var(--rule);
    }}

    .profile-section > h2,
    .profile-section > p {{
      transform: translateY(-2mm);
    }}

    section:last-child {{ border-bottom: 0; }}

    h2 {{
      margin: 1px 0 0;
      color: var(--accent);
      font-size: 10px;
      line-height: 1.2;
      font-weight: 750;
      letter-spacing: 0.08em;
      text-transform: uppercase;
    }}

    p {{ margin: 0; }}

    .item {{ margin-bottom: 10px; }}
    .item:last-child {{ margin-bottom: 0; }}

    .topline {{
      display: flex;
      justify-content: space-between;
      gap: 12px;
      align-items: baseline;
      margin-bottom: 2px;
    }}

    h3 {{
      margin: 0;
      font-size: 12.2px;
      line-height: 1.25;
      font-weight: 720;
    }}

    .date {{
      color: var(--muted);
      font-size: 10.4px;
      white-space: nowrap;
    }}

    .sub {{
      color: var(--muted);
      font-size: 10.6px;
      margin-bottom: 4px;
    }}

    ul {{
      margin: 4px 0 0;
      padding-left: 14px;
    }}

    li {{
      margin: 2px 0;
      padding-left: 1px;
    }}

    .project h3 {{ color: var(--ink); }}

    .tags {{
      color: var(--muted);
      font-size: 10.2px;
      margin: 2px 0 3px;
    }}

    .skills {{
      display: grid;
      gap: 5px;
    }}

    .skill-row strong {{ color: var(--ink); }}

    @media print {{
      body {{ background: #fff; }}
      main {{
        width: auto;
        min-height: auto;
        margin: 0;
        padding: 0;
      }}
    }}
  </style>
</head>
<body>
  <main>
    <header>
      <div>
        <h1>{html_escape(profile_display_name(profile))}</h1>
        <div class="headline">{html_escape(AI_RESUME_HEADLINE)}</div>
        <div class="contact">{html_escape(profile_contact_line(profile)).replace(" | ", "<br>")}</div>
      </div>
      <img class="photo" src="{html_escape(photo_filename)}" alt="Profile portrait">
    </header>

    <section class="profile-section">
      <h2>Profile</h2>
      <p>{html_escape(tailored_profile(focus, job))}</p>
    </section>

    <section>
      <h2>Education</h2>
      <div>
        <div class="item">
          <h3>{html_escape(profile.get("school") or "Your School")}</h3>
          <div class="sub">{html_escape(profile.get("degree") or "Your Degree / Program")}</div>
        </div>
      </div>
    </section>

    <section>
      <h2>Selected Projects</h2>
      <div>{''.join(project_html)}
      </div>
    </section>

    <section>
      <h2>Experience</h2>
      <div>{''.join(experience_html)}
      </div>
    </section>

    <section>
      <h2>Skills</h2>
      <div class="skills">
        {skill_html}
      </div>
    </section>
  </main>
</body>
</html>
""",
        encoding="utf-8",
    )


def chrome_path() -> str | None:
    candidates = [
        os.environ.get("CHROME_PATH", ""),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return None


def render_html_pdf(html_path: Path, output_path: Path) -> bool:
    import subprocess
    import tempfile

    browser = chrome_path()
    if not browser:
        return False
    with tempfile.TemporaryDirectory(prefix="sg-career-chrome-") as user_data_dir:
        command = [
            browser,
            "--headless=new",
            "--disable-gpu",
            "--no-pdf-header-footer",
            f"--user-data-dir={user_data_dir}",
            f"--print-to-pdf={output_path}",
            html_path.resolve().as_uri(),
        ]
        try:
            completed = subprocess.run(command, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=30)
        except (subprocess.SubprocessError, OSError):
            return False
    return completed.returncode == 0 and output_path.exists() and output_path.stat().st_size > 10000


def write_resume_pdf_fallback(markdown_text: str, output_path: Path, job: dict) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    normal = ParagraphStyle("normal", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.6, leading=11, spaceAfter=3)
    small = ParagraphStyle("small", parent=normal, fontSize=7.8, leading=10, textColor=colors.HexColor("#4f5b55"))
    heading = ParagraphStyle("heading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor("#214c42"), spaceBefore=6, spaceAfter=4)
    title = ParagraphStyle("title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=21, textColor=colors.HexColor("#202521"), spaceAfter=2)
    subtitle = ParagraphStyle("subtitle", parent=normal, fontName="Helvetica-Bold", fontSize=9.2, leading=11, textColor=colors.HexColor("#2f6f5e"))

    profile = load_profile()
    name = profile_display_name(profile)
    contact_line = profile_contact_line(profile)
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=14 * mm, leftMargin=14 * mm, topMargin=12 * mm, bottomMargin=12 * mm)
    story = [
        Paragraph(ascii_pdf_text(name), title),
        Paragraph(AI_RESUME_HEADLINE, subtitle),
        Paragraph(ascii_pdf_text(contact_line), small),
    ]
    profile_photo_path = current_profile_photo_path()
    if profile_photo_path.exists():
        story.extend([Spacer(1, 4), Image(str(profile_photo_path), width=24 * mm, height=30 * mm)])

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith(f"# {name}") or line.startswith("Service &") or line.startswith("AI-Enabled") or line == contact_line:
            continue
        if line.startswith("## "):
            story.append(Paragraph(escape(ascii_pdf_text(line[3:].upper())), heading))
            continue
        if line.startswith("### "):
            story.append(Paragraph(f"<b>{escape(ascii_pdf_text(line[4:]))}</b>", normal))
            continue
        if line.startswith("**") and "** | " in line:
            story.append(Paragraph(f"<b>{escape(ascii_pdf_text(line.replace('**', '')))}</b>", small))
            continue
        if line.startswith("- "):
            story.append(Paragraph(f"&bull; {escape(ascii_pdf_text(line[2:]))}", normal))
            continue
        story.append(Paragraph(escape(ascii_pdf_text(line)), normal))
    doc.build(story)


def write_resume_pdf(markdown_text: str, output_path: Path, job: dict) -> None:
    import shutil

    html_path = output_path.with_suffix(".html")
    photo_name = "profile-photo.jpg"
    profile_photo_path = current_profile_photo_path()
    if profile_photo_path.exists():
        shutil.copyfile(profile_photo_path, output_path.parent / photo_name)
    render_resume_html(job, html_path, photo_name)
    if not render_html_pdf(html_path, output_path):
        write_resume_pdf_fallback(markdown_text, output_path, job)


def write_cover_pdf(markdown_text: str, output_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    normal = ParagraphStyle("normal", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, spaceAfter=9)
    title = ParagraphStyle("title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#202521"), spaceAfter=16)
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm)
    story = [Paragraph("Cover Letter", title)]
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("# ") or line.startswith("---") or line.startswith("Target URL:") or line.startswith("Score:"):
            continue
        story.append(Paragraph(escape(ascii_pdf_text(line)), normal))
        if line == "Dear Hiring Team,":
            story.append(Spacer(1, 4))
    doc.build(story)


def materials_need_refresh(job: dict) -> bool:
    resume_path = job.get("resume_path") or ""
    cover_path = job.get("cover_letter_path") or ""
    if not resume_path or not cover_path:
        return True
    resume = Path(resume_path)
    cover = Path(cover_path)
    return resume.suffix.lower() != ".pdf" or cover.suffix.lower() != ".pdf" or not resume.exists() or not cover.exists()


def make_drafts(job: dict) -> tuple[str, str]:
    date = today()
    job_dir = current_workspace_dir() / "applications" / date / sanitize_filename(f"{job['company']} - {job['position']}")
    job_dir.mkdir(parents=True, exist_ok=True)
    resume_md_path = job_dir / "tailored-resume.md"
    resume_pdf_path = job_dir / "tailored-resume.pdf"
    cover_md_path = job_dir / "cover-letter.md"
    cover_pdf_path = job_dir / "cover-letter.pdf"
    reference_path = job_dir / "job-reference.md"

    resume_markdown = build_tailored_resume_markdown(job)
    cover_markdown = build_cover_letter_markdown(job)
    resume_md_path.write_text(resume_markdown, encoding="utf-8")
    cover_md_path.write_text(cover_markdown, encoding="utf-8")
    reference_path.write_text(
        f"""# Job Reference

Company: {job['company']}
Position: {job['position']}
URL: {job['url']}
Score: {job['score']}/5.0
Resume source: {current_resume_path()}

## Match Notes

{job.get('match_notes') or ''}

## JD

{(job.get('jd_text') or '').strip()}
""",
        encoding="utf-8",
    )
    write_resume_pdf(resume_markdown, resume_pdf_path, job)
    write_cover_pdf(cover_markdown, cover_pdf_path)
    return str(resume_pdf_path), str(cover_pdf_path)


def upsert_job(payload: dict) -> dict:
    company = clean_company_name(payload.get("company") or "") or "Unknown Company"
    position = (payload.get("position") or "").strip() or "Unknown Position"
    source = (payload.get("source") or "Manual").strip()
    external_job_id = (payload.get("external_job_id") or "").strip() or None
    url = canonical_job_url(source, (payload.get("url") or "").strip(), external_job_id)
    jd_text = (payload.get("jd_text") or payload.get("JD") or "").strip()
    job_type = (payload.get("job_type") or "").strip()
    region = active_region_code(payload.get("region") or payload.get("source_region"))
    city = (payload.get("city") or active_region_context(region).get("city") or REGION_CONFIGS[region]["default_city"]).strip()
    source_region = active_region_code(payload.get("source_region") or region)
    location = (payload.get("location") or city or REGION_CONFIGS[region]["search_location"]).strip()

    if not url:
        raise ValueError("Job URL is required.")
    if not jd_text:
        jd_text = "JD not pasted yet. Preserve URL and update JD before drafting."

    metadata = job_metadata(position, jd_text, job_type, region, company, source)
    application_deadline = (
        normalize_application_deadline(str(payload.get("application_deadline") or ""))
        or extract_application_deadline(jd_text)
    )
    for key in ["salary_min", "salary_max", "salary_currency", "salary_period", "salary_text", "salary_fit"]:
        if key in payload and payload.get(key) not in {None, ""}:
            metadata[key] = payload[key]
    score, flags, match_notes = score_job(company, position, jd_text, source)
    hard_blocked = any(flag in flags for flag in ["citizen_or_pr_only", "local_only", "clearance_required"])
    initial_status = "Recommended" if score >= 3.0 and not hard_blocked else "New"
    batch_date = today() if initial_status == "Recommended" else None
    recommended_date = batch_date
    jd_hash = hashlib.sha256(jd_text.encode("utf-8")).hexdigest()
    stamp = now_iso()
    name = f"{company} - {position}"

    resume_path = None
    cover_path = None

    with get_db() as conn:
        existing = conn.execute("select id from jobs where url = ?", (url,)).fetchone()
        if existing:
            conn.execute(
                """
                update jobs set
                    company=?,
                    position=?,
                    name=?,
                    source=?,
                    external_job_id=coalesce(?, external_job_id),
                    location=?,
                    region=?,
                    city=?,
                    source_region=?,
                    job_type=?,
                    employment_type=?,
                    conversion_opportunity=?,
                    salary_min=?,
                    salary_max=?,
                    salary_currency=?,
                    salary_period=?,
                    salary_text=?,
                    salary_fit=?,
                    conversion_signal=?,
                    visa_sponsorship_signal=?,
                    language_signal=?,
                    pathway_score=?,
                    pathway_evidence_json=?,
                    application_deadline=coalesce(?, application_deadline),
                    jd_text=?,
                    jd_hash=?,
                    score=?,
                    eligibility_flags=?,
                    match_notes=?,
                    status=case when status='New' and ? then 'Recommended' else status end,
                    batch_date=case when status='New' and ? then coalesce(batch_date, ?) else batch_date end,
                    recommended_date=case when status='New' and ? then coalesce(recommended_date, ?) else recommended_date end,
                    last_checked_at=?,
                    updated_at=?,
                    resume_path=coalesce(?, resume_path),
                    cover_letter_path=coalesce(?, cover_letter_path)
                where id=?
                """,
                (
                    company,
                    position,
                    name,
                    source,
                    external_job_id,
                    location,
                    region,
                    city,
                    source_region,
                    job_type,
                    metadata["employment_type"],
                    metadata["conversion_opportunity"],
                    metadata["salary_min"],
                    metadata["salary_max"],
                    metadata["salary_currency"],
                    metadata["salary_period"],
                    metadata["salary_text"],
                    metadata["salary_fit"],
                    metadata["conversion_signal"],
                    metadata["visa_sponsorship_signal"],
                    metadata["language_signal"],
                    metadata["pathway_score"],
                    json.dumps(metadata["pathway_evidence_json"], ensure_ascii=False),
                    application_deadline or None,
                    jd_text,
                    jd_hash,
                    score,
                    json.dumps(flags),
                    match_notes,
                    initial_status == "Recommended",
                    initial_status == "Recommended",
                    today(),
                    initial_status == "Recommended",
                    today(),
                    stamp,
                    stamp,
                    resume_path,
                    cover_path,
                    existing["id"],
                ),
            )
            job_id = existing["id"]
        else:
            conn.execute(
                """
                insert into jobs(
                    company, position, name, source, url, external_job_id, location, region, city, source_region,
                    job_type, employment_type, conversion_opportunity, salary_min, salary_max,
                    salary_currency, salary_period, salary_text, salary_fit,
                    conversion_signal, visa_sponsorship_signal, language_signal, pathway_score, pathway_evidence_json,
                    application_deadline,
                    jd_text, jd_hash,
                    score, status, eligibility_flags, match_notes, found_date, batch_date,
                    recommended_date, last_checked_at, resume_path, cover_letter_path, created_at, updated_at
                )
                values(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    company,
                    position,
                    name,
                    source,
                    url,
                    external_job_id,
                    location,
                    region,
                    city,
                    source_region,
                    job_type,
                    metadata["employment_type"],
                    metadata["conversion_opportunity"],
                    metadata["salary_min"],
                    metadata["salary_max"],
                    metadata["salary_currency"],
                    metadata["salary_period"],
                    metadata["salary_text"],
                    metadata["salary_fit"],
                    metadata["conversion_signal"],
                    metadata["visa_sponsorship_signal"],
                    metadata["language_signal"],
                    metadata["pathway_score"],
                    json.dumps(metadata["pathway_evidence_json"], ensure_ascii=False),
                    application_deadline or None,
                    jd_text,
                    jd_hash,
                    score,
                    initial_status,
                    json.dumps(flags),
                    match_notes,
                    today(),
                    batch_date,
                    recommended_date,
                    stamp,
                    resume_path,
                    cover_path,
                    stamp,
                    stamp,
                ),
            )
            job_id = conn.execute("select last_insert_rowid()").fetchone()[0]
    return get_job(job_id)


def get_job(job_id: int) -> dict:
    with get_db() as conn:
        row = conn.execute("select * from jobs where id = ?", (job_id,)).fetchone()
    if not row:
        raise KeyError(f"Job {job_id} not found.")
    return row_to_dict(row)


def translate_text_to_zh(text: str) -> str:
    text = re.sub(r"\s+", " ", text or "").strip()
    if not text:
        return ""
    text = text[:7000]
    chunks = [text[i:i + 900] for i in range(0, len(text), 900)]
    translated: list[str] = []
    for chunk in chunks:
        params = urlencode({"client": "gtx", "sl": "auto", "tl": "zh-CN", "dt": "t", "q": chunk})
        raw = http_get(f"https://translate.googleapis.com/translate_a/single?{params}", timeout=20)
        data = json.loads(raw)
        translated.append("".join(part[0] for part in data[0] if part and part[0]))
    suffix = "\n\n（中文内容由在线翻译生成；请以原始 JD 为准。）"
    if len(text) >= 7000:
        suffix = "\n\n（原 JD 较长，此处翻译前 7000 个字符；请以原始链接和英文 JD 为准。）"
    return "\n\n".join(translated).strip() + suffix


def fallback_chinese_jd(job: dict) -> str:
    excerpt = clean_text(job.get("jd_text") or "")[:1600]
    return f"""岗位：{job.get("position") or "-"}
公司：{job.get("company") or "-"}
来源：{job.get("source") or "-"}
岗位链接：{job.get("url") or "-"}

匹配说明：
{job.get("match_notes") or "暂无。"}

自动翻译服务暂时不可用。以下是原始 JD 摘要，方便先判断岗位：
{excerpt}
"""


def ensure_job_translation(job_id: int) -> dict:
    job = get_job(job_id)
    if (job.get("jd_cn_text") or "").strip():
        return job
    source_text = f"{job.get('position')}\n{job.get('company')}\n{job.get('jd_text') or ''}"
    try:
        translated = translate_text_to_zh(source_text)
    except Exception:
        translated = fallback_chinese_jd(job)
    with get_db() as conn:
        conn.execute("update jobs set jd_cn_text=?, updated_at=? where id=?", (translated, now_iso(), job_id))
    return get_job(job_id)


def list_jobs(params: dict[str, list[str]]) -> list[dict]:
    status = (params.get("status") or [""])[0]
    date_filter = (params.get("date") or [""])[0]
    region_filter = (params.get("region") or [""])[0]
    city_filter = (params.get("city") or [""])[0].strip()
    query = "select * from jobs"
    clauses = []
    values: list[str] = []
    if status:
        clauses.append("status = ?")
        values.append(status)
    if date_filter:
        clauses.append("(batch_date = ? or found_date = ? or applied_date = ?)")
        values.extend([date_filter, date_filter, date_filter])
    if region_filter:
        code = active_region_code(region_filter)
        clauses.append("region = ?")
        values.append(code)
        if city_filter and code == "CN":
            clauses.append("(city = ? or location like ? or coalesce(city, '') = '')")
            values.extend([city_filter, f"%{city_filter}%"])
    if clauses:
        query += " where " + " and ".join(clauses)
    user_state_filter = bool(status and status not in {"New", "Recommended"})
    query += f" order by score desc, updated_at desc limit {5000 if user_state_filter or date_filter else 500}"
    with get_db() as conn:
        rows = conn.execute(query, values).fetchall()
        pathway_rows = []
        deadline_rows = []
        state_rows = []
        if region_filter and not status and not date_filter:
            code = active_region_code(region_filter)
            deadline_rows = conn.execute(
                """
                select * from jobs
                where region=?
                  and status in ('New', 'Recommended')
                  and application_deadline between ? and ?
                order by application_deadline, score desc
                limit 100
                """,
                (
                    code,
                    today(),
                    (dt.date.today() + dt.timedelta(days=7)).strftime(DATE_FMT),
                ),
            ).fetchall()
            pathway_query = """
                select * from jobs
                where region=?
                  and score>=2.5 and score<3.0
                  and pathway_score>=3.4
                  and employment_type in ('Internship', 'Graduate')
                order by pathway_score desc, score desc, updated_at desc
                limit 100
            """
            pathway_rows = conn.execute(pathway_query, (code,)).fetchall()
            state_clauses = ["region=?", "status not in ('New', 'Recommended')"]
            state_values: list[str] = [code]
            if city_filter and code == "CN":
                state_clauses.append("(city = ? or location like ? or coalesce(city, '') = '')")
                state_values.extend([city_filter, f"%{city_filter}%"])
            state_rows = conn.execute(
                f"select * from jobs where {' and '.join(state_clauses)} order by updated_at desc limit 5000",
                state_values,
            ).fetchall()
    jobs = [row_to_dict(row) for row in rows]
    seen_ids = {job.get("id") for job in jobs}
    for job in (row_to_dict(row) for row in state_rows):
        if job.get("id") not in seen_ids:
            jobs.append(job)
            seen_ids.add(job.get("id"))
    for job in (row_to_dict(row) for row in deadline_rows):
        if job.get("id") not in seen_ids:
            jobs.append(job)
            seen_ids.add(job.get("id"))
    jobs.extend(
        job for job in (row_to_dict(row) for row in pathway_rows)
        if job.get("id") not in seen_ids and is_pathway_recommendation_candidate(job)
    )
    return jobs


def clean_ai_detection_text(text: str) -> str:
    lowered = (text or "").lower()
    for pattern in AI_NOISE_PATTERNS:
        lowered = re.sub(pattern, " ", lowered, flags=re.I)
    return re.sub(r"\s+", " ", lowered).strip()


def keyword_hits(text: str, keywords: list[str]) -> list[str]:
    return [keyword for keyword in keywords if has_keyword(text, keyword)]


def is_actionable_job_title(title: str) -> bool:
    normalized = clean_text(title).lower().strip(" -|:")
    if not normalized:
        return False
    if re.match(r"^(view all|see all|browse all|explore all)\b", normalized):
        return False
    if re.match(r"^(careers?|jobs?)\s+at\b", normalized):
        return False
    return normalized not in {
        "career",
        "careers",
        "job",
        "jobs",
        "job search",
        "search jobs",
        "open roles",
        "open positions",
        "opportunities",
        "join us",
        "developer center",
        "design center",
        "resource center",
    }


def is_actionable_company_name(company: str, position: str = "") -> bool:
    normalized = re.sub(r"[^a-z0-9]+", " ", clean_text(company).lower()).strip()
    position_normalized = re.sub(r"[^a-z0-9]+", " ", clean_text(position).lower()).strip()
    if not normalized or normalized == position_normalized:
        return False
    return normalized not in {
        "linkedin company",
        "indeed company",
        "internsg company",
        "glints nodeflair startups",
        "mainland public search",
        "public search",
        "jobsdb",
    }


def is_pathway_recommendation_candidate(job: dict) -> bool:
    employment = job.get("employment_type") or detect_employment_type(
        job.get("position") or "",
        job.get("jd_text") or "",
        job.get("job_type") or "",
    )
    score = float(job.get("score") or 0)
    pathway_score = float(job.get("pathway_score") or 0)
    conversion = job.get("conversion_signal") or "unknown"
    visa = job.get("visa_sponsorship_signal") or "unknown"
    language = job.get("language_signal") or "unknown"
    has_pathway_signal = conversion in {"strong", "possible"} or visa == "possible" or language == "chinese_friendly_possible"
    return employment in {"Internship", "Graduate"} and score >= 2.5 and pathway_score >= 3.4 and has_pathway_signal


def job_listing_freshness(job: dict) -> dict:
    checked_age = days_since_date(job.get("last_checked_at") or job.get("updated_at"))
    found_age = days_since_date(job.get("found_date") or job.get("batch_date") or job.get("recommended_date"))
    source = clean_text(job.get("source") or "").lower()
    if checked_age == 0:
        return {"status": "verified", "label": "今日已验证", "adjustment": 0.0, "checked_age_days": checked_age}
    if checked_age is not None and checked_age <= 7:
        return {"status": "verified", "label": "本周已验证", "adjustment": 0.0, "checked_age_days": checked_age}
    if found_age is not None and found_age <= 7:
        return {"status": "recent", "label": "近期新发现", "adjustment": 0.0, "checked_age_days": checked_age}
    if checked_age is None:
        return {"status": "unknown", "label": "有效性未知", "adjustment": -0.18, "checked_age_days": None}
    if checked_age <= 14:
        return {"status": "aging", "label": "建议确认", "adjustment": -0.08, "checked_age_days": checked_age}
    if checked_age <= 30:
        return {"status": "aging", "label": "较早岗位", "adjustment": -0.18, "checked_age_days": checked_age}
    if "linkedin" in source:
        return {"status": "likely_closed", "label": "可能已下架", "adjustment": -1.0, "checked_age_days": checked_age}
    return {"status": "verify", "label": "需确认有效", "adjustment": -0.32, "checked_age_days": checked_age}


def is_recommendation_available(job: dict) -> bool:
    blocked_statuses = NOTION_APPLICATION_STATUSES | {"Dropped", "Closed"}
    hard_flags = {"citizen_or_pr_only", "local_only", "clearance_required"}
    region = job.get("source_region") or job.get("region") or ""
    city = job.get("city") or (REGION_CONFIGS.get(region) or {}).get("default_city") or ""
    source = clean_text(job.get("source") or "").lower()
    location_validator = explicit_ats_location_matches_region if "ats" in source else location_matches_region
    location_is_valid = not region or location_validator(job.get("location") or "", region, city)
    freshness = job_listing_freshness(job)
    return (
        job.get("status") not in blocked_statuses
        and freshness["status"] != "likely_closed"
        and (float(job.get("score") or 0) >= 3.0 or is_pathway_recommendation_candidate(job))
        and is_actionable_job_title(job.get("position") or job.get("title") or "")
        and is_actionable_company_name(job.get("company") or "", job.get("position") or job.get("title") or "")
        and location_is_valid
        and not hard_flags.intersection(set(job.get("eligibility_flags") or []))
    )


def ai_relevance_details(job: dict, matching_text: str | None = None) -> dict:
    title_text = clean_ai_detection_text(job.get("position", ""))
    jd_text = matching_text if matching_text is not None else job_matching_text(job)
    combined = f"{title_text} {jd_text}"
    title_hits = keyword_hits(title_text, AI_EXPLICIT_KEYWORDS)
    jd_hits = keyword_hits(jd_text, AI_EXPLICIT_KEYWORDS)
    domain_hits = keyword_hits(jd_text, AI_DOMAIN_KEYWORDS)
    adjacent_hits = keyword_hits(combined, AI_ADJACENT_KEYWORDS)
    role_hits = keyword_hits(combined, AI_ROLE_ANCHORS)

    title_signal = bool(title_hits)
    jd_signal = bool(jd_hits or domain_hits) and bool(role_hits)
    adjacent_signal = bool(adjacent_hits) and bool(role_hits) and any(
        has_keyword(title_text, term) for term in adjacent_hits
    )
    is_ai_related = title_signal or jd_signal or adjacent_signal

    relevance = 0.0
    if title_hits:
        relevance += 2.0 + min(len(title_hits) * 0.35, 1.0)
    if jd_hits:
        relevance += min(len(jd_hits) * 0.28, 1.1)
    if domain_hits:
        relevance += min(len(domain_hits) * 0.34, 1.0)
    if adjacent_hits:
        relevance += min(len(adjacent_hits) * 0.22, 0.7)
    if role_hits:
        relevance += min(len(role_hits) * 0.12, 0.6)
    relevance += min(float(job.get("score") or 0) / 5.0, 1.0)

    evidence_terms = []
    for term in [*title_hits, *domain_hits, *jd_hits, *adjacent_hits, *role_hits]:
        if term not in evidence_terms:
            evidence_terms.append(term)
    notes = "AI relevance: " + ", ".join(evidence_terms[:8]) if evidence_terms else "AI relevance: related signal pending"
    return {
        "is_ai_related": is_ai_related,
        "ai_relevance": round(relevance, 2),
        "ai_match_notes": notes,
    }


def list_ai_jobs(params: dict[str, list[str]]) -> list[dict]:
    try:
        limit = int((params.get("limit") or ["20"])[0])
    except ValueError:
        limit = 20
    limit = max(1, min(20, limit))
    region = active_region_code((params.get("region") or [""])[0] or None)
    city = (params.get("city") or [active_region_context(region).get("city") or ""])[0]
    watched = watched_company_keys(region)
    dismissed = dismissed_company_keys(region)
    candidates: list[dict] = []
    for job in list_jobs({"region": [region], "city": [city]}):
        if not is_recommendation_available(job):
            continue
        if is_watched_company_job(job, watched) or is_watched_company_job(job, dismissed):
            continue
        ai_details = ai_relevance_details(job)
        if not ai_details["is_ai_related"]:
            continue
        job.update(ai_details)
        candidates.append(job)

    candidates = apply_preference_scores_to_jobs(candidates, region)
    current_date = today()
    candidates.sort(
        key=lambda item: (
            item.get("batch_date") == current_date or item.get("recommended_date") == current_date,
            float(item.get("ai_relevance") or 0),
            float(item.get("rank_score") or item.get("score") or 0),
            item.get("updated_at") or "",
        ),
        reverse=True,
    )
    selected = collapse_duplicate_job_groups(candidates)[:limit]
    return compact_job_payloads(selected) if request_uses_compact_jobs(params) else selected


def career_direction_by_id(direction_id: str) -> dict | None:
    return next((item for item in CAREER_DIRECTIONS if item["id"] == direction_id), None)


def job_matching_text(job: dict) -> str:
    raw_jd_text = job.get("jd_text") or ""
    preamble_lines = raw_jd_text.splitlines()
    if preamble_lines and "official career match" in preamble_lines[0].lower():
        first_content_line = next(
            (index + 1 for index, line in enumerate(preamble_lines[:8]) if index >= 4 and not line.strip()),
            min(5, len(preamble_lines)),
        )
        raw_jd_text = "\n".join(preamble_lines[first_content_line:])
    raw_jd_text = re.sub(r"(?im)^\s*source\s+query\s*:.*$", " ", raw_jd_text)
    jd_text = clean_text(raw_jd_text)
    source = (job.get("source") or "").lower()
    lowered = jd_text.lower()
    if "internsg" in source:
        marker = lowered.find("job description")
        if marker >= 0:
            jd_text = jd_text[marker + len("job description"):]
        lowered_detail = jd_text.lower()
        end_markers = [
            lowered_detail.find(value)
            for value in [
                "you can discuss this job",
                "share this page",
                "previous job listing",
                "related posts",
                "related job searches",
                "copyright ©",
            ]
            if lowered_detail.find(value) >= 0
        ]
        if end_markers:
            jd_text = jd_text[:min(end_markers)]
    elif "linkedin" in source:
        markers = [
            lowered.find(value)
            for value in ["about the job", "job summary", "position "]
            if lowered.find(value) >= 0
        ]
        if markers:
            jd_text = jd_text[min(markers):]
    return clean_ai_detection_text(
        f"{job.get('position', '')} {jd_text}"
    )


def active_preference_direction_ids() -> tuple[list[str], str]:
    preferences = get_career_preferences()
    selected = [item for item in preferences["selected_directions"] if career_direction_by_id(item)]
    if selected:
        return selected, "user_selected"
    active_resume = get_active_resume_version()
    analysis = latest_resume_analysis(active_resume.get("id")) if active_resume else latest_resume_analysis()
    if analysis:
        suggested = [
            item["id"]
            for item in analysis.get("directions", [])
            if item.get("id") and item.get("score", 0) > 0
        ][:3]
        if suggested:
            return suggested, "resume_analysis"
    context_selected = [
        item for item in active_region_context().get("target_directions", [])
        if career_direction_by_id(item)
    ]
    if context_selected:
        return context_selected, "user_context"
    return [], "base_score"


def direction_match_for_job(job: dict, direction: dict, matching_text: str | None = None) -> tuple[float, list[str]]:
    text = matching_text if matching_text is not None else job_matching_text(job)
    hits = [keyword for keyword in direction["keywords"] if has_keyword(text, keyword)]
    if not hits:
        return 0.0, []
    title_text = clean_ai_detection_text(job.get("position", ""))
    title_hits = [keyword for keyword in hits if has_keyword(title_text, keyword)]
    generic_keywords = DIRECTION_GENERIC_KEYWORDS.get(direction["id"], set())
    specific_hits = [keyword for keyword in hits if keyword not in generic_keywords]
    if (
        direction["id"] == "ai-product"
        and PURE_TECHNICAL_TITLE_PATTERN.search(title_text)
        and not PRODUCT_FACING_TITLE_PATTERN.search(title_text)
    ):
        return 0.0, []
    if direction["id"] == "ai-product" and not PRODUCT_FACING_TITLE_PATTERN.search(title_text) and not specific_hits:
        return 0.0, []
    if not title_hits and NON_TARGET_FUNCTION_TITLE_PATTERN.search(title_text):
        return 0.0, []
    if not title_hits and not specific_hits:
        return 0.0, []
    score = min(1.0, len(hits) / 4)
    if title_hits:
        score = min(1.0, score + 0.25)
    elif specific_hits:
        score = min(1.0, score + 0.1)
    return score, hits[:8]


def region_fit_for_job(job: dict, region: str, watched_companies: set[str], context: dict | None = None) -> dict:
    code = active_region_code(region)
    context = context or active_region_context(code)
    city = (context.get("city") or REGION_CONFIGS[code]["default_city"]).lower()
    job_region = active_region_code(job.get("region") or job.get("source_region") or code)
    location_text = f"{job.get('location', '')} {job.get('city', '')}".lower()
    region_fit = 1.0 if job_region == code else 0.0
    location_match = bool(city and city in location_text) or bool(REGION_CONFIGS[code]["label"].lower() in location_text)
    if code == "SG" and "singapore" in location_text:
        location_match = True
    company_text = f"{job.get('company') or ''} {job.get('name') or ''}"
    company_boost = 0.35 if any(company_text_has_term(company_text, term) for term in watched_companies) else 0.0
    location_reason = REGION_CONFIGS[code]["label"]
    if location_match:
        location_reason = f"Matches {context.get('city') or REGION_CONFIGS[code]['default_city']}"
    elif region_fit:
        location_reason = f"Within {REGION_CONFIGS[code]['label']}"
    return {
        "region_fit": round(region_fit, 2),
        "location_match": location_match,
        "location_reason": location_reason,
        "company_boost": company_boost,
        "work_auth_fit": 1.0 if not {"citizen_or_pr_only", "local_only", "clearance_required"}.intersection(set(job.get("eligibility_flags") or [])) else 0.0,
    }


def employment_preference_for_job(job: dict, context: dict) -> dict:
    priority = context.get("employment_priority") or "unspecified"
    employment_type = job.get("employment_type") or "Unknown"
    flags = set(job.get("eligibility_flags") or [])
    boost = 0.0
    label = ""
    if priority == "internship":
        if employment_type == "Internship":
            boost, label = 0.28, "实习优先"
        elif employment_type == "Graduate":
            boost, label = 0.12, "Graduate 可考虑"
        elif employment_type == "Full-time":
            boost, label = -0.08, "正式工次优先"
    elif priority == "full_time":
        if employment_type == "Full-time":
            boost, label = 0.24, "正式工优先"
        elif employment_type == "Graduate":
            boost, label = 0.12, "Graduate 可考虑"
        elif employment_type == "Internship":
            boost, label = -0.08, "实习次优先"
    elif priority == "both" and employment_type in {"Internship", "Graduate", "Full-time"}:
        boost, label = 0.08, "类型匹配"
    if "experience_too_high" in flags and priority in {"internship", "both", "unspecified"}:
        boost -= 0.35
        label = "年限偏高"
    return {
        "employment_boost": round(boost, 2),
        "employment_fit_label": label,
    }


def salary_preference_for_job(job: dict, context: dict, strong_match_score: float) -> dict:
    preferred = context.get("salary_preferred")
    minimum = context.get("salary_min")
    target = preferred or minimum
    if not target:
        return {"salary_adjustment": 0.0, "salary_fit": job.get("salary_fit") or "unknown", "salary_fit_label": "薪资未设置偏好"}
    if not job.get("salary_max"):
        return {"salary_adjustment": 0.0, "salary_fit": "unknown", "salary_fit_label": "薪资未知"}

    context_currency = (context.get("salary_currency") or "").upper()
    job_currency = (job.get("salary_currency") or "").upper()
    if context_currency and job_currency and context_currency != job_currency:
        return {"salary_adjustment": 0.0, "salary_fit": "unknown", "salary_fit_label": "薪资币种待确认"}

    job_min_monthly = salary_to_monthly(job.get("salary_min"), job.get("salary_period"))
    job_max_monthly = salary_to_monthly(job.get("salary_max"), job.get("salary_period"))
    context_monthly = salary_to_monthly(target, context.get("salary_period") or "monthly")
    preferred_monthly = salary_to_monthly(preferred, context.get("salary_period") or "monthly") if preferred else context_monthly
    if not job_max_monthly or not context_monthly:
        return {"salary_adjustment": 0.0, "salary_fit": "unknown", "salary_fit_label": "薪资待确认"}

    if job_max_monthly < context_monthly:
        label = "薪资偏低 · 其他匹配强" if strong_match_score >= 4.0 else "薪资偏低"
        return {"salary_adjustment": -0.16, "salary_fit": "low", "salary_fit_label": label}
    if preferred_monthly and job_min_monthly and job_min_monthly >= preferred_monthly:
        return {"salary_adjustment": 0.12, "salary_fit": "strong", "salary_fit_label": "薪资达偏好"}
    return {"salary_adjustment": 0.06, "salary_fit": "match", "salary_fit_label": "薪资可接受"}


def conversion_signal_label(value: str) -> str:
    return {
        "strong": "可转正",
        "possible": "可能可转正",
        "none": "明确无转正",
        "unknown": "转正未知",
    }.get(value or "unknown", "转正未知")


def visa_signal_label(value: str) -> str:
    return {
        "possible": "工签可能",
        "unclear": "工签待确认",
        "unlikely": "工签风险",
        "unknown": "工签未知",
    }.get(value or "unknown", "工签未知")


def language_signal_label(value: str) -> str:
    return {
        "chinese_friendly_possible": "中文友好可能",
        "english_first": "英文为主",
        "unknown": "语言未知",
    }.get(value or "unknown", "语言未知")


def job_decision_summary(job: dict) -> str:
    employment_label = {
        "Internship": "实习岗位",
        "Graduate": "毕业生岗位",
        "Full-time": "正式岗位",
        "Contract": "合同岗位",
    }.get(job.get("employment_type") or "Unknown", "岗位方向")
    direction_mismatch = bool(job.get("direction_mismatch_adjustment"))
    matched = job.get("matched_directions") or []
    lead = employment_label
    if matched and not direction_mismatch:
        direction_label = clean_text(matched[0].get("label") or "")[:24]
        if direction_label:
            lead = f"{lead}，与你的 {direction_label} 方向匹配"

    conversion = job.get("conversion_signal") or ("strong" if job.get("conversion_opportunity") else "unknown")
    visa = job.get("visa_sponsorship_signal") or "unknown"
    language = job.get("language_signal") or "unknown"
    positive: list[str] = []
    cautions: list[str] = []

    if conversion == "strong":
        positive.append("明确可转正")
    elif conversion == "possible":
        positive.append("有转正信号")
    elif conversion == "none":
        cautions.append("明确无转正")

    if visa == "possible":
        positive.append("有工签可能")
    elif visa == "unlikely":
        cautions.append("工签风险较高")

    if language == "chinese_friendly_possible":
        positive.append("中文友好可能")

    freshness = job.get("listing_freshness_status") or ""
    if freshness == "likely_closed":
        cautions.insert(0, "岗位可能已下架")
    elif freshness in {"verify", "unknown"}:
        cautions.append("招聘状态需确认")
    if direction_mismatch:
        cautions.insert(0, "方向匹配较弱")
    muted = job.get("user_tag_mutes") or []
    if muted:
        labels = "、".join(clean_text(item.get("label") or item.get("id") or "") for item in muted[:2])
        if labels:
            cautions.insert(0, f"命中少看项 {labels}")
    if job.get("employment_type") == "Internship" and conversion not in {"strong", "possible", "none"}:
        cautions.append("转正需确认")
    if visa not in {"possible", "unlikely"}:
        cautions.append("工签需确认")

    segments = [lead]
    if positive:
        segments.append("、".join(positive[:2]))
    if cautions:
        segments.append("、".join(list(dict.fromkeys(cautions))[:2]))
    summary = "；".join(segments) + "。"
    if len(summary) <= 72:
        return summary
    return summary[:71].rstrip("，；、。 ") + "。"


def job_tag_label(tag_id: str) -> str:
    return USER_JOB_TAG_LABELS.get(tag_id, tag_id)


def source_tag_for_job(source: str) -> str:
    lowered = (source or "").lower()
    if "company" in lowered or "ats" in lowered or "官网" in lowered or "careers@gov" in lowered:
        return "source_official"
    if "linkedin" in lowered:
        return "source_linkedin"
    if "internsg" in lowered:
        return "source_internsg"
    if "internship.sg" in lowered:
        return "source_internship_sg"
    if "cultjobs" in lowered:
        return "source_cultjobs"
    if any(value in lowered for value in ["创业", "startup", "glints", "nodeflair", "wellfound"]):
        return "source_startup"
    if "jobstreet" in lowered:
        return "source_jobstreet"
    if "indeed" in lowered:
        return "source_indeed"
    if "mycareersfuture" in lowered:
        return "source_mycareersfuture"
    if "google jobs" in lowered:
        return "source_google_jobs"
    return ""


def freshness_tag_for_job(job: dict) -> str:
    dates = [job.get("found_date"), job.get("batch_date"), job.get("recommended_date"), job.get("applied_date")]
    parsed: list[dt.date] = []
    for value in dates:
        if not value:
            continue
        try:
            parsed.append(dt.datetime.strptime(str(value)[:10], DATE_FMT).date())
        except ValueError:
            continue
    if not parsed:
        return ""
    age_days = max(0, (dt.date.today() - max(parsed)).days)
    if age_days <= 1:
        return "fresh_today"
    if age_days <= 7:
        return "fresh_recent"
    return "fresh_stale"


def content_tag_ids_for_job(
    job: dict,
    matching_text: str | None = None,
    direction_matches: dict[str, tuple[float, list[str]]] | None = None,
) -> set[str]:
    text = matching_text if matching_text is not None else job_matching_text(job)
    title = clean_ai_detection_text(job.get("position") or "")
    conflicting_function = bool(NON_TARGET_FUNCTION_TITLE_PATTERN.search(title))
    tags: set[str] = set()
    direction_matches = direction_matches if direction_matches is not None else {}
    direction_scores = {}
    for direction_id in ["ai-product", "ux-product-design", "user-research", "service-design", "product-ops", "growth-content"]:
        if direction_id not in direction_matches:
            direction_matches[direction_id] = direction_match_for_job(job, career_direction_by_id(direction_id), text)
        direction_scores[direction_id] = direction_matches[direction_id][0]
    title_has_ai_signal = any(
        has_keyword(title, keyword)
        for keyword in ["ai", "genai", "generative ai", "llm", "machine learning", "chatbot"]
    )
    role_specific_ai_signal = bool(re.search(
        r"\b(support|develop|design|test|evaluate|implement|manage|analy[sz]e|work\s+(?:on|with)|use|apply)\b"
        r".{0,80}\b(ai|genai|generative ai|llm|machine learning|chatbot)\b",
        text,
        flags=re.I,
    ))
    if ai_relevance_details(job, text)["is_ai_related"] and (
        not conflicting_function or title_has_ai_signal or role_specific_ai_signal
    ):
        tags.add("ai_related")
    product_role_terms = [
        "product manager",
        "product management",
        "product design",
        "product marketing",
        "product development",
        "product owner",
        "product operations",
        "product strategy",
        "product roadmap",
    ]
    if has_keyword(title, "product") or (
        not conflicting_function
        and any(has_keyword(text, keyword) for keyword in product_role_terms)
    ):
        tags.add("product_related")
    if any(direction_scores[item] > 0 for item in ["ux-product-design", "user-research", "service-design"]):
        tags.add("ux_related")
    if direction_scores["product-ops"] > 0:
        tags.add("operations_related")
    if direction_scores["growth-content"] > 0:
        tags.add("marketing_related")
    return tags


def job_tag_ids_for_preferences(
    job: dict,
    region: str,
    company_profile: dict | None = None,
    matching_text: str | None = None,
    direction_matches: dict[str, tuple[float, list[str]]] | None = None,
) -> set[str]:
    tags: set[str] = set()
    employment = job.get("employment_type") or "Unknown"
    employment_map = {
        "Internship": "internship",
        "Graduate": "graduate",
        "Full-time": "full_time",
        "Contract": "contract",
    }
    if employment in employment_map:
        tags.add(employment_map[employment])

    conversion = job.get("conversion_signal") or ("strong" if job.get("conversion_opportunity") else "unknown")
    if conversion == "strong":
        tags.add("conversion_strong")
    elif conversion == "possible":
        tags.add("conversion_possible")
    elif conversion == "none":
        tags.add("conversion_none")

    visa = job.get("visa_sponsorship_signal") or "unknown"
    if visa == "possible":
        tags.add("visa_possible")
    elif visa == "unclear":
        tags.add("visa_unclear")
    elif visa == "unlikely":
        tags.add("visa_unlikely")

    language = job.get("language_signal") or "unknown"
    if language == "chinese_friendly_possible":
        tags.add("chinese_friendly")
    elif language == "english_first":
        tags.add("english_first")

    profile = company_profile or company_catalog_match_for_job(job.get("company") or "", region) or {}
    company_group = profile.get("company_group")
    if company_group in COMPANY_GROUP_VALUES:
        tags.add(f"company_{company_group}")
    if profile.get("sponsorship_signal") == "possible" and visa != "unlikely":
        tags.add("visa_possible")
    if "chinese" in str(profile.get("language_signal") or "").lower():
        tags.add("chinese_friendly")

    salary_fit = job.get("salary_fit") or "unknown"
    if salary_fit in {"strong", "match"}:
        tags.add("salary_match")
    elif salary_fit == "low":
        tags.add("salary_low")
    elif salary_fit == "unknown":
        tags.add("salary_unknown")

    source_tag = source_tag_for_job(job.get("source") or "")
    if source_tag:
        tags.add(source_tag)
    freshness_tag = freshness_tag_for_job(job)
    if freshness_tag:
        tags.add(freshness_tag)
    if "experience_too_high" in set(job.get("eligibility_flags") or []):
        tags.add("high_experience")
    tags.update(content_tag_ids_for_job(job, matching_text, direction_matches))
    return tags


def user_tag_preference_for_job(
    job: dict,
    context: dict,
    region: str,
    company_profile: dict | None = None,
    matching_text: str | None = None,
    direction_matches: dict[str, tuple[float, list[str]]] | None = None,
) -> dict:
    all_tags = job_tag_ids_for_preferences(job, region, company_profile, matching_text, direction_matches)
    preferred = [tag for tag in (context.get("preferred_job_tags") or []) if tag in USER_JOB_TAG_VALUES]
    muted = [tag for tag in (context.get("muted_job_tags") or []) if tag in USER_JOB_TAG_VALUES]
    preferred_hits = [tag for tag in preferred if tag in all_tags]
    muted_hits = [tag for tag in muted if tag in all_tags]
    lightweight_prefixes = ("source_", "fresh_")
    lightweight_tags = {"salary_match", "salary_unknown", "salary_low"}
    high_risk_mutes = {"visa_unlikely", "conversion_none", "high_experience"}
    employment_mutes = {"internship", "graduate", "full_time", "contract"}
    preferred_score = sum(
        0.06 if tag.startswith(lightweight_prefixes) or tag in lightweight_tags else 0.16
        for tag in preferred_hits
    )
    muted_score = sum(
        0.42 if tag in high_risk_mutes else 0.28 if tag in employment_mutes else 0.2
        for tag in muted_hits
    )
    adjustment = min(0.65, preferred_score) - min(1.2, muted_score)
    tag_priority = preferred_score - (muted_score * 1.25)
    return {
        "job_tag_ids": sorted(all_tags),
        "user_tag_matches": [{"id": tag, "label": job_tag_label(tag)} for tag in preferred_hits[:8]],
        "user_tag_mutes": [{"id": tag, "label": job_tag_label(tag)} for tag in muted_hits[:6]],
        "user_tag_adjustment": round(adjustment, 2),
        "user_tag_priority": round(tag_priority, 2),
    }


def pathway_preference_for_job(job: dict, context: dict, region: str, company_profile: dict | None = None) -> dict:
    conversion = job.get("conversion_signal") or ("strong" if job.get("conversion_opportunity") else "unknown")
    visa = job.get("visa_sponsorship_signal") or "unknown"
    language = job.get("language_signal") or "unknown"
    company_profile = company_profile or company_catalog_match_for_job(job.get("company") or "", region) or {}
    company_group = company_profile.get("company_group") or "other"
    company_visa_possible = company_profile.get("sponsorship_signal") == "possible" and visa != "unlikely"
    base_pathway = float(job.get("pathway_score") or 0)
    if not base_pathway:
        base_pathway = pathway_score_from_signals(job.get("employment_type") or "Unknown", conversion, visa, language, company_profile)

    adjustment = max(-0.34, min(0.42, (base_pathway - 2.5) * 0.12))
    tags: list[str] = []
    questions: list[str] = []
    if (context.get("career_goal") or "") == "sg_internship_to_fulltime":
        if job.get("employment_type") == "Internship":
            adjustment += 0.12
        if conversion in {"strong", "possible"}:
            adjustment += 0.12
        if visa == "possible" or company_visa_possible:
            adjustment += 0.12
        if visa == "unlikely":
            adjustment -= 0.24
    if context.get("conversion_priority") == "high":
        if conversion in {"strong", "possible"}:
            adjustment += 0.14
        elif conversion == "none":
            adjustment -= 0.16
    if context.get("sponsorship_priority") == "high":
        if visa == "possible" or company_visa_possible:
            adjustment += 0.14
        elif visa == "unlikely":
            adjustment -= 0.26
    if context.get("language_preference") in {"chinese_friendly", "bilingual"}:
        if language == "chinese_friendly_possible" or "chinese" in str(company_profile.get("language_signal") or "").lower():
            adjustment += 0.12
        elif context.get("language_preference") == "chinese_friendly" and language == "english_first":
            adjustment -= 0.04
    if company_group in set(context.get("preferred_company_groups") or []):
        adjustment += 0.1

    if conversion in {"strong", "possible"}:
        tags.append(conversion_signal_label(conversion))
    elif conversion == "none":
        tags.append("明确无转正")
    elif job.get("employment_type") == "Internship":
        tags.append("转正待确认")
        questions.append("问 HR：该实习是否有转正或 return offer 机制？")
    if visa in {"possible", "unclear", "unlikely"}:
        tags.append(visa_signal_label(visa))
    elif company_profile.get("sponsorship_signal") == "possible":
        tags.append("工签可能待确认")
    elif region == "SG":
        tags.append("工签待确认")
        questions.append("问 HR：若转为正式岗位，公司是否支持 EP 或 S Pass？")
    if language == "chinese_friendly_possible" or "chinese" in str(company_profile.get("language_signal") or "").lower():
        tags.append("中文友好可能")
    elif context.get("language_preference") in {"chinese_friendly", "bilingual"} and language == "unknown":
        questions.append("面试确认：团队日常协作中是否可以使用中文？")
    if company_group == "greater_china":
        tags.append("大中华背景")
    elif company_group == "ai_startup":
        tags.append("AI/初创")
    elif company_group == "sg_anchor":
        tags.append("本地大厂")

    evidence = list(job.get("pathway_evidence_json") or [])[:5]
    if company_profile.get("recommend_reason"):
        evidence.append(f"公司理由: {company_profile['recommend_reason']}")
    if not evidence:
        evidence.append("暂无明确转正/工签证据，请打开原 JD 人工确认。")
    return {
        "pathway_score": round(base_pathway, 2),
        "pathway_adjustment": round(max(-0.5, min(0.65, adjustment)), 2),
        "pathway_tags": list(dict.fromkeys(tags))[:6],
        "pathway_questions": list(dict.fromkeys(questions))[:4],
        "evidence": evidence[:6],
        "conversion_signal_label": conversion_signal_label(conversion),
        "visa_sponsorship_label": visa_signal_label(visa),
        "language_signal_label": language_signal_label(language),
        "company_group": company_group,
    }


def bounded_display_score(value: float) -> float:
    return round(max(0.0, min(5.0, float(value or 0))), 1)


def rank_job_with_preferences(
    job: dict,
    direction_ids: list[str],
    weights: dict,
    region: str | None = None,
    watched_companies: set[str] | None = None,
    context: dict | None = None,
) -> dict:
    matched = []
    boost = 0.0
    matching_text = job_matching_text(job)
    direction_matches: dict[str, tuple[float, list[str]]] = {}
    for direction_id in direction_ids:
        direction = career_direction_by_id(direction_id)
        if not direction:
            continue
        match_score, hits = direction_match_for_job(job, direction, matching_text)
        direction_matches[direction_id] = (match_score, hits)
        if match_score <= 0:
            continue
        weight = float(weights.get(direction_id, 1.0) or 1.0)
        boost += min(0.35, match_score * 0.28 * weight)
        matched.append({"id": direction_id, "label": direction["label"], "keywords": hits, "score": round(match_score, 2)})
    boost = round(min(0.8, boost), 2)
    out = dict(job)
    direction_mismatch_adjustment = (
        -0.55
        if direction_ids
        and not matched
        and NON_TARGET_FUNCTION_TITLE_PATTERN.search(out.get("position") or "")
        else 0.0
    )
    code = active_region_code(region or job.get("region"))
    context = context or active_region_context(code)
    region_fit = region_fit_for_job(out, code, watched_companies or set(), context)
    out["base_score"] = float(job.get("score") or 0)
    out["preference_boost"] = boost
    out["region_fit"] = region_fit["region_fit"]
    out["location_match"] = region_fit["location_match"]
    out["location_reason"] = region_fit["location_reason"]
    out["company_boost"] = region_fit["company_boost"]
    out["work_auth_fit"] = region_fit["work_auth_fit"]
    detected_employment = detect_employment_type(out.get("position") or "", out.get("jd_text") or "", out.get("job_type") or "")
    if detected_employment != "Unknown":
        out["employment_type"] = detected_employment
    employment = employment_preference_for_job(out, context)
    strong_match_score = out["base_score"] + boost + out["company_boost"] + employment["employment_boost"]
    salary = salary_preference_for_job(out, context, strong_match_score)
    company_profile = company_catalog_match_for_job(out.get("company") or "", code) or {}
    pathway = pathway_preference_for_job(out, context, code, company_profile)
    freshness = job_listing_freshness(out)
    conversion_boost = 0.0
    if out.get("conversion_opportunity") and context.get("employment_priority") in {"internship", "both", "unspecified"}:
        conversion_boost = 0.14
    out.update(employment)
    out.update(salary)
    out.update(pathway)
    out["listing_freshness_status"] = freshness["status"]
    out["listing_freshness_label"] = freshness["label"]
    out["listing_checked_age_days"] = freshness["checked_age_days"]
    out["freshness_adjustment"] = freshness["adjustment"]
    out["pathway_candidate"] = is_pathway_recommendation_candidate(out)
    tag_preference = user_tag_preference_for_job(
        out,
        context,
        code,
        company_profile,
        matching_text,
        direction_matches,
    )
    out.update(tag_preference)
    out["conversion_boost"] = round(conversion_boost, 2)
    out["direction_mismatch_adjustment"] = direction_mismatch_adjustment
    out["rank_score"] = round(
        out["base_score"]
        + boost
        + direction_mismatch_adjustment
        + out["company_boost"]
        + out["employment_boost"]
        + out["salary_adjustment"]
        + conversion_boost
        + out["pathway_adjustment"]
        + out["user_tag_adjustment"]
        + out["freshness_adjustment"],
        2,
    )
    out["fit_score"] = bounded_display_score(out["rank_score"])
    out["score_breakdown"] = {
        "base": round(out["base_score"], 2),
        "direction": boost,
        "direction_mismatch": direction_mismatch_adjustment,
        "company": out["company_boost"],
        "employment": out["employment_boost"],
        "salary": out["salary_adjustment"],
        "conversion": out["conversion_boost"],
        "pathway": out["pathway_adjustment"],
        "user_tags": out["user_tag_adjustment"],
        "freshness": out["freshness_adjustment"],
    }
    out["matched_directions"] = matched
    out["fit_reasons"] = [f"{item['label']}: {', '.join(item['keywords'][:4])}" for item in matched]
    if out.get("user_tag_matches"):
        out["fit_reasons"].append("你选的标签: " + "、".join(item["label"] for item in out["user_tag_matches"][:4]))
    if out.get("user_tag_mutes"):
        out["fit_reasons"].append("少看标签: " + "、".join(item["label"] for item in out["user_tag_mutes"][:3]))
    if out.get("pathway_tags"):
        out["fit_reasons"].append("留新加坡路径: " + "、".join(out["pathway_tags"][:4]))
    if out.get("pathway_candidate") and out.get("base_score", 0) < 3.0:
        out["fit_reasons"].append("留新路径补充候选")
    if out["company_boost"]:
        out["fit_reasons"].append("Watched company")
    if out.get("employment_fit_label"):
        out["fit_reasons"].append(out["employment_fit_label"])
    if out.get("salary_fit_label") and out["salary_fit_label"] != "薪资未设置偏好":
        out["fit_reasons"].append(out["salary_fit_label"])
    if conversion_boost:
        out["fit_reasons"].append("可转正机会")
    if out["location_reason"]:
        out["fit_reasons"].append(out["location_reason"])
    risk_reasons = []
    if direction_mismatch_adjustment:
        risk_reasons.append("方向偏离")
    if out.get("user_tag_mutes"):
        risk_reasons.append("少看标签: " + "、".join(item["label"] for item in out["user_tag_mutes"][:3]))
    if out.get("employment_fit_label") == "年限偏高":
        risk_reasons.append(out["employment_fit_label"])
    if out.get("salary_fit") == "low" and out.get("salary_fit_label"):
        risk_reasons.append(out["salary_fit_label"])
    if out.get("listing_freshness_status") in {"aging", "verify", "unknown", "likely_closed"}:
        risk_reasons.append(out.get("listing_freshness_label") or "需确认有效")
    ordered_reasons = list(dict.fromkeys([*risk_reasons, *out["fit_reasons"]]))
    out["recommendation_reason"] = " · ".join(ordered_reasons[:4]) or "按基础评分推荐，建议打开 JD 确认转正和工签信息。"
    deadline = application_deadline_status(out.get("application_deadline"))
    out["deadline_status"] = deadline["code"]
    out["deadline_label"] = deadline["label"]
    out["deadline_days_remaining"] = deadline["days_remaining"]
    if out.get("status") == "Apply Queue":
        queue = queue_decision(out)
        out["queue_priority"] = queue["priority"]
        out["queue_priority_label"] = queue["label"]
        out["queue_reason"] = queue["reason"]
    if out.get("status") in {"Applied", "Follow Up"}:
        followup = followup_decision(out)
        out["followup_priority"] = followup["priority"]
        out["followup_priority_label"] = followup["label"]
        out["followup_reason"] = followup["reason"]
    out["decision_summary"] = job_decision_summary(out)
    return out


def apply_preference_scores_to_jobs(jobs: list[dict], region: str | None = None) -> list[dict]:
    direction_ids, direction_source = active_preference_direction_ids()
    preferences = get_career_preferences()
    region_code = active_region_code(region) if region else None
    watched = watched_company_keys(region_code) if region_code else set()
    dismissed = dismissed_company_keys(region_code) if region_code else set()
    watched_by_region: dict[str, set[str]] = {}
    dismissed_by_region: dict[str, set[str]] = {}
    context_by_region: dict[str, dict] = {}
    ranked = []
    for job in jobs:
        job_region = region_code or active_region_code(job.get("region") or job.get("source_region") or None)
        if region_code:
            watched_for_job = watched
            dismissed_for_job = dismissed
        else:
            if job_region not in watched_by_region:
                watched_by_region[job_region] = watched_company_keys(job_region)
                dismissed_by_region[job_region] = dismissed_company_keys(job_region)
            watched_for_job = watched_by_region[job_region]
            dismissed_for_job = dismissed_by_region[job_region]
        if job_region not in context_by_region:
            context_by_region[job_region] = active_region_context(job_region)
        item = rank_job_with_preferences(
            job,
            direction_ids,
            preferences["direction_weights"],
            job_region,
            watched_for_job,
            context_by_region[job_region],
        )
        item["direction_source"] = direction_source
        item["company_watched_by_user"] = is_watched_company_job(item, watched_for_job)
        item["company_hidden_by_watchlist"] = is_watched_company_job(item, dismissed_for_job)
        item["supplemental_candidate"] = (
            is_recommendation_available(item)
            and not item["company_watched_by_user"]
            and not item["company_hidden_by_watchlist"]
        )
        ranked.append(item)
    return ranked


COMPACT_JOB_FIELDS = {
    "id", "company", "position", "source", "status", "url",
    "location", "region", "city", "source_region",
    "employment_type", "conversion_opportunity",
    "salary_min", "salary_max", "salary_currency", "salary_period", "salary_fit", "salary_fit_label",
    "conversion_signal", "visa_sponsorship_signal", "language_signal",
    "found_date", "batch_date", "recommended_date", "updated_at", "last_checked_at", "applied_date",
    "last_followup_at", "followup_count",
    "score", "base_score", "fit_score", "rank_score", "pathway_score", "pathway_tags",
    "user_tag_matches", "user_tag_mutes", "decision_summary", "recommendation_reason",
    "source_count", "alternate_links", "alternate_sources", "dedupe_key",
    "supplemental_candidate", "company_watched_by_user", "company_hidden_by_watchlist",
    "listing_freshness_status", "listing_freshness_label",
    "deadline_status", "deadline_label", "deadline_days_remaining",
    "queue_priority", "queue_priority_label", "queue_reason",
    "followup_priority", "followup_priority_label", "followup_reason", "next_step",
    "ai_relevance", "ai_match_notes",
}
WORKBENCH_JOB_FIELDS = {
    "id", "company", "position", "source", "status",
    "employment_type", "found_date", "updated_at", "applied_date",
    "fit_score", "rank_score", "pathway_score", "pathway_tags",
    "user_tag_matches", "user_tag_mutes",
    "salary_fit", "salary_fit_label", "salary_min", "salary_max", "salary_currency", "salary_period",
    "decision_summary", "recommendation_reason", "source_count", "supplemental_candidate",
    "company_watched_by_user", "company_hidden_by_watchlist", "next_step",
    "listing_freshness_status", "listing_freshness_label",
    "application_deadline", "deadline_status", "deadline_label", "deadline_days_remaining",
    "queue_priority", "queue_priority_label", "queue_reason",
    "followup_priority", "followup_priority_label", "followup_reason", "last_followup_at", "followup_count",
}


def request_uses_compact_jobs(params: dict[str, list[str]] | None) -> bool:
    value = ((params or {}).get("compact") or [""])[0]
    return str(value).strip().lower() in {"1", "true", "yes", "on"}


def compact_job_payload(job: dict) -> dict:
    out = {key: value for key, value in job.items() if key in COMPACT_JOB_FIELDS}
    out.setdefault("dedupe_key", job_dedupe_key(job))
    return out


def compact_job_payloads(jobs: list[dict]) -> list[dict]:
    return [compact_job_payload(job) for job in jobs]


def workbench_job_payload(job: dict) -> dict:
    out = {key: value for key, value in job.items() if key in WORKBENCH_JOB_FIELDS}
    alternate_ids = [
        {"id": item.get("id")}
        for item in (job.get("alternate_links") or [])
        if item.get("id") is not None
    ]
    if alternate_ids:
        out["alternate_links"] = alternate_ids
    return out


def workbench_job_payloads(jobs: list[dict]) -> list[dict]:
    return [workbench_job_payload(job) for job in jobs]


def list_jobs_payload(params: dict[str, list[str]]) -> list[dict]:
    region = (params.get("region") or [""])[0] or None
    jobs = apply_preference_scores_to_jobs(list_jobs(params), region)
    return compact_job_payloads(jobs) if request_uses_compact_jobs(params) else jobs


def job_payload(job: dict) -> dict:
    return apply_preference_scores_to_jobs([job], job.get("region") or None)[0]


def recommendation_payload_from_ranked_jobs(
    ranked_jobs: list[dict],
    region: str,
    limit: int,
    context: dict,
    direction_ids: list[str],
    direction_source: str,
    exclude_keywords: list[str],
    compact: bool = False,
) -> dict:
    candidates = []
    for job in ranked_jobs:
        if not is_recommendation_available(job):
            continue
        if job.get("company_watched_by_user") or job.get("company_hidden_by_watchlist"):
            continue
        combined = f"{job.get('company', '')} {job.get('position', '')} {job.get('jd_text', '')}".lower()
        if any(has_keyword(combined, keyword) for keyword in exclude_keywords):
            continue
        ranked = dict(job)
        ranked["company_watched_by_user"] = False
        ranked["company_hidden_by_watchlist"] = False
        candidates.append(ranked)

    current_date = today()
    has_tag_preferences = bool(context.get("preferred_job_tags") or context.get("muted_job_tags"))
    candidates.sort(
        key=lambda item: (
            not bool(item.get("user_tag_mutes")) if has_tag_preferences else True,
            not bool(item.get("direction_mismatch_adjustment")),
            deadline_recommendation_priority(item),
            item.get("batch_date") == current_date or item.get("recommended_date") == current_date,
            float(item.get("user_tag_priority") or 0) if has_tag_preferences else 0,
            float(item.get("user_tag_adjustment") or 0) if has_tag_preferences else 0,
            float(item.get("rank_score") or 0),
            float(item.get("base_score") or 0),
            item.get("updated_at") or "",
        ),
        reverse=True,
    )
    candidates = collapse_duplicate_job_groups(candidates)
    preferred = [tag for tag in (context.get("preferred_job_tags") or []) if tag in USER_JOB_TAG_VALUES]
    muted = [tag for tag in (context.get("muted_job_tags") or []) if tag in USER_JOB_TAG_VALUES]
    selected_jobs = candidates[:limit]
    if compact:
        selected_jobs = compact_job_payloads(selected_jobs)
    return {
        "date": current_date,
        "region": region,
        "region_label": REGION_CONFIGS[region]["label"],
        "jobs": selected_jobs,
        "active_direction_ids": direction_ids,
        "direction_source": direction_source,
        "limit": limit,
        "tag_scope": {
            "preferred_tags": [{"id": tag, "label": job_tag_label(tag)} for tag in preferred],
            "muted_tags": [{"id": tag, "label": job_tag_label(tag)} for tag in muted],
            "matched_jobs": sum(1 for item in candidates if item.get("user_tag_matches")),
            "muted_jobs": sum(1 for item in candidates if item.get("user_tag_mutes")),
            "effective": has_tag_preferences,
        },
    }


def list_today_recommendations(params: dict[str, list[str]] | None = None) -> dict:
    params = params or {}
    try:
        limit = int((params.get("limit") or ["20"])[0])
    except ValueError:
        limit = 20
    limit = max(1, min(200, limit))
    region = active_region_code((params.get("region") or [""])[0] or None)
    city = (params.get("city") or [active_region_context(region).get("city") or ""])[0]
    context = active_region_context(region)
    direction_ids, direction_source = active_preference_direction_ids()
    preferences = get_career_preferences()
    primary_jobs = list_jobs({"region": [region], "city": [city]})
    fresh_jobs = list_jobs({"date": [today()], "region": [region], "city": [city]})
    merged_jobs: list[dict] = []
    seen_jobs: set[str] = set()
    for job in [*primary_jobs, *fresh_jobs]:
        key = str(job.get("id") or canonical_job_url(job.get("source") or "", job.get("url") or "", job.get("external_job_id")))
        if key and key in seen_jobs:
            continue
        if key:
            seen_jobs.add(key)
        merged_jobs.append(job)
    ranked_jobs = apply_preference_scores_to_jobs(merged_jobs, region)
    return recommendation_payload_from_ranked_jobs(
        ranked_jobs,
        region,
        limit,
        context,
        direction_ids,
        direction_source,
        preferences["exclude_keywords"],
        request_uses_compact_jobs(params),
    )


def daily_status(region: str | None = None) -> dict:
    code = active_region_code(region)
    current_date = today()
    latest = latest_scan_run(current_date, code)
    successful = latest_successful_scan(current_date, code)
    return {
        "date": current_date,
        "region": code,
        "region_label": REGION_CONFIGS[code]["label"],
        "has_successful_scan": successful is not None,
        "latest_run": latest,
        "latest_successful_run": successful,
        "auto_run_mode": "open_app_once_per_day",
    }


def days_since_date(value: str | None, reference_date: dt.date | None = None) -> int | None:
    if not value:
        return None
    try:
        parsed = dt.datetime.strptime(str(value)[:10], DATE_FMT).date()
    except ValueError:
        return None
    return max(0, ((reference_date or dt.date.today()) - parsed).days)


def number_like(value, default=0) -> float:
    try:
        return float(value or default)
    except (TypeError, ValueError):
        return float(default)


def next_step_for_job(job: dict) -> str:
    status = job.get("status") or ""
    if status == "Apply Queue":
        deadline = application_deadline_status(job.get("application_deadline"))
        if deadline["code"] == "expired":
            return "截止日期已过，请先确认岗位是否仍开放。"
        if deadline["code"] == "today":
            return "今天截止，优先完成投递。"
        if deadline["code"] in {"urgent", "soon"}:
            return f"{deadline['label']}，建议优先投递。"
        if job.get("queue_reason"):
            return str(job["queue_reason"])
        return "今天可以打开填表助手，最终提交前人工确认。"
    if status == "Applied":
        return followup_decision(job)["reason"]
    if status == "Follow Up":
        return followup_decision(job)["reason"]
    if status == "Recommended":
        return "先看转正、工签和语言证据，再决定是否加入队列。"
    return "打开岗位详情确认下一步。"


def followup_decision(job: dict, reference_date: dt.date | None = None) -> dict:
    status = job.get("status") or ""
    if status == "Follow Up":
        return {
            "priority": "followup",
            "label": "今天跟进",
            "reason": "已标记需要跟进，今天发送一条简短消息。",
        }
    if status != "Applied":
        return {"priority": "", "label": "", "reason": ""}
    days = days_since_date(job.get("applied_date"), reference_date)
    if days is None or days < 3:
        return {
            "priority": "waiting",
            "label": "等待反馈",
            "reason": "刚完成投递，先给招聘方一点处理时间。",
        }
    followup_days = days_since_date(job.get("last_followup_at"), reference_date)
    followup_count = int(job.get("followup_count") or 0)
    if followup_days is not None:
        if followup_days < 7:
            return {
                "priority": "waiting",
                "label": "等待反馈",
                "reason": f"等待反馈：第 {max(1, followup_count)} 次跟进已发送，{7 - followup_days} 天后再判断。",
            }
        if followup_count < 2:
            return {
                "priority": "followup",
                "label": "今天跟进",
                "reason": "上次跟进已过 7 天，可以发送第二次简短跟进。",
            }
        return {
            "priority": "archive",
            "label": "建议归档",
            "reason": f"已跟进 {followup_count} 次仍无回复，建议最后确认后暂停。",
        }
    if days <= 14:
        return {
            "priority": "followup",
            "label": "今天跟进",
            "reason": f"已投 {days} 天，今天适合发送第一次简短跟进。",
        }
    return {
        "priority": "archive",
        "label": "建议归档",
        "reason": f"已投 {days} 天仍无更新，建议最后确认一次后归档。",
    }


def application_action_bucket(job: dict) -> str:
    decision = followup_decision(job)
    return "stale" if decision["priority"] == "archive" else decision["priority"]


def scan_run_summary_text(run: dict | None) -> str:
    if not run:
        return "今天还没有扫描记录。"
    failures = len(run.get("failures_json") or [])
    status = run.get("status") or "pending"
    label = {
        "pending": "等待",
        "running": "扫描中",
        "success": "成功",
        "partial": "部分成功",
        "limited": "受限",
        "failed": "失败",
        "interrupted": "已中断",
    }.get(status, status)
    new_count = int(run.get("new_count") or 0)
    updated_count = int(run.get("updated_count") or 0)
    duplicate_count = int(run.get("duplicate_count") or 0)
    if new_count or updated_count or duplicate_count:
        return (
            f"{label}：{new_count} 条新发现，{updated_count} 条更新，"
            f"合并 {duplicate_count} 条重复，失败/受限 {failures} 条。"
        )
    return f"{label}：抓到 {run.get('scanned_count') or 0} 条，保存/更新 {run.get('saved_count') or 0} 条，失败/受限 {failures} 条。"


def scan_overview(scan_payload: dict) -> dict:
    run = scan_payload.get("run") or {}
    status = run.get("status") or "pending"
    run_sources = run.get("sources") or []
    source_rows = {
        item.get("source"): {
            "source": item.get("source"),
            "mode": item.get("mode") or scan_source_mode(item.get("source") or ""),
            "status": "pending",
            "scanned_count": 0,
            "saved_count": 0,
            "new_count": 0,
            "updated_count": 0,
            "duplicate_count": 0,
            "failure_count": 0,
        }
        for item in (scan_payload.get("expected_source_details") or [])
        if item.get("source")
    }
    for source in run_sources:
        name = SCAN_SOURCE_NAME_ALIASES.get(source.get("source"), source.get("source"))
        if not name or (name in RETIRED_AUTO_SCAN_SOURCES and name not in source_rows):
            continue
        source_rows[name] = {
            "source": name,
            "mode": source.get("mode") or scan_source_mode(name),
            "status": source.get("status") or "pending",
            "scanned_count": source.get("scanned_count") or 0,
            "saved_count": source.get("saved_count") or 0,
            "new_count": source.get("new_count") or 0,
            "updated_count": source.get("updated_count") or 0,
            "duplicate_count": source.get("duplicate_count") or 0,
            "failure_count": source.get("failure_count") or 0,
        }
    sources = list(source_rows.values())
    limited_sources = [
        source.get("source")
        for source in sources
        if source.get("status") in {"failed", "limited", "partial"} or number_like(source.get("failure_count")) > 0
    ]
    latest_success = scan_payload.get("latest_successful_run") or {}
    return {
        "status": status,
        "summary": scan_run_summary_text(run),
        "running": bool(scan_payload.get("running")),
        "source_count": len(sources),
        "limited_count": len(limited_sources),
        "failure_count": len(run.get("failures_json") or []),
        "new_count": run.get("new_count") or 0,
        "updated_count": run.get("updated_count") or 0,
        "duplicate_count": run.get("duplicate_count") or 0,
        "last_success": latest_success.get("finished_at") or latest_success.get("started_at"),
        "sources": sources,
    }


def workbench_actions(
    summary_payload: dict,
    recommendations: list[dict],
    queue_jobs: list[dict],
    followups: list[dict],
    scan_payload: dict,
    stale_applications: list[dict] | None = None,
) -> list[dict]:
    actions: list[dict] = []
    context = load_user_context()
    if not context.get("resume_analyzed"):
        actions.append({
            "kind": "resume",
            "title": "先分析简历",
            "body": "让系统按你的经历预选方向，后续推荐会更稳。",
            "view": "fit",
            "priority": 95,
        })
    if recommendations:
        top = recommendations[:3]
        actions.append({
            "kind": "recommendations",
            "title": f"先看 {len(recommendations)} 个每日推荐岗位",
            "body": "这些岗位已避开关注公司刷屏，并按标签、留新路径和分数排序。",
            "view": "today",
            "priority": 90,
            "job_ids": [job.get("id") for job in top if job.get("id")],
        })
    if queue_jobs:
        priority_counts = {"today": 0, "next": 0, "review": 0}
        for job in queue_jobs:
            priority = str(job.get("queue_priority") or queue_decision(job)["priority"])
            priority_counts[priority] = priority_counts.get(priority, 0) + 1
        title = (
            f"今天先投 {priority_counts['today']} 个岗位"
            if priority_counts["today"]
            else f"投递队列里有 {len(queue_jobs)} 个岗位"
        )
        actions.append({
            "kind": "queue",
            "title": title,
            "body": f"共 {len(queue_jobs)} 个待投；{priority_counts['review']} 个建议先确认方向或有效性。",
            "view": "queue",
            "priority": 82,
        })
    if followups:
        actions.append({
            "kind": "followup",
            "title": f"{len(followups)} 个岗位需要跟进",
            "body": "已投递超过 3 天的岗位可以做一次轻量 follow-up。",
            "view": "tracker",
            "priority": 78,
        })
    if stale_applications:
        actions.append({
            "kind": "stale",
            "title": f"{len(stale_applications)} 个长期无回复岗位待整理",
            "body": "做最后一次确认，仍无进展的可以暂停，保留记录但不占每日注意力。",
            "view": "tracker",
            "priority": 72,
        })
    run = scan_payload.get("run") or {}
    if not run:
        actions.append({
            "kind": "scan",
            "title": "今天还没有扫描",
            "body": "需要新机会时再启动扫描，受限来源不会阻塞主流程。",
            "view": "today",
            "priority": 64,
        })
    elif run.get("status") in {"failed", "limited", "partial"}:
        actions.append({
            "kind": "scan_limited",
            "title": "扫描部分受限",
            "body": "主流程仍可用；需要时展开扫描详情查看来源。",
            "view": "today",
            "priority": 58,
        })
    if not actions and summary_payload.get("today_applied"):
        actions.append({
            "kind": "done",
            "title": "今天已经有投递记录",
            "body": "可以复盘已投岗位，或继续看 Top 推荐。",
            "view": "tracker",
            "priority": 40,
        })
    return sorted(actions, key=lambda item: item.get("priority", 0), reverse=True)[:5]


def is_watched_company_job(job: dict, watched_terms: set[str]) -> bool:
    company_text = f"{job.get('company') or ''} {job.get('name') or ''}"
    from_watched_source = clean_text(job.get("source") or "").lower() == "关注公司公开来源"
    has_watched_alternate_source = any(
        clean_text(item.get("source") or "").lower() == "关注公司公开来源"
        for item in (job.get("alternate_links") or [])
        if isinstance(item, dict)
    ) or any(
        clean_text(source).lower() == "关注公司公开来源"
        for source in (job.get("alternate_sources") or [])
    )
    return from_watched_source or has_watched_alternate_source or (
        bool(watched_terms)
        and any(company_text_has_term(company_text, term) for term in watched_terms)
    )


def diversified_workbench_recommendations(jobs: list[dict], watched_terms: set[str], limit: int = 20) -> list[dict]:
    selected: list[dict] = []
    overflow: list[dict] = []
    company_counts: dict[str, int] = {}
    source_counts: dict[str, int] = {}
    source_limit = max(3, int(limit * 0.3 + 0.999))
    for job in jobs:
        company_key = normalize_company_phrase(job.get("company") or "")
        source_key = clean_text(job.get("source") or "Unknown").lower()
        is_watched = is_watched_company_job(job, watched_terms)
        if is_watched:
            continue
        if company_counts.get(company_key, 0) >= 2 or source_counts.get(source_key, 0) >= source_limit:
            overflow.append(job)
            continue
        selected.append(job)
        company_counts[company_key] = company_counts.get(company_key, 0) + 1
        source_counts[source_key] = source_counts.get(source_key, 0) + 1
        if len(selected) >= limit:
            return selected
    for job in overflow:
        if len(selected) >= limit:
            break
        company_key = normalize_company_phrase(job.get("company") or "")
        if company_counts.get(company_key, 0) >= 2:
            continue
        selected.append(job)
        company_counts[company_key] = company_counts.get(company_key, 0) + 1
    return selected


def job_discovery_age_days(job: dict) -> int | None:
    for key in ["found_date", "batch_date", "recommended_date"]:
        age = days_since_date(job.get(key))
        if age is not None:
            return age
    return None


def workbench_recommendation_bucket(
    jobs: list[dict],
    watched_terms: set[str],
    limit: int = 20,
    *,
    max_age_days: int | None = None,
    today_only: bool = False,
    exclude_ids: set[int] | None = None,
    exclude_dedupe_keys: set[str] | None = None,
    exclude_statuses: set[str] | None = None,
    exclude_company_terms: set[str] | None = None,
) -> list[dict]:
    exclude_ids = exclude_ids or set()
    exclude_dedupe_keys = exclude_dedupe_keys or set()
    exclude_statuses = exclude_statuses or set()
    exclude_company_terms = exclude_company_terms or set()
    candidates: list[dict] = []
    for job in jobs:
        if job.get("id") in exclude_ids or job_dedupe_key(job) in exclude_dedupe_keys or job.get("status") in exclude_statuses:
            continue
        if is_watched_company_job(job, exclude_company_terms):
            continue
        age = job_discovery_age_days(job)
        if today_only and age != 0:
            continue
        if max_age_days is not None and (age is None or age > max_age_days):
            continue
        candidates.append(job)
    if max_age_days is not None and not today_only:
        candidates.sort(
            key=lambda item: (
                not bool(item.get("user_tag_mutes")),
                float(item.get("rank_score") or item.get("score") or 0),
                float(item.get("user_tag_priority") or 0),
                float(item.get("pathway_score") or 0),
                float(item.get("base_score") or item.get("score") or 0),
                item.get("updated_at") or "",
            ),
            reverse=True,
        )
    return diversified_workbench_recommendations(candidates, set(), limit)


def workbench_payload(params: dict[str, list[str]] | None = None) -> dict:
    params = params or {}
    region = active_region_code((params.get("region") or [""])[0] or None)
    city = (params.get("city") or [active_region_context(region).get("city") or ""])[0]
    context = active_region_context(region)
    direction_ids, direction_source = active_preference_direction_ids()
    preferences = get_career_preferences()
    jobs = apply_preference_scores_to_jobs(list_jobs({"region": [region], "city": [city]}), region)
    recommendations = recommendation_payload_from_ranked_jobs(
        jobs,
        region,
        120,
        context,
        direction_ids,
        direction_source,
        preferences["exclude_keywords"],
    )
    summary_payload = summary(region, include_ai=False)
    watched_terms = watched_company_keys(region)
    queue_jobs = [job for job in jobs if job.get("status") == "Apply Queue"]
    today_date = today()
    today_applied = [job for job in jobs if job.get("status") == "Applied" and job.get("applied_date") == today_date]
    followups = [job for job in jobs if application_action_bucket(job) == "followup"]
    stale_applications = [job for job in jobs if application_action_bucket(job) == "stale"]
    scan_payload = scan_status_payload(region=region)
    recommendation_jobs = recommendations.get("jobs", [])
    watched_company_jobs = diversified_workbench_recommendations(
        [job for job in recommendation_jobs if is_watched_company_job(job, watched_terms)],
        set(),
        6,
    )
    today_new_recommendations = workbench_recommendation_bucket(
        recommendation_jobs,
        watched_terms,
        20,
        today_only=True,
        exclude_company_terms=watched_terms,
    )
    today_new_ids = {job.get("id") for job in today_new_recommendations if job.get("id")}
    today_new_dedupe_keys = {job_dedupe_key(job) for job in today_new_recommendations}
    weekly_unqueued_recommendations = workbench_recommendation_bucket(
        recommendation_jobs,
        watched_terms,
        20,
        max_age_days=6,
        exclude_ids=today_new_ids,
        exclude_dedupe_keys=today_new_dedupe_keys,
        exclude_statuses={"Apply Queue"},
        exclude_company_terms=watched_terms,
    )
    today_discovered_jobs = collapse_duplicate_job_groups(
        [job for job in jobs if job_discovery_age_days(job) == 0]
    )
    today_actionable_jobs = [
        job for job in recommendation_jobs
        if job_discovery_age_days(job) == 0 and not is_watched_company_job(job, watched_terms)
    ]
    discovery_summary = {
        "today_discovered": len(today_discovered_jobs),
        "today_actionable": len(today_actionable_jobs),
        "today_shown": len(today_new_recommendations),
    }
    top_recommendations = today_new_recommendations or weekly_unqueued_recommendations
    queue_preview = sorted(queue_jobs, key=queue_job_sort_key)[:5]
    followup_preview = sorted(
        followups,
        key=lambda job: (
            int(job.get("followup_count") or 0),
            job.get("last_followup_at") or job.get("applied_date") or job.get("updated_at") or "",
        ),
    )[:5]
    recommendation_sections = [
        {
            "id": "today_new",
            "label": "今天新发现",
            "count": len(today_new_recommendations),
            "limit": 20,
        },
        {
            "id": "weekly_unqueued",
            "label": "近一周未投",
            "count": len(weekly_unqueued_recommendations),
            "limit": 20,
        },
    ]
    compact_recommendations = {
        **recommendations,
        "jobs": [],
    }
    return {
        "date": today_date,
        "region": region,
        "city": city,
        "summary": summary_payload,
        "today_actions": workbench_actions(summary_payload, top_recommendations, queue_jobs, followups, scan_payload, stale_applications),
        "top_recommendations": workbench_job_payloads(top_recommendations),
        "today_new_recommendations": workbench_job_payloads(today_new_recommendations),
        "weekly_unqueued_recommendations": workbench_job_payloads(weekly_unqueued_recommendations),
        "discovery_summary": discovery_summary,
        "recommendation_sections": recommendation_sections,
        "watched_company_jobs": workbench_job_payloads(watched_company_jobs),
        "queue_preview": [workbench_job_payload({**job, "next_step": next_step_for_job(job)}) for job in queue_preview],
        "today_applied": [workbench_job_payload({**job, "next_step": next_step_for_job(job)}) for job in today_applied[:5]],
        "followups": [workbench_job_payload({**job, "next_step": next_step_for_job(job)}) for job in followup_preview],
        "followup_count": len(followups),
        "stale_application_count": len(stale_applications),
        "scan_overview": scan_overview(scan_payload),
        "active_context": active_region_context(region),
        "recommendations": compact_recommendations,
    }


def run_daily_scan(force: bool = False, triggered_by: str = "auto_open", async_mode: bool = False, region: str | None = None) -> dict:
    code = active_region_code(region)
    current_date = today()
    existing = latest_successful_scan(current_date, code)
    if existing and not force:
        return {
            "date": current_date,
            "region": code,
            "skipped": True,
            "reason": "A successful or partial scan already exists for today.",
            "scan_run": existing,
            "status": daily_status(code),
        }
    if async_mode:
        started = start_scan_async(triggered_by=triggered_by, forced=force, region=code)
        return {
            "date": current_date,
            "region": code,
            "skipped": False,
            "async": True,
            **started,
            "status": daily_status(code),
        }
    result = scan_sources(triggered_by=triggered_by, forced=force, region=code)
    return {
        "date": current_date,
        "region": code,
        "skipped": False,
        "result": result,
        "status": daily_status(code),
    }


def summary(region: str | None = None, include_ai: bool = True) -> dict:
    current_date = today()
    region = active_region_code(region)
    with get_db() as conn:
        rows = conn.execute(
            """
            select
                count(*) as total,
                sum(case when status='Recommended' then 1 else 0 end) as recommended,
                sum(case when status='Apply Queue' then 1 else 0 end) as apply_queue,
                sum(case when status='Applied' then 1 else 0 end) as applied,
                sum(case when status='Watch' then 1 else 0 end) as watch,
                sum(case when status='Dropped' then 1 else 0 end) as dropped,
                sum(case when score >= 3.0 then 1 else 0 end) as above_threshold
            from jobs
            where region=?
            """
            ,
            (region,),
        ).fetchone()
        today_rows = conn.execute(
            """
            select
                sum(case when batch_date=? then 1 else 0 end) as today_recommended,
                sum(case when applied_date=? then 1 else 0 end) as today_applied,
                sum(case when status='Apply Queue' and batch_date=? then 1 else 0 end) as today_queue
            from jobs
            where region=?
            """,
            (current_date, current_date, current_date, region),
        ).fetchone()
    if include_ai:
        ai_recommended = len(list_ai_jobs({"limit": ["20"], "region": [region]}))
    else:
        latest_run = latest_scan_run(current_date, region) or {}
        ai_recommended = latest_run.get("ai_recommended_count") or 0
    return {
        "date": current_date,
        "region": region,
        "region_label": REGION_CONFIGS[region]["label"],
        "resume_path": str(current_resume_path()),
        "total": rows["total"] or 0,
        "recommended": rows["recommended"] or 0,
        "apply_queue": rows["apply_queue"] or 0,
        "applied": rows["applied"] or 0,
        "watch": rows["watch"] or 0,
        "dropped": rows["dropped"] or 0,
        "above_threshold": rows["above_threshold"] or 0,
        "today_recommended": today_rows["today_recommended"] or 0,
        "today_applied": today_rows["today_applied"] or 0,
        "today_queue": today_rows["today_queue"] or 0,
        "daily_target": 15,
        "recommendation_target": 20,
        "ai_recommendation_target": 20,
        "ai_recommended": ai_recommended,
    }


def set_decision(job_id: int, decision: str, notes: str = "") -> dict:
    decision_map = {
        "Apply": "Apply Queue",
        "Watch": "Watch",
        "Drop": "Dropped",
        "Pause": "Closed",
        "FollowUpSent": "Applied",
        "Restore": "Recommended",
    }
    if decision not in decision_map:
        raise ValueError("Unsupported job decision.")

    job = get_job(job_id)
    status = decision_map[decision]
    stamp = now_iso()
    applied_date = None
    resume_path = job.get("resume_path")
    cover_path = job.get("cover_letter_path")
    if decision == "Restore":
        status = "Recommended" if float(job.get("score") or 0) >= 3.0 or is_pathway_recommendation_candidate(job) else "New"
    if decision == "FollowUpSent":
        with get_db() as conn:
            conn.execute(
                """
                update jobs set
                    status='Applied',
                    last_followup_at=?,
                    followup_count=coalesce(followup_count, 0) + 1,
                    updated_at=?
                where id=?
                """,
                (today(), stamp, job_id),
            )
        return get_job(job_id)
    if decision == "Apply":
        hard_flags = {"citizen_or_pr_only", "local_only", "clearance_required"}
        blocked = hard_flags.intersection(set(job.get("eligibility_flags") or []))
        if blocked:
            raise ValueError("Hard local eligibility flag detected: " + ", ".join(sorted(blocked)))
        if materials_need_refresh(job):
            resume_path, cover_path = make_drafts(job)
        with get_db() as conn:
            exists = conn.execute("select id from applications where job_id = ?", (job_id,)).fetchone()
            if not exists:
                conn.execute(
                    """
                    insert into applications(job_id, status, resume_path, cover_letter_path, notes, created_at, updated_at)
                    values(?, 'Drafted', ?, ?, ?, ?, ?)
                    """,
                    (job_id, resume_path, cover_path, notes, stamp, stamp),
                )
            else:
                conn.execute(
                    """
                    update applications set
                        resume_path=coalesce(?, resume_path),
                        cover_letter_path=coalesce(?, cover_letter_path),
                        updated_at=?
                    where job_id=?
                    """,
                    (resume_path, cover_path, stamp, job_id),
                )
    with get_db() as conn:
        conn.execute(
            """
            update jobs set
                status=?,
                decision=?,
                applied_date=coalesce(?, applied_date),
                resume_path=coalesce(?, resume_path),
                cover_letter_path=coalesce(?, cover_letter_path),
                updated_at=?
            where id=?
            """,
            (status, decision, applied_date, resume_path, cover_path, stamp, job_id),
        )
    return get_job(job_id)


def confirm_applied(job_id: int) -> dict:
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            update jobs set status='Applied', decision='Apply', applied_date=?, updated_at=?
            where id=?
            """,
            (today(), stamp, job_id),
        )
        conn.execute(
            """
            update applications set status='Submitted', submitted_at=?, updated_at=?
            where job_id=?
            """,
            (stamp, stamp, job_id),
        )
    return get_job(job_id)


def watchlist(region: str | None = None) -> list[dict]:
    code = active_region_code(region)
    city = active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]
    with get_db() as conn:
        rows = conn.execute(
            """
            select * from watch_companies
            where region=? and status='Watch'
            order by priority desc, company
            """,
            (code,),
        ).fetchall()
    return enrich_company_items([row_to_dict(row) for row in rows], code, city)


def add_watch_company(payload: dict) -> dict:
    code = active_region_code(payload.get("region"))
    company = (payload.get("company") or "").strip()
    if not company:
        raise ValueError("Company name is required.")
    url = validate_http_url(payload.get("url") or "")
    focus = (payload.get("focus") or "Company career page").strip()
    source = (payload.get("source") or "Company Site").strip()
    city_tags = payload.get("city_tags") or payload.get("city_tags_json") or []
    if isinstance(city_tags, str):
        city_tags = [item.strip() for item in re.split(r"[,，\n]+", city_tags) if item.strip()]
    if not city_tags:
        city_tags = [active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]]
    aliases = json_list(payload.get("aliases") or payload.get("aliases_json"))
    aliases = company_alias_values(company, {"aliases": aliases})
    company_type = (payload.get("company_type") or "Company").strip()
    priority = int(payload.get("priority") or 70)
    notes = (payload.get("notes") or "").strip()
    user_added = 1 if payload.get("user_added", True) else 0
    stamp = now_iso()
    with get_db() as conn:
        duplicate_url = conn.execute(
            "select * from watch_companies where region=? and lower(url)=lower(?) and lower(company)<>lower(?)",
            (code, url, company),
        ).fetchone()
        if duplicate_url:
            raise ValueError(f"This career URL is already tracked for {duplicate_url['company']}.")
        conn.execute(
            """
            insert into watch_companies(
                company, source, url, focus, region, city_tags_json, company_type,
                aliases_json, user_added, priority, notes, last_checked_at, status
            )
            values(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, null, 'Watch')
            on conflict(region, company) do update set
                source=excluded.source,
                url=excluded.url,
                focus=excluded.focus,
                city_tags_json=excluded.city_tags_json,
                company_type=excluded.company_type,
                aliases_json=excluded.aliases_json,
                user_added=max(watch_companies.user_added, excluded.user_added),
                priority=excluded.priority,
                notes=excluded.notes,
                status='Watch'
            """,
            (
                company,
                source,
                url,
                focus,
                code,
                json.dumps(city_tags, ensure_ascii=False),
                company_type,
                json.dumps(aliases, ensure_ascii=False),
                user_added,
                priority,
                notes,
            ),
        )
        row = conn.execute(
            "select * from watch_companies where region=? and company=?",
            (code, company),
        ).fetchone()
    return row_to_dict(row)


def dismiss_watch_company(payload: dict) -> dict:
    code = active_region_code(payload.get("region"))
    company = (payload.get("company") or "").strip()
    if not company:
        raise ValueError("Company name is required.")
    catalog_item = company_catalog_item(company, code) or {}
    url = validate_http_url(payload.get("url") or catalog_item.get("url") or catalog_item.get("official_careers_url") or "")
    focus = (payload.get("focus") or catalog_item.get("focus") or "Temporarily hidden company").strip()
    source = (payload.get("source") or catalog_item.get("source") or "Company Site").strip()
    city_tags = payload.get("city_tags") or catalog_item.get("city_tags") or [active_region_context(code).get("city") or REGION_CONFIGS[code]["default_city"]]
    aliases = json_list(payload.get("aliases") or payload.get("aliases_json") or catalog_item.get("aliases") or [])
    aliases = company_alias_values(company, {"aliases": aliases})
    company_type = (payload.get("company_type") or catalog_item.get("company_type") or "Company").strip()
    priority = int(payload.get("priority") or catalog_item.get("priority") or 50)
    stamp = now_iso()
    with get_db() as conn:
        conn.execute(
            """
            insert into watch_companies(
                company, source, url, focus, region, city_tags_json, company_type,
                aliases_json, user_added, priority, notes, last_checked_at, status
            )
            values(?, ?, ?, ?, ?, ?, ?, ?, 0, ?, '', ?, 'Dropped')
            on conflict(region, company) do update set
                source=excluded.source,
                url=excluded.url,
                focus=excluded.focus,
                city_tags_json=excluded.city_tags_json,
                company_type=excluded.company_type,
                aliases_json=excluded.aliases_json,
                priority=excluded.priority,
                last_checked_at=excluded.last_checked_at,
                status='Dropped'
            """,
            (
                company,
                source,
                url,
                focus,
                code,
                json.dumps(city_tags, ensure_ascii=False),
                company_type,
                json.dumps(aliases, ensure_ascii=False),
                priority,
                stamp,
            ),
        )
        row = conn.execute(
            "select * from watch_companies where region=? and company=?",
            (code, company),
        ).fetchone()
    return row_to_dict(row)


def update_watch_company(company_id: int, payload: dict) -> dict:
    with get_db() as conn:
        existing = conn.execute("select * from watch_companies where id=?", (company_id,)).fetchone()
        if not existing:
            raise KeyError(f"Watch company {company_id} not found.")
        current = row_to_dict(existing)
        company = (payload.get("company") or current["company"]).strip()
        url = validate_http_url(payload.get("url") or current["url"])
        city_tags = payload.get("city_tags") or current.get("city_tags_json") or []
        if isinstance(city_tags, str):
            city_tags = [item.strip() for item in re.split(r"[,，\n]+", city_tags) if item.strip()]
        aliases = json_list(payload.get("aliases") or payload.get("aliases_json") or current.get("aliases_json"))
        aliases = company_alias_values(company, {"aliases": aliases})
        conn.execute(
            """
            update watch_companies set
                company=?,
                source=?,
                url=?,
                focus=?,
                city_tags_json=?,
                aliases_json=?,
                company_type=?,
                priority=?,
                notes=?,
                status=coalesce(?, status)
            where id=?
            """,
            (
                company,
                payload.get("source") or current["source"],
                url,
                payload.get("focus") or current["focus"],
                json.dumps(city_tags, ensure_ascii=False),
                json.dumps(aliases, ensure_ascii=False),
                payload.get("company_type") or current.get("company_type") or "Company",
                int(payload.get("priority") or current.get("priority") or 70),
                payload.get("notes") if payload.get("notes") is not None else current.get("notes") or "",
                payload.get("status"),
                company_id,
            ),
        )
        row = conn.execute("select * from watch_companies where id=?", (company_id,)).fetchone()
    return row_to_dict(row)


def delete_watch_company(company_id: int) -> dict:
    with get_db() as conn:
        existing = conn.execute("select * from watch_companies where id=?", (company_id,)).fetchone()
        if not existing:
            raise KeyError(f"Watch company {company_id} not found.")
        if existing["user_added"]:
            conn.execute("delete from watch_companies where id=?", (company_id,))
        else:
            conn.execute("update watch_companies set status='Dropped' where id=?", (company_id,))
    return {"ok": True, "id": company_id}


def generate_report(region: str | None = None) -> dict:
    current_date = today()
    region = active_region_code(region)
    jobs = list_jobs({"date": [current_date], "region": [region]})
    scan_run = latest_scan_run(current_date, region)
    scan_failures = (scan_run or {}).get("failures_json") or []
    counts = {
        "searched_count": len(jobs),
        "recommended_count": sum(1 for job in jobs if job["status"] == "Recommended"),
        "drafted_count": sum(1 for job in jobs if job.get("resume_path")),
        "apply_queue_count": sum(1 for job in jobs if job["status"] == "Apply Queue"),
        "applied_count": sum(1 for job in jobs if job["status"] == "Applied" or job.get("applied_date") == current_date),
        "watch_count": sum(1 for job in jobs if job["status"] == "Watch"),
        "drop_count": sum(1 for job in jobs if job["status"] == "Dropped"),
    }
    report_path = current_workspace_dir() / "reports" / f"{current_date}-{region}.md"
    lines = [
        f"# Career Copilot Daily Report - {REGION_CONFIGS[region]['label']} - {current_date}",
        "",
        "## Summary",
        f"- Searched/imported: {counts['searched_count']}",
        f"- Recommended: {counts['recommended_count']}",
        f"- Drafted: {counts['drafted_count']}",
        f"- Apply queue: {counts['apply_queue_count']}",
        f"- Applied: {counts['applied_count']}",
        f"- Watch: {counts['watch_count']}",
        f"- Dropped: {counts['drop_count']}",
    ]
    if scan_run:
        lines.extend(
            [
                f"- Latest scan status: {scan_run.get('status')} at {scan_run.get('finished_at') or scan_run.get('started_at')}",
                f"- Scan failures: {len(scan_failures)}",
            ]
        )
    if scan_failures:
        lines.extend(["", "## Scan Failures"])
        for failure in scan_failures[:20]:
            if isinstance(failure, dict):
                lines.append(f"- {failure.get('source', '-')}: {failure.get('error', failure)}")
            else:
                lines.append(f"- {failure}")
    lines.extend(["", "## Jobs"])
    for job in jobs:
        lines.extend(
            [
                f"### {job['company']} - {job['position']}",
                f"- Status: {job['status']}",
                f"- Score: {job['score']}/5.0",
                f"- URL: {job['url']}",
                f"- Date: found {job['found_date']}, batch {job.get('batch_date') or '-'}, applied {job.get('applied_date') or '-'}",
                f"- Flags: {', '.join(job.get('eligibility_flags') or []) or 'none'}",
                "",
            ]
        )
    report_path.write_text("\n".join(lines), encoding="utf-8")

    with get_db() as conn:
        conn.execute(
            """
            insert into daily_reports(
                date, searched_count, recommended_count, drafted_count, apply_queue_count,
                applied_count, watch_count, drop_count, failures_json, report_markdown_path, updated_at
            )
            values(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            on conflict(date) do update set
                searched_count=excluded.searched_count,
                recommended_count=excluded.recommended_count,
                drafted_count=excluded.drafted_count,
                apply_queue_count=excluded.apply_queue_count,
                applied_count=excluded.applied_count,
                watch_count=excluded.watch_count,
                drop_count=excluded.drop_count,
                failures_json=excluded.failures_json,
                report_markdown_path=excluded.report_markdown_path,
                updated_at=excluded.updated_at
            """,
            (
                current_date,
                counts["searched_count"],
                counts["recommended_count"],
                counts["drafted_count"],
                counts["apply_queue_count"],
                counts["applied_count"],
                counts["watch_count"],
                counts["drop_count"],
                json.dumps(scan_failures, ensure_ascii=False),
                str(report_path),
                now_iso(),
            ),
        )
    return {"date": current_date, "path": str(report_path), "counts": counts, "markdown": report_path.read_text(encoding="utf-8")}


def notion_schema() -> dict:
    return {
        "required": [
            {"name": "Name / Company", "type": "title", "purpose": "已加入投递的公司 + 岗位名称。"},
            {"name": "Status / Stage", "type": "status", "purpose": "To apply, Applied, Offer, Rejected 等投递阶段。"},
            {"name": "URL / Link", "type": "url", "purpose": "原始岗位页面链接，永远保留。"},
            {"name": "position", "type": "rich_text", "purpose": "原始岗位标题。"},
            {"name": "JD", "type": "rich_text", "purpose": "岗位描述摘要或完整 JD 前段。"},
        ],
        "recommended": [
            "Company",
            "Source",
            "Score",
            "Timeline Date",
            "Batch Date",
            "Found Date",
            "Recommended Date",
            "Applied Date",
            "Decision",
            "Eligibility Flags",
            "Resume Path",
            "Cover Letter Path",
            "Notes",
            "Drop Reason",
            "Last Checked",
        ],
    }


def load_notion_config() -> dict:
    ensure_dirs()
    path = current_notion_config_path()
    config = {"token": "", "database_id": "", "updated_at": ""}
    if path.exists():
        try:
            stored = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(stored, dict):
                config.update({
                    "token": str(stored.get("token") or ""),
                    "database_id": str(stored.get("database_id") or ""),
                    "updated_at": str(stored.get("updated_at") or ""),
                })
        except json.JSONDecodeError:
            pass
    if not auth_required():
        config["token"] = config["token"] or os.environ.get("NOTION_TOKEN", "")
        config["database_id"] = config["database_id"] or os.environ.get("NOTION_DATABASE_ID", "")
    return config


def save_notion_config(payload: dict) -> dict:
    ensure_dirs()
    current = load_notion_config()
    if "token" in payload:
        current["token"] = str(payload.get("token") or "").strip()
    if "database_id" in payload:
        current["database_id"] = str(payload.get("database_id") or "").strip()
    current["updated_at"] = now_iso()
    current_notion_config_path().write_text(json.dumps(current, ensure_ascii=False, indent=2), encoding="utf-8")
    return notion_status()


def notion_status() -> dict:
    config = load_notion_config()
    user_config_exists = current_notion_config_path().exists()
    return {
        "token_configured": bool(config.get("token")),
        "database_id_configured": bool(config.get("database_id")),
        "source": "user" if user_config_exists else ("env" if not auth_required() and (os.environ.get("NOTION_TOKEN") or os.environ.get("NOTION_DATABASE_ID")) else "none"),
        "env_file": ".env.local" if not auth_required() and (APP_DIR / ".env.local").exists() else "",
        "updated_at": config.get("updated_at") or "",
    }


def notion_request(method: str, path: str, payload: dict | None = None) -> dict:
    token = load_notion_config().get("token")
    if not token:
        raise ValueError("还没有配置你的 Notion token。可以先只保存在 Job Assistant，或在 Notion 页填入自己的配置。")
    body = None if payload is None else json.dumps(payload).encode("utf-8")
    request = urllib.request.Request(
        f"https://api.notion.com/v1{path}",
        data=body,
        method=method,
        headers={
            "Authorization": f"Bearer {token}",
            "Notion-Version": NOTION_VERSION,
            "Content-Type": "application/json",
        },
    )
    try:
        with urllib.request.urlopen(request, timeout=30) as response:
            raw = response.read().decode("utf-8")
            return json.loads(raw or "{}")
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        raise ValueError(f"Notion API error {exc.code}: {detail}") from exc


def rich_text(value: str, limit: int = 1900) -> dict:
    return {"rich_text": [{"type": "text", "text": {"content": (value or "")[:limit]}}]}


def date_property(value: str | None) -> dict:
    if not value:
        return {"date": None}
    return {"date": {"start": value}}


def timeline_date(job: dict) -> str | None:
    if job.get("status") in NOTION_APPLICATION_STATUSES:
        return (
            job.get("applied_date")
            or (job.get("updated_at") or "")[:10]
            or job.get("batch_date")
            or job.get("found_date")
        )
    return (
        job.get("applied_date")
        or job.get("batch_date")
        or job.get("recommended_date")
        or job.get("found_date")
    )


def first_notion_property(database_properties: dict, notion_type: str, preferred: list[str]) -> str | None:
    for name in preferred:
        if database_properties.get(name, {}).get("type") == notion_type:
            return name
    for name, value in database_properties.items():
        if value.get("type") == notion_type:
            return name
    return None


def notion_property_map(database_properties: dict) -> dict:
    return {
        "title": first_notion_property(database_properties, "title", ["Name", "Company"]),
        "status": first_notion_property(database_properties, "status", ["Status", "Stage"]),
        "url": first_notion_property(database_properties, "url", ["URL", "Link"]),
        "position": first_notion_property(database_properties, "rich_text", ["position", "Position"]),
        "jd": "JD" if database_properties.get("JD", {}).get("type") == "rich_text" else None,
    }


def notion_stage_name(status: str, status_property: dict | None) -> str:
    options = []
    if status_property and status_property.get("type") == "status":
        options = [option.get("name") for option in status_property.get("status", {}).get("options", [])]
    if status in options:
        return status
    stage_fallback = {
        "Apply Queue": "To apply",
        "Drafted": "To apply",
        "Applied": "Applied",
        "Follow Up": "Applied",
        "Interview": "Applied",
        "Rejected": "Rejected",
        "Offer": "Offer",
        "Dropped": "No Answer",
        "Closed": "No Answer",
    }
    fallback = stage_fallback.get(status, status)
    if fallback in options:
        return fallback
    return options[0] if options else fallback


def notion_exact_properties(job: dict) -> dict:
    flags = ", ".join(job.get("eligibility_flags") or [])
    return {
        "Name": {"title": [{"type": "text", "text": {"content": job["name"][:200]}}]},
        "Status": {"status": {"name": job["status"]}},
        "URL": {"url": job["url"]},
        "position": rich_text(job["position"]),
        "JD": rich_text(job.get("jd_text") or ""),
        "Company": rich_text(job["company"]),
        "Source": {"select": {"name": job["source"]}},
        "Score": {"number": float(job["score"])},
        "Timeline Date": date_property(timeline_date(job)),
        "Batch Date": date_property(job.get("batch_date")),
        "Found Date": date_property(job.get("found_date")),
        "Recommended Date": date_property(job.get("recommended_date")),
        "Applied Date": date_property(job.get("applied_date")),
        "Decision": {"select": {"name": job.get("decision") or "None"}},
        "Eligibility Flags": {"multi_select": [{"name": flag} for flag in job.get("eligibility_flags") or []]},
        "Resume Path": rich_text(job.get("resume_path") or ""),
        "Cover Letter Path": rich_text(job.get("cover_letter_path") or ""),
        "Notes": rich_text(f"{job.get('match_notes') or ''}\nFlags: {flags}".strip()),
        "Last Checked": date_property((job.get("last_checked_at") or "")[:10] or None),
    }


def notion_database_properties(database_id: str) -> dict:
    result = notion_request("GET", f"/databases/{database_id}")
    return result.get("properties") or {}


def notion_payload_properties(job: dict, database_properties: dict) -> dict:
    aliases = notion_property_map(database_properties)
    missing = [label for label in ["title", "status", "url", "position"] if not aliases.get(label)]
    if missing:
        raise ValueError(f"Notion 数据库缺少可同步字段：{', '.join(missing)}。至少需要 title、status、url、rich_text 岗位列。")

    props = {
        aliases["title"]: {"title": [{"type": "text", "text": {"content": job["name"][:200]}}]},
        aliases["status"]: {"status": {"name": notion_stage_name(job["status"], database_properties.get(aliases["status"]))}},
        aliases["url"]: {"url": job["url"]},
        aliases["position"]: rich_text(job["position"]),
    }
    if aliases.get("jd"):
        props[aliases["jd"]] = rich_text(job.get("jd_text") or "")

    available = set(database_properties.keys())
    for name, value in notion_exact_properties(job).items():
        if name in available and name not in props:
            props[name] = value
    return props


def find_notion_page_by_url(database_id: str, url: str, url_property: str = "URL") -> str | None:
    result = notion_request(
        "POST",
        f"/databases/{database_id}/query",
        {
            "page_size": 1,
            "filter": {"property": url_property, "url": {"equals": url}},
        },
    )
    results = result.get("results") or []
    if not results:
        return None
    return results[0].get("id")


def create_notion_page(database_id: str, props: dict) -> str | None:
    result = notion_request(
        "POST",
        "/pages",
        {
            "parent": {"database_id": database_id},
            "properties": props,
        },
    )
    return result.get("id")


def is_notion_application_job(job: dict) -> bool:
    return job.get("status") in NOTION_APPLICATION_STATUSES or bool(job.get("applied_date"))


def list_notion_application_jobs() -> list[dict]:
    return [job for job in list_jobs({}) if is_notion_application_job(job)]


def sync_notion(job_id: int | None = None) -> dict:
    database_id = load_notion_config().get("database_id")
    if not database_id:
        raise ValueError("还没有配置你的 Notion database ID。可以先只保存在 Job Assistant，或在 Notion 页填入自己的配置。")

    if job_id:
        candidates = [get_job(job_id)]
    else:
        candidates = list_jobs({})
    jobs = [job for job in candidates if is_notion_application_job(job)]
    skipped = len(candidates) - len(jobs)

    database_properties = notion_database_properties(database_id)
    url_property = notion_property_map(database_properties).get("url") or "URL"
    synced = 0
    failures = []
    with get_db() as conn:
        for job in jobs:
            try:
                page_id = job.get("notion_page_id") or find_notion_page_by_url(database_id, job["url"], url_property)
                props = notion_payload_properties(job, database_properties)
                if page_id:
                    try:
                        notion_request("PATCH", f"/pages/{page_id}", {"properties": props})
                    except ValueError as exc:
                        if "archived" not in str(exc).lower():
                            raise
                        page_id = create_notion_page(database_id, props)
                else:
                    page_id = create_notion_page(database_id, props)
                if page_id:
                    conn.execute("update jobs set notion_page_id=?, updated_at=? where id=?", (page_id, now_iso(), job["id"]))
                synced += 1
            except Exception as exc:
                failures.append({"job_id": job["id"], "name": job["name"], "error": str(exc)})
    return {"synced": synced, "failures": failures, "total": len(jobs), "skipped": skipped}


def json_response(handler: SimpleHTTPRequestHandler, data: dict | list, status: HTTPStatus = HTTPStatus.OK) -> None:
    body = json.dumps(data, ensure_ascii=False).encode("utf-8")
    accepts_gzip = "gzip" in (handler.headers.get("Accept-Encoding", "").lower())
    compressed = accepts_gzip and len(body) >= 1024
    if compressed:
        body = gzip.compress(body, compresslevel=5)
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Vary", "Accept-Encoding")
    if compressed:
        handler.send_header("Content-Encoding", "gzip")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def read_json(handler: SimpleHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length", "0"))
    if length == 0:
        return {}
    raw = handler.rfile.read(length).decode("utf-8")
    return json.loads(raw or "{}")


def decode_header_param(value: str) -> str:
    try:
        return value.encode("latin-1").decode("utf-8")
    except UnicodeError:
        return value


def parse_multipart_file(handler: SimpleHTTPRequestHandler, field_name: str = "resume") -> tuple[str, bytes, str]:
    content_type = handler.headers.get("Content-Type", "")
    match = re.search(r"boundary=(?P<boundary>[^;]+)", content_type)
    if not match:
        raise ValueError("Expected multipart form upload.")
    boundary = match.group("boundary").strip().strip('"').encode("utf-8")
    length = int(handler.headers.get("Content-Length", "0"))
    raw = handler.rfile.read(length)
    for part in raw.split(b"--" + boundary):
        part = part.strip(b"\r\n")
        if not part or part == b"--":
            continue
        if part.endswith(b"--"):
            part = part[:-2].strip(b"\r\n")
        if b"\r\n\r\n" not in part:
            continue
        header_blob, body = part.split(b"\r\n\r\n", 1)
        headers = header_blob.decode("latin-1", errors="ignore")
        disposition = next((line for line in headers.split("\r\n") if line.lower().startswith("content-disposition:")), "")
        name_match = re.search(r'name="([^"]+)"', disposition)
        filename_match = re.search(r'filename="([^"]*)"', disposition)
        if not name_match or name_match.group(1) != field_name or not filename_match:
            continue
        filename = decode_header_param(filename_match.group(1))
        mime_match = re.search(r"content-type:\s*([^\r\n]+)", headers, flags=re.I)
        mime_type = mime_match.group(1).strip() if mime_match else ""
        if body.endswith(b"\r\n"):
            body = body[:-2]
        return filename, body, mime_type
    raise ValueError("No resume file was uploaded.")


def open_local_path(payload: dict) -> dict:
    raw_path = (payload.get("path") or "").strip()
    mode = payload.get("mode") or "folder"
    if not raw_path:
        raise ValueError("缺少要打开的材料路径。")

    target = Path(raw_path).expanduser().resolve()
    allowed_roots = [APP_DIR.resolve(), APP_DIR.parent.resolve()]
    with get_db() as conn:
        known_material = conn.execute(
            """
            select 1 from jobs
            where resume_path = ? or cover_letter_path = ?
            limit 1
            """,
            (str(target), str(target)),
        ).fetchone()
    if not known_material and not any(target == root or root in target.parents for root in allowed_roots):
        raise ValueError("为了安全，只能打开求职助手目录或已记录的简历材料。")
    if not target.exists():
        raise FileNotFoundError(f"找不到这个文件：{target}")

    if os.name == "nt":
        if mode == "file":
            os.startfile(str(target))  # type: ignore[attr-defined]
            opened = target
        else:
            folder = target if target.is_dir() else target.parent
            os.startfile(str(folder))  # type: ignore[attr-defined]
            opened = folder
    else:
        folder = target if target.is_dir() or mode == "folder" else target
        import subprocess

        subprocess.Popen(["xdg-open", str(folder)])
        opened = folder
    return {"opened": str(opened), "mode": mode}


def get_application(job_id: int) -> dict | None:
    with get_db() as conn:
        row = conn.execute("select * from applications where job_id=?", (job_id,)).fetchone()
        return row_to_dict(row) if row else None


def draft_custom_questions(job: dict, profile: dict, existing: list | None = None) -> list[dict]:
    existing = existing or []
    if existing:
        return existing
    role = f"{job.get('company')} - {job.get('position')}"
    return [
        {
            "question": "Why are you interested in this role?",
            "draft_answer": (
                f"I am interested in {role} because it connects with my human-centred design, "
                "service design, user research, prototyping, and AI-assisted workflow experience. "
                "I would like to contribute evidence-led design thinking while learning from the team."
            ),
            "status": "needs_confirmation",
            "source": "profile_and_job",
        },
        {
            "question": "Work authorisation / sponsorship",
            "draft_answer": profile.get("work_authorisation") or "",
            "status": "needs_confirmation",
            "source": "profile",
        },
        {
            "question": "Availability",
            "draft_answer": profile.get("availability") or "",
            "status": "needs_confirmation",
            "source": "profile",
        },
    ]


def update_application_assist(job_id: int, payload_path: Path | None, result_path: Path | None, status: str, questions: list[dict]) -> None:
    stamp = now_iso()
    with get_db() as conn:
        exists = conn.execute("select id from applications where job_id=?", (job_id,)).fetchone()
        if exists:
            conn.execute(
                """
                update applications set
                    custom_questions_json=?,
                    assist_payload_path=?,
                    assist_result_path=?,
                    assist_status=?,
                    assist_updated_at=?,
                    updated_at=?
                where job_id=?
                """,
                (
                    json.dumps(questions, ensure_ascii=False),
                    str(payload_path) if payload_path else None,
                    str(result_path) if result_path else None,
                    status,
                    stamp,
                    stamp,
                    job_id,
                ),
            )
        else:
            conn.execute(
                """
                insert into applications(
                    job_id, status, custom_questions_json, assist_payload_path,
                    assist_result_path, assist_status, assist_updated_at, notes, created_at, updated_at
                )
                values(?, 'Drafted', ?, ?, ?, ?, ?, '', ?, ?)
                """,
                (
                    job_id,
                    json.dumps(questions, ensure_ascii=False),
                    str(payload_path) if payload_path else None,
                    str(result_path) if result_path else None,
                    status,
                    stamp,
                    stamp,
                    stamp,
                ),
            )


def apply_assist(job_id: int) -> dict:
    job = get_job(job_id)
    if job.get("status") not in {"Apply Queue", "Drafted"}:
        raise ValueError("Apply assist is only available after a job is added to the apply queue.")
    if materials_need_refresh(job):
        resume_path, cover_path = make_drafts(job)
        with get_db() as conn:
            conn.execute(
                "update jobs set resume_path=?, cover_letter_path=?, updated_at=? where id=?",
                (resume_path, cover_path, now_iso(), job_id),
            )
        job = get_job(job_id)

    profile = load_profile()
    application = get_application(job_id)
    questions = draft_custom_questions(job, profile, (application or {}).get("custom_questions_json"))
    source = (job.get("source") or "").lower()
    supported = "linkedin" in source or "jobstreet" in source

    if not supported:
        webbrowser.open(job["url"])
        update_application_assist(job_id, None, None, "opened_manual", questions)
        return {
            "status": "opened_manual",
            "message": "This source is not adapted yet. The job page was opened; use the generated materials manually.",
            "job": get_job(job_id),
            "custom_questions": questions,
        }

    if importlib.util.find_spec("playwright") is None:
        update_application_assist(job_id, None, None, "playwright_missing", questions)
        raise ValueError("Playwright is not installed. Run: python -m pip install -r requirements.txt && python -m playwright install chromium")

    stamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    payload_path = current_apply_assist_dir() / f"job-{job_id}-{stamp}.json"
    result_path = current_apply_assist_dir() / f"job-{job_id}-{stamp}-result.json"
    log_path = current_apply_assist_dir() / f"job-{job_id}-{stamp}.log"
    payload = {
        "job": job,
        "profile": profile,
        "custom_questions": questions,
        "browser_profile_dir": str(current_browser_profile_dir()),
        "result_path": str(result_path),
        "review_required": True,
    }
    payload_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    update_application_assist(job_id, payload_path, result_path, "launched", questions)

    script_path = APP_DIR / "scripts" / "browser_apply_assist.py"
    stdout = log_path.open("a", encoding="utf-8")
    creationflags = subprocess.CREATE_NEW_PROCESS_GROUP if os.name == "nt" else 0
    subprocess.Popen(
        [sys.executable, str(script_path), str(payload_path)],
        cwd=str(APP_DIR),
        stdout=stdout,
        stderr=subprocess.STDOUT,
        creationflags=creationflags,
    )
    stdout.close()
    return {
        "status": "launched",
        "message": "Browser assist launched. Review every field in the visible browser before submitting.",
        "payload_path": str(payload_path),
        "result_path": str(result_path),
        "log_path": str(log_path),
        "custom_questions": questions,
        "job": get_job(job_id),
    }


class CareerHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(PUBLIC_DIR), **kwargs)

    def end_headers(self) -> None:
        self.send_header("Cache-Control", "no-store, max-age=0")
        super().end_headers()

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        params = parse_qs(parsed.query)
        previous_user_id = getattr(REQUEST_CONTEXT, "user_id", None)
        sync_after_request = False
        try:
            REQUEST_CONTEXT.user_id = user_id_for_request(self, parsed.path)
            ensure_cloud_state_loaded()
            sync_after_request = should_sync_after_request("GET", parsed.path)
            if parsed.path == "/api/health":
                json_response(self, health_payload())
            elif parsed.path == "/api/auth/config":
                json_response(self, auth_config_payload())
            elif parsed.path == "/api/summary":
                json_response(self, summary())
            elif parsed.path == "/api/workbench":
                json_response(self, workbench_payload(params))
            elif parsed.path == "/api/regions":
                json_response(self, regions_payload())
            elif parsed.path == "/api/profile-options":
                json_response(self, profile_options_payload((params.get("region") or [""])[0] or None))
            elif parsed.path == "/api/user-context":
                json_response(self, load_user_context())
            elif parsed.path == "/api/company-catalog":
                json_response(self, company_catalog((params.get("region") or [""])[0] or None, (params.get("city") or [""])[0] or None))
            elif parsed.path == "/api/company-jobs":
                company_id = int((params.get("company_id") or ["0"])[0] or 0)
                json_response(
                    self,
                    company_jobs_payload(
                        (params.get("company") or [""])[0],
                        (params.get("region") or [""])[0] or None,
                        (params.get("city") or [""])[0] or None,
                        company_id or None,
                    ),
                )
            elif parsed.path == "/api/daily/status":
                json_response(self, daily_status((params.get("region") or [""])[0] or None))
            elif parsed.path == "/api/profile":
                json_response(self, load_profile())
            elif parsed.path == "/api/career-fit":
                json_response(self, career_fit())
            elif parsed.path == "/api/recommendations/today":
                json_response(self, list_today_recommendations(params))
            elif parsed.path == "/api/scan/status":
                json_response(self, scan_status_payload(region=(params.get("region") or [""])[0] or None))
            elif re.match(r"^/api/scan-runs/\d+$", parsed.path):
                scan_run_id = int(parsed.path.rstrip("/").split("/")[-1])
                json_response(self, scan_status_payload(scan_run_id))
            elif parsed.path == "/api/jobs":
                json_response(self, list_jobs_payload(params))
            elif parsed.path == "/api/jobs/ai":
                json_response(self, list_ai_jobs(params))
            elif parsed.path.startswith("/api/jobs/"):
                translate_match = re.match(r"^/api/jobs/(\d+)/translate$", parsed.path)
                if translate_match:
                    json_response(self, job_payload(ensure_job_translation(int(translate_match.group(1)))))
                else:
                    job_id = int(parsed.path.rstrip("/").split("/")[-1])
                    json_response(self, job_payload(get_job(job_id)))
            elif parsed.path == "/api/watchlist":
                json_response(self, watchlist((params.get("region") or [""])[0] or None))
            elif parsed.path == "/api/notion-schema":
                json_response(self, notion_schema())
            elif parsed.path == "/api/notion-status":
                json_response(self, notion_status())
            elif parsed.path == "/api/notion-config":
                json_response(self, notion_status())
            elif parsed.path == "/api/report/today":
                json_response(self, generate_report())
            else:
                super().do_GET()
        except AuthError as exc:
            json_response(self, {"error": str(exc), "auth_required": True}, HTTPStatus.UNAUTHORIZED)
        except Exception as exc:
            json_response(self, {"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        finally:
            if sync_after_request:
                safe_sync_cloud_state(f"GET {parsed.path}")
            if previous_user_id is None:
                try:
                    delattr(REQUEST_CONTEXT, "user_id")
                except AttributeError:
                    pass
            else:
                REQUEST_CONTEXT.user_id = previous_user_id

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        previous_user_id = getattr(REQUEST_CONTEXT, "user_id", None)
        sync_after_request = False
        try:
            REQUEST_CONTEXT.user_id = user_id_for_request(self, parsed.path)
            ensure_cloud_state_loaded()
            sync_after_request = should_sync_after_request("POST", parsed.path)
            if parsed.path == "/api/jobs":
                json_response(self, job_payload(upsert_job(read_json(self))), HTTPStatus.CREATED)
                return

            if parsed.path == "/api/profile":
                json_response(self, save_profile(read_json(self)))
                return

            if parsed.path == "/api/user-context":
                json_response(self, save_user_context(read_json(self)))
                return

            if parsed.path == "/api/watchlist":
                json_response(self, add_watch_company(read_json(self)), HTTPStatus.CREATED)
                return

            if parsed.path == "/api/watchlist/dismiss":
                json_response(self, dismiss_watch_company(read_json(self)), HTTPStatus.CREATED)
                return

            if parsed.path == "/api/resumes":
                filename, content, mime_type = parse_multipart_file(self, "resume")
                json_response(self, save_uploaded_resume(filename, content, mime_type), HTTPStatus.CREATED)
                return

            if parsed.path == "/api/career-fit/analyze":
                payload = read_json(self)
                json_response(self, {"analysis": analyze_resume_version(payload.get("resume_version_id"), payload.get("mode") or "local"), "career_fit": career_fit()})
                return

            if parsed.path == "/api/daily/run":
                payload = read_json(self)
                json_response(
                    self,
                    run_daily_scan(
                        bool(payload.get("force")),
                        payload.get("triggered_by") or "auto_open",
                        bool(payload.get("async")),
                        payload.get("region"),
                    ),
                )
                return

            if parsed.path == "/api/scan/async":
                payload = read_json(self)
                json_response(
                    self,
                    start_scan_async(
                        payload.get("triggered_by") or "manual",
                        bool(payload.get("force", True)),
                        payload.get("region"),
                    ),
                    HTTPStatus.ACCEPTED,
                )
                return

            if parsed.path == "/api/scan":
                payload = read_json(self) if int(self.headers.get("Content-Length", "0") or 0) else {}
                json_response(self, scan_sources(region=payload.get("region")))
                return

            if parsed.path == "/api/notion-config":
                json_response(self, save_notion_config(read_json(self)))
                return

            if parsed.path == "/api/open-path":
                json_response(self, open_local_path(read_json(self)))
                return

            decision_match = re.match(r"^/api/jobs/(\d+)/decision$", parsed.path)
            if decision_match:
                payload = read_json(self)
                job = set_decision(int(decision_match.group(1)), payload.get("decision", ""), payload.get("notes", ""))
                json_response(self, job_payload(job))
                return

            applied_match = re.match(r"^/api/jobs/(\d+)/confirm-applied$", parsed.path)
            if applied_match:
                json_response(self, job_payload(confirm_applied(int(applied_match.group(1)))))
                return

            assist_match = re.match(r"^/api/jobs/(\d+)/apply-assist$", parsed.path)
            if assist_match:
                json_response(self, apply_assist(int(assist_match.group(1))))
                return

            translate_match = re.match(r"^/api/jobs/(\d+)/translate$", parsed.path)
            if translate_match:
                json_response(self, job_payload(ensure_job_translation(int(translate_match.group(1)))))
                return

            notion_sync_match = re.match(r"^/api/notion/sync(?:/(\d+))?$", parsed.path)
            if notion_sync_match:
                job_id = int(notion_sync_match.group(1)) if notion_sync_match.group(1) else None
                json_response(self, sync_notion(job_id))
                return

            json_response(self, {"error": "Not found"}, HTTPStatus.NOT_FOUND)
        except AuthError as exc:
            json_response(self, {"error": str(exc), "auth_required": True}, HTTPStatus.UNAUTHORIZED)
        except Exception as exc:
            json_response(self, {"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        finally:
            if sync_after_request:
                safe_sync_cloud_state(f"POST {parsed.path}")
            if previous_user_id is None:
                try:
                    delattr(REQUEST_CONTEXT, "user_id")
                except AttributeError:
                    pass
            else:
                REQUEST_CONTEXT.user_id = previous_user_id

    def do_PUT(self) -> None:
        parsed = urlparse(self.path)
        previous_user_id = getattr(REQUEST_CONTEXT, "user_id", None)
        sync_after_request = False
        try:
            REQUEST_CONTEXT.user_id = user_id_for_request(self, parsed.path)
            ensure_cloud_state_loaded()
            sync_after_request = should_sync_after_request("PUT", parsed.path)
            if parsed.path == "/api/career-fit/preferences":
                json_response(self, {"preferences": save_career_preferences(read_json(self)), "career_fit": career_fit()})
                return
            if parsed.path == "/api/user-context":
                json_response(self, save_user_context(read_json(self)))
                return
            watch_match = re.match(r"^/api/watchlist/(\d+)$", parsed.path)
            if watch_match:
                json_response(self, update_watch_company(int(watch_match.group(1)), read_json(self)))
                return
            json_response(self, {"error": "Not found"}, HTTPStatus.NOT_FOUND)
        except AuthError as exc:
            json_response(self, {"error": str(exc), "auth_required": True}, HTTPStatus.UNAUTHORIZED)
        except Exception as exc:
            json_response(self, {"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        finally:
            if sync_after_request:
                safe_sync_cloud_state(f"PUT {parsed.path}")
            if previous_user_id is None:
                try:
                    delattr(REQUEST_CONTEXT, "user_id")
                except AttributeError:
                    pass
            else:
                REQUEST_CONTEXT.user_id = previous_user_id

    def do_DELETE(self) -> None:
        parsed = urlparse(self.path)
        previous_user_id = getattr(REQUEST_CONTEXT, "user_id", None)
        sync_after_request = False
        try:
            REQUEST_CONTEXT.user_id = user_id_for_request(self, parsed.path)
            ensure_cloud_state_loaded()
            sync_after_request = should_sync_after_request("DELETE", parsed.path)
            watch_match = re.match(r"^/api/watchlist/(\d+)$", parsed.path)
            if watch_match:
                json_response(self, delete_watch_company(int(watch_match.group(1))))
                return
            json_response(self, {"error": "Not found"}, HTTPStatus.NOT_FOUND)
        except AuthError as exc:
            json_response(self, {"error": str(exc), "auth_required": True}, HTTPStatus.UNAUTHORIZED)
        except Exception as exc:
            json_response(self, {"error": str(exc)}, HTTPStatus.BAD_REQUEST)
        finally:
            if sync_after_request:
                safe_sync_cloud_state(f"DELETE {parsed.path}")
            if previous_user_id is None:
                try:
                    delattr(REQUEST_CONTEXT, "user_id")
                except AttributeError:
                    pass
            else:
                REQUEST_CONTEXT.user_id = previous_user_id


def main() -> None:
    load_env_files()
    setup_db()
    server = ThreadingHTTPServer((APP_HOST, APP_PORT), CareerHandler)
    display_host = "127.0.0.1" if APP_HOST in {"0.0.0.0", "::"} else APP_HOST
    print(f"Job Assistant running at http://{display_host}:{APP_PORT}")
    if display_host != APP_HOST:
        print(f"Listening on {APP_HOST}:{APP_PORT}")
    print(f"Database: {current_db_path()}")
    server.serve_forever()


if __name__ == "__main__":
    main()
